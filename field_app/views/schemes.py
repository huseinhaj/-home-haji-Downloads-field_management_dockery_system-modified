"""
Scheme of Work and Lesson Plan generation views.
"""
import json
import re
from io import BytesIO

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render
from django.utils import timezone

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from .ai_utils import client, model_name
from .forms import SchemeOfWorkForm
from .models import (
    StudentTeacher, Subject, SchemeOfWork, LessonPlan,
)


@login_required
def generate_scheme_view(request):
    from .models import EducationLevel, ClassLevel, StudentApplication

    form = SchemeOfWorkForm()
    education_levels = EducationLevel.objects.all().order_by('order')

    # Build inline JSON maps so template needs no AJAX for dropdowns
    import json as _json
    classes_by_level = {}
    for cl in ClassLevel.objects.select_related('education_level').order_by('education_level', 'order'):
        classes_by_level.setdefault(cl.education_level_id, []).append({'id': cl.id, 'name': cl.name})

    primary_ids = list(
        EducationLevel.objects.filter(name__icontains='primary').values_list('id', flat=True)
    )
    subjects_by_level = {
        lvl.id: list(
            Subject.objects.filter(level='primary' if lvl.id in primary_ids else 'secondary')
            .order_by('name').values('id', 'name')
        )
        for lvl in education_levels
    }

    student = None
    school = None
    if request.user.is_authenticated:
        try:
            student = StudentTeacher.objects.select_related('selected_school').get(user=request.user)
            school = student.selected_school
            if not school:
                app = StudentApplication.objects.filter(
                    student=student, status='approved'
                ).select_related('school').first()
                if app:
                    school = app.school
        except StudentTeacher.DoesNotExist:
            pass

    return render(request, 'field_app/generate_scheme.html', {
        'form': form,
        'education_levels': education_levels,
        'classes_by_level_json': _json.dumps(classes_by_level),
        'subjects_by_level_json': _json.dumps(subjects_by_level),
        'student': student,
        'school': school,
    })


@login_required
def ajax_generate_scheme(request):
    """API ya AI kuzalisha Scheme of Work kwa format ya KitabuSmart"""
    if client is None:
        return JsonResponse({'success': False, 'error': 'Huduma ya AI haitumiki. Ufunguo wa API (GOOGLE_API_KEY) haujawekwa. Wasiliana na msimamizi.'}, status=503)
    if request.method == 'POST':
        try:
            data = json.loads(request.body)

            education_level = data.get('education_level')
            class_name = data.get('class_name')
            subject = data.get('subject')
            term = data.get('term')
            year = data.get('year')
            syllabus = data.get('syllabus', 'New Syllabus')
            total_weeks = data.get('total_weeks', 12)
            periods_per_week = data.get('periods_per_week', 8)
            start_date = data.get('start_date')
            end_date = data.get('end_date')
            teacher_name = data.get('teacher_name')
            school_name = data.get('school_name')
            reference_source = data.get('reference_source', '')
            breaks = data.get('breaks', [])

            breaks_text = ""
            if breaks:
                breaks_text = "\nBreaks (holidays/exams) to respect:\n"
                for b in breaks:
                    breaks_text += f"- {b.get('name', 'Break')}: {b.get('start', '')} to {b.get('end', '')}\n"

            prompt = f"""
You are an AI assistant for Tanzanian teachers. Generate a complete Scheme of Work following EXACTLY the SEQUIP/TIE Tanzania revised format (2023).

Input details:
- Education Level: {education_level}
- Class/Form: {class_name}
- Subject: {subject}
- Term: {term} {year}
- Syllabus: {syllabus}
- Total Weeks: {total_weeks} weeks
- Periods per Week: {periods_per_week}
- Start Date: {start_date}
- End Date: {end_date}
- Teacher: {teacher_name}
- School: {school_name}
- Reference Source: {reference_source}
{breaks_text}
The output MUST be a JSON list of objects. Each object must have exactly these 12 keys:
"Main Competence", "Specific Competences", "Main Learning Activities", "Specific Learning Activities", "Month", "Week", "Number of Periods", "Teaching and Learning Methods", "Teaching and Learning Resources", "Assessment Tools", "References", "Remarks"

Requirements:
1. Distribute content across {total_weeks} weeks, respecting any breaks (skip weeks that fall on breaks).
2. For each week, assign appropriate Month (e.g., MAY, JUNE, JULY, AUGUST, SEPTEMBER, OCTOBER).
3. "Week" column should be like "1st", "2nd", "3rd", etc.
4. "Number of Periods" should be {periods_per_week} for normal weeks.
5. "Specific Learning Activities" are obtained by DECONSTRUCTING the Main Learning Activity into smaller measurable steps (easy to difficult).
6. "References" must follow APA style version 7: Author, Year, Title, Publisher. E.g. "Ministry of Education (2023). {subject} Syllabus Form I-IV. TIE, Dar es Salaam."
7. "Main Competence" should be numbered like "1.0 Demonstrate mastery of {subject} fundamental principles".
8. "Teaching and Learning Methods" should include CBC-aligned methods (group discussion, brainstorming, Q&A, role play, investigation, etc.).
9. "Remarks" should address: Strength of the activity, Weakness/challenges, and Way forward. Leave blank for future teacher to fill.
10. Content must be realistic for {education_level} {class_name} {subject} in Tanzania.
11. Return ONLY valid JSON, no extra text.
Example row:
{{"Main Competence": "1.0 Demonstrate understanding of {subject} concepts", "Specific Competences": "Explain key concepts of the topic", "Main Learning Activities": "Describe the concept of {subject}", "Specific Learning Activities": "1. Explain the concept\\n2. Analyse branches\\n3. Evaluate importance", "Month": "MAY", "Week": "1st", "Number of Periods": {periods_per_week}, "Teaching and Learning Methods": "Group discussion, Q&A, Brainstorming", "Teaching and Learning Resources": "Textbook, Charts, Chalkboard", "Assessment Tools": "Oral questions, Quiz", "References": "TIE (2023). {subject} Syllabus. TIE, Dar es Salaam.", "Remarks": ""}}
"""

            response = client.models.generate_content(model=model_name, contents=prompt)
            response_text = response.text

            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                json_data = json_match.group()
            else:
                json_data = response_text

            try:
                scheme_data = json.loads(json_data)
            except Exception as e:
                print("JSON parse error:", e)
                scheme_data = []

            saved_id = None
            try:
                student = StudentTeacher.objects.get(user=request.user)
                school = student.selected_school
                if school and scheme_data:
                    level_map = {
                        'primary school': 'primary',
                        'ordinary level': 'ordinary',
                        'advanced level': 'advanced',
                    }
                    edu_level = level_map.get(education_level.lower(), 'ordinary')
                    subj_obj = Subject.objects.filter(name__iexact=subject).first()
                    if subj_obj:
                        start_dt = None
                        end_dt = None
                        if start_date:
                            try:
                                from datetime import date as _dt
                                start_dt = _dt.fromisoformat(start_date)
                            except Exception:
                                pass
                        if end_date:
                            try:
                                from datetime import date as _dt
                                end_dt = _dt.fromisoformat(end_date)
                            except Exception:
                                pass
                        scheme_obj, _ = SchemeOfWork.objects.update_or_create(
                            student=student,
                            subject=subj_obj,
                            term=term,
                            year=int(year),
                            defaults={
                                'school': school,
                                'education_level': edu_level,
                                'class_name': class_name,
                                'syllabus': syllabus,
                                'total_weeks': int(total_weeks),
                                'periods_per_week': int(periods_per_week),
                                'start_date': start_dt,
                                'end_date': end_dt,
                                'teacher_name': teacher_name,
                                'reference_source': reference_source,
                                'breaks': breaks,
                                'scheme_data': scheme_data,
                                'generated_by_ai': True,
                            }
                        )
                        saved_id = scheme_obj.id
            except Exception as save_err:
                print(f"Scheme save error (non-fatal): {save_err}")

            return JsonResponse({'success': True, 'data': scheme_data, 'saved_id': saved_id})

        except Exception as e:
            import traceback
            traceback.print_exc()
            raw = str(e)
            print(f"[AI SCHEME ERROR] {raw}")
            if 'PERMISSION_DENIED' in raw or 'suspended' in raw.lower() or '403' in raw:
                user_msg = (
                    "Huduma ya AI imesimamishwa kwa sababu ya ufunguo wa API uliosimamishwa. "
                    "Tafadhali wasiliana na msimamizi ili upate ufunguo mpya wa Google AI."
                )
            elif 'quota' in raw.lower() or '429' in raw or 'rate' in raw.lower() or 'exhausted' in raw.lower() or 'resource' in raw.lower():
                user_msg = f"Kikomo cha matumizi: {raw[:200]}"
            elif 'API_KEY' in raw or 'api_key' in raw.lower():
                user_msg = "Ufunguo wa API ya AI haujawekwa. Wasiliana na msimamizi."
            else:
                user_msg = f"Hitilafu: {raw[:200]}"
            return JsonResponse({'success': False, 'error': user_msg}, status=500)

    return JsonResponse({'success': False}, status=400)


@login_required
def download_scheme_pdf(request):
    """Generate PDF ya Scheme of Work — landscape A4, proper cell wrapping"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)

    data = json.loads(request.body)
    scheme_data = data.get('scheme_data') or []
    subject      = data.get('subject', '')
    class_name   = data.get('class_name', '')
    term         = data.get('term', '')
    year         = data.get('year', '')
    syllabus     = data.get('syllabus', '')
    teacher_name = data.get('teacher_name', '')
    school_name  = data.get('school_name', '')
    total_weeks  = data.get('total_weeks', '')

    NAVY   = colors.HexColor('#0A2B5E')
    GOLD   = colors.HexColor('#C8900A')
    STRIPE = colors.HexColor('#EBF0FB')
    BORDER = colors.HexColor('#9BAAC4')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4),
                            rightMargin=18, leftMargin=18, topMargin=22, bottomMargin=22)
    elements = []

    cell_style = ParagraphStyle('SchCell', fontName='Helvetica', fontSize=7,
                                leading=10, wordWrap='LTR', spaceAfter=0, spaceBefore=0)
    hdr_style  = ParagraphStyle('SchHdr', fontName='Helvetica-Bold', fontSize=7,
                                leading=10, textColor=colors.white, wordWrap='LTR',
                                alignment=1, spaceAfter=0, spaceBefore=0)

    elements.append(Paragraph("SCHEME OF WORK",
        ParagraphStyle('ST', fontName='Helvetica-Bold', fontSize=14, alignment=1,
                       textColor=NAVY, spaceAfter=3)))
    elements.append(Paragraph(
        f"{subject}  |  {class_name}  |  Term {term} {year}  |  {syllabus}",
        ParagraphStyle('SS', fontSize=9, alignment=1, spaceAfter=2)))
    elements.append(Paragraph(
        f"Teacher: {teacher_name}    |    School: {school_name}    |    Total Weeks: {total_weeks}",
        ParagraphStyle('SI', fontSize=8, alignment=1, textColor=colors.grey, spaceAfter=10)))

    if scheme_data:
        headers = list(scheme_data[0].keys())

        WIDTH_MAP = {
            'Main Competence': 72, 'Specific Competences': 72,
            'Main Learning Activities': 68, 'Specific Learning Activities': 76,
            'Month': 38, 'Week': 34, 'Number of Periods': 32,
            'Teaching and Learning Methods': 64,
            'Teaching and Learning Resources': 64,
            'Assessment Tools': 54, 'References': 64, 'Remarks': 68,
        }
        TOTAL = 806
        col_widths = []
        for h in headers:
            w = WIDTH_MAP.get(h)
            if w is None:
                for k, v in WIDTH_MAP.items():
                    if k.lower() in h.lower() or h.lower() in k.lower():
                        w = v; break
            col_widths.append(w if w else TOTAL // len(headers))

        scale = TOTAL / sum(col_widths)
        col_widths = [w * scale for w in col_widths]

        table_data = [[Paragraph(h, hdr_style) for h in headers]]
        for row in scheme_data:
            table_data.append([
                Paragraph(str(row.get(h, '') or ''), cell_style) for h in headers
            ])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        ts = [
            ('BACKGROUND',  (0, 0), (-1, 0),  NAVY),
            ('ALIGN',       (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN',      (0, 0), (-1, -1), 'TOP'),
            ('GRID',        (0, 0), (-1, -1), 0.35, BORDER),
            ('TOPPADDING',  (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0,0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING',(0, 0), (-1, -1), 4),
            ('LINEBELOW',   (0, 0), (-1, 0),  1.2, GOLD),
        ]
        for i in range(1, len(table_data)):
            bg = STRIPE if i % 2 == 0 else colors.white
            ts.append(('BACKGROUND', (0, i), (-1, i), bg))
        tbl.setStyle(TableStyle(ts))
        elements.append(tbl)

    doc.build(elements)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = (
        f'attachment; filename="Scheme_of_Work_{subject}_{class_name}.pdf"')
    return response


@login_required
def download_scheme_word(request):
    """Export Scheme of Work as Word (.docx)"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as _json

    data = _json.loads(request.body)
    scheme_data = data.get('scheme_data', [])
    subject = data.get('subject', '')
    class_name = data.get('class_name', '')
    term = data.get('term', '')
    year = data.get('year', '')
    teacher_name = data.get('teacher_name', '')
    school_name = data.get('school_name', '')

    doc = Document()
    doc.core_properties.title = f"Scheme of Work — {subject}"

    title = doc.add_heading(f"{subject} — {class_name} — Term {term} {year}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph(f"Teacher: {teacher_name}    School: {school_name}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if scheme_data:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        headers = list(scheme_data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        WIDTH_CM = {
            'Main Competence': 2.6, 'Specific Competences': 2.6,
            'Main Learning Activities': 2.4, 'Specific Learning Activities': 2.8,
            'Month': 1.3, 'Week': 1.2, 'Number of Periods': 1.1,
            'Teaching and Learning Methods': 2.2,
            'Teaching and Learning Resources': 2.2,
            'Assessment Tools': 1.9, 'References': 2.2, 'Remarks': 2.5,
        }
        TOTAL_CM = 24.0
        col_cms = []
        for h in headers:
            w = WIDTH_CM.get(h)
            if w is None:
                for k, v in WIDTH_CM.items():
                    if k.lower() in h.lower() or h.lower() in k.lower():
                        w = v; break
            col_cms.append(w if w else TOTAL_CM / len(headers))
        scale = TOTAL_CM / sum(col_cms)
        col_cms = [w * scale for w in col_cms]

        from docx.shared import Cm
        section = doc.sections[0]
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = section.right_margin = Cm(1.2)
        section.top_margin  = section.bottom_margin = Cm(1.2)

        for i, cell in enumerate(table.rows[0].cells):
            cell.width = Cm(col_cms[i])

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            p = hdr_cells[i].paragraphs[0]
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0A2B5E')
            tcPr.append(shd)

        for ri, row in enumerate(scheme_data):
            row_cells = table.add_row().cells
            fill_hex = 'EBF0FB' if ri % 2 == 0 else 'FFFFFF'
            for i, h in enumerate(headers):
                row_cells[i].width = Cm(col_cms[i])
                row_cells[i].text = str(row.get(h, '') or '')
                p = row_cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(8)
                tc = row_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"Scheme_{subject}_{class_name}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


@login_required
def download_lesson_plan_pdf(request):
    """Export Lesson Plan as PDF — proper cell wrapping, navy/gold theme"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    import json as _json

    data   = _json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form   = data.get('form_data', {})

    NAVY   = colors.HexColor('#0A2B5E')
    GOLD   = colors.HexColor('#C8900A')
    LIGHT  = colors.HexColor('#EEF1F6')
    STRIPE = colors.HexColor('#F4F7FF')
    BORDER = colors.HexColor('#9BAAC4')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    elements = []

    normal = ParagraphStyle('LP_N', fontName='Helvetica', fontSize=9, leading=13,
                             wordWrap='LTR', spaceAfter=3)
    section_hdr = ParagraphStyle('LP_H', fontName='Helvetica-Bold', fontSize=10,
                                  textColor=NAVY, spaceBefore=10, spaceAfter=4,
                                  borderPad=4, borderColor=GOLD, borderWidth=0,
                                  leftIndent=0)
    cell_s  = ParagraphStyle('LP_C',  fontName='Helvetica',      fontSize=8,  leading=11, wordWrap='LTR')
    hdr_s   = ParagraphStyle('LP_CH', fontName='Helvetica-Bold', fontSize=8,  leading=11,
                              textColor=colors.white, wordWrap='LTR', alignment=1)
    label_s = ParagraphStyle('LP_L',  fontName='Helvetica-Bold', fontSize=8,  leading=11,
                              textColor=NAVY)

    elements.append(Paragraph("LESSON PLAN",
        ParagraphStyle('LP_T', fontName='Helvetica-Bold', fontSize=16, alignment=1,
                       textColor=NAVY, spaceAfter=2)))
    elements.append(Paragraph(
        f"{form.get('subject','')}  |  {form.get('class_name','')}  |  "
        f"Term {form.get('term','')} {form.get('year','')}",
        ParagraphStyle('LP_S', fontSize=9, alignment=1, textColor=colors.grey, spaceAfter=10)))

    def P(txt, st=normal): return Paragraph(str(txt or ''), st)
    meta_rows = [
        [P('Teacher',   label_s), P(form.get('teacher_name','')),
         P('Subject',   label_s), P(form.get('subject',''))],
        [P('Class',     label_s), P(form.get('class_name','')),
         P('Term / Year',label_s),P(f"Term {form.get('term','')} {form.get('year','')}")],
        [P('Topic',     label_s), P(form.get('topic','')),
         P('Subtopic',  label_s), P(form.get('subtopic',''))],
        [P('Duration',  label_s), P(f"{form.get('duration','')} minutes"),
         P('Date',      label_s), P(str(timezone.now().date()))],
        [P('Students',  label_s),
         P(f"{form.get('total_students','')} total / {form.get('present_students','')} present"),
         P(''), P('')],
    ]
    meta_tbl = Table(meta_rows, colWidths=[72, 188, 72, 191])
    meta_tbl.setStyle(TableStyle([
        ('BACKGROUND',   (0, 0), (0, -1), LIGHT),
        ('BACKGROUND',   (2, 0), (2, -1), LIGHT),
        ('GRID',         (0, 0), (-1, -1), 0.4, BORDER),
        ('VALIGN',       (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING',   (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
        ('LEFTPADDING',  (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('LINEBELOW',    (0, -1), (-1, -1), 1.2, GOLD),
    ]))
    elements.append(meta_tbl)
    elements.append(Spacer(1, 10))

    for label, key in [('Main Competence', 'main_competence'),
                       ('Specific Competence', 'specific_competence'),
                       ('Previous Knowledge', 'previous_knowledge')]:
        val = lesson.get(key, '')
        if val:
            elements.append(Paragraph(f"<b>{label}:</b>  {val}", normal))

    for label, key in [('Learning Objectives', 'learning_objectives'),
                       ('Teaching Methods',    'teaching_methods'),
                       ('Teaching Resources',  'teaching_resources')]:
        items = lesson.get(key, [])
        if items:
            elements.append(Paragraph(label, section_hdr))
            for item in items:
                elements.append(Paragraph(
                    f"<bullet>•</bullet> {item}",
                    ParagraphStyle('LP_B', fontName='Helvetica', fontSize=9,
                                   leading=13, leftIndent=14, wordWrap='LTR')))

    ld = lesson.get('lesson_development', [])
    if ld:
        elements.append(Paragraph("Lesson Development (IDDR Model)", section_hdr))
        ld_headers = ['Time', 'Stage (IDDR)', 'Methods', 'Teacher Activities', 'Student Activities', 'Assessment Criteria']
        ld_data = [[Paragraph(h, hdr_s) for h in ld_headers]]
        for i, stage in enumerate(ld):
            bg = colors.white if i % 2 == 0 else STRIPE
            ld_data.append([
                Paragraph(str(stage.get('time', '') or ''), cell_s),
                Paragraph(str(stage.get('stage', stage.get('phase', '')) or ''), cell_s),
                Paragraph(str(stage.get('methods', '') or ''), cell_s),
                Paragraph(str(stage.get('teacher_activities', '') or ''), cell_s),
                Paragraph(str(stage.get('student_activities', '') or ''), cell_s),
                Paragraph(str(stage.get('assessment_criteria', '') or ''), cell_s),
            ])
        ld_tbl = Table(ld_data, colWidths=[38, 72, 72, 118, 118, 105], repeatRows=1)
        ld_ts = [
            ('BACKGROUND',   (0, 0), (-1, 0),  NAVY),
            ('LINEBELOW',    (0, 0), (-1, 0),  1.2, GOLD),
            ('GRID',         (0, 0), (-1, -1), 0.4, BORDER),
            ('VALIGN',       (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING',   (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
            ('LEFTPADDING',  (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ]
        for i in range(1, len(ld_data)):
            bg = STRIPE if i % 2 == 0 else colors.white
            ld_ts.append(('BACKGROUND', (0, i), (-1, i), bg))
        ld_tbl.setStyle(TableStyle(ld_ts))
        elements.append(ld_tbl)

    remarks = lesson.get('remarks', '')
    if remarks:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("Remarks", section_hdr))
        if isinstance(remarks, dict):
            for label, key in [('1. Strength', 'strength'), ('2. Weakness', 'weakness'), ('3. Way Forward', 'way_forward')]:
                val = remarks.get(key, '') or '...............................................'
                elements.append(Paragraph(f"<b>{label}:</b>  {val}", normal))
        else:
            elements.append(Paragraph(str(remarks), normal))

    doc.build(elements)
    buffer.seek(0)
    safe_name = (f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}"
                 .replace(' ', '_'))
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.pdf"'
    return response


@login_required
def download_lesson_plan_word(request):
    """Export Lesson Plan as Word (.docx)"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as _json

    data = _json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form = data.get('form_data', {})

    doc = Document()
    doc.add_heading('LESSON PLAN', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_table = doc.add_table(rows=5, cols=4)
    meta_table.style = 'Table Grid'
    rows = meta_table.rows
    def _set(r, c, text, bold=False):
        cell = rows[r].cells[c]
        cell.text = text
        if bold and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    _set(0,0,'Teacher:', True); _set(0,1, form.get('teacher_name',''))
    _set(0,2,'Subject:', True); _set(0,3, form.get('subject',''))
    _set(1,0,'Class:', True);   _set(1,1, form.get('class_name',''))
    _set(1,2,'Term/Year:', True);_set(1,3, f"Term {form.get('term','')} {form.get('year','')}")
    _set(2,0,'Topic:', True);   _set(2,1, form.get('topic',''))
    _set(2,2,'Subtopic:', True);_set(2,3, form.get('subtopic',''))
    _set(3,0,'Duration:', True);_set(3,1, f"{form.get('duration','')} min")
    _set(3,2,'Date:', True);    _set(3,3, str(timezone.now().date()))
    _set(4,0,'Students:', True);_set(4,1, f"{form.get('total_students','')} / {form.get('present_students','')}")

    doc.add_paragraph()

    for label, key in [('Main Competence', 'main_competence'),
                       ('Specific Competence', 'specific_competence'),
                       ('Previous Knowledge', 'previous_knowledge')]:
        val = lesson.get(key, '')
        if val:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)

    for label, key in [('Learning Objectives', 'learning_objectives'),
                       ('Teaching Methods', 'teaching_methods'),
                       ('Teaching Resources', 'teaching_resources')]:
        items = lesson.get(key, [])
        if items:
            doc.add_heading(label, level=2)
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

    ld = lesson.get('lesson_development', [])
    if ld:
        doc.add_heading('Lesson Development (IDDR Model)', level=2)
        ld_table = doc.add_table(rows=1, cols=6)
        ld_table.style = 'Table Grid'
        hdr = ld_table.rows[0].cells
        for i, h in enumerate(['Time', 'Stage (IDDR)', 'Methods', 'Teacher Activities', 'Student Activities', 'Assessment Criteria']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for stage in ld:
            row = ld_table.add_row().cells
            row[0].text = stage.get('time', '')
            row[1].text = stage.get('stage', stage.get('phase', ''))
            row[2].text = stage.get('methods', '')
            row[3].text = stage.get('teacher_activities', '')
            row[4].text = stage.get('student_activities', '')
            row[5].text = stage.get('assessment_criteria', '')

    remarks = lesson.get('remarks', '')
    if remarks:
        doc.add_heading('Remarks', level=2)
        if isinstance(remarks, dict):
            for label, key in [('1. Strength', 'strength'), ('2. Weakness', 'weakness'), ('3. Way Forward', 'way_forward')]:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(remarks.get(key, '') or '...............................................')
        else:
            p = doc.add_paragraph()
            p.add_run('Remarks: ').bold = True
            p.add_run(str(remarks))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


@login_required
def lesson_plan_view(request):
    """Display lesson plan generator form"""
    from .models import EducationLevel, ClassLevel, StudentApplication

    import json as _json
    education_levels = EducationLevel.objects.all().order_by('order')

    classes_by_level = {}
    for cl in ClassLevel.objects.select_related('education_level').order_by('education_level', 'order'):
        classes_by_level.setdefault(cl.education_level_id, []).append({'id': cl.id, 'name': cl.name})

    primary_ids = list(
        EducationLevel.objects.filter(name__icontains='primary').values_list('id', flat=True)
    )
    subjects_by_level = {
        lvl.id: list(
            Subject.objects.filter(level='primary' if lvl.id in primary_ids else 'secondary')
            .order_by('name').values('id', 'name')
        )
        for lvl in education_levels
    }

    student = None
    school = None
    if request.user.is_authenticated:
        try:
            student = StudentTeacher.objects.select_related('selected_school').get(user=request.user)
            school = student.selected_school
            if not school:
                app = StudentApplication.objects.filter(
                    student=student, status='approved'
                ).select_related('school').first()
                if app:
                    school = app.school
        except StudentTeacher.DoesNotExist:
            pass

    return render(request, 'field_app/lesson_plan.html', {
        'education_levels': education_levels,
        'classes_by_level_json': _json.dumps(classes_by_level),
        'subjects_by_level_json': _json.dumps(subjects_by_level),
        'student': student,
        'school': school,
    })


@login_required
def ajax_generate_lessonplan(request):
    """Generate lesson plan using AI"""
    if client is None:
        return JsonResponse({'success': False, 'error': 'Huduma ya AI haitumiki. Ufunguo wa API (GOOGLE_API_KEY) haujawekwa. Wasiliana na msimamizi.'}, status=503)
    if request.method == 'POST':
        try:
            import json as json_module
            import re as re_module

            data = json_module.loads(request.body)

            education_level = data.get('education_level', '')
            class_name = data.get('class_name', '')
            subject = data.get('subject', '')
            subject_id = data.get('subject_id', '')
            topic = data.get('topic', '')
            subtopic = data.get('subtopic', '')
            term = data.get('term', 'I')
            year = data.get('year', 2026)
            duration = data.get('duration', 40)
            total_students = data.get('total_students', '')
            present_students = data.get('present_students', '')
            learning_objectives = data.get('learning_objectives', '')
            teaching_methods = data.get('teaching_methods', '')
            reference_source = data.get('reference_source', '')
            teacher_name = data.get('teacher_name', '')

            prompt = f"""
You are an AI assistant for Tanzanian teachers. Generate a detailed LESSON PLAN following the SEQUIP/TIE Tanzania IDDR revised format (2023).

The IDDR model stages are:
- I = Introduction: Engage learners, activate prior knowledge, arouse curiosity, use questions/activities related to new content.
- D = Competence Development: Guide learners to build competence through reading, discussion, investigation, videos, group work, presentation.
- D = Design: Deepen learning — apply knowledge in real-life/other contexts. Role play, problem solving, creating models.
- R = Realisation: Assess and evaluate student achievement using portfolios, performance assessments, quizzes, peer/self-assessment.

Input Details:
- Education Level: {education_level}
- Class/Form: {class_name}
- Subject: {subject}
- Topic: {topic}
- Subtopic: {subtopic}
- Term: {term}
- Year: {year}
- Duration: {duration} minutes
- Total Students: {total_students}
- Present Students: {present_students}
- Learning Objectives: {learning_objectives}
- Teaching Methods: {teaching_methods}
- Reference Source: {reference_source}

Allocate time across the 4 IDDR stages proportionally within {duration} minutes.
Introduction ≈ 15% | Competence Development ≈ 40% | Design ≈ 30% | Realisation ≈ 15%

Output MUST be ONLY valid JSON. Do not include any other text. Use this exact structure:
{{
    "lesson_title": "{subject} - {topic}",
    "date": "today's date",
    "main_competence": "Main competence from syllabus",
    "specific_competence": "Specific competence from syllabus",
    "previous_knowledge": "What students already know related to this topic",
    "learning_objectives": ["By end of lesson, students will be able to...", "..."],
    "teaching_methods": ["Group discussion", "Q&A", "Brainstorming", "Role play"],
    "teaching_resources": ["Textbook", "Chalkboard", "Charts", "Realia"],
    "lesson_development": [
        {{"time": "X min", "stage": "Introduction (I)", "methods": "...", "teacher_activities": "...", "student_activities": "...", "assessment_criteria": "..."}},
        {{"time": "X min", "stage": "Competence Development (D)", "methods": "...", "teacher_activities": "...", "student_activities": "...", "assessment_criteria": "..."}},
        {{"time": "X min", "stage": "Design (D)", "methods": "...", "teacher_activities": "...", "student_activities": "...", "assessment_criteria": "..."}},
        {{"time": "X min", "stage": "Realisation (R)", "methods": "...", "teacher_activities": "...", "student_activities": "...", "assessment_criteria": "..."}}
    ],
    "remarks": {{
        "strength": "What went well in this lesson",
        "weakness": "What did not go well / challenges",
        "way_forward": "What will be done to improve"
    }}
}}
"""

            response = client.models.generate_content(model=model_name, contents=prompt)
            response_text = response.text

            cleaned_text = re_module.sub(r'```json\s*', '', response_text)
            cleaned_text = re_module.sub(r'```\s*', '', cleaned_text)
            cleaned_text = cleaned_text.strip()

            start_idx = cleaned_text.find('{')
            end_idx = cleaned_text.rfind('}')

            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_data = cleaned_text[start_idx:end_idx + 1]
                lesson_data = json_module.loads(json_data)
            else:
                lesson_data = {
                    "lesson_title": f"{subject} - {topic}",
                    "date": timezone.now().strftime('%Y-%m-%d'),
                    "main_competence": f"Demonstrate understanding of {topic}",
                    "specific_competence": f"Explain key concepts of {topic}",
                    "previous_knowledge": "Basic knowledge from previous lessons",
                    "learning_objectives": [f"Define {topic}", f"Explain {topic}", f"Apply {topic} in real-life situations"],
                    "teaching_methods": ["Group discussion", "Q&A", "Brainstorming", "Role play"],
                    "teaching_resources": ["Chalkboard", "Textbook", "Charts"],
                    "lesson_development": [
                        {"time": f"{max(5, int(int(duration)*0.15))} min", "stage": "Introduction (I)", "methods": "Q&A, Brainstorming", "teacher_activities": f"Use questions to activate prior knowledge about {topic}", "student_activities": "Respond to questions and share what they know", "assessment_criteria": "Active participation and relevant responses"},
                        {"time": f"{max(10, int(int(duration)*0.40))} min", "stage": "Competence Development (D)", "methods": "Group discussion, Guided discovery", "teacher_activities": f"Guide students through group discussions to explore {topic}", "student_activities": "Conduct discussions, explore content, present findings", "assessment_criteria": "Accuracy of explanations and active participation"},
                        {"time": f"{max(8, int(int(duration)*0.30))} min", "stage": "Design (D)", "methods": "Role play, Problem solving", "teacher_activities": f"Organise activities where students apply {topic} in real-life scenarios", "student_activities": "Apply knowledge creatively to solve real-life problems", "assessment_criteria": "Correct application and creativity demonstrated"},
                        {"time": f"{max(5, int(int(duration)*0.15))} min", "stage": "Realisation (R)", "methods": "Quiz, Peer assessment", "teacher_activities": "Administer assessment activities and provide feedback", "student_activities": "Complete assessment tasks and self/peer assess", "assessment_criteria": "Achievement of learning objectives"}
                    ],
                    "remarks": {
                        "strength": "",
                        "weakness": "",
                        "way_forward": ""
                    }
                }

            saved_id = None
            save_error_msg = None
            try:
                student = StudentTeacher.objects.get(user=request.user)
                school = student.selected_school
                if not school:
                    save_error_msg = 'Student has no selected school'
                else:
                    level_map = {
                        'primary school': 'primary',
                        'ordinary level': 'ordinary',
                        'advanced level': 'advanced',
                    }
                    edu_level = level_map.get((education_level or '').lower(), 'ordinary')

                    subj_obj = None
                    if subject_id:
                        try:
                            subj_obj = Subject.objects.get(id=int(subject_id))
                        except (Subject.DoesNotExist, ValueError):
                            pass
                    if not subj_obj and subject:
                        subj_obj = Subject.objects.filter(name__iexact=subject).first()

                    if not subj_obj:
                        save_error_msg = f'Subject not found: id={subject_id} name={subject}'
                    else:
                        lp_date = timezone.now().date()
                        lo_list = lesson_data.get('learning_objectives', [])
                        if isinstance(lo_list, str):
                            lo_list = [x.strip() for x in lo_list.split('\n') if x.strip()]
                        tm_list = lesson_data.get('teaching_methods', [])
                        if isinstance(tm_list, str):
                            tm_list = [x.strip() for x in tm_list.split('\n') if x.strip()]
                        lp_obj = LessonPlan.objects.create(
                            student=student,
                            school=school,
                            subject=subj_obj,
                            education_level=edu_level,
                            class_name=class_name,
                            term=term,
                            year=int(year),
                            topic=topic,
                            subtopic=subtopic or '',
                            date=lp_date,
                            duration=int(duration) if duration else 40,
                            total_students=int(total_students) if total_students else 0,
                            present_students=int(present_students) if present_students else 0,
                            teacher_name=teacher_name or student.full_name,
                            reference_source=reference_source or '',
                            main_competence=lesson_data.get('main_competence', ''),
                            specific_competence=lesson_data.get('specific_competence', ''),
                            previous_knowledge=lesson_data.get('previous_knowledge', ''),
                            learning_objectives=lo_list,
                            teaching_methods=tm_list,
                            teaching_resources=lesson_data.get('teaching_resources', []),
                            lesson_development=lesson_data.get('lesson_development', []),
                            remarks=lesson_data.get('remarks', ''),
                            generated_by_ai=True,
                        )
                        saved_id = lp_obj.id
            except StudentTeacher.DoesNotExist:
                save_error_msg = f'No StudentTeacher for user {request.user}'
            except Exception as save_err:
                import traceback as _tb
                save_error_msg = str(save_err)
                print(f"Lesson Plan save error: {save_error_msg}")
                print(_tb.format_exc())

            if save_error_msg:
                print(f"[LessonPlan] NOT saved — {save_error_msg}")

            return JsonResponse({'success': True, 'data': lesson_data, 'saved_id': saved_id})

        except Exception as e:
            import traceback
            traceback.print_exc()
            raw = str(e)
            print(f"[AI LESSON ERROR] {raw}")
            if 'PERMISSION_DENIED' in raw or 'suspended' in raw.lower() or '403' in raw:
                user_msg = (
                    "Huduma ya AI imesimamishwa kwa sababu ya ufunguo wa API uliosimamishwa. "
                    "Tafadhali wasiliana na msimamizi ili upate ufunguo mpya wa Google AI."
                )
            elif 'quota' in raw.lower() or '429' in raw or 'rate' in raw.lower() or 'exhausted' in raw.lower() or 'resource' in raw.lower():
                user_msg = f"Kikomo cha matumizi: {raw[:200]}"
            elif 'API_KEY' in raw or 'api_key' in raw.lower():
                user_msg = "Ufunguo wa API ya AI haujawekwa. Wasiliana na msimamizi."
            else:
                user_msg = f"Hitilafu ya AI: {raw[:300]}"
            return JsonResponse({'success': False, 'error': user_msg}, status=500)

    return JsonResponse({'success': False}, status=400)


@login_required
def ajax_load_saved_scheme(request):
    """Load most recent saved SchemeOfWork for this student from DB — no AI call."""
    try:
        student = StudentTeacher.objects.get(user=request.user)
        scheme = (SchemeOfWork.objects
                  .filter(student=student)
                  .select_related('subject')
                  .order_by('-updated_at')
                  .first())
        if not scheme or not scheme.scheme_data:
            return JsonResponse({'success': False, 'error': 'Hakuna mpango uliohifadhiwa. Tengeneza kwanza.'}, status=404)
        return JsonResponse({
            'success': True,
            'data': scheme.scheme_data,
            'saved_id': scheme.id,
            'meta': {
                'subject': scheme.subject.name,
                'class_name': scheme.class_name,
                'term': scheme.term,
                'year': scheme.year,
                'teacher_name': scheme.teacher_name,
                'school_name': scheme.school.name if scheme.school else '',
                'total_weeks': scheme.total_weeks,
                'syllabus': scheme.syllabus,
                'updated_at': scheme.updated_at.strftime('%d %b %Y, %H:%M'),
            }
        })
    except StudentTeacher.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Wasifu wa mwanafunzi haupatikani.'}, status=404)


@login_required
def ajax_load_saved_lessonplan(request):
    """Load most recent saved LessonPlan for this student from DB — no AI call."""
    try:
        student = StudentTeacher.objects.get(user=request.user)
        lp = (LessonPlan.objects
              .filter(student=student)
              .select_related('subject')
              .order_by('-created_at')
              .first())
        if not lp or not lp.lesson_development:
            return JsonResponse({'success': False, 'error': 'Hakuna mpango wa somo uliohifadhiwa. Tengeneza kwanza.'}, status=404)
        lesson_data = {
            'lesson_title': f"{lp.subject.name} - {lp.topic}",
            'main_competence': lp.main_competence,
            'specific_competence': lp.specific_competence,
            'previous_knowledge': lp.previous_knowledge,
            'learning_objectives': lp.learning_objectives,
            'teaching_methods': lp.teaching_methods,
            'teaching_resources': lp.teaching_resources,
            'lesson_development': lp.lesson_development,
            'remarks': lp.remarks,
        }
        form_data = {
            'subject': lp.subject.name,
            'class_name': lp.class_name,
            'term': lp.term,
            'year': lp.year,
            'topic': lp.topic,
            'subtopic': lp.subtopic,
            'teacher_name': lp.teacher_name,
            'duration': lp.duration,
            'total_students': lp.total_students,
            'present_students': lp.present_students,
        }
        return JsonResponse({
            'success': True,
            'data': lesson_data,
            'form_data': form_data,
            'saved_id': lp.id,
            'updated_at': lp.created_at.strftime('%d %b %Y, %H:%M'),
        })
    except StudentTeacher.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Wasifu wa mwanafunzi haupatikani.'}, status=404)
