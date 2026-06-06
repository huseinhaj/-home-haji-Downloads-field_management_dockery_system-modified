import csv
import io
import json
import re
import secrets
import string
from io import BytesIO
from datetime import datetime, timedelta

from django.conf import settings
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import authenticate, login, logout, get_user_model
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.gis.db.models.functions import Distance
from django.contrib.gis.geos import Point
from django.core.cache import cache
from django.core.exceptions import PermissionDenied
from django.core.files.base import ContentFile
from django.core.mail import send_mail
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Count, Case, When, Value, BooleanField, F, Q, Prefetch, Max
from django.db import models
from django.db.models.functions import Greatest
from django.http import HttpResponse, HttpResponseNotAllowed, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from geopy.distance import geodesic

from .ai_utils import client, model_name
from .decorators import board_login_required, assessor_login_required
from .forms import (
    CustomLoginForm, StudentRegistrationForm, StudentTeacherForm,
    LogbookForm, AssessorLoginForm, BulkAssignForm, RegionFieldInputForm
)
from .forms import SchemeOfWorkForm
from .models import (
    Assessor, School, SchoolAssignment, StudentTeacher,
    StudentAssessment, SchoolAssessment, SchoolRequirement,
    StudentApplication, Region, RegionPin, SchoolPin,
    District, Subject, SchoolSubjectCapacity,
    LogbookEntry, ApprovalLetter, AcademicYear, SchemeOfWork,
    BoardMember, BoardComment, LessonPlan, MonthlyReport,
    DistrictAllocation, SchoolAllocation, SchoolHeadRequest,
)

User = get_user_model()

# =========================
# CACHE HELPERS
# =========================

def _cached_active_year():
    """Cache the active AcademicYear for 5 minutes."""
    key = 'active_academic_year'
    yr = cache.get(key)
    if yr is None:
        yr = AcademicYear.objects.filter(is_active=True).first()
        cache.set(key, yr, 300)
    return yr

def _cached_subjects(student):
    """Return all subjects grouped by level. Always show all so teacher can record any subject."""
    key = 'all_subjects_list'
    subs = cache.get(key)
    if subs is None:
        subs = list(Subject.objects.order_by('level', 'name'))
        cache.set(key, subs, 600)
    return subs


def _cached_today_logbook(student, school, today):
    """Cache today's logbook entry for 30s to avoid get_or_create hit on every page load."""
    key = f'logbook_today_{student.id}_{today}'
    entry = cache.get(key)
    if entry is None:
        entry, _ = LogbookEntry.objects.get_or_create(
            student=student, date=today,
            defaults={'school': school, 'morning_check_in': timezone.now()}
        )
        cache.set(key, entry, 30)
    return entry


def _invalidate_today_logbook(student, today):
    cache.delete(f'logbook_today_{student.id}_{today}')


def _cached_schools_by_district(district_id):
    """Cache schools za district fulani kwa dakika 10 - inatumiwa mara nyingi kwenye ombi/DEO views."""
    key = f'schools_district_{district_id}'
    schools = cache.get(key)
    if schools is None:
        schools = list(
            School.objects.filter(district_id=district_id)
            .select_related('district__region')
            .order_by('name')
        )
        cache.set(key, schools, 600)
    return schools


def _cached_all_districts():
    """Cache districts zote kwa dakika 15 - data haitabadilika mara kwa mara."""
    key = 'all_districts_list'
    districts = cache.get(key)
    if districts is None:
        districts = list(District.objects.select_related('region').order_by('name'))
        cache.set(key, districts, 900)
    return districts


def _cached_all_regions():
    """Cache regions zote kwa dakika 30."""
    key = 'all_regions_list'
    regions = cache.get(key)
    if regions is None:
        regions = list(Region.objects.order_by('name'))
        cache.set(key, regions, 1800)
    return regions

# =========================
# HELPER FUNCTIONS
# =========================
# =========================
# EMAIL TEMPLATES - Add this at the top of views.py
# =========================

def get_assessor_email_template(assessor, school, temp_password, is_new_account, assignments_count, login_url):
    """Generate beautiful HTML email for assessor - Tanzania Teacher Colleges"""
    
    if is_new_account:
        credential_html = f"""
        <div style="background-color: #fef9e6; border-left: 4px solid #f59e0b; padding: 20px; margin: 20px 0; border-radius: 8px;">
            <h3 style="margin-top: 0; color: #d97706;">🆕 AKAUNTI MPYA YA ASSESSOR</h3>
            <p style="margin: 10px 0;"><strong>📧 Email:</strong> {assessor.email}</p>
            <p style="margin: 10px 0;"><strong>🔑 Nywila ya Muda:</strong> <code style="background-color: #fff3cd; padding: 4px 8px; border-radius: 4px; font-size: 16px;">{temp_password}</code></p>
            <p style="margin: 10px 0; color: #856404;">⚠️ Badilisha nywila yako mara tu baada ya kuingia mara ya kwanza</p>
        </div>
        """
    else:
        credential_html = f"""
        <div style="background-color: #e3f2fd; border-left: 4px solid #2196f3; padding: 20px; margin: 20px 0; border-radius: 8px;">
            <h3 style="margin-top: 0; color: #1976d2;">🔄 MWAKA MPYA WA MASOMO {assessor.current_academic_year.year if assessor.current_academic_year else '2024/2025'}</h3>
            <p style="margin: 10px 0;"><strong>📧 Email:</strong> {assessor.email}</p>
            <p style="margin: 10px 0;"><strong>🔑 Nywila Mpya:</strong> <code style="background-color: #fff; padding: 4px 8px; border-radius: 4px; font-size: 16px;">{temp_password}</code></p>
            <p style="margin: 10px 0; color: #d32f2f;">🔐 Nywila yako imebadilishwa kwa mwaka mpya wa masomo</p>
        </div>
        """
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Tanzania Teacher Colleges - Field Placement</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: #333;
                margin: 0;
                padding: 0;
                background-color: #f0f2f5;
            }}
            .container {{
                max-width: 600px;
                margin: 20px auto;
                background: white;
                border-radius: 16px;
                overflow: hidden;
                box-shadow: 0 4px 12px rgba(0,0,0,0.1);
            }}
            .header {{
                background: linear-gradient(135deg, #2c3e50 0%, #3498db 100%);
                color: white;
                padding: 30px;
                text-align: center;
            }}
            .header h1 {{
                margin: 0;
                font-size: 24px;
            }}
            .header p {{
                margin: 8px 0 0;
                opacity: 0.9;
            }}
            .content {{
                padding: 30px;
            }}
            .info-card {{
                background: #f8f9fa;
                border-radius: 12px;
                padding: 20px;
                margin: 20px 0;
                border: 1px solid #e9ecef;
            }}
            .button {{
                display: inline-block;
                background: linear-gradient(135deg, #2c3e50 0%, #3498db 100%);
                color: white;
                text-decoration: none;
                padding: 12px 30px;
                border-radius: 8px;
                margin: 20px 0;
                font-weight: bold;
                text-align: center;
            }}
            .footer {{
                background: #f8f9fa;
                padding: 20px;
                text-align: center;
                font-size: 11px;
                color: #6c757d;
                border-top: 1px solid #e9ecef;
            }}
            @media only screen and (max-width: 600px) {{
                .content {{
                    padding: 20px;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>🏫 Tanzania Teacher Colleges</h1>
                <p>Field Placement Management System</p>
            </div>
            
            <div class="content">
                <h2>Dear {assessor.full_name},</h2>
                
                <p>You have been assigned as a <strong>Field Placement Assessor</strong> for the academic year <strong>{assessor.current_academic_year.year if assessor.current_academic_year else '2024/2025'}</strong>.</p>
                
                <div class="info-card">
                    <h3 style="margin-top: 0;">📋 Assignment Details</h3>
                    <p><strong>🏫 College/School:</strong> {school.name}</p>
                    <p><strong>📍 District:</strong> {school.district.name}</p>
                    <p><strong>🗺️ Region:</strong> {school.district.region.name}</p>
                    <p><strong>📅 Assignment Date:</strong> {timezone.now().strftime('%d/%m/%Y')}</p>
                    <p><strong>👥 Student Teachers:</strong> {assignments_count}</p>
                </div>
                
                {credential_html}
                
                <div style="text-align: center;">
                    <a href="{login_url}" class="button" style="color: white; text-decoration: none;">
                        🔐 LOGIN TO YOUR DASHBOARD
                    </a>
                </div>
                
                <div class="info-card" style="background: #e7f3ff;">
                    <h3 style="margin-top: 0;">✅ After Login You Can:</h3>
                    <ul>
                        <li>📊 View assigned college/school details</li>
                        <li>👨‍🎓 See list of student teachers</li>
                        <li>📝 Track teaching practice logbooks</li>
                        <li>📋 Submit assessment reports</li>
                        <li>📈 Monitor student progress</li>
                    </ul>
                </div>
                
                <p style="margin-top: 30px;">Best regards,<br>
                <strong>Field Placement Coordination Unit</strong><br>
                Tanzania Teacher Colleges</p>
                
                <p style="font-size: 11px; color: #999; margin-top: 20px;">
                    📧 This is an automated message. Please do not reply.
                </p>
            </div>
            
            <div class="footer">
                <p>© {timezone.now().year} Tanzania Teacher Colleges - Field Placement System</p>
                <p>📍 Empowering Future Educators | Tanzania</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    text_content = f"""
TANZANIA TEACHER COLLEGES - FIELD PLACEMENT SYSTEM
{'='*60}

Dear {assessor.full_name},

You have been assigned as a Field Placement Assessor.

ASSIGNMENT DETAILS:
• College/School: {school.name}
• District: {school.district.name}
• Region: {school.district.region.name}
• Assignment Date: {timezone.now().strftime('%d/%m/%Y')}
• Student Teachers: {assignments_count}

{'='*60}
LOGIN CREDENTIALS:
Login URL: {login_url}
Email: {assessor.email}
Password: {temp_password if is_new_account else 'Use your existing password (reset for new year)'}

After login you can:
1. View assigned college/school details
2. See list of student teachers
3. Track teaching practice logbooks
4. Submit assessment reports

IMPORTANT:
• Change your password immediately after first login
• Contact your college coordinator if you face issues

Best regards,
Field Placement Coordination Unit
Tanzania Teacher Colleges
"""
    
    return html_content, text_content
# =========================
# UPDATED BULK ASSIGNMENT FUNCTION WITH HTML EMAIL
# =========================

def process_bulk_assignment_with_academic_year(assessor_ids, school_ids, assessment_date, request):
    """
    Create assignments ONLY - NO NEW CREDENTIALS UNLESS NEW ACADEMIC YEAR
    FIXED: No username field - uses email as USERNAME_FIELD
    WITH BEAUTIFUL HTML EMAIL
    """
    
    print(f"\n{'='*60}")
    print(f"📊 BULK ASSIGNMENT PROCESS STARTED")
    print(f"{'='*60}")
    print(f"👥 Assessors: {len(assessor_ids)}")
    print(f"🏫 Schools: {len(school_ids)}")
    print(f"📅 Assessment Date: {assessment_date}")
    
    # Get current academic year
    current_academic_year = get_current_academic_year()
    
    print(f"📚 Current Academic Year: {current_academic_year.year}")
    
    # Get all assessors and schools
    assessors = Assessor.objects.filter(id__in=assessor_ids).select_related('user')
    schools = School.objects.filter(id__in=school_ids).select_related('district', 'district__region')
    
    print(f"✅ Found {assessors.count()} assessors")
    print(f"✅ Found {schools.count()} schools")
    
    # Process each assessor
    email_results = []
    new_accounts_count = 0
    new_year_resets = 0
    assignments_created = 0
    email_sent_count = 0
    
    for assessor in assessors:
        print(f"\n{'─'*50}")
        print(f"👤 Processing: {assessor.full_name}")
        print(f"📧 Email: {assessor.email}")
        print(f"📚 Current Academic Year in DB: {assessor.current_academic_year}")
        
        # Validate email
        if not assessor.email or '@' not in assessor.email:
            email_results.append({
                'assessor': assessor.full_name,
                'status': '❌ Skipped - Invalid email',
                'email': assessor.email or 'No email',
                'credentials': 'N/A',
                'is_new': False,
                'is_new_year': False
            })
            print(f"❌ Invalid email - skipped")
            continue
        
        temp_password = None
        credential_action = ""
        is_new = False
        is_new_year = False
        send_email = False
        
        # ========== LOGIC 1: New assessor (no user account) ==========
        if not assessor.user:
            print("🆕 SCENARIO 1: New assessor - CREATING ACCOUNT")
            
            try:
                temp_password = generate_random_password()
                credential_action = "New account created"
                is_new = True
                new_accounts_count += 1
                send_email = True

                print(f"🔐 Generated password: {temp_password}")

                # Check if user with this email already exists
                existing_user = User.objects.filter(email=assessor.email).first()
                if existing_user:
                    user = existing_user
                    user.set_password(temp_password)
                    user.save()
                    print(f"♻️ Reusing existing user account: {assessor.email}")
                else:
                    user = User.objects.create_user(
                        email=assessor.email,
                        password=temp_password,
                        is_staff=False,
                        is_active=True
                    )
                    print(f"✅ Account created with email: {assessor.email}")

                assessor.user = user
                assessor.current_academic_year = current_academic_year
                assessor.save()

                print(f"✅ Academic year set: {current_academic_year.year}")

            except Exception as e:
                print(f"❌ ACCOUNT CREATION FAILED: {e}")
                email_results.append({
                    'assessor': assessor.full_name,
                    'status': f'❌ Account creation failed: {str(e)[:100]}',
                    'email': assessor.email,
                    'credentials': 'FAILED',
                    'is_new': False,
                    'is_new_year': False,
                    'error': str(e)[:100]
                })
                continue
        
        # ========== LOGIC 2: Existing assessor ==========
        elif assessor.user:
            print(f"🔄 SCENARIO 2: Existing assessor - CHECKING ACADEMIC YEAR")
            
            if not assessor.current_academic_year or assessor.current_academic_year != current_academic_year:
                print(f"📅 New academic year detected - NEEDS NEW CREDENTIALS")
                
                try:
                    temp_password = generate_random_password()
                    credential_action = f"New credentials for {current_academic_year.year}"
                    is_new_year = True
                    new_year_resets += 1
                    send_email = True
                    
                    print(f"🔐 New password generated: {temp_password}")
                    
                    assessor.user.set_password(temp_password)
                    assessor.user.save()
                    
                    assessor.current_academic_year = current_academic_year
                    assessor.save()
                    
                    print(f"✅ Password reset for new academic year: {current_academic_year.year}")
                    
                except Exception as e:
                    print(f"❌ PASSWORD RESET FAILED: {e}")
                    email_results.append({
                        'assessor': assessor.full_name,
                        'status': f'❌ Password reset failed: {str(e)[:100]}',
                        'email': assessor.email,
                        'credentials': 'FAILED',
                        'is_new': False,
                        'is_new_year': False,
                        'error': str(e)[:100]
                    })
                    continue
                print(f"✅ Already has credentials for {current_academic_year.year}")
                credential_action = f"Already has credentials for {current_academic_year.year}"
                send_email = False
        
        # ========== CREATE ASSIGNMENTS FOR THIS ASSESSOR ==========
        assignments_for_this_assessor = 0
        skipped_assignments = 0
        
        for school in schools:
            try:
                print(f"\n📝 Processing school: {school.name} (ID: {school.id})")
                
                assignment, created = SchoolAssessment.objects.get_or_create(
                    assessor=assessor,
                    school=school,
                    academic_year=current_academic_year,
                    defaults={
                        'assigned_date': timezone.now().date(),
                        'assessment_date': assessment_date,
                        'is_completed': False,
                        'supervisor': request.user if request.user.is_authenticated else None
                    }
                )
                
                if created:
                    assignments_created += 1
                    assignments_for_this_assessor += 1
                    print(f"✅ NEW assignment created: {assessor.full_name} -> {school.name}")
                    print(f"   Assignment ID: {assignment.id}")

                    approved_students = StudentTeacher.objects.filter(
                        selected_school=school,
                        approval_status='approved'
                    )

                    student_assessments_created = 0
                    for student in approved_students:
                        sa, sa_created = StudentAssessment.objects.get_or_create(
                            assessor=assessor,
                            student=student,
                            school=school,
                            academic_year=current_academic_year,
                            defaults={
                                'assessment_date': assessment_date,
                                'status': 'pending'
                            }
                        )
                        if sa_created:
                            student_assessments_created += 1
                            print(f"   ✓ Created student assessment: {student.full_name}")

                    if student_assessments_created > 0:
                        print(f"   📊 Total student assessments: {student_assessments_created}")
                    else:
                        print(f"   ℹ️ No new student assessments needed")
                else:
                    skipped_assignments += 1
                    print(f"⚠️ SKIPPED: Assignment already exists for {assessor.full_name} -> {school.name}")
                
            except Exception as e:
                print(f"❌ Assignment failed for {school.name}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        print(f"\n📊 This assessor: {assignments_for_this_assessor} NEW assignments, {skipped_assignments} SKIPPED")
        
        # ========== SEND BEAUTIFUL HTML EMAIL ==========
        if send_email and temp_password:
            try:
                login_url = request.build_absolute_uri(reverse('assessor_login'))
                
                # Build schools list for email
                assigned_schools_list = ""
                school_counter = 0
                for school in schools:
                    if SchoolAssessment.objects.filter(
                        assessor=assessor,
                        school=school,
                        academic_year=current_academic_year
                    ).exists():
                        school_counter += 1
                        assigned_schools_list += f"{school_counter}. {school.name} ({school.district.name})\n"
                
                # Create beautiful HTML email
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Field Placement Credentials</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #333; margin: 0; padding: 20px; background: #f0f2f5; }}
        .container {{ max-width: 550px; margin: 0 auto; background: white; border-radius: 20px; overflow: hidden; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; text-align: center; }}
        .header h1 {{ margin: 0; font-size: 24px; }}
        .header p {{ margin: 8px 0 0; opacity: 0.9; }}
        .content {{ padding: 30px; }}
        .greeting {{ font-size: 20px; font-weight: 600; margin-bottom: 20px; }}
        .credential-box {{ background: #fef9e6; border-left: 4px solid #f59e0b; padding: 20px; margin: 20px 0; border-radius: 12px; }}
        .schools-box {{ background: #e7f3ff; padding: 20px; margin: 20px 0; border-radius: 12px; }}
        .button {{ display: block; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; text-decoration: none; padding: 14px 24px; border-radius: 50px; text-align: center; margin: 24px 0; font-weight: 600; }}
        .footer {{ background: #f8f9fa; padding: 20px; text-align: center; font-size: 11px; color: #888; border-top: 1px solid #e9ecef; }}
        code {{ background: #fff3cd; padding: 4px 8px; border-radius: 6px; font-size: 14px; }}
        @media only screen and (max-width: 480px) {{ .content {{ padding: 20px; }} }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🎓 Mfumo wa Ufuatiliaji wa Walimu Wanafunzi</h1>
            <p>Wizara ya Elimu, Sayansi na Teknolojia</p>
        </div>
        <div class="content">
            <div class="greeting">👋 Dear {assessor.full_name},</div>
            <p>You have been assigned as a <strong>Field Placement Assessor</strong> for <strong>{current_academic_year.year}</strong>.</p>
            <div class="credential-box">
                <h3 style="margin: 0 0 15px 0;">🔐 YOUR LOGIN CREDENTIALS</h3>
                <p><strong>📧 Email:</strong> {assessor.email}</p>
                <p><strong>🔑 Password:</strong> <code>{temp_password}</code></p>
                <p style="margin: 15px 0 0 0; color: #856404;">⚠️ Change password after first login</p>
            </div>
            <div class="schools-box">
                <h3 style="margin: 0 0 15px 0;">🏫 ASSIGNED SCHOOLS ({assignments_for_this_assessor})</h3>
                <pre style="background: white; padding: 12px; border-radius: 8px; margin: 0; font-size: 14px;">{assigned_schools_list if assigned_schools_list else 'No new assignments'}</pre>
            </div>
            <a href="{login_url}" class="button">🔐 LOGIN TO YOUR DASHBOARD</a>
            <div style="background: #fff3e0; padding: 15px; border-radius: 12px;">
                <h3 style="margin: 0 0 8px 0;">✅ After Login You Can:</h3>
                <ul style="margin: 0; padding-left: 20px;">
                    <li>View assigned school details</li>
                    <li>See list of students</li>
                    <li>Track logbook entries</li>
                    <li>Submit assessment reports</li>
                </ul>
            </div>
        </div>
        <div class="footer">
            <p>Wizara ya Elimu, Sayansi na Teknolojia — IMS v2.1.0</p>
            <p>📧 Ujumbe huu umetumwa kiotomatiki. Tafadhali usijibu.</p>
        </div>
    </div>
</body>
</html>"""
                
                # Plain text fallback
                text_content = f"""
FIELD PLACEMENT ASSESSOR CREDENTIALS & ASSIGNMENTS
{'='*60}

Dear {assessor.full_name},

{'NEW ACCOUNT CREATED FOR YOU' if is_new else f'NEW CREDENTIALS FOR {current_academic_year.year}'}

ACADEMIC YEAR: {current_academic_year.year}

YOUR LOGIN DETAILS:
• Login URL: {login_url}
• Email: {assessor.email}
• Password: {temp_password}

YOUR ASSIGNMENTS ({assignments_for_this_assessor} schools):
{assigned_schools_list if assigned_schools_list else 'No new assignments created'}

IMPORTANT INSTRUCTIONS:
1. This is a temporary password
2. Change it immediately after first login
3. Login to see your assigned schools and students

Best regards,
Kitengo cha Uratibu wa Mafunzo ya Uwanjani
Wizara ya Elimu, Sayansi na Teknolojia
"""
                
                subject = f'🎓 Field Placement Credentials & Assignments - {current_academic_year.year}'
                
                # Send HTML email
                send_mail(
                    subject=subject,
                    message=text_content,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[assessor.email],
                    html_message=html_content,
                    fail_silently=False,
                )
                
                email_sent_count += 1
                print(f"✅ HTML email sent successfully to: {assessor.email}")
                
                email_results.append({
                    'assessor': assessor.full_name,
                    'email': assessor.email,
                    'status': f'✅ HTML Credentials sent & {assignments_for_this_assessor} new assignments',
                    'credentials': temp_password,
                    'assignments_count': assignments_for_this_assessor,
                    'credential_action': credential_action,
                    'academic_year': current_academic_year.year,
                    'is_new': is_new,
                    'is_new_year': is_new_year,
                })
                
            except Exception as e:
                print(f"❌ EMAIL SEND FAILED: {e}")
                import traceback
                traceback.print_exc()
                
                email_results.append({
                    'assessor': assessor.full_name,
                    'email': assessor.email,
                    'status': f'⚠️ Email failed - {assignments_for_this_assessor} assignments created',
                    'credentials': temp_password,
                    'assignments_count': assignments_for_this_assessor,
                    'is_new': is_new,
                    'is_new_year': is_new_year,
                    'error': str(e)[:100],
                    'note': 'MANUALLY SHARE THESE CREDENTIALS'
                })
            status_msg = f'ℹ️ {assignments_for_this_assessor} new assignments created'
            if not send_email:
                status_msg += ' (no email - already has credentials)'
            elif assignments_for_this_assessor == 0:
                status_msg = f'⚠️ No new assignments - all {len(schools)} schools already assigned'
            
            email_results.append({
                'assessor': assessor.full_name,
                'email': assessor.email,
                'status': status_msg,
                'credentials': temp_password if temp_password else 'Existing credentials',
                'assignments_count': assignments_for_this_assessor,
                'credential_action': credential_action,
                'is_new': False,
                'is_new_year': False,
            })
    
    # ========== FINAL STATISTICS ==========
    sent_count = email_sent_count
    failed_count = len([r for r in email_results if '❌' in r.get('status', '')])
    warning_count = len([r for r in email_results if '⚠️' in r.get('status', '')])
    
    print(f"\n{'='*60}")
    print(f"✅ BULK ASSIGNMENT PROCESS COMPLETE!")
    print(f"{'='*60}")
    print(f"📊 New Accounts Created: {new_accounts_count}")
    print(f"📊 Password Resets for New Year: {new_year_resets}")
    print(f"📧 Emails Sent Successfully: {sent_count}")
    print(f"⚠️  Emails Failed: {failed_count + warning_count}")
    print(f"📝 TOTAL NEW ASSIGNMENTS CREATED: {assignments_created}")
    print(f"📚 Academic Year: {current_academic_year.year}")
    print(f"{'='*60}")
    
    return {
        'total_assessors': len(assessor_ids),
        'total_schools': len(school_ids),
        'email_results': email_results,
        'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
        'academic_year': current_academic_year.year,
        'sent_count': sent_count,
        'failed_count': failed_count,
        'warning_count': warning_count,
        'new_accounts_count': new_accounts_count,
        'new_year_resets': new_year_resets,
        'assignments_created': assignments_created,
        'note': f'Successfully created {assignments_created} new assignments for {current_academic_year.year}'
    }
def get_or_create_student_profile(user):
    """Hakikisha kila user ana StudentTeacher profile. Cached for 2 minutes."""
    cache_key = f'student_profile_{user.pk}'
    profile = cache.get(cache_key)
    if profile is not None:
        return profile
    try:
        profile = StudentTeacher.objects.select_related(
            'selected_school', 'selected_school__district', 'selected_school__district__region'
        ).get(user=user)
    except StudentTeacher.DoesNotExist:
        email_username = user.email.split('@')[0] if user.email else user.username
        profile = StudentTeacher.objects.create(
            user=user,
            full_name=email_username,
            phone_number='Not provided'
        )
    cache.set(cache_key, profile, 120)
    return profile


def invalidate_student_cache(student):
    """Call after saving student profile changes."""
    cache.delete(f'student_profile_{student.user_id}')
    cache.delete(f'student_subjects_{student.id}')

def is_assessor(user):
    """Check if user is an assessor"""
    return hasattr(user, 'assessor')

def generate_random_password(length=12):
    """Generate random password for new assessors"""
    alphabet = string.ascii_letters + string.digits + "@#$%"
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def get_current_academic_year():
    """Get current academic year - FIXED: Jan-Dec = current_year/current_year+1"""
    current_date = timezone.now().date()
    current_year = current_date.year
    
    # LOGIC MPYA: Mwaka mzima (Jan-Dec) ni academic year current_year/current_year+1
    # Kwa mfano: Jan 2026 hadi Dec 2026 = 2026/2027
    academic_year_string = f"{current_year}/{current_year + 1}"
    
    print(f"🔍 Academic year based on {current_date}: {academic_year_string}")
    
    # Get OR create academic year
    try:
        academic_year = AcademicYear.objects.get(year=academic_year_string)
        print(f"✅ Found existing academic year: {academic_year.year}")
    except AcademicYear.DoesNotExist:
        print(f"⚠️ Academic year not found, creating: {academic_year_string}")
        academic_year = AcademicYear.objects.create(
            year=academic_year_string,
            is_active=True
        )
        # Set only this one as active
        AcademicYear.objects.exclude(id=academic_year.id).update(is_active=False)
        print(f"✅ Created new academic year: {academic_year.year}")
    
    # Double-check it's active
    if not academic_year.is_active:
        academic_year.is_active = True
        academic_year.save()
        print(f"🔧 Activated academic year: {academic_year.year}")
    
    return academic_year
# =========================
# AUTHENTICATION VIEWS
# =========================

def register(request):
    if request.method == 'POST':
        form = StudentRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.email = form.cleaned_data['email']
            user.set_password(form.cleaned_data['password1'])
            user.save()

            full_name = form.cleaned_data['full_name']
            phone_number = form.cleaned_data['phone_number']
            StudentTeacher.objects.create(user=user, full_name=full_name, phone_number=phone_number)

            messages.success(request, 'Account created successfully. Please login.')
            return redirect('login')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = StudentRegistrationForm()

    return render(request, 'field_app/registration/register.html', {
        'form': form,
        'hide_navbar': True
    })

def login_view(request):
    """Login view for STUDENTS ONLY"""
    
    # SPECIAL CASE: If assessor wants to login as student
    if 'assessor_logout' in request.GET and request.user.is_authenticated:
        logout(request)
        messages.info(request, "Logged out from assessor account. You can now login as student.")
        return redirect('login')
    
    # If already logged in, check user type
    if request.user.is_authenticated:
        try:
            assessor = Assessor.objects.get(user=request.user)
            return render(request, 'field_app/registration/login.html', {
                'assessor_warning': True,
                'assessor_name': assessor.full_name,
                'assessor_email': assessor.email,
                'logout_url': f"{request.path}?assessor_logout=true",
                'assessor_login_url': reverse('assessor_login'),
            })
        except Assessor.DoesNotExist:
            get_or_create_student_profile(request.user)
            return redirect('dashboard')
    
    if request.method == 'POST':
        form = CustomLoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()

            # Staff/Admin — must use admin login, not student form
            if user.is_staff:
                messages.error(request,
                    "Akaunti hii ni ya msimamizi. Tafadhali tumia ukurasa wa msimamizi."
                    if request.session.get('lang') == 'sw' else
                    "This is an admin account. Please use the admin login page."
                )
                return redirect('login')

            # Assessor — must use assessor login
            try:
                Assessor.objects.get(user=user)
                messages.warning(request,
                    "Akaunti hii ni ya mkaguzi. Tumia ukurasa wa mkaguzi kuingia."
                    if request.session.get('lang') == 'sw' else
                    "This is an assessor account. Please use the assessor login page."
                )
                return redirect('assessor_login')
            except Assessor.DoesNotExist:
                pass

            # Board member (DEO / HT / REO) — must use board login
            try:
                BoardMember.objects.get(user=user)
                messages.warning(request,
                    "Akaunti hii ni ya bodi. Tafadhali tumia ukurasa wa bodi kuingia."
                    if request.session.get('lang') == 'sw' else
                    "This is a board account. Please use the board login page."
                )
                return redirect('board_login')
            except BoardMember.DoesNotExist:
                pass

            login(request, user, backend='field_app.backends.EmailBackend')
            get_or_create_student_profile(user)
            messages.success(request,
                "Umeingia mfumo." if request.session.get('lang') == 'sw' else "Login successful."
            )
            return redirect('dashboard')
        else:
            messages.error(request, 'Invalid credentials')
    else:
        form = CustomLoginForm()

    return render(request, 'field_app/registration/login.html', {
        'form': form,
        'hide_navbar': True
    })

# views.py - Badilisha logout_view kwa hii

def logout_view(request):
    """Logout na upeleke kwenye login page sahihi kulingana na aina ya mtumiaji"""
    redirect_to = 'login'
    if request.user.is_authenticated:
        if request.user.is_staff:
            redirect_to = 'login_page'
        else:
            try:
                Assessor.objects.get(user=request.user)
                redirect_to = 'assessor_login'
            except Assessor.DoesNotExist:
                try:
                    BoardMember.objects.get(user=request.user)
                    redirect_to = 'board_login'
                except BoardMember.DoesNotExist:
                    redirect_to = 'login'
    logout(request)
    return redirect(redirect_to)
# =========================
# ASSESSOR LOGIN VIEW
# =========================

# views.py - FIX ASSESSOR LOGIN VIEW

# views.py - COMPLETE FIX FOR ASSESSOR LOGIN

def assessor_login(request):
    """Simple and fixed assessor login"""
    
    print(f"\n🔐 ASSESSOR LOGIN STARTED - Method: {request.method}")
    
    # Already logged in as assessor? Go to dashboard
    if request.user.is_authenticated:
        try:
            assessor = Assessor.objects.get(user=request.user)
            print(f"✅ Already logged in as: {assessor.full_name}")
            return redirect('assessor_dashboard')
        except Assessor.DoesNotExist:
            pass
    
    # Handle POST request (login attempt)
    if request.method == 'POST':
        email = request.POST.get('email', '').strip().lower()
        password = request.POST.get('password', '')
        
        print(f"📧 Login attempt: {email}")
        
        if not email or not password:
            messages.error(request, 'Please enter both email and password.')
            return render(request, 'field_app/assessor_login.html')
        
        # Authenticate the user
        user = authenticate(request, username=email, password=password)
        
        if user is None:
            # Try finding user by email
            try:
                user = User.objects.get(email__iexact=email)
                if user.check_password(password):
                    print(f"✅ Password check passed")
                    messages.error(request, 'Invalid email or password.')
                    return render(request, 'field_app/assessor_login.html')
            except User.DoesNotExist:
                messages.error(request, 'No account found with this email.')
                return render(request, 'field_app/assessor_login.html')
        
        # Check if user is an assessor
        try:
            assessor = Assessor.objects.get(user=user)
            print(f"✅ User is assessor: {assessor.full_name}")
            
            # Verify email matches
            if assessor.email.lower() != email.lower():
                messages.error(request, 
                    f'Email mismatch. This assessor is registered with: {assessor.email}'
                )
                return render(request, 'field_app/assessor_login.html')
            
            # LOGIN SUCCESSFUL
            login(request, user, backend='field_app.backends.EmailBackend')
            print(f"✅ Login successful, redirecting to dashboard")

            messages.success(request, f'Welcome Assessor {assessor.full_name}!')
            return redirect('assessor_dashboard')
            
        except Assessor.DoesNotExist:
            # Check if assessor exists with this email but different user
            try:
                assessor = Assessor.objects.get(email__iexact=email)
                print(f"⚠️ Assessor found but not linked: {assessor.email}")
                
                # Link assessor to this user
                assessor.user = user
                assessor.save()

                # Login
                login(request, user, backend='field_app.backends.EmailBackend')

                messages.success(request, f'Welcome Assessor {assessor.full_name}!')
                return redirect('assessor_dashboard')
                
            except Assessor.DoesNotExist:
                messages.error(request, 
                    'This email is not registered as an assessor. '
                    'Please use the student login page.'
                )
                return render(request, 'field_app/assessor_login.html')  # 🔥 ADDED THIS RETURN
    
    # GET request or failed login
    return render(request, 'field_app/assessor_login.html')
# =========================
# DASHBOARD VIEWS
# =========================


@login_required
def help_page(request):
    """Help and user guide page — shows relevant section based on user role."""
    user = request.user
    if user.is_staff:
        role = 'admin'
    elif Assessor.objects.filter(user=user).exists():
        role = 'assessor'
    elif BoardMember.objects.filter(user=user).exists():
        role = 'board'
    else:
        role = 'student'
    return render(request, 'field_app/help.html', {'role': role})


@login_required
def dashboard(request):
    """Student dashboard"""

    if request.user.is_staff:
        return redirect('admin_dashboard')

    try:
        Assessor.objects.get(user=request.user)
        return redirect('assessor_dashboard')
    except Assessor.DoesNotExist:
        pass

    try:
        BoardMember.objects.get(user=request.user)
        return redirect('board_home')
    except BoardMember.DoesNotExist:
        pass
    
    student = get_or_create_student_profile(request.user)
    current_year = _cached_active_year()

    pinned_regions = Region.objects.none()
    if current_year:
        pinned_region_ids = RegionPin.objects.filter(
            academic_year=current_year,
            is_pinned=True
        ).values_list('region_id', flat=True)
        pinned_regions = Region.objects.filter(id__in=pinned_region_ids)
    
    assessors = []
    if student.selected_school:
        qs = SchoolAssessment.objects.filter(school=student.selected_school)
        if current_year:
            qs = qs.filter(academic_year=current_year)
        for assessment in qs.select_related('assessor'):
            assessors.append({
                'assessor': assessment.assessor,
                'assignment_date': assessment.assessment_date,
                'is_completed': assessment.is_completed,
            })
    
    applications = []
    approved_applications_count = 0
    pending_applications_count = 0
    has_approved_applications = False
    school_has_completed_quota = False
    can_download_group_letter = False
    approved_students_count = 0
    group_letter_quota = 5
    
    if student:
        applications = StudentApplication.objects.filter(student=student).select_related('subject', 'school')
        approved_applications_count = applications.filter(status='approved').count()
        pending_applications_count = applications.filter(status='pending').count()
        has_approved_applications = approved_applications_count > 0

        # Auto-sync selected_school to match the actual application's school
        # Priority: approved application first, then pending
        canonical_app = (
            applications.filter(status='approved').first() or
            applications.filter(status='pending').first()
        )
        if canonical_app and student.selected_school_id != canonical_app.school_id:
            old_school = student.selected_school
            student.selected_school = canonical_app.school
            student.save(update_fields=['selected_school'])
            if old_school:
                School.objects.filter(id=old_school.id).update(current_students=Greatest(F('current_students') - 1, 0))
            School.objects.filter(id=canonical_app.school_id).update(current_students=F('current_students') + 1)
            invalidate_student_cache(student)

        if student.selected_school:
            school = student.selected_school

            approved_students_count = StudentApplication.objects.filter(
                school=school,
                status='approved'
            ).count()

            school_has_completed_quota = approved_students_count >= group_letter_quota
            can_download_group_letter = school_has_completed_quota and has_approved_applications

    logbook_entries = []
    if student:
        logbook_entries = LogbookEntry.objects.filter(
            student=student
        ).select_related('subject_taught').order_by('-date')[:5]

    board_comments = []
    if student:
        board_comments = BoardComment.objects.filter(
            student=student
        ).select_related('board_member').order_by('-created_at')[:5]
        BoardComment.objects.filter(student=student, is_read=False).update(is_read=True)

    # Monthly reports visible to the student — find their entry in each report
    my_monthly_reports = []
    if student and student.selected_school:
        district = student.selected_school.district
        reports = MonthlyReport.objects.filter(district=district).order_by('-year', '-month')[:6]
        for report in reports:
            content = report.ai_content or {}
            student_entry = None
            # Search primary and secondary lists for this student by name
            for level_key in ('primary', 'secondary'):
                for entry in content.get(level_key, []):
                    if entry.get('jina', '').strip().lower() == student.full_name.strip().lower():
                        student_entry = entry
                        break
                if student_entry:
                    break
            my_monthly_reports.append({
                'report': report,
                'entry': student_entry,   # None if student had no logbook that month
            })

    return render(request, 'field_app/dashboard.html', {
        'regions': pinned_regions,
        'current_year': current_year,
        'student': student,
        'applications': applications,
        'approved_applications_count': approved_applications_count,
        'pending_applications_count': pending_applications_count,
        'has_approved_applications': has_approved_applications,
        'school_has_completed_quota': school_has_completed_quota,
        'can_download_group_letter': can_download_group_letter,
        'approved_students_count': approved_students_count,
        'group_letter_quota': group_letter_quota,
        'logbook_entries': logbook_entries,
        'assessors': assessors,
        'board_comments': board_comments,
        'my_monthly_reports': my_monthly_reports,
    })

@login_required
def student_monthly_report(request, report_id):
    """Show a monthly report to the student who is in that district."""
    student = get_or_create_student_profile(request.user)
    report = get_object_or_404(MonthlyReport, id=report_id)

    # Ensure the student belongs to this district
    if not student.selected_school or student.selected_school.district_id != report.district_id:
        messages.error(request, 'Huna ruhusa ya kuona ripoti hii.')
        return redirect('dashboard')

    content = report.ai_content or {}
    student_entry = None
    for level_key in ('primary', 'secondary'):
        for entry in content.get(level_key, []):
            if entry.get('jina', '').strip().lower() == student.full_name.strip().lower():
                student_entry = entry
                break
        if student_entry:
            break

    return render(request, 'field_app/student_monthly_report.html', {
        'report': report,
        'student': student,
        'student_entry': student_entry,
        'muhtasari': content.get('muhtasari', ''),
        'hitimisho': content.get('hitimisho', ''),
    })


# views.py - SAHIHISHA SEHEMU YA ASSESSOR DASHBOARD

@assessor_login_required
def assessor_dashboard(request):
    """Dashboard ya Assessor - FIXED FIELD ERROR"""
    try:
        assessor = Assessor.objects.get(user=request.user)
    except Assessor.DoesNotExist:
        messages.error(request, "You are not registered as an assessor.")
        return redirect('dashboard')
    
    current_year = get_current_academic_year()
    
    print(f"\n🔍 ASSESSOR DASHBOARD: {assessor.full_name}")
    
    # Get assignments for this assessor
    school_assignments = SchoolAssessment.objects.filter(
        assessor=assessor,
        academic_year=current_year
    ).select_related('school', 'school__district', 'school__district__region')
    
    schools_data = []
    total_students_all_schools = 0
    
    for assignment in school_assignments:
        school = assignment.school
        
        print(f"\n🏫 Processing school: {school.name}")
        
        # 🔴 FIX 1: Get students via TWO separate queries
        from django.db.models import Q
        
        # Method 1: Students who selected this school
        students_via_selected = StudentTeacher.objects.filter(
            selected_school=school
        )
        
        # Method 2: Students with approved applications for this school
        approved_app_student_ids = StudentApplication.objects.filter(
            school=school,
            status='approved'
        ).values_list('student_id', flat=True).distinct()
        
        students_via_applications = StudentTeacher.objects.filter(
            id__in=approved_app_student_ids
        )
        
        # Combine both querysets
        student_ids = set()
        
        for student in students_via_selected:
            student_ids.add(student.id)
        
        for student in students_via_applications:
            student_ids.add(student.id)
        
        # Get all unique students
        all_students = StudentTeacher.objects.filter(
            id__in=list(student_ids)
        ).select_related('user')
        
        print(f"   Found {len(student_ids)} unique students")
        
        # Get student assessments for this school by this assessor
        student_assessments = StudentAssessment.objects.filter(
            assessor=assessor,
            school=school,
            academic_year=current_year
        ).select_related('student')
        
        # Create assessment map
        assessment_map = {}
        for sa in student_assessments:
            if sa.student:
                assessment_map[sa.student.id] = sa
        
        # Prepare detailed student data
        students_data = []
        for student in all_students:
            # Get student's approved applications for this school
            approved_apps = StudentApplication.objects.filter(
                student=student,
                school=school,
                status='approved'
            ).select_related('subject')
            
            approved_subjects = [app.subject.name for app in approved_apps]
            
            # Check if student selected this school
            has_selected_school = (student.selected_school == school)
            
            # Get assessment
            assessment = assessment_map.get(student.id)
            
            # Logbook stats for this student
            logbook_entries = LogbookEntry.objects.filter(student=student)
            logbook_total = logbook_entries.count()
            logbook_verified = logbook_entries.filter(is_location_verified=True).count()
            logbook_this_week = logbook_entries.filter(
                date__gte=timezone.now().date() - timedelta(days=7)
            ).count()
            logbook_last = logbook_entries.order_by('-date').first()

            students_data.append({
                'student': student,
                'has_selected_school': has_selected_school,
                'has_approved_application': approved_apps.exists(),
                'approved_subjects': approved_subjects,
                'approved_apps_count': approved_apps.count(),
                'assessment': assessment,
                'is_completed': assessment.score is not None if assessment else False,
                'score': assessment.score if assessment else None,
                'email': student.user.email if student.user else "No email",
                'phone': student.phone_number or "Not provided",
                'logbook_total': logbook_total,
                'logbook_verified': logbook_verified,
                'logbook_this_week': logbook_this_week,
                'logbook_last_date': logbook_last.date if logbook_last else None,
            })
            
            # Debug print
            if has_selected_school or approved_apps.exists():
                print(f"   👤 {student.full_name}:")
                if has_selected_school:
                    print(f"      ✅ Selected this school")
                if approved_apps.exists():
                    print(f"      ✅ Approved subjects: {', '.join(approved_subjects)}")
        
        students_count = len(students_data)
        total_students_all_schools += students_count
        
        # Assessment counts
        completed_student_assessments = len([s for s in students_data if s['is_completed']])
        pending_student_assessments = students_count - completed_student_assessments
        
        # Get other assessors
        other_assessors = []
        if current_year:
            other_assignments = SchoolAssessment.objects.filter(
                school=school,
                academic_year=current_year
            ).exclude(assessor=assessor).select_related('assessor')
            
            for other_assignment in other_assignments:
                other_assessors.append({
                    'name': other_assignment.assessor.full_name,
                    'email': other_assignment.assessor.email,
                })
        
        schools_data.append({
            'school': school,
            'assignment': assignment,
            'students': students_data,
            'students_count': students_count,
            'other_assessors': other_assessors,
            'completed_student_assessments': completed_student_assessments,
            'pending_student_assessments': pending_student_assessments,
            'academic_year': current_year.year if current_year else "Not Set",
        })
    
    total_completed = sum(d['completed_student_assessments'] for d in schools_data)

    return render(request, 'field_app/assessor_dashboard.html', {
        'assessor': assessor,
        'schools_data': schools_data,
        'total_schools': school_assignments.count(),
        'total_students': total_students_all_schools,
        'total_completed': total_completed,
        'current_year': current_year,
    })
@login_required
def select_region(request):
    """Show ONLY regions that are NOT pinned for current academic year"""
    current_year = get_current_academic_year()
    
    if not current_year:
        messages.error(request, "No active academic year found!")
        return redirect('dashboard')
    
    # Get PINNED region IDs (regions to HIDE from students)
    pinned_region_ids = RegionPin.objects.filter(
        academic_year=current_year,
        is_pinned=True  # Pinned = HIDDEN
    ).values_list('region_id', flat=True)
    
    # Show ONLY regions that are NOT pinned
    available_regions = Region.objects.exclude(
        id__in=pinned_region_ids
    ).order_by('name')
    
    print(f"📅 Academic Year: {current_year.year}")
    print(f"🔒 Pinned (Hidden) Regions: {pinned_region_ids.count()}")
    print(f"✅ Available Regions: {available_regions.count()}")
    
    # Debug: Print region names
    for region in available_regions:
        print(f"   - {region.name}")
    
    return render(request, 'field_app/select_region.html', {
        'regions': available_regions,
        'current_year': current_year,
        'pinned_count': pinned_region_ids.count(),
        'available_count': available_regions.count(),
    })

@login_required
def select_district(request, region_id):
    region = get_object_or_404(Region, id=region_id)
    districts = District.objects.filter(region=region)
    current_year = _cached_active_year()

    # Fetch allocations for all districts in this region at once
    alloc_qs = DistrictAllocation.objects.filter(district__region=region)
    if current_year:
        alloc_qs = alloc_qs.filter(academic_year=current_year)
    alloc_map = {a.district_id: a for a in alloc_qs}

    for district in districts:
        district.school_count = School.objects.filter(district=district).count()
        alloc = alloc_map.get(district.id)
        district.allocation = alloc  # may be None if DEO hasn't set it yet

    request.session['selected_region_id'] = region.id

    return render(request, 'field_app/select_district.html', {
        'districts': districts,
        'region': region,
    })

@login_required
def select_school(request, district_id):
    district = get_object_or_404(District, id=district_id)
    current_year = _cached_active_year()
    
    # Get pinned schools
    pinned_school_ids = []
    if current_year:
        pinned_school_ids = list(SchoolPin.objects.filter(
            academic_year=current_year, is_pinned=True
        ).values_list('school_id', flat=True))
    
    # Get parameters
    search_query = request.GET.get('q', '')
    selected_level = request.GET.get('level', 'Secondary')
    
    # Get schools
    schools_qs = School.objects.filter(district=district, level=selected_level)
    if search_query:
        schools_qs = schools_qs.filter(name__icontains=search_query)
    
    # Fetch district allocation and per-school quotas for display
    da_qs = DistrictAllocation.objects.filter(district=district)
    if current_year:
        da_qs = da_qs.filter(academic_year=current_year)
    district_alloc = da_qs.first()
    school_alloc_map = {}
    if district_alloc:
        school_alloc_map = {
            sa.school_id: sa
            for sa in SchoolAllocation.objects.filter(district_allocation=district_alloc)
        }

    schools = []
    for school in schools_qs:
        school.is_pinned = school.id in pinned_school_ids
        school.is_selectable = (not school.is_pinned) and (school.current_students < school.capacity)
        school.occupancy_percentage = round((school.current_students / school.capacity) * 100) if school.capacity > 0 else 0
        sa = school_alloc_map.get(school.id)
        school.deo_quota = sa.quota if sa else None
        school.deo_filled = sa.filled if sa else None
        school.deo_remaining = sa.remaining if sa else None
        schools.append(school)

    total_schools = len(schools)
    pinned_schools_count = sum(1 for s in schools if s.is_pinned)
    available_schools_count = sum(1 for s in schools if s.is_selectable)
    full_schools_count = total_schools - pinned_schools_count - available_schools_count
    
    # ========== SIMPLE SESSION LOGIC ==========
    selected_school = None
    temp_selected_id = request.session.get('temp_selected_school_id')
    if temp_selected_id:
        try:
            selected_school = School.objects.get(id=temp_selected_id)
        except School.DoesNotExist:
            request.session.pop('temp_selected_school_id', None)
    
    # ========== HANDLE POST ==========
    if request.method == 'POST':
        action = request.POST.get('action')
        school_id = request.POST.get('school_id')
        
        if action == 'select' and school_id:
            school = get_object_or_404(School, id=school_id)
            request.session['temp_selected_school_id'] = school.id
            messages.info(request, f'Umchagua {school.name}. Bonyeza Thibitisha au Ghairi.')
            return redirect(f"{request.path}?level={selected_level}&q={search_query}")
        
        elif action == 'confirm':
            if temp_selected_id:
                try:
                    school = School.objects.get(id=temp_selected_id)
                    student = get_or_create_student_profile(request.user)

                    is_changing_school = student.selected_school and student.selected_school.id != school.id

                    if is_changing_school:
                        # Block if student has any approved application — cannot change school after approval
                        approved_apps = StudentApplication.objects.filter(student=student, status='approved')
                        if approved_apps.exists():
                            messages.error(request,
                                'Huwezi kubadili shule. Ombi lako limeshaidhinishwa tayari.'
                                if request.session.get('lang') == 'sw' else
                                'You cannot change school. You already have an approved application.'
                            )
                            request.session.pop('temp_selected_school_id', None)
                            return redirect('dashboard')

                        # Delete pending applications at old school before switching
                        StudentApplication.objects.filter(student=student, status='pending').delete()
                        School.objects.filter(id=student.selected_school.id).update(current_students=Greatest(F('current_students') - 1, 0))

                    elif student.selected_school:
                        School.objects.filter(id=student.selected_school.id).update(current_students=Greatest(F('current_students') - 1, 0))

                    student.selected_school = school
                    student.save()
                    invalidate_student_cache(student)
                    School.objects.filter(id=school.id).update(current_students=F('current_students') + 1)

                    request.session.pop('temp_selected_school_id', None)
                    messages.success(request, f'Shule imethibitishwa: {school.name}')
                    return redirect('select_subjects', school_id=school.id)
                except School.DoesNotExist:
                    messages.error(request, 'Shule haipo')
                messages.error(request, 'Hakuna shule iliyochaguliwa')
            return redirect(f"{request.path}?level={selected_level}&q={search_query}")
        
        elif action == 'cancel':
            request.session.pop('temp_selected_school_id', None)
            messages.success(request, 'Umeghairi uchaguzi wa shule')
            return redirect(f"{request.path}?level={selected_level}&q={search_query}")
    
    return render(request, 'field_app/select_school.html', {
        'district': district,
        'schools': schools,
        'selected_school': selected_school,
        'query': search_query,
        'selected_level': selected_level,
        'total_schools': total_schools,
        'pinned_schools_count': pinned_schools_count,
        'available_schools_count': available_schools_count,
        'full_schools_count': full_schools_count,
        'current_year': current_year,
        'district_alloc': district_alloc,
    })
@login_required
def select_subjects(request, school_id):
    school = get_object_or_404(School, id=school_id)
    subject_capacities = SchoolSubjectCapacity.objects.filter(school=school).select_related('subject')
    
    student = get_or_create_student_profile(request.user)
    
    existing_applications = StudentApplication.objects.filter(
        student=student, 
        school=school
    ).select_related('subject')
    
    applied_subject_ids = {app.subject.id for app in existing_applications}

    if request.method == 'POST':
        subject_id = request.POST.get('subject_id')
        action = request.POST.get('action')

        if not subject_id:
            messages.error(request, "No subject selected.")
            return redirect('select_subjects', school_id=school.id)

        try:
            subject = Subject.objects.get(id=subject_id)
        except Subject.DoesNotExist:
            messages.error(request, "Subject does not exist.")
            return redirect('select_subjects', school_id=school.id)

        try:
            capacity = SchoolSubjectCapacity.objects.get(school=school, subject=subject)
        except SchoolSubjectCapacity.DoesNotExist:
            messages.error(request, f"{subject.name} is not available at this school.")
            return redirect('select_subjects', school_id=school.id)

        if action == 'apply':
            existing_application = StudentApplication.objects.filter(
                student=student,
                subject=subject,
                school=school
            ).first()

            # Enforce one school, one subject rule
            any_other_application = StudentApplication.objects.filter(
                student=student
            ).exclude(subject=subject, school=school).exists()

            if existing_application:
                messages.info(request, f"Tayari uliomba {subject.name}." if request.session.get('lang') == 'sw' else f"You have already applied for {subject.name}.")
            elif any_other_application:
                messages.error(request,
                    "Una ombi tayari katika shule/somo lingine. Unaweza kuomba shule moja na somo moja tu. Ghairi ombi lako la kwanza kwanza."
                    if request.session.get('lang') == 'sw' else
                    "You already have an application elsewhere. Only one school and one subject is allowed. Cancel your existing application first."
                )
            elif capacity.current_students >= capacity.max_students:
                messages.error(request, f"{subject.name} is already full.")
            else:
                StudentApplication.objects.create(
                    student=student,
                    subject=subject,
                    school=school,
                    status='pending'
                )
                messages.success(request,
                    f"✅ Ombi la {subject.name} limetumwa! Linasubiri idhini ya Admin."
                    if request.session.get('lang') == 'sw' else
                    f"✅ Application for {subject.name} submitted successfully! Waiting for Admin approval."
                )

        elif action == 'cancel_application':
            application = StudentApplication.objects.filter(
                student=student,
                subject=subject,
                school=school
            ).first()

            if application:
                application.delete()
                messages.success(request, f"Application for {subject.name} cancelled.")
            else:
                messages.error(request, f"Cannot cancel application for {subject.name}.")

        return redirect('select_subjects', school_id=school.id)

    return render(request, 'field_app/select_subjects.html', {
        'school': school,
        'subject_capacities': subject_capacities,
        'existing_applications': existing_applications,
        'applied_subject_ids': applied_subject_ids,
    })

@login_required
def apply_for_subject(request, subject_id, school_id):
    subject = get_object_or_404(Subject, id=subject_id)
    school = get_object_or_404(School, id=school_id)
    
    student = get_or_create_student_profile(request.user)
    
    existing_application = StudentApplication.objects.filter(
        student=student,
        subject=subject,
        school=school
    ).first()

    if existing_application:
        messages.info(request, f"You have already applied for {subject.name}")
        return redirect('select_subjects', school_id=school.id)

    # Enforce one school, one subject rule
    any_other_application = StudentApplication.objects.filter(
        student=student
    ).exclude(subject=subject, school=school).exists()
    if any_other_application:
        messages.error(request,
            "Una ombi tayari katika shule/somo lingine. Unaweza kuomba shule moja na somo moja tu."
            if request.session.get('lang') == 'sw' else
            "You already have an application elsewhere. Only one school and one subject is allowed."
        )
        return redirect('select_subjects', school_id=school.id)

    try:
        capacity = SchoolSubjectCapacity.objects.get(school=school, subject=subject)
        if capacity.current_students >= capacity.max_students:
            messages.error(request, f"{subject.name} is already full at {school.name}")
            return redirect('select_subjects', school_id=school.id)
    except SchoolSubjectCapacity.DoesNotExist:
        messages.error(request, f"{subject.name} is not available at {school.name}")
        return redirect('select_subjects', school_id=school.id)

    StudentApplication.objects.create(
        student=student,
        subject=subject,
        school=school,
        status='pending'
    )

    messages.success(request, f"Application for {subject.name} submitted successfully! Waiting for approval.")
    return redirect('dashboard')

# =========================
# LOGBOOK VIEWS
# =========================

@login_required
def submit_logbook(request):
    student = get_or_create_student_profile(request.user)
    today = timezone.now().date()
    
    # Check if weekend
    if today.weekday() >= 5:  # Saturday (5) or Sunday (6)
        messages.info(request, "Hakuna kazi ya uwanjani wikendi. Rudi tena Jumatatu.")
        return redirect('dashboard')
    
    if not student.selected_school:
        messages.error(request, "Lazima uchague shule kabla ya kujaza logbook.")
        return redirect('select_region')
    
    school = student.selected_school

    logbook_entry = _cached_today_logbook(student, school, today)
    
    if request.method == 'POST':
        form = LogbookForm(request.POST, instance=logbook_entry)

        latitude = request.POST.get('latitude')
        longitude = request.POST.get('longitude')
        is_location_verified = request.POST.get('is_location_verified', 'false') == 'true'

        days_swahili = {0: 'Jumatatu', 1: 'Jumanne', 2: 'Jumatano', 3: 'Alhamisi', 4: 'Ijumaa'}

        def _render_logbook(extra=None):
            subjects = _cached_subjects(student)
            ctx = {
                'form': form, 'student': student, 'logbook_entry': logbook_entry,
                'today': today, 'today_name': days_swahili.get(today.weekday(), 'Leo'),
                'school': school, 'subjects': subjects,
            }
            if extra:
                ctx.update(extra)
            return render(request, 'field_app/logbook.html', ctx)

        if not is_location_verified:
            messages.error(request, "Thibiti eneo lako kwanza kabla ya kuwasilisha logbook.")
            return _render_logbook({'location_error': True})

        if not latitude or not longitude:
            messages.error(request, "Eneo halipatikani. Washa GPS na ujaribu tena.")
            return _render_logbook({'location_error': True})

        try:
            lat = float(latitude)
            lng = float(longitude)
            logbook_entry.latitude = lat
            logbook_entry.longitude = lng
            logbook_entry.location_address = request.POST.get('location_address', '')

            if school and school.latitude and school.longitude:
                from geopy.distance import geodesic
                distance_m = geodesic((lat, lng), (school.latitude, school.longitude)).meters
                school_verified = distance_m <= 1000
            else:
                school_verified = (-11.8 <= lat <= -1.0) and (29.3 <= lng <= 40.5)

            logbook_entry.is_location_verified = school_verified
            logbook_entry.is_at_school = school_verified

        except (ValueError, TypeError) as e:
            messages.error(request, f"Hitilafu ya eneo: {e}")
            return _render_logbook()

        if form.is_valid():
            entry = form.save(commit=False)

            import json as _json
            try:
                entry.lessons_data = _json.loads(request.POST.get('lessons_data', '[]'))
            except (ValueError, TypeError):
                entry.lessons_data = []

            entry.latitude = logbook_entry.latitude
            entry.longitude = logbook_entry.longitude
            entry.is_location_verified = logbook_entry.is_location_verified
            entry.is_at_school = logbook_entry.is_at_school
            entry.location_address = logbook_entry.location_address
            entry.save()
            _invalidate_today_logbook(student, today)

            # Send email in background thread — never block the response
            import threading
            def _notify(school_id, student_name, entry_id, verified):
                try:
                    from .models import SchoolAssessment, AcademicYear
                    from django.core.mail import send_mail as _send
                    from django.conf import settings as _s
                    yr = AcademicYear.objects.filter(is_active=True).first()
                    asgn = SchoolAssessment.objects.filter(
                        school_id=school_id, academic_year=yr
                    ).select_related('assessor').first()
                    if asgn and asgn.assessor.email:
                        loc = "GPS Verified" if verified else "Not Verified"
                        _send(
                            subject=f"[IMS] Logbook Submitted — {student_name}",
                            message=(
                                f"Student {student_name} amewasilisha logbook.\n"
                                f"Eneo: {loc}\n\nIngia IMS kuona maelezo."
                            ),
                            from_email=_s.DEFAULT_FROM_EMAIL,
                            recipient_list=[asgn.assessor.email],
                            fail_silently=True,
                        )
                except Exception:
                    pass
            threading.Thread(
                target=_notify,
                args=(school.id, student.full_name, entry.id, entry.is_location_verified),
                daemon=True,
            ).start()

            if entry.is_location_verified:
                messages.success(request, "✅ Logbook imesajiliwa kikamilifu!")
            else:
                messages.warning(request, "⚠️ Logbook imesajiliwa. Eneo halikuthibitishwa.")

            return redirect('logbook_history')
        else:
            messages.error(request, "Tafadhali kagua makosa yaliyomo kwenye fomu.")
    else:
        form = LogbookForm(instance=logbook_entry)

    subjects = _cached_subjects(student)
    days_swahili = {0: 'Jumatatu', 1: 'Jumanne', 2: 'Jumatano', 3: 'Alhamisi', 4: 'Ijumaa'}

    return render(request, 'field_app/logbook.html', {
        'form': form,
        'student': student,
        'logbook_entry': logbook_entry,
        'today': today,
        'today_name': days_swahili.get(today.weekday(), 'Leo'),
        'school': school,
        'subjects': subjects,
    })
@login_required
def logbook_history(request):
    student = get_or_create_student_profile(request.user)

    week_filter = request.GET.get('week')
    month_filter = request.GET.get('month')

    entries = LogbookEntry.objects.filter(student=student).select_related('subject_taught', 'school')

    if week_filter:
        try:
            year, week = map(int, week_filter.split('-W'))
            start_date = datetime.strptime(f'{year}-W{week}-1', "%Y-W%W-%w").date()
            entries = entries.filter(date__range=[start_date, start_date + timedelta(days=6)])
        except ValueError:
            messages.error(request, "Tarehe ya wiki si sahihi.")

    elif month_filter:
        try:
            year, month = map(int, month_filter.split('-'))
            entries = entries.filter(date__year=year, date__month=month)
        except ValueError:
            messages.error(request, "Tarehe ya mwezi si sahihi.")

    else:
        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        entries = entries.filter(date__range=[start_of_week, start_of_week + timedelta(days=4)])

    return render(request, 'field_app/logbook_history.html', {
        'entries': entries.order_by('-date'),
        'student': student,
    })

@login_required
def download_logbook_pdf(request, period=None, period_type=None):
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm

    student = get_or_create_student_profile(request.user)
    period_value = period or period_type or 'week'
    today = timezone.now().date()

    if period_value == 'today':
        entries = LogbookEntry.objects.filter(student=student, date=today)
        filename = f"logbook_{today}.pdf"
        title = f"Logbook — {today}"
    elif period_value == 'week':
        start_of_week = today - timedelta(days=today.weekday())
        end_of_week = start_of_week + timedelta(days=4)
        entries = LogbookEntry.objects.filter(student=student, date__range=[start_of_week, end_of_week])
        filename = f"logbook_wiki_{start_of_week}.pdf"
        title = f"Logbook — Wiki {start_of_week} hadi {end_of_week}"
    elif period_value == 'month':
        start_of_month = today.replace(day=1)
        next_month = today.replace(day=28) + timedelta(days=4)
        end_of_month = next_month - timedelta(days=next_month.day)
        entries = LogbookEntry.objects.filter(student=student, date__range=[start_of_month, end_of_month])
        filename = f"logbook_mwezi_{today.year}_{today.month:02d}.pdf"
        title = f"Logbook — Mwezi {today.month}/{today.year}"
    else:
        entries = LogbookEntry.objects.filter(student=student)
        filename = f"logbook_zote_{today}.pdf"
        title = "Logbook Zote"

    entries = entries.order_by('date')

    NAVY = rl_colors.HexColor('#0A2B5E')
    GOLD = rl_colors.HexColor('#C8900A')
    LIGHT = rl_colors.HexColor('#EEF1F6')
    WHITE = rl_colors.white

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
    )

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=14, textColor=NAVY, spaceAfter=4)
    s_sub = ParagraphStyle('sub', fontName='Helvetica', fontSize=9, textColor=rl_colors.HexColor('#4A5568'), spaceAfter=2)
    s_head = ParagraphStyle('head', fontName='Helvetica-Bold', fontSize=10, textColor=WHITE)
    s_label = ParagraphStyle('label', fontName='Helvetica-Bold', fontSize=8, textColor=NAVY)
    s_body = ParagraphStyle('body', fontName='Helvetica', fontSize=8, textColor=rl_colors.HexColor('#1A1A2E'), leading=11)
    s_small = ParagraphStyle('small', fontName='Helvetica', fontSize=7.5, textColor=rl_colors.HexColor('#4A5568'), leading=10)

    school_name = student.selected_school.name if student.selected_school else '—'
    district_name = (student.selected_school.district.name if student.selected_school and student.selected_school.district else '—')
    current_year = _cached_active_year()
    year_label = str(current_year) if current_year else '—'

    story = []

    # ── Cover header ──────────────────────────────────────────────────────────
    story.append(Paragraph("WIZARA YA ELIMU, SAYANSI NA TEKNOLOJIA", ParagraphStyle('gov', fontName='Helvetica-Bold', fontSize=11, textColor=NAVY, alignment=1)))
    story.append(Paragraph("Mfumo wa Ufuatiliaji wa Walimu Wanafunzi (IMS)", ParagraphStyle('gov2', fontName='Helvetica', fontSize=9, textColor=rl_colors.HexColor('#4A5568'), alignment=1, spaceAfter=6)))
    story.append(HRFlowable(width="100%", thickness=2, color=GOLD, spaceAfter=6))
    story.append(Paragraph(title, s_title))

    info_data = [
        [Paragraph('<b>Jina la Mwanafunzi:</b>', s_label), Paragraph(student.full_name, s_body),
         Paragraph('<b>Shule:</b>', s_label), Paragraph(school_name, s_body)],
        [Paragraph('<b>Wilaya:</b>', s_label), Paragraph(district_name, s_body),
         Paragraph('<b>Mwaka wa Masomo:</b>', s_label), Paragraph(year_label, s_body)],
        [Paragraph('<b>Tarehe ya Chapisha:</b>', s_label), Paragraph(str(today), s_body),
         Paragraph('<b>Idadi ya Siku:</b>', s_label), Paragraph(str(entries.count()), s_body)],
    ]
    info_tbl = Table(info_data, colWidths=[3.5*cm, 5.5*cm, 3.5*cm, 5.5*cm])
    info_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), LIGHT),
        ('BOX', (0,0), (-1,-1), 0.5, NAVY),
        ('INNERGRID', (0,0), (-1,-1), 0.3, rl_colors.HexColor('#CBD5E0')),
        ('PADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 10))

    # ── Per-entry ─────────────────────────────────────────────────────────────
    day_names = {'Monday':'Jumatatu','Tuesday':'Jumanne','Wednesday':'Jumatano',
                 'Thursday':'Alhamisi','Friday':'Ijumaa','Saturday':'Jumamosi','Sunday':'Jumapili'}

    for entry in entries:
        day_sw = day_names.get(entry.date.strftime('%A'), entry.date.strftime('%A'))
        gps_status = "✓ Imehakikiwa" if entry.is_location_verified else "✗ Haijahakikiwa"
        gps_color = rl_colors.HexColor('#10B981') if entry.is_location_verified else rl_colors.HexColor('#F59E0B')

        # Day header row
        day_hdr = Table(
            [[Paragraph(f"{day_sw} — {entry.date}", s_head),
              Paragraph(f"GPS: {gps_status}", ParagraphStyle('gps', fontName='Helvetica-Bold', fontSize=8, textColor=gps_color))]],
            colWidths=[13*cm, 5*cm]
        )
        day_hdr.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), NAVY),
            ('PADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (1,0), (1,0), 'RIGHT'),
        ]))
        story.append(day_hdr)

        # Lesson periods
        lessons = entry.lessons_data if entry.lessons_data else []
        if not lessons:
            # Fallback to legacy fields
            if entry.morning_activity or entry.afternoon_activity:
                legacy_data = [
                    [Paragraph('<b>Shughuli za Asubuhi</b>', s_label),
                     Paragraph(entry.morning_activity or '—', s_body)],
                    [Paragraph('<b>Shughuli za Mchana</b>', s_label),
                     Paragraph(entry.afternoon_activity or '—', s_body)],
                ]
                leg_tbl = Table(legacy_data, colWidths=[4*cm, 14*cm])
                leg_tbl.setStyle(TableStyle([
                    ('BOX', (0,0), (-1,-1), 0.5, rl_colors.HexColor('#CBD5E0')),
                    ('INNERGRID', (0,0), (-1,-1), 0.3, rl_colors.HexColor('#CBD5E0')),
                    ('BACKGROUND', (0,0), (0,-1), LIGHT),
                    ('PADDING', (0,0), (-1,-1), 4),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(leg_tbl)
            else:
                story.append(Paragraph("Hakuna data ya masomo.", s_small))
        else:
            for lesson in lessons:
                period_num = lesson.get('period', '?')
                subj = lesson.get('subject', '—')
                cls = lesson.get('class', '—')
                main_topic = lesson.get('main_topic', '—')
                subtopic = lesson.get('subtopic', '') or '—'
                enrolled = lesson.get('enrolled', '—')
                present = lesson.get('present', '—')
                activity = lesson.get('activity_type', '—')
                methods = lesson.get('methods', '') or '—'
                aids = lesson.get('teaching_aids', '') or '—'
                intro = lesson.get('introduction', '') or '—'
                development = lesson.get('development', '') or '—'
                conclusion = lesson.get('conclusion', '') or '—'
                assessment = lesson.get('assessment', '') or '—'
                homework = lesson.get('homework', '') or '—'

                # Period header
                period_hdr = Table(
                    [[Paragraph(f"Kipindi {period_num}  |  {subj}  |  Darasa: {cls}  |  Waliojumuishwa: {enrolled}  |  Waliopo: {present}  |  Aina: {activity}",
                               ParagraphStyle('ph', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE))]],
                    colWidths=[18*cm]
                )
                period_hdr.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), GOLD),
                    ('PADDING', (0,0), (-1,-1), 5),
                ]))
                story.append(period_hdr)

                lesson_rows = [
                    [Paragraph('<b>Mada Kuu</b>', s_label), Paragraph(main_topic, s_body),
                     Paragraph('<b>Mada Ndogo</b>', s_label), Paragraph(subtopic, s_body)],
                    [Paragraph('<b>Mbinu</b>', s_label), Paragraph(methods, s_body),
                     Paragraph('<b>Vifaa vya Kufundishia</b>', s_label), Paragraph(aids, s_body)],
                    [Paragraph('<b>Utangulizi</b>', s_label), Paragraph(intro, s_small),
                     Paragraph('<b>Hitimisho</b>', s_label), Paragraph(conclusion, s_small)],
                    [Paragraph('<b>Uendelezaji</b>', s_label), Paragraph(development, s_small),
                     Paragraph('<b>Tathmini</b>', s_label), Paragraph(assessment, s_small)],
                    [Paragraph('<b>Kazi ya Nyumbani</b>', s_label), Paragraph(homework, s_small),
                     Paragraph('', s_label), Paragraph('', s_body)],
                ]
                lesson_tbl = Table(lesson_rows, colWidths=[3.5*cm, 5.5*cm, 3.5*cm, 5.5*cm])
                lesson_tbl.setStyle(TableStyle([
                    ('BOX', (0,0), (-1,-1), 0.5, rl_colors.HexColor('#CBD5E0')),
                    ('INNERGRID', (0,0), (-1,-1), 0.3, rl_colors.HexColor('#CBD5E0')),
                    ('BACKGROUND', (0,0), (0,-1), LIGHT),
                    ('BACKGROUND', (2,0), (2,-1), LIGHT),
                    ('PADDING', (0,0), (-1,-1), 4),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(lesson_tbl)

        # Other activities / challenges / reflection
        extra_rows = []
        if entry.other_activities:
            extra_rows.append([Paragraph('<b>Shughuli Nyingine</b>', s_label), Paragraph(entry.other_activities, s_body)])
        if entry.challenges_faced:
            extra_rows.append([Paragraph('<b>Changamoto</b>', s_label), Paragraph(entry.challenges_faced, s_body)])
        if entry.lessons_learned:
            extra_rows.append([Paragraph('<b>Tafakuri</b>', s_label), Paragraph(entry.lessons_learned, s_body)])
        if entry.supervisor_remarks:
            extra_rows.append([Paragraph('<b>Maoni ya Msimamizi</b>', ParagraphStyle('sr', fontName='Helvetica-Bold', fontSize=8, textColor=NAVY)),
                               Paragraph(entry.supervisor_remarks, s_body)])
        if extra_rows:
            extra_tbl = Table(extra_rows, colWidths=[4*cm, 14*cm])
            extra_tbl.setStyle(TableStyle([
                ('BOX', (0,0), (-1,-1), 0.5, rl_colors.HexColor('#CBD5E0')),
                ('INNERGRID', (0,0), (-1,-1), 0.3, rl_colors.HexColor('#CBD5E0')),
                ('BACKGROUND', (0,0), (0,-1), LIGHT),
                ('PADDING', (0,0), (-1,-1), 4),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ]))
            story.append(extra_tbl)

        story.append(Spacer(1, 10))

    if not entries.exists():
        story.append(Spacer(1, 20))
        story.append(Paragraph("Hakuna rekodi za logbook kwa kipindi kilichochaguliwa.", s_sub))

    # Footer
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceBefore=10, spaceAfter=4))
    story.append(Paragraph("© Wizara ya Elimu, Sayansi na Teknolojia — IMS v2.1.0",
                           ParagraphStyle('footer', fontName='Helvetica', fontSize=7, textColor=rl_colors.HexColor('#4A5568'), alignment=1)))

    doc.build(story)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
@login_required
def logbook_download_options(request):
    """Page for choosing download options"""
    student = get_or_create_student_profile(request.user)
    
    total_entries = LogbookEntry.objects.filter(student=student).count()
    this_week_entries = LogbookEntry.objects.filter(
        student=student,
        date__gte=timezone.now().date() - timedelta(days=7)
    ).count()
    
    return render(request, 'field_app/logbook_download.html', {
        'student': student,
        'total_entries': total_entries,
        'this_week_entries': this_week_entries,
    })

# =========================
# ADMIN VIEWS
# =========================

@staff_member_required
def admin_dashboard(request):
    pending_applications = StudentApplication.objects.filter(status='pending').select_related('student', 'subject', 'school')

    total_applications = StudentApplication.objects.count()
    approved_applications = StudentApplication.objects.filter(status='approved').count()
    rejected_applications = StudentApplication.objects.filter(status='rejected').count()
    pending_count = StudentApplication.objects.filter(status='pending').count()

    paginator = Paginator(pending_applications, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    schools = School.objects.annotate(
        current_count=Count('studentteacher'),
        is_full=Case(
            When(capacity__lte=F('current_students'), then=Value(True)),
            default=Value(False),
            output_field=BooleanField()
        )
    )

    total_assessors = Assessor.objects.count()
    active_assessors = Assessor.objects.filter(is_active=True).count()

    total_school_assignments = SchoolAssessment.objects.count()
    completed_assessments = SchoolAssessment.objects.filter(is_completed=True).count()

    recent_assignments = SchoolAssessment.objects.select_related(
        'assessor', 'school'
    ).order_by('-assessment_date')[:10]

    # --- Chart data ---
    # 1. Applications by status (donut chart)
    chart_app_status = {
        'labels': ['Approved', 'Pending', 'Rejected'],
        'data': [approved_applications, pending_count, rejected_applications],
    }

    # 2. Students by region (bar chart)
    region_qs = Region.objects.annotate(
        student_count=Count('district__school__studentteacher', distinct=True)
    ).order_by('-student_count')[:10]
    chart_regions = {
        'labels': list(region_qs.values_list('name', flat=True)),
        'data': list(region_qs.values_list('student_count', flat=True)),
    }

    # 3. Logbook submissions – last 14 days (line chart)
    today = timezone.now().date()
    start_day = today - timedelta(days=13)
    from django.db.models.functions import TruncDate
    counts_by_date = {
        row['date']: row['count']
        for row in LogbookEntry.objects.filter(date__range=(start_day, today))
        .values('date').annotate(count=Count('id'))
    }
    logbook_dates = []
    logbook_counts = []
    for i in range(13, -1, -1):
        day = today - timedelta(days=i)
        logbook_dates.append(day.strftime('%b %d'))
        logbook_counts.append(counts_by_date.get(day, 0))
    chart_logbook = {
        'labels': logbook_dates,
        'data': logbook_counts,
    }

    # 4. Assessment coverage (schools with vs without active assessor assignment)
    total_schools = School.objects.count()
    schools_with_assessor = SchoolAssessment.objects.values('school').distinct().count()
    schools_without_assessor = max(0, total_schools - schools_with_assessor)
    chart_coverage = {
        'labels': ['With Assessor', 'Without Assessor'],
        'data': [schools_with_assessor, schools_without_assessor],
    }

    context = {
        'pending_applications': pending_applications,
        'schools': schools,
        'total_applications': total_applications,
        'approved_applications': approved_applications,
        'rejected_applications': rejected_applications,
        'page_obj': page_obj,
        'total_assessors': total_assessors,
        'active_assessors': active_assessors,
        'total_school_assignments': total_school_assignments,
        'completed_assessments': completed_assessments,
        'recent_assignments': recent_assignments,
        'total_schools': total_schools,
        # Charts (JSON strings for template)
        'chart_app_status_json': json.dumps(chart_app_status),
        'chart_regions_json': json.dumps(chart_regions),
        'chart_logbook_json': json.dumps(chart_logbook),
        'chart_coverage_json': json.dumps(chart_coverage),
    }

    return render(request, 'field_app/admin_dashboard.html', context)


@staff_member_required
def download_backup(request):
    """Admin: export all IMS data as JSON and send as file download."""
    import json as _json
    from django.core import serializers as _ser
    from django.apps import apps as _apps
    from datetime import datetime as _dt

    BACKUP_MODELS = [
        'StudentTeacher', 'School', 'District', 'Region', 'Subject',
        'StudentApplication', 'LogbookEntry', 'SchemeOfWork', 'LessonPlan',
        'Assessor', 'SchoolAssessment', 'AcademicYear',
        'BoardMember', 'BoardComment', 'MonthlyReport',
    ]

    all_objects = []
    for name in BACKUP_MODELS:
        try:
            model = _apps.get_model('field_app', name)
            qs = model.objects.all()
            if qs.exists():
                all_objects.extend(_json.loads(_ser.serialize('json', qs)))
        except Exception:
            pass

    filename = f"ims_backup_{_dt.now().strftime('%Y%m%d_%H%M%S')}.json"
    payload = _json.dumps(all_objects, ensure_ascii=False, indent=2, default=str)
    response = HttpResponse(payload, content_type='application/json; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@staff_member_required
def approve_application(request, application_id):
    application = get_object_or_404(StudentApplication, id=application_id)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'approve':
            application.status = 'approved'
            application.approved_by = request.user
            application.approval_date = timezone.now()
            application.save()
            
            application.student.subjects.add(application.subject)
            
            try:
                capacity = SchoolSubjectCapacity.objects.get(
                    school=application.school, 
                    subject=application.subject
                )
                capacity.current_students = F('current_students') + 1
                capacity.save()
            except SchoolSubjectCapacity.DoesNotExist:
                SchoolSubjectCapacity.objects.create(
                    school=application.school,
                    subject=application.subject,
                    current_students=1,
                    max_students=5
                )
            
            messages.success(request, f"Application for {application.subject.name} approved successfully!")
            
        elif action == 'reject':
            application.status = 'rejected'
            application.approved_by = request.user
            application.approval_date = timezone.now()
            application.save()
            messages.success(request, f"Application for {application.subject.name} rejected.")
        
        return redirect('admin_dashboard')
    
    return render(request, 'field_app/approve_application.html', {'application': application})

# =========================
# ASSESSOR MANAGEMENT VIEWS
# =========================

@staff_member_required
def assign_assessor(request):
    """Assign single assessor to school"""
    if request.method == 'POST':
        assessor_id = request.POST.get('assessor_id')
        school_id = request.POST.get('school_id')
        
        assessor = get_object_or_404(Assessor, id=assessor_id)
        school = get_object_or_404(School, id=school_id)
        
        existing = SchoolAssessment.objects.filter(assessor=assessor, school=school).first()
        if existing:
            messages.warning(request, f"Assessor {assessor.full_name} is already assigned to {school.name}")
            if not assessor.email:
                messages.error(request, 
                    f"Assessor {assessor.full_name} has no email address! "
                    f"Cannot send credentials."
                )
                return redirect('assign_assessor')
            
            temp_password = None
            is_new_account = False
            
            if not assessor.user:
                temp_password = generate_random_password()

                try:
                    existing_user = User.objects.filter(email=assessor.email).first()
                    if existing_user:
                        user = existing_user
                        user.set_password(temp_password)
                        user.save()
                    else:
                        user = User.objects.create_user(
                            email=assessor.email,
                            password=temp_password,
                            is_staff=False,
                            is_superuser=False,
                            is_active=True
                        )
                    assessor.user = user
                    assessor.save()
                    is_new_account = True

                except Exception as e:
                    messages.error(request, f"Failed to create user account: {str(e)}")
                    return redirect('assign_assessor')
            
            school_assessment = SchoolAssessment.objects.create(
                assessor=assessor,
                school=school,
                assessment_date=timezone.now().date()
            )
            
            students = StudentTeacher.objects.filter(
                selected_school=school,
                approval_status='approved'
            )
            
            student_assessments_created = 0
            for student in students:
                StudentAssessment.objects.create(
                    assessor=assessor,
                    student=student,
                    school=school,
                    assessment_date=timezone.now().date()
                )
                student_assessments_created += 1
            
            try:
                login_url = request.build_absolute_uri(reverse('assessor_login'))
                
                subject = f'Field Placement Assessor Assignment - {school.name}'
                
                if is_new_account:
                    password_info = f"""
                    NEW ACCOUNT CREATED FOR YOU:
                    
                    Login Email: {assessor.email}
                    Temporary Password: {temp_password}
                    
                    Please change your password immediately after first login.
                    """
                    password_info = f"""
                    USE YOUR EXISTING ACCOUNT:
                    
                    Login Email: {assessor.email}
                    
                    If you forgot your password, use 'Forgot Password' on login page.
                    """
                
                message = f"""
                FIELD PLACEMENT ASSESSOR ASSIGNMENT
                {'=' * 50}
                
                Dear {assessor.full_name},
                
                You have been assigned as a Field Placement Assessor.
                
                ASSIGNMENT DETAILS:
                • School: {school.name}
                • District: {school.district.name} 
                • Region: {school.district.region.name}
                • Assignment Date: {timezone.now().strftime('%d/%m/%Y')}
                • Number of Students: {student_assessments_created}
                
                YOUR LOGIN CREDENTIALS:
                {password_info}
                
                LOGIN URL: {login_url}
                
                AFTER LOGIN, YOU CAN:
                1. View assigned school details
                2. See list of students assigned to you
                3. Track student progress
                4. Submit assessment reports
                5. Monitor logbook entries
                
                IMPORTANT:
                • Login using your email address
                • First-time users must change password
                • Contact administrator if you face issues
                
                Best regards,
                Kitengo cha Uratibu wa Mafunzo ya Uwanjani
                Wizara ya Elimu, Sayansi na Teknolojia

                Ujumbe huu umetumwa kiotomatiki. Tafadhali usijibu.
                """
                
                send_mail(
                    subject=subject,
                    message=message,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[assessor.email],
                    fail_silently=False,
                )
                
                if is_new_account:
                    messages.success(request, 
                        f"✅ Assessor {assessor.full_name} assigned successfully!<br>"
                        f"• Email sent to: {assessor.email}<br>"
                        f"• Temporary password: {temp_password}<br>"
                        f"• Assigned to: {school.name}<br>"
                        f"• Students: {student_assessments_created}"
                    )
                    messages.success(request,
                        f"✅ Assessor {assessor.full_name} assigned successfully!<br>"
                        f"• Email sent to: {assessor.email}<br>"
                        f"• Assigned to: {school.name}<br>"
                        f"• Students: {student_assessments_created}"
                    )
                
                print(f"📧 Email sent to assessor {assessor.email}")
                
            except Exception as e:
                error_msg = str(e)
                print(f"❌ Email failed for {assessor.email}: {error_msg}")
                
                if is_new_account:
                    messages.warning(request,
                        f"⚠️ Assessor assigned but email failed!<br>"
                        f"• Assessor: {assessor.full_name}<br>"
                        f"• School: {school.name}<br>"
                        f"• ERROR: {error_msg}<br>"
                        f"• <strong>MANUAL CREDENTIALS:</strong><br>"
                        f"Email: {assessor.email}<br>"
                        f"Password: {temp_password}"
                    )
                    messages.warning(request,
                        f"⚠️ Assessor assigned but email failed!<br>"
                        f"• Assessor: {assessor.full_name}<br>"
                        f"• School: {school.name}<br>"
                        f"• ERROR: {error_msg}"
                    )
        
        return redirect('admin_dashboard')
    
    # GET REQUEST
    assessors = Assessor.objects.filter(is_active=True).order_by('full_name')
    schools = School.objects.all().order_by('name')
    
    assessors_with_email = []
    assessors_without_email = []
    
    for assessor in assessors:
        if assessor.email and '@' in assessor.email:
            assessors_with_email.append(assessor)
            assessors_without_email.append(assessor)
    
    return render(request, 'field_app/assign_assessor.html', {
        'assessors_with_email': assessors_with_email,
        'assessors_without_email': assessors_without_email,
        'schools': schools,
    })

@staff_member_required
def bulk_assign_assessors(request):
    """Bulk assign assessors to schools - WITH ACADEMIC YEAR LOGIC"""
    
    if request.method == 'GET':
        # Get current academic year
        current_year = get_current_academic_year()
        if not current_year:
            messages.error(request, "⚠️ Hakuna mwaka wa masomo unaofanya kazi! Fungua mwaka mpya kwanza.")
            return redirect('admin_dashboard')
        
        # Get ALL assessors
        all_assessors = Assessor.objects.all().order_by('full_name')
        
        # Check each assessor's status for current year
        for assessor in all_assessors:
            if not assessor.user:
                # Assessor mpya kabisa - hana akaunti
                assessor.needs_new_credentials = True
                assessor.year_status = "Hakuna akaunti"
            elif not assessor.current_academic_year or assessor.current_academic_year != current_year:
                # Assessor ana akaunti lakini ni mwaka mpya
                assessor.needs_new_credentials = True
                assessor.year_status = f"Mpya kwa {current_year.year}"
            else:
                # Assessor ana akaunti na ni mwaka huohuo - shule tu zinaongezwa
                assessor.needs_new_credentials = False
                assessor.year_status = f"Tayari kwa {current_year.year}"
            
            assessor.schools_assigned = SchoolAssessment.objects.filter(
                assessor=assessor,
                academic_year=current_year
            ).count()
        
        # Search and pagination for schools
        search_query = request.GET.get('q', '')
        page_number = request.GET.get('page', 1)
        schools_per_page = 50
        
        schools_qs = School.objects.all().order_by('name')
        
        if search_query:
            schools_qs = schools_qs.filter(
                Q(name__icontains=search_query) |
                Q(district__name__icontains=search_query) |
                Q(district__region__name__icontains=search_query)
            )
        
        paginator = Paginator(schools_qs, schools_per_page)
        
        try:
            page_obj = paginator.page(page_number)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)
        
        schools_on_page = list(page_obj.object_list)
        all_filtered_schools = list(schools_qs.values_list('id', flat=True))
        
        # Get stats for schools on current page
        school_ids_on_page = [school.id for school in schools_on_page]
        
        # Count assessors per school for current year
        assignment_counts = SchoolAssessment.objects.filter(
            school_id__in=school_ids_on_page,
            academic_year=current_year
        ).values('school_id').annotate(assessors_count=Count('assessor_id'))
        
        assignment_dict = {item['school_id']: item['assessors_count'] for item in assignment_counts}
        
        # Count students per school
        student_counts = StudentTeacher.objects.filter(
            selected_school_id__in=school_ids_on_page,
            approval_status='approved'
        ).values('selected_school_id').annotate(student_count=Count('id'))
        
        student_count_dict = {item['selected_school_id']: item['student_count'] for item in student_counts}
        
        for school in schools_on_page:
            school.assessors_count = assignment_dict.get(school.id, 0)
            school.student_count = student_count_dict.get(school.id, 0)
        
        # Get assessors without valid emails
        assessors_no_email = Assessor.objects.filter(
            Q(email__isnull=True) | Q(email='')
        ).values_list('full_name', flat=True)
        
        default_date = (timezone.now() + timedelta(days=1)).strftime('%Y-%m-%d')
        
        return render(request, 'field_app/bulk_assign_assessors.html', {
            'all_assessors': all_assessors,
            'schools': schools_on_page,
            'page_obj': page_obj,
            'assessors_no_email': list(assessors_no_email),
            'total_assessors': all_assessors.count(),
            'total_schools': School.objects.count(),
            'available_assessors': all_assessors.exclude(
                Q(email__isnull=True) | Q(email='')
            ).count(),
            'default_date': default_date,
            'total_approved_students': StudentTeacher.objects.filter(
                approval_status='approved'
            ).count(),
            'search_query': search_query,
            'all_filtered_schools': all_filtered_schools,
            'current_year': current_year,
        })
    
    elif request.method == 'POST':
        # Get form data
        assessor_ids = request.POST.getlist('assessors')
        selected_schools_from_checkboxes = request.POST.getlist('schools[]')
        selected_schools_from_hidden = request.POST.get('selected_schools', '')
        assessment_date_str = request.POST.get('assessment_date', '')
        
        # Combine school selections
        school_ids = []
        if selected_schools_from_checkboxes:
            school_ids.extend(selected_schools_from_checkboxes)
        
        if selected_schools_from_hidden:
            hidden_ids = [sid.strip() for sid in selected_schools_from_hidden.split(',') if sid.strip()]
            school_ids.extend(hidden_ids)
        
        school_ids = list(set(school_ids))
        
        # Validate inputs
        if not assessor_ids:
            messages.error(request, "❌ Tafadhali chagua assessor mmoja au zaidi.")
            return redirect('bulk_assign_assessors')
        
        if not school_ids:
            messages.error(request, "❌ Tafadhali chagua shule moja au zaidi.")
            return redirect('bulk_assign_assessors')
        
        # Parse date
        try:
            assessment_date = datetime.strptime(assessment_date_str, '%Y-%m-%d').date()
        except ValueError:
            assessment_date = timezone.now().date()
        
        # Process assignment
        try:
            results = process_bulk_assignment_with_academic_year(
                assessor_ids, school_ids, assessment_date, request
            )
            
            request.session['bulk_assignment_results'] = results
            return redirect('bulk_assignment_results')
            
        except Exception as e:
            messages.error(request, f"❌ Hitilafu: {str(e)}")
            import traceback
            traceback.print_exc()
            return redirect('bulk_assign_assessors')
    
        return HttpResponseNotAllowed(['GET', 'POST'])

# views.py - FIX BULK ASSIGNMENT FUNCTION

# views.py - FIX BULK ASSIGNMENT BUG

@staff_member_required
def bulk_assignment_results(request):
    """Show results of bulk assignment with credentials"""
    results = request.session.get('bulk_assignment_results')
    
    if not results:
        messages.info(request, "No assignment results found.")
        return redirect('admin_dashboard')
    
    # Clear session after showing results
    if 'bulk_assignment_results' in request.session:
        del request.session['bulk_assignment_results']
    
    return render(request, 'field_app/bulk_assignment_results.html', {
        'results': results
    })

# =========================
# ASSESSOR ASSESSMENT VIEWS
# =========================

@assessor_login_required
def assessor_student_detail(request, school_id):
    """Assessor aone details za wanafunzi wa shule maalum"""
    try:
        assessor = Assessor.objects.get(user=request.user)
    except Assessor.DoesNotExist:
        messages.error(request, "You are not registered as an assessor.")
        return redirect('dashboard')
    
    school = get_object_or_404(School, id=school_id)
    
    # Check assignment
    school_assignment = SchoolAssessment.objects.filter(
        assessor=assessor,
        school=school
    ).first()
    
    if not school_assignment:
        messages.error(request, "You are not assigned to this school.")
        return redirect('assessor_dashboard')
    
    # ========== FIX: GET STUDENTS FROM BOTH SOURCES ==========
    
    # Method 1: Students who selected this school directly
    students_selected = StudentTeacher.objects.filter(
        selected_school=school,
        approval_status='approved'
    )
    
    # Method 2: Students with approved applications for this school
    approved_app_student_ids = StudentApplication.objects.filter(
        school=school,
        status='approved'
    ).values_list('student_id', flat=True).distinct()
    
    students_with_apps = StudentTeacher.objects.filter(
        id__in=approved_app_student_ids
    )
    
    # Combine both querysets
    student_ids = set()
    for student in students_selected:
        student_ids.add(student.id)
    for student in students_with_apps:
        student_ids.add(student.id)
    
    # Get all unique students
    students = StudentTeacher.objects.filter(
        id__in=list(student_ids)
    ).select_related('user')
    
    print(f"📊 Found {students.count()} students for school {school.name}")
    for student in students:
        print(f"   - {student.full_name}")
    
    # Get assessments
    student_assessments = StudentAssessment.objects.filter(
        assessor=assessor,
        school=school
    ).select_related('student')
    
    # Get other assessors
    other_assessors_assessments = SchoolAssessment.objects.filter(
        school=school
    ).exclude(assessor=assessor).select_related('assessor')
    
    other_assessors = [oa.assessor for oa in other_assessors_assessments]
    
    return render(request, 'field_app/assessor_student_detail.html', {
        'assessor': assessor,
        'school': school,
        'students': students,  # ← SASA ITAKUWA NA DATA
        'student_assessments': student_assessments,
        'school_assignment': school_assignment,
        'other_assessors': other_assessors,
    })
@assessor_login_required
def assessor_student_assessment(request, student_id):
    """Assessor assess specific student"""
    try:
        assessor = Assessor.objects.get(user=request.user)
    except Assessor.DoesNotExist:
        messages.error(request, "You are not registered as an assessor.")
        return redirect('dashboard')
    
    student = get_object_or_404(StudentTeacher, id=student_id)
    
    school_assignment = SchoolAssessment.objects.filter(
        assessor=assessor,
        school=student.selected_school
    ).first()
    
    if not school_assignment:
        messages.error(request, "You are not assigned to assess this student.")
        return redirect('assessor_dashboard')
    
    student_assessment, created = StudentAssessment.objects.get_or_create(
        assessor=assessor,
        student=student,
        school=student.selected_school,
        defaults={
            'assessment_date': timezone.now().date()
        }
    )
    
    if request.method == 'POST':
        student_assessment.attendance_score = request.POST.get('attendance_score')
        student_assessment.participation_score = request.POST.get('participation_score')
        student_assessment.teaching_skills_score = request.POST.get('teaching_skills_score')
        student_assessment.lesson_planning_score = request.POST.get('lesson_planning_score')
        student_assessment.classroom_management_score = request.POST.get('classroom_management_score')
        student_assessment.overall_score = request.POST.get('overall_score')
        student_assessment.comments = request.POST.get('comments')
        student_assessment.is_completed = True
        student_assessment.completed_date = timezone.now()
        student_assessment.save()
        
        messages.success(request, f"Assessment for {student.full_name} submitted successfully!")
        return redirect('assessor_student_detail', school_id=student.selected_school.id)
    
    logbook_entries = LogbookEntry.objects.filter(
        student=student
    ).order_by('-date')[:20]
    
    approved_subjects = student.subjects.all()
    
    return render(request, 'field_app/assessor_student_assessment.html', {
        'assessor': assessor,
        'student': student,
        'student_assessment': student_assessment,
        'logbook_entries': logbook_entries,
        'approved_subjects': approved_subjects,
        'school_assignment': school_assignment,
    })


@assessor_login_required
def assessor_add_logbook_remark(request, entry_id):
    """Assessor anaandika maoni kwenye logbook maalum ya mwanafunzi."""
    if request.method != 'POST':
        return HttpResponseNotAllowed(['POST'])
    try:
        assessor = Assessor.objects.get(user=request.user)
    except Assessor.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Hautambuliki kama msimamizi.'}, status=403)

    entry = get_object_or_404(LogbookEntry, id=entry_id)
    # Verify assessor is assigned to this student's school
    if not SchoolAssessment.objects.filter(assessor=assessor, school=entry.student.selected_school).exists():
        return JsonResponse({'ok': False, 'error': 'Huna ruhusa ya kuandika maoni hapa.'}, status=403)

    remarks = request.POST.get('remarks', '').strip()
    if not remarks:
        return JsonResponse({'ok': False, 'error': 'Maoni hayawezi kuwa matupu.'}, status=400)

    entry.supervisor_remarks = remarks
    entry.save(update_fields=['supervisor_remarks'])
    return JsonResponse({'ok': True, 'message': 'Maoni yamehifadhiwa.'})


# =========================
# STUDENT LIST VIEWS
# =========================

@staff_member_required
def student_list(request):
    students = StudentTeacher.objects.all().select_related('user', 'selected_school')
    
    school_filter = request.GET.get('school')
    if school_filter:
        students = students.filter(selected_school__name__icontains=school_filter)
    
    status_filter = request.GET.get('status')
    if status_filter:
        students = students.filter(approval_status=status_filter)
    
    return render(request, 'field_app/student_list.html', {'students': students})

@staff_member_required
def approve_student(request, student_id):
    student = get_object_or_404(StudentTeacher, id=student_id)
    
    if request.method == 'POST':
        student.approval_status = 'approved'
        student.approval_date = timezone.now()
        student.save()
        invalidate_student_cache(student)
        messages.success(request, f'Student {student.full_name} approved successfully!')
        return redirect('student_list')
    
    return render(request, 'field_app/approve_student.html', {'student': student})

# =========================
# LETTER DOWNLOAD VIEWS
# =========================

@login_required
def download_individual_letter(request):
    student = get_or_create_student_profile(request.user)
    
    approved_applications = StudentApplication.objects.filter(
        student=student, 
        status='approved'
    )
    
    if not approved_applications.exists():
        messages.error(request, "You don't have any approved applications to download a letter.")
        return redirect('dashboard')

    buffer = BytesIO()
    p = canvas.Canvas(buffer)
    
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, 800, "INDIVIDUAL FIELD PLACEMENT APPROVAL LETTER")
    
    p.setFont("Helvetica", 12)
    p.drawString(100, 770, f"Student Name: {student.full_name}")
    p.drawString(100, 750, f"Student ID: {student.id}")
    p.drawString(100, 730, f"Phone: {student.phone_number}")
    p.drawString(100, 710, f"Email: {student.user.email}")
    
    if student.selected_school:
        p.drawString(100, 680, f"Assigned School: {student.selected_school.name}")
        p.drawString(100, 660, f"School District: {student.selected_school.district.name}")
        p.drawString(100, 640, f"School Region: {student.selected_school.district.region.name}")
    
    p.drawString(100, 610, "Approved Teaching Subjects:")
    y_position = 590
    for application in approved_applications:
        p.drawString(120, y_position, f"✓ {application.subject.name} at {application.school.name}")
        y_position -= 20
        if application.approval_date:
            p.drawString(140, y_position, f"Approved on: {application.approval_date.strftime('%Y-%m-%d')}")
            y_position -= 20
    
    p.drawString(100, 530, "This letter confirms that the above student has been approved")
    p.drawString(100, 510, "for field placement teaching practice.")
    p.drawString(100, 490, f"Generated on: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
    
    p.showPage()
    p.save()
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="individual_approval_{student.full_name}.pdf"'
    return response

@login_required
def download_group_letter(request):
    student = get_or_create_student_profile(request.user)
    
    if not student.selected_school:
        messages.error(request, "Huna shule uliyochagua.")
        return redirect('dashboard')
    
    school = student.selected_school
    
    group_letter_quota = 5
    
    approved_students_count = StudentApplication.objects.filter(
        school=school,
        status='approved'
    ).count()
    
    student_has_approved_application = StudentApplication.objects.filter(
        student=student,
        school=school,
        status='approved'
    ).exists()
    
    if approved_students_count < group_letter_quota:
        messages.error(request, 
            f"Bado hatujafikia idadi ya wanafunzi {group_letter_quota} walioidhinishwa. " 
            f"Kwa sasa kuna {approved_students_count}/{group_letter_quota}."
        )
        return redirect('dashboard')
    
    if not student_has_approved_application:
        messages.error(request, 
            "Huwezi kupata barua ya kikundi kwa sababu huna maombi yaliyoidhinishwa kwenye shule hii."
        )
        return redirect('dashboard')
    
    buffer = BytesIO()
    p = canvas.Canvas(buffer)
    
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, 800, "BARUA YA UTHIBITISHO WA KIKUNDI")
    p.drawString(100, 780, "Taasisi ya Ualimu Tanzania")
    
    p.setFont("Helvetica", 12)
    p.drawString(100, 750, f"Jina la Shule: {school.name}")
    p.drawString(100, 730, f"Wilaya: {school.district.name}")
    p.drawString(100, 710, f"Mkoa: {school.district.region.name}")
    p.drawString(100, 690, f"Idadi ya Wanafunzi Inayohitajika: {group_letter_quota}")
    p.drawString(100, 670, f"Wanafunzi Walioidhinishwa: {approved_students_count}")
    
    p.drawString(100, 640, "Orodha ya Wanafunzi Walioidhinishwa:")
    y_position = 620
    
    approved_applications = StudentApplication.objects.filter(
        school=school,
        status='approved'
    ).select_related('student').distinct()
    
    for idx, application in enumerate(approved_applications, 1):
        student_name = application.student.full_name
        subject_name = application.subject.name
        p.drawString(120, y_position, f"{idx}. {student_name} - {subject_name}")
        y_position -= 20
        if y_position < 100:
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = 780
    
    p.drawString(100, y_position - 40, "Barua hii inathibitisha kuwa shule imefikia idadi ya wanafunzi 5")
    p.drawString(100, y_position - 60, "wa kufanya mafunzo ya ualimu kwenye uwanja kama kikundi.")
    p.drawString(100, y_position - 80, f"Imetolewa tarehe: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
    
    p.showPage()
    p.save()
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="barua_kikundi_{school.name}.pdf"'
    
    messages.success(request, "Barua ya kikundi imepakuliwa kikamilifu!")
    return response

# =========================
# AJAX AND API VIEWS
# =========================

@staff_member_required
@csrf_exempt
def ajax_search_schools(request):
    """AJAX endpoint for searching ALL schools"""
    search_query = request.GET.get('q', '').strip()
    
    if not search_query or len(search_query) < 2:
        return JsonResponse({'results': [], 'count': 0, 'error': 'Search term too short'})
    
    # Search in all schools
    schools = School.objects.filter(
        Q(name__icontains=search_query) |
        Q(district__name__icontains=search_query) |
        Q(district__region__name__icontains=search_query)
    ).select_related('district', 'district__region')[:100]
    
    results = []
    for school in schools:
        student_count = StudentTeacher.objects.filter(
            selected_school=school,
            approval_status='approved'
        ).count()
        
        assessor_count = SchoolAssessment.objects.filter(school=school).count()
        
        results.append({
            'id': school.id,
            'name': school.name,
            'district': school.district.name,
            'region': school.district.region.name,
            'level': school.level,
            'students': student_count,
            'assessors': assessor_count,
            'capacity': school.capacity,
            'current_students': school.current_students,
        })
    
    return JsonResponse({
        'results': results,
        'count': len(results),
        'search_term': search_query
    })

@staff_member_required
@csrf_exempt
def assessor_details_api(request, assessor_id):
    """API endpoint for assessor details"""
    if request.method == 'GET':
        assessor = get_object_or_404(Assessor, id=assessor_id)
        
        school_assignments = SchoolAssessment.objects.filter(assessor=assessor)
        schools_data = []
        for assignment in school_assignments:
            schools_data.append({
                'name': assignment.school.name,
                'district': assignment.school.district.name,
                'level': assignment.school.level,
                'assessment_date': assignment.assessment_date.strftime('%Y-%m-%d'),
            })
        
        data = {
            'id': assessor.id,
            'full_name': assessor.full_name,
            'email': assessor.email,
            'phone_number': assessor.phone_number,
            'is_active': assessor.is_active,
            'has_account': bool(assessor.user),
            'schools_count': len(schools_data),
            'schools': schools_data,
        }
        
        return JsonResponse(data)
    return JsonResponse({'error': 'Invalid method'}, status=405)

# =========================
# REGION PINNING VIEWS
# =========================
@staff_member_required
def region_pinning_view(request):
    """
    Pin (hide) regions from students for a specific academic year.
    Regions entered in the form will be HIDDEN from students.
    """
    if request.method == 'POST':
        form = RegionFieldInputForm(request.POST)
        if form.is_valid():
            year_name = form.cleaned_data['academic_year']
            
            # Get regions to hide from form
            regions_data = form.cleaned_data['regions_to_hide']
            
            # Handle both string and list input
            if isinstance(regions_data, str):
                regions_to_hide_names = [
                    name.strip().lower() for name in regions_data.split(',') if name.strip()
                ]
                regions_to_hide_names = [name.strip().lower() for name in regions_data if name.strip()]
            
            print(f"📋 Regions to HIDE: {regions_to_hide_names}")
            
            # Validate regions exist
            existing_region_names = list(Region.objects.values_list('name', flat=True))
            existing_region_names_lower = [r.lower() for r in existing_region_names]
            
            invalid_regions = []
            valid_regions = []
            
            for region_name in regions_to_hide_names:
                if region_name in existing_region_names_lower:
                    # Find original case
                    for er in existing_region_names:
                        if er.lower() == region_name:
                            valid_regions.append(er)
                            break
                    invalid_regions.append(region_name)
            
            if invalid_regions:
                messages.error(
                    request,
                    f"❌ These regions don't exist: {', '.join(invalid_regions)}\n"
                    f"Available regions: {', '.join(existing_region_names[:20])}"
                )
                # Show current status
                current_year = get_current_academic_year()
                if current_year:
                    current_pins = RegionPin.objects.filter(
                        academic_year=current_year,
                        is_pinned=True
                    ).select_related('region')
                    currently_hidden = [pin.region.name for pin in current_pins]
                    currently_hidden = []
                
                return render(request, 'field_app/pin_regions_form.html', {
                    'form': form,
                    'current_year': current_year if 'current_year' in locals() else None,
                    'currently_hidden_regions': currently_hidden if 'currently_hidden' in locals() else [],
                    'total_regions': Region.objects.count(),
                })
            
            # Get or create academic year
            year, created = AcademicYear.objects.get_or_create(
                year=year_name,
                defaults={'is_active': True}
            )
            
            # Set this year as active
            if not year.is_active:
                year.is_active = True
                year.save()
                AcademicYear.objects.exclude(id=year.id).update(is_active=False)
                print(f"📅 Activated academic year: {year.year}")
            
            # ========== FIX: Clear old pins ==========
            # Delete ALL region pins for this academic year
            region_pins_deleted = RegionPin.objects.filter(academic_year=year).delete()
            
            # Delete ALL school pins for this academic year
            school_pins_deleted = SchoolPin.objects.filter(academic_year=year).delete()
            
            print(f"🗑️ Deleted {region_pins_deleted[0]} region pins, {school_pins_deleted[0]} school pins")
            
            # ========== FIX: Create RegionPin objects for ALL regions ==========
            all_regions = Region.objects.all()
            region_pins_to_create = []
            pinned_region_names = []
            visible_region_names = []
            
            for region in all_regions:
                # is_pinned = True if region should be HIDDEN from students
                is_pinned = region.name in valid_regions  # Use original case for comparison
                
                region_pins_to_create.append(RegionPin(
                    academic_year=year,
                    region=region,
                    is_pinned=is_pinned
                ))
                
                if is_pinned:
                    pinned_region_names.append(region.name)
                    visible_region_names.append(region.name)
            
            # Bulk create region pins
            if region_pins_to_create:
                RegionPin.objects.bulk_create(region_pins_to_create)
                print(f"✅ Created {len(region_pins_to_create)} region pins")
                print(f"   🔒 Pinned (Hidden): {len(pinned_region_names)}")
                print(f"   ✅ Visible: {len(visible_region_names)}")
            
            # ========== FIX: Pin schools ONLY in hidden regions ==========
            if pinned_region_names:
                # Get all schools in hidden regions
                schools_to_pin = School.objects.filter(
                    district__region__name__in=pinned_region_names
                ).select_related('district__region')
                
                school_pins_to_create = []
                for school in schools_to_pin:
                    school_pins_to_create.append(SchoolPin(
                        academic_year=year,
                        school=school,
                        is_pinned=True,
                        pin_reason='region_restricted',
                        notes=f"Region {school.district.region.name} is restricted for {year.year}"
                    ))
                
                if school_pins_to_create:
                    SchoolPin.objects.bulk_create(school_pins_to_create)
                    print(f"✅ Created {len(school_pins_to_create)} school pins for hidden regions")
                    print(f"ℹ️ No schools found in hidden regions")
                print(f"ℹ️ No hidden regions - no school pins created")
                school_pins_to_create = []
            
            # Store summary in session for success page
            request.session['pinning_summary'] = {
                'academic_year': year.year,
                'pinned_regions': pinned_region_names,
                'visible_regions': visible_region_names,
                'pinned_regions_count': len(pinned_region_names),
                'visible_regions_count': len(visible_region_names),
                'schools_pinned_count': len(school_pins_to_create),
                'is_new_year': created,
                'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Success message
            if len(pinned_region_names) == 0:
                message = f"✅ All {len(visible_region_names)} regions are now VISIBLE to students for {year.year}!"
                message = (
                    f"✅ Success for {year.year}!\n"
                    f"🔒 HIDDEN Regions ({len(pinned_region_names)}): {', '.join(pinned_region_names[:5])}"
                    f"{'...' if len(pinned_region_names) > 5 else ''}\n"
                    f"✅ VISIBLE Regions ({len(visible_region_names)}): {', '.join(visible_region_names[:5])}"
                    f"{'...' if len(visible_region_names) > 5 else ''}\n"
                    f"🏫 Schools Hidden: {len(school_pins_to_create)}"
                )
            
            messages.success(request, message)
            return redirect('pinning_success')
        else:
            messages.error(request, f"Please correct the errors below: {form.errors}")
    else:
        form = RegionFieldInputForm()
    
    # GET request or form errors - show current status
    current_year = get_current_academic_year()
    currently_hidden = []
    
    if current_year:
        current_pins = RegionPin.objects.filter(
            academic_year=current_year,
            is_pinned=True
        ).select_related('region')
        currently_hidden = [pin.region.name for pin in current_pins]
    
    return render(request, 'field_app/pin_regions_form.html', {
        'form': form,
        'current_year': current_year,
        'currently_hidden_regions': currently_hidden,
        'total_regions': Region.objects.count(),
        'hidden_count': len(currently_hidden),
        'visible_count': Region.objects.count() - len(currently_hidden),
    })
@login_required
def pinning_success_view(request):
    """Display summary after pinning regions"""
    summary = request.session.get('pinning_summary', {})
    
    context = {
        'summary': summary,
        'has_summary': bool(summary),
    }
    
    # Clear from session after displaying
    if 'pinning_summary' in request.session:
        del request.session['pinning_summary']
    
    return render(request, 'field_app/pinning_success.html', context)
# =========================
# PROFILE VIEWS
# =========================

@login_required
def profile_create(request):
    student = get_or_create_student_profile(request.user)
    
    if request.method == 'POST':
        form = StudentTeacherForm(request.POST, instance=student)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profile updated successfully!')
            return redirect('dashboard')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = StudentTeacherForm(instance=student)

    return render(request, 'field_app/profile_create.html', {'form': form})

# =========================
# MISCELLANEOUS VIEWS
# =========================

@login_required
def get_subjects(request, school_id):
    subject_caps = SchoolSubjectCapacity.objects.filter(school_id=school_id).select_related('subject')
    data = [
        {
            'id': sc.subject.id,
            'name': sc.subject.name,
            'current': sc.current_students,
            'max': sc.max_students
        }
        for sc in subject_caps
    ]
    return JsonResponse(data, safe=False)

@login_required
def confirm_school_selection(request, district_id):
    if request.method != 'POST':
        return redirect('select_school', district_id=district_id)

    school_id = request.POST.get('school_id')
    if not school_id:
        messages.error(request, "No school selected.")
        return redirect('select_school', district_id=district_id)

    district = get_object_or_404(District, id=district_id)
    school = get_object_or_404(School, id=school_id)

    school.refresh_from_db()
    if school.current_students >= school.capacity:
        messages.error(request, f"{school.name} is at full capacity!")
        return redirect('select_school', district_id=district_id)

    old_school_id = request.session.get('selected_school_id')
    if old_school_id and old_school_id != school.id:
        old_school = School.objects.filter(id=old_school_id).first()
        if old_school:
            old_school.current_students = Greatest(F('current_students') - 1, 0)
            old_school.save()

    request.session['selected_school_id'] = school.id
    school.current_students = F('current_students') + 1
    school.save()
    school.refresh_from_db()

    student = get_or_create_student_profile(request.user)
    student.selected_school = school
    student.save()
    invalidate_student_cache(student)

    messages.success(request, f"You have successfully selected {school.name}.")
    return redirect('dashboard')

@login_required
def my_assessors(request):
    """Wanafunzi waone assessors wao kwa mwaka huu wa masomo"""
    student = get_or_create_student_profile(request.user)
    
    if not student.selected_school:
        messages.error(request, "You need to select a school first to see your assessors.")
        return redirect('select_region')
    
    school = student.selected_school
    
    # 🔴 FIX: Get current academic year
    current_year = get_current_academic_year()
    
    print(f"\n🔍 MY ASSESSORS DEBUG:")
    print(f"   Student: {student.full_name}")
    print(f"   School: {school.name} (ID: {school.id})")
    print(f"   Current Academic Year: {current_year.year if current_year else 'None'}")
    
    if current_year:
        school_assessments = SchoolAssessment.objects.filter(
            school=school,
            academic_year=current_year
        ).select_related('assessor')
    else:
        school_assessments = SchoolAssessment.objects.none()
    
    assessors_data = []
    for assessment in school_assessments:
        assessors_data.append({
            'assessor': assessment.assessor,
            'assessment_date': assessment.assessment_date,
            'is_completed': assessment.is_completed,
        })
        print(f"   - Assessor: {assessment.assessor.full_name}, Completed: {assessment.is_completed}")
    
    return render(request, 'field_app/my_assessors.html', {
        'student': student,
        'school': school,
        'assessors_data': assessors_data,
        'current_year': current_year,
    })

@staff_member_required
def assessor_list(request):
    """List all assessors with their credentials"""
    assessors = Assessor.objects.filter(is_active=True).select_related('user')
    
    for assessor in assessors:
        assessor.assigned_schools = SchoolAssessment.objects.filter(
            assessor=assessor
        ).select_related('school')
        
        assessor.schools_count = assessor.assigned_schools.count()
        
        if assessor.user:
            assessor.has_account = True
            assessor.login_email = assessor.user.email
            assessor.password_info = "Use existing password"
            assessor.has_account = False
            assessor.login_email = assessor.email or "No email"
            assessor.password_info = "No account created yet"
    
    return render(request, 'field_app/assessor_list.html', {
        'assessors': assessors
    })

@staff_member_required
def reset_assessors_for_new_year(request):
    """Reset all assessors for new academic year"""
    if request.method == 'POST':
        current_year = get_current_academic_year()
        if not current_year:
            messages.error(request, "No active academic year found!")
            return redirect('admin_dashboard')
        
        assessors = Assessor.objects.filter(is_active=True, user__isnull=False)
        
        results = []
        for assessor in assessors:
            if assessor.email and '@' in assessor.email:
                temp_password = generate_random_password()
                
                assessor.user.set_password(temp_password)
                assessor.user.save()
                
                assessor.current_academic_year = current_year
                assessor.save()
                
                results.append({
                    'name': assessor.full_name,
                    'email': assessor.email,
                    'password': temp_password,
                    'status': '✅ Password reset for new year'
                })
        
        request.session['new_year_credentials'] = results
        
        messages.success(request, 
            f"✅ Reset {len(results)} assessors for new academic year {current_year.year}"
        )
        return redirect('new_year_credentials')
    
    return render(request, 'field_app/reset_new_year.html')

@staff_member_required
def new_year_credentials(request):
    """Show credentials after new year reset"""
    results = request.session.get('new_year_credentials', [])
    
    if not results:
        messages.info(request, "No credentials found. Please reset assessors first.")
        return redirect('reset_assessors_for_new_year')
    
    return render(request, 'field_app/new_year_credentials.html', {
        'results': results,
        'total': len(results)
    })
# =========================
# SCHOOL CHANGE FUNCTIONALITY
# =========================




# =========================
# CHANGE SCHOOL - PERFORMANCE OPTIMIZED
# =========================



# =========================
# CHANGE SCHOOL - COMPLETE FIXED VERSION
# =========================

@login_required
def change_school(request):
    """Mwanafunzi anaweza kubadili shule ndani ya wiki moja - WITH COMPLETE DB UPDATE"""
    student = get_or_create_student_profile(request.user)
    
    # Check if student has selected a school
    if not student.selected_school:
        messages.error(request, "Hujachagua shule yoyote bado.")
        return redirect('select_region')
    
    # Set initial selection date if not set
    if not student.initial_school_selection_date:
        student.initial_school_selection_date = timezone.now()
        student.save()
        invalidate_student_cache(student)
    
    # Calculate days passed
    days_passed = (timezone.now() - student.initial_school_selection_date).days
    CAN_CHANGE_DAYS = 7
    MAX_CHANGES = 3

    # Block if student has an approved application — placement is confirmed
    has_approved_application = StudentApplication.objects.filter(
        student=student, status='approved'
    ).exists()

    can_change = (
        not has_approved_application and
        days_passed <= CAN_CHANGE_DAYS and
        student.school_change_count < MAX_CHANGES
    )
    cant_change_reason = None
    if has_approved_application:
        cant_change_reason = 'approved'
    elif days_passed > CAN_CHANGE_DAYS:
        cant_change_reason = 'expired'
    elif student.school_change_count >= MAX_CHANGES:
        cant_change_reason = 'limit'

    # Get districts in current school's region (only same region for faster loading)
    current_district = student.selected_school.district
    current_region = current_district.region

    # Get all districts in the same region
    districts_in_region = District.objects.filter(region=current_region).select_related('region')

    remaining_days = max(0, CAN_CHANGE_DAYS - days_passed)
    remaining_changes = max(0, MAX_CHANGES - student.school_change_count)
    
    # ========== HANDLE POST REQUEST DIRECTLY (if not using AJAX) ==========
    if request.method == 'POST' and not request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        new_school_id = request.POST.get('new_school_id')
        
        if not new_school_id:
            messages.error(request, "Tafadhali chagua shule mpya.")
            return redirect('change_school')
        
        # Process the change
        try:
            new_school = School.objects.select_related('district', 'district__region').get(id=new_school_id)
            old_school = student.selected_school
            
            # Validate
            if not can_change:
                messages.error(request, f"Huwezi kubadili shule tena. Umebadili tayari {student.school_change_count} mara.")
                return redirect('change_school')
            
            # Check capacity
            if new_school.current_students >= new_school.capacity:
                messages.error(request, f"Shule {new_school.name} imejaa. Hakuna nafasi.")
                return redirect('change_school')
            
            # Check if pinned
            current_year = get_current_academic_year()
            is_pinned = SchoolPin.objects.filter(
                school=new_school,
                academic_year=current_year,
                is_pinned=True
            ).exists()
            
            if is_pinned:
                messages.error(request, f"Shule {new_school.name} haipatikani kwa sasa.")
                return redirect('change_school')
            
            # ========== UPDATE DATABASE ==========
            print(f"\n🔄 CHANGING SCHOOL FOR: {student.full_name}")
            print(f"   OLD: {old_school.name} (ID: {old_school.id})")
            print(f"   NEW: {new_school.name} (ID: {new_school.id})")
            
            # 1. Decrease old school counter
            School.objects.filter(id=old_school.id).update(current_students=Greatest(F('current_students') - 1, 0))
            print(f"   ✅ Decreased old school: {old_school.name}")
            
            # 2. Increase new school counter
            School.objects.filter(id=new_school.id).update(current_students=F('current_students') + 1)
            print(f"   ✅ Increased new school: {new_school.name}")
            
            # 3. Update student record
            student.selected_school = new_school
            student.school_change_count = F('school_change_count') + 1
            student.last_school_change_date = timezone.now()
            student.save()
            invalidate_student_cache(student)
            
            # Refresh student to get updated count
            student.refresh_from_db()
            
            # 4. Delete pending applications for old school
            deleted_count, _ = StudentApplication.objects.filter(
                student=student,
                school=old_school,
                status='pending'
            ).delete()
            print(f"   ✅ Deleted {deleted_count} pending applications for old school")
            
            # 5. Clear any session cache
            if 'selected_school_id' in request.session:
                del request.session['selected_school_id']
            
            # 6. Clear cache for API
            cache_key = f'schools_district_{old_school.district_id}'
            cache.delete(cache_key)
            cache_key = f'schools_district_{new_school.district_id}'
            cache.delete(cache_key)
            
            print(f"   ✅ School change COMPLETE!")
            
            messages.success(
                request,
                f"✅ Umefanikiwa kubadili shule!\n"
                f"Shule mpya: {new_school.name}\n"
                f"Umesalia na {MAX_CHANGES - student.school_change_count} nafasi za kubadilisha."
            )
            
            return redirect('dashboard')
            
        except School.DoesNotExist:
            messages.error(request, "Shule haipatikani.")
            return redirect('change_school')
        except Exception as e:
            print(f"❌ ERROR: {e}")
            import traceback
            traceback.print_exc()
            messages.error(request, f"Hitilafu: {str(e)}")
            return redirect('change_school')
    
    # Render the template for GET request
    return render(request, 'field_app/change_school.html', {
        'student': student,
        'current_school': student.selected_school,
        'current_region': current_region,
        'districts': districts_in_region,
        'days_passed': days_passed,
        'remaining_days': remaining_days,
        'remaining_changes': remaining_changes,
        'max_change_days': CAN_CHANGE_DAYS,
        'max_changes': MAX_CHANGES,
        'initial_selection_date': student.initial_school_selection_date,
        'can_change': can_change,
        'cant_change_reason': cant_change_reason,
    })


@login_required
def api_confirm_change_school(request):
    """API endpoint to confirm school change - SIMPLE WORKING VERSION"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        new_school_id = data.get('school_id')
    except:
        return JsonResponse({'error': 'Invalid data'}, status=400)
    
    if not new_school_id:
        return JsonResponse({'error': 'School ID required'}, status=400)
    
    student = get_or_create_student_profile(request.user)
    
    # Check if student can change school
    if not student.selected_school:
        return JsonResponse({'error': 'No school selected'}, status=400)

    # Block if student has an approved application — placement is confirmed
    if StudentApplication.objects.filter(student=student, status='approved').exists():
        return JsonResponse({
            'error': 'Huwezi kubadili shule. Ombi lako limeshaidhinishwa. Wasiliana na msimamizi.'
        }, status=403)

    if not student.initial_school_selection_date:
        student.initial_school_selection_date = timezone.now()
        student.save()
        invalidate_student_cache(student)

    days_passed = (timezone.now() - student.initial_school_selection_date).days
    if days_passed > 7:
        return JsonResponse({'error': 'Change window expired (only 7 days)'}, status=400)

    if student.school_change_count >= 3:
        return JsonResponse({'error': 'Maximum 3 changes allowed'}, status=400)
    
    try:
        new_school = School.objects.get(id=new_school_id)
        old_school = student.selected_school
        
        # Check capacity
        if new_school.current_students >= new_school.capacity:
            return JsonResponse({'error': f'Shule {new_school.name} imejaa'}, status=400)
        
        # Check if pinned
        current_year = get_current_academic_year()
        is_pinned = SchoolPin.objects.filter(
            school=new_school,
            academic_year=current_year,
            is_pinned=True
        ).exists()
        
        if is_pinned:
            return JsonResponse({'error': f'Shule {new_school.name} haipatikani'}, status=400)
        
        # ========== SIMPLE DATABASE UPDATE ==========
        old_school.current_students = max(0, old_school.current_students - 1)
        old_school.save()

        new_school.current_students += 1
        new_school.save()
        
        # Update student
        student.selected_school = new_school
        student.school_change_count += 1
        student.last_school_change_date = timezone.now()
        student.save()
        invalidate_student_cache(student)
        
        # Delete pending applications for old school
        StudentApplication.objects.filter(
            student=student,
            school=old_school,
            status='pending'
        ).delete()
        
        print(f"✅ SUCCESS: {student.full_name} changed from {old_school.name} to {new_school.name}")
        
        return JsonResponse({
            'success': True,
            'message': f'Successfully changed to {new_school.name}',
            'new_school': {
                'id': new_school.id,
                'name': new_school.name,
                'district': new_school.district.name,
                'region': new_school.district.region.name,
                'current_students': new_school.current_students,
                'capacity': new_school.capacity,
            },
            'remaining_changes': 3 - student.school_change_count,
            'remaining_days': max(0, 7 - days_passed),
        })
        
    except School.DoesNotExist:
        return JsonResponse({'error': 'School not found'}, status=404)
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)
@login_required
def api_get_schools_for_change(request):
    """API endpoint for getting schools by district - FAST with caching"""
    
    student = get_or_create_student_profile(request.user)
    
    if not student.selected_school:
        return JsonResponse({'error': 'No school selected'}, status=400)
    
    # Get parameters
    district_id = request.GET.get('district_id')
    level = request.GET.get('level', 'Secondary')
    search = request.GET.get('search', '').strip()
    page = int(request.GET.get('page', 1))
    per_page = 12  # Schools per page
    
    # Validate district
    if not district_id:
        return JsonResponse({'error': 'District ID required'}, status=400)
    
    try:
        district = District.objects.get(id=district_id)
    except District.DoesNotExist:
        return JsonResponse({'error': 'District not found'}, status=404)
    
    # Get current academic year for pinned schools
    current_year = get_current_academic_year()
    
    # Get pinned school IDs (cache for 5 minutes)
    cache_key = f'pinned_schools_{current_year.id if current_year else "none"}'
    pinned_school_ids = cache.get(cache_key)
    if pinned_school_ids is None and current_year:
        pinned_school_ids = list(SchoolPin.objects.filter(
            academic_year=current_year,
            is_pinned=True
        ).values_list('school_id', flat=True))
        cache.set(cache_key, pinned_school_ids, 300)  # Cache for 5 minutes
    elif not current_year:
        pinned_school_ids = []
    
    # Base queryset - only schools in selected district
    schools_qs = School.objects.filter(
        district_id=district_id,
        level=level
    ).exclude(id=student.selected_school.id)
    
    # Exclude pinned schools
    if pinned_school_ids:
        schools_qs = schools_qs.exclude(id__in=pinned_school_ids)
    
    # Exclude full schools
    schools_qs = schools_qs.filter(current_students__lt=F('capacity'))
    
    # Apply search filter
    if search:
        schools_qs = schools_qs.filter(
            Q(name__icontains=search) |
            Q(district__name__icontains=search) |
            Q(district__region__name__icontains=search)
        )
    
    # Select related for efficiency
    schools_qs = schools_qs.select_related('district', 'district__region')
    
    # Count total (for pagination)
    total_count = schools_qs.count()
    
    # Apply pagination
    start = (page - 1) * per_page
    end = start + per_page
    schools = schools_qs[start:end]
    
    # Prepare data
    schools_data = []
    for school in schools:
        # Calculate occupancy
        if school.capacity > 0:
            occupancy = round((school.current_students / school.capacity) * 100)
            occupancy = 0
        
        schools_data.append({
            'id': school.id,
            'name': school.name,
            'district': school.district.name,
            'region': school.district.region.name,
            'level': school.level,
            'current_students': school.current_students,
            'capacity': school.capacity,
            'available_spots': school.capacity - school.current_students,
            'occupancy_percentage': occupancy,
            'is_available': school.current_students < school.capacity,
        })
    
    return JsonResponse({
        'success': True,
        'schools': schools_data,
        'total': total_count,
        'page': page,
        'total_pages': (total_count + per_page - 1) // per_page,
        'has_next': end < total_count,
        'has_previous': page > 1,
    })        
@staff_member_required
def toggle_region_pin(request, region_id):
    """Toggle region pin (hide/unhide) from admin interface"""
    if request.method == 'POST':
        current_year = get_current_academic_year()
        region = get_object_or_404(Region, id=region_id)
        
        # Get or create pin
        region_pin, created = RegionPin.objects.get_or_create(
            academic_year=current_year,
            region=region,
            defaults={'is_pinned': False}
        )
        
        # Toggle the pin status
        region_pin.is_pinned = not region_pin.is_pinned
        region_pin.save()
        
        # Also pin/unpin all schools in this region
        schools_in_region = School.objects.filter(district__region=region)
        
        if region_pin.is_pinned:
            # Pin all schools (hide)
            for school in schools_in_region:
                SchoolPin.objects.update_or_create(
                    academic_year=current_year,
                    school=school,
                    defaults={
                        'is_pinned': True,
                        'pin_reason': 'region_restricted',
                        'notes': f"Region {region.name} is restricted for {current_year.year}"
                    }
                )
            status = "HIDDEN"
            # Unpin all schools (show)
            SchoolPin.objects.filter(
                academic_year=current_year,
                school__in=schools_in_region
            ).delete()
            status = "VISIBLE"
        
        messages.success(
            request, 
            f"✅ Region '{region.name}' is now {status} to students!"
        )
        
        return redirect('manage_regions')
    
    return redirect('manage_regions')
@staff_member_required
def change_academic_year(request):
    """Change active academic year from admin interface"""
    if request.method == 'POST':
        year_id = request.POST.get('academic_year_id')
        
        if not year_id:
            messages.error(request, "Please select an academic year")
            return redirect('manage_regions')
        
        try:
            new_year = AcademicYear.objects.get(id=year_id)
            
            # Deactivate all years
            AcademicYear.objects.all().update(is_active=False)
            
            # Activate selected year
            new_year.is_active = True
            new_year.save()
            
            messages.success(
                request,
                f"✅ Academic year changed to {new_year.year}\n"
                f"⚠️ Remember to set region pins for this new year!"
            )
            
        except AcademicYear.DoesNotExist:
            messages.error(request, "Academic year not found")
        
        return redirect('manage_regions')
    
    return redirect('manage_regions')
@staff_member_required
def create_academic_year(request):
    """Create new academic year from admin interface"""
    if request.method == 'POST':
        year_name = request.POST.get('year_name', '').strip()
        
        if not year_name:
            messages.error(request, "Please enter academic year (e.g., 2027/2028)")
            return redirect('manage_regions')
        
        # Validate format
        if '/' not in year_name:
            messages.error(request, "Use format YYYY/YYYY (e.g., 2027/2028)")
            return redirect('manage_regions')
        
        # Check if exists
        if AcademicYear.objects.filter(year=year_name).exists():
            messages.warning(request, f"Academic year {year_name} already exists!")
            # Create new year (not active by default)
            AcademicYear.objects.create(
                year=year_name,
                is_active=False
            )
            messages.success(request, f"✅ Academic year {year_name} created successfully!")
        
        return redirect('manage_regions')
    
    return redirect('manage_regions')
@staff_member_required
def reset_all_region_pins(request):
    """Make all regions visible (unpin everything)"""
    if request.method == 'POST':
        current_year = get_current_academic_year()
        
        # Delete all pins for current year
        region_deleted = RegionPin.objects.filter(academic_year=current_year).delete()
        school_deleted = SchoolPin.objects.filter(academic_year=current_year).delete()
        
        messages.success(
            request,
            f"✅ All {Region.objects.count()} regions are now VISIBLE to students!\n"
            f"Deleted {region_deleted[0]} region pins and {school_deleted[0]} school pins"
        )
        
        return redirect('manage_regions')
    
    return redirect('manage_regions')
    
@staff_member_required
def manage_regions(request):
    """Main management page for regions and academic years"""
    current_year = get_current_academic_year()
    all_years = AcademicYear.objects.all().order_by('-year')
    all_regions = Region.objects.all().order_by('name')
    
    # Get pin status for each region
    hidden_count = 0
    for region in all_regions:
        pin = RegionPin.objects.filter(
            academic_year=current_year, 
            region=region
        ).first()
        region.is_pinned = pin.is_pinned if pin else False
        if region.is_pinned:
            hidden_count += 1
    
    context = {
        'current_year': current_year,
        'all_years': all_years,
        'all_regions': all_regions,
        'hidden_count': hidden_count,
        'visible_count': all_regions.count() - hidden_count,
        'total_regions': all_regions.count(),
    }
    
    return render(request, 'field_app/manage_regions.html', context)       
# =========================
# ASSESSOR PASSWORD RESET VIEWS - ONGEZA HIZI MWISHONI MWA VIEWS.PY
# =========================

def assessor_password_reset(request):
    """Reset password for assessor - sends new temporary password via email"""

    if request.method == 'POST':
        email = request.POST.get('email', '').strip().lower()

        if not email:
            messages.error(request, 'Please enter your email address.')
            return redirect('assessor_login')

        try:
            assessor = Assessor.objects.get(email__iexact=email)
        except Assessor.DoesNotExist:
            messages.error(request, f'Hakuna assessor aliyepatikana na email: {email}')
            return redirect('assessor_login')

        # If assessor has no user account, create one automatically
        if not assessor.user:
            try:
                User = get_user_model()
                existing_user = User.objects.filter(email__iexact=email).first()
                if existing_user:
                    assessor.user = existing_user
                else:
                    new_user = User.objects.create_user(
                        email=email,
                        password=None,
                        is_staff=False,
                        is_active=True,
                    )
                    assessor.user = new_user
                assessor.save()
            except Exception as create_err:
                messages.error(request,
                    f'Imeshindwa kuunda akaunti: {str(create_err)[:100]}. '
                    f'Wasiliana na msimamizi.'
                )
                return redirect('assessor_login')

        # Generate new temp password and save it
        temp_password = generate_random_password()
        assessor.user.set_password(temp_password)
        assessor.user.save()

        # Send email
        login_url = request.build_absolute_uri(reverse('assessor_login'))
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Password Reset - Field Placement System</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 500px; margin: 0 auto; padding: 20px; }}
                .header {{ background: #2c3e50; color: white; padding: 20px; text-align: center; }}
                .content {{ padding: 20px; }}
                .password {{ background: #fef9e6; padding: 15px; border-left: 4px solid #f59e0b; margin: 20px 0; }}
                .footer {{ text-align: center; padding: 20px; font-size: 12px; color: #666; }}
                code {{ background: #f0f0f0; padding: 4px 8px; border-radius: 4px; font-size: 16px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>🔐 Password Reset</h2>
                    <p>Mfumo wa Ufuatiliaji wa Walimu Wanafunzi (IMS)</p>
                </div>
                <div class="content">
                    <p>Dear <strong>{assessor.full_name}</strong>,</p>
                    <p>You requested to reset your password. Here are your new login credentials:</p>
                    <div class="password">
                        <p><strong>📧 Email:</strong> {assessor.email}</p>
                        <p><strong>🔑 New Password:</strong> <code>{temp_password}</code></p>
                    </div>
                    <p><strong>⚠️ Important:</strong> Please change this password immediately after logging in.</p>
                    <p style="margin-top: 20px;">
                        <a href="{login_url}" style="background: #2c3e50; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px;">
                            Login to Dashboard
                        </a>
                    </p>
                </div>
                <div class="footer">
                    <p>This is an automated message. Please do not reply.</p>
                    <p>© {timezone.now().year} Wizara ya Elimu, Sayansi na Teknolojia — IMS v2.1.0</p>
                </div>
            </div>
        </body>
        </html>
        """
        text_content = (
            f"PASSWORD RESET - Field Placement System\n"
            f"{'='*50}\n\n"
            f"Dear {assessor.full_name},\n\n"
            f"Your new login credentials:\n"
            f"Email: {assessor.email}\n"
            f"Password: {temp_password}\n\n"
            f"Login URL: {login_url}\n\n"
            f"IMPORTANT: Change this password immediately after logging in.\n"
        )

        try:
            send_mail(
                subject='🔐 Badilisha Nywila - Field Placement System',
                message=text_content,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[assessor.email],
                html_message=html_content,
                fail_silently=False,
            )
            messages.success(request,
                f'✅ Nywila mpya ya muda imetumwa kwa {assessor.email}. '
                f'Angalia inbox na spam folder yako.'
            )
        except Exception as email_error:
            # Email failed — password is saved, show it on screen so admin can relay it.
            import logging
            logging.getLogger('django').error(
                f'Assessor password reset email failed for {assessor.email}: {email_error}'
            )
            messages.warning(request,
                f'Nywila mpya imewekwa lakini email imeshindwa kutumwa kwa {assessor.email}. '
                f'Nywila mpya ya muda ni: {temp_password} — toa kwa mkono kwa {assessor.full_name}. '
                f'(Hitilafu: {str(email_error)[:100]})'
            )

        return redirect('assessor_login')

    # GET request - redirect to login page
    return redirect('assessor_login')


def assessor_password_reset_done(request):
    """Show success message after password reset"""
    messages.success(request, 
        'Password reset email has been sent. Please check your inbox.'
    )
    return redirect('assessor_login')                     
# ============================================================
# ADD THIS AT THE END OF YOUR views.py FILE
# ============================================================

# =========================
# LOGIN PAGE WITH DYNAMIC DATA FROM DATABASE
# =========================

def login_page(request):
    """
    Custom login page that displays dynamic data from database:
    - Total partner colleges
    - Total students placed
    - Current academic year
    - Notices and announcements
    - List of partner colleges
    """
    
    # Get current academic year
    current_academic_year_obj = get_current_academic_year()
    current_academic_year = current_academic_year_obj.year if current_academic_year_obj else "2025/2026"
    
    # ========== TOTAL COLLEGES (Partner Schools) ==========
    # Use School model since that's what you have
    total_colleges = School.objects.filter(is_active=True).count()
    
    # If you have a College model, uncomment this:
    # total_colleges = College.objects.count()
    
    # ========== TOTAL STUDENTS PLACED ==========
    # Count students who have been approved and assigned to schools
    total_students_placed = StudentTeacher.objects.filter(
        approval_status='approved',
        selected_school__isnull=False
    ).count()
    
    # Alternative: Count from StudentApplication with approved status
    # total_students_placed = StudentApplication.objects.filter(
    #     status='approved'
    # ).values('student').distinct().count()
    
    # ========== NOTICES / ANNOUNCEMENTS ==========
    # You need to create a Notice model first (see below)
    # For now, using hardcoded fallback or get from database if model exists
    
    notices = []
    try:
        # If you have Notice model, use this:
        from .models import Notice  # Uncomment if you have Notice model
        notices = Notice.objects.filter(is_active=True).order_by('-created_at')[:5]
    except (ImportError, AttributeError):
        # Fallback notices if model doesn't exist yet
        notices = [
            {'title': 'Field placement applications for 2025/2026 are now open', 'date': timezone.now()},
            {'title': 'Deadline for application submission is 30th June 2025', 'date': timezone.now()},
            {'title': 'All students must complete logbook entries daily', 'date': timezone.now()},
        ]
    
    # ========== PARTNER COLLEGES LIST ==========
    # Get schools that have capacity and are active
    partner_colleges = School.objects.filter(
        is_active=True
    ).select_related('district', 'district__region').order_by('name')[:10]
    
    # If you need different format for the template
    colleges_data = []
    for college in partner_colleges:
        colleges_data.append({
            'name': college.name,
            'region': college.district.region.name if college.district else 'Tanzania',
        })
    
    # If you have a separate College model, use this:
    # colleges_data = College.objects.all()[:10]
    
    context = {
        'total_colleges': total_colleges,
        'total_students_placed': total_students_placed,
        'current_academic_year': current_academic_year,
        'notices': notices,
        'colleges': colleges_data,
        'hide_navbar': True,  # To hide navbar on login page
    }
    
    return render(request, 'field_app/registration/login.html', context)


# =========================
# NOTICE MODEL - ADD THIS TO YOUR models.py
# =========================
"""
Add this to your models.py file:

class Notice(models.Model):
    title = models.CharField(max_length=200)
    content = models.TextField(blank=True)
    is_active = models.BooleanField(default=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
"""


# =========================
# ALTERNATIVE: If you don't want to create Notice model yet
# Use this simpler version with hardcoded notices but dynamic stats
# =========================

def login_page_simple(request):
    """
    Simplified login page - stats from database, notices hardcoded
    Use this if you don't want to create Notice model yet
    """
    
    # Get current academic year
    current_academic_year_obj = get_current_academic_year()
    current_academic_year = current_academic_year_obj.year if current_academic_year_obj else "2025/2026"
    
    # Total partner colleges (schools with capacity > 0)
    total_colleges = School.objects.filter(
        is_active=True,
        capacity__gt=0
    ).count()
    
    # Total students placed
    total_students_placed = StudentTeacher.objects.filter(
        approval_status='approved',
        selected_school__isnull=False
    ).count()
    
    # Partner colleges list
    partner_colleges = School.objects.filter(
        is_active=True
    ).select_related('district__region').order_by('name')[:10]
    
    # Hardcoded notices (will be shown even without database)
    notices = [
        {'title': 'Field placement applications for 2025/2026 are now open', 'date': 'March 2025'},
        {'title': 'Deadline for application submission is 30th June 2025', 'date': 'March 2025'},
        {'title': 'All students must complete logbook entries daily', 'date': 'March 2025'},
        {'title': 'Contact your academic advisor for placement inquiries', 'date': 'March 2025'},
    ]
    
    context = {
        'total_colleges': total_colleges if total_colleges > 0 else 6,
        'total_students_placed': total_students_placed if total_students_placed > 0 else 2400,
        'current_academic_year': current_academic_year,
        'notices': notices,
        'colleges': partner_colleges,
        'hide_navbar': True,
    }
    
    return render(request, 'field_app/registration/login.html', context)


# =========================
# UPDATE URLS.PY - Add this to your urls.py
# =========================
"""
In your main urls.py or app urls.py, add:

from django.urls import path
from your_app.views import login_page  # or login_page_simple

urlpatterns = [
    # ... other URLs ...
    path('', login_page, name='login_page'),  # Or use 'login_page_simple'
    path('login/', login_page, name='login'),  # If you want to replace default login
]
"""  
def homepage(request):
    """Homepage that includes meta tag for Google verification"""
    return render(request, 'field_app/base.html')  
# views.py


@login_required
def generate_scheme_view(request):
    from .models import EducationLevel

    form = SchemeOfWorkForm()
    education_levels = EducationLevel.objects.all().order_by('order')

    return render(request, 'field_app/generate_scheme.html', {
        'form': form,
        'education_levels': education_levels,
    })
@login_required
def ajax_generate_scheme(request):
    """API ya AI kuzalisha Scheme of Work kwa format ya KitabuSmart"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            # Extract form data
            education_level = data.get('education_level')
            class_name = data.get('class_name')
            subject = data.get('subject')
            term = data.get('term')
            year = data.get('year')
            syllabus = data.get('syllabus', 'New Syllabus')
            total_weeks = data.get('total_weeks', 12)
            periods_per_week = data.get('periods_per_week', 8)
            start_date = data.get('start_date')
            end_date = data.get('end_date')
            teacher_name = data.get('teacher_name')
            school_name = data.get('school_name')
            reference_source = data.get('reference_source', '')
            breaks = data.get('breaks', [])
            
            # Build breaks text for prompt
            breaks_text = ""
            if breaks:
                breaks_text = "\nBreaks (holidays/exams) to respect:\n"
                for b in breaks:
                    breaks_text += f"- {b.get('name', 'Break')}: {b.get('start', '')} to {b.get('end', '')}\n"
            
            # Prompt ya AI kufuata format ya KitabuSmart
            prompt = f"""
You are an AI assistant for Tanzanian teachers. Generate a complete Scheme of Work following EXACTLY the KitabuSmart format.

Input details:
- Education Level: {education_level}
- Class: {class_name}
- Subject: {subject}
- Term: {term} {year}
- Syllabus: {syllabus}
- Total Weeks: {total_weeks} weeks
- Periods per Week: {periods_per_week}
- Start Date: {start_date}
- End Date: {end_date}
- Teacher: {teacher_name}
- School: {school_name}
- Reference Source: {reference_source}
{breaks_text}
The output MUST be a JSON list of objects. Each object must have exactly these 12 keys (column names):
"Main Competence", "Specific Competence", "Learning Activities", "Specific Learning Activities", "Month", "Week", "Periods", "Reference", "Teaching & Learning Methods", "Teaching & Learning Resources", "Assessment Tools", "Remarks"

Requirements:
1. Distribute content across {total_weeks} weeks, respecting any breaks (skip weeks that fall on breaks).
2. For each week, assign appropriate Month (e.g., MAY, JUNE, JULY, AUGUST, SEPTEMBER, OCTOBER).
3. "Week" column should be like "1st", "2nd", "3rd", etc.
4. "Periods" should be {periods_per_week} for normal weeks.
5. "Reference" should include {reference_source} with page numbers (e.g., "ENGLISH_iv-vii.pdf, page 10-15").
6. "Main Competence" should be numbered like "1.0 Demonstrate mastery of BASIC MATHEMATICS fundamental principles".
7. Content must be realistic for {education_level} {class_name} {subject} in Tanzania.
8. After the last week, add a row with remarks about examination preparation if needed.
9. Return ONLY valid JSON, no extra text.
Example row:
{{"Main Competence": "1.0 Demonstrate mastery of concepts", "Specific Competence": "Understand numbers", "Learning Activities": "Group discussion", "Specific Learning Activities": "Define numbers", "Month": "MAY", "Week": "1st", "Periods": 8, "Reference": "book.pdf, page 5-10", "Teaching & Learning Methods": "Think-Pair-Share", "Teaching & Learning Resources": "Charts", "Assessment Tools": "Quizzes", "Remarks": "Emphasize basics"}}
"""
            
            # ========== CALL GEMINI AI ==========
            print("🤖 Calling Gemini AI...")
            
            response = client.models.generate_content(model=model_name, contents=prompt)
            response_text = response.text
            
            print(f"✅ Gemini response received ({len(response_text)} characters)")
            
            # Extract JSON
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                json_data = json_match.group()
            else:
                json_data = response_text
            
            try:
                scheme_data = json.loads(json_data)
            except Exception as e:
                print("JSON parse error:", e)
                scheme_data = []

            # Save to database
            saved_id = None
            try:
                student = StudentTeacher.objects.get(user=request.user)
                school = student.selected_school
                if school and scheme_data:
                    level_map = {
                        'primary school': 'primary',
                        'ordinary level': 'ordinary',
                        'advanced level': 'advanced',
                    }
                    edu_level = level_map.get(education_level.lower(), 'ordinary')
                    subj_obj = Subject.objects.filter(name__iexact=subject).first()
                    if subj_obj:
                        start_dt = None
                        end_dt = None
                        if start_date:
                            try:
                                from datetime import date as _dt
                                start_dt = _dt.fromisoformat(start_date)
                            except Exception:
                                pass
                        if end_date:
                            try:
                                from datetime import date as _dt
                                end_dt = _dt.fromisoformat(end_date)
                            except Exception:
                                pass
                        scheme_obj, _ = SchemeOfWork.objects.update_or_create(
                            student=student,
                            subject=subj_obj,
                            term=term,
                            year=int(year),
                            defaults={
                                'school': school,
                                'education_level': edu_level,
                                'class_name': class_name,
                                'syllabus': syllabus,
                                'total_weeks': int(total_weeks),
                                'periods_per_week': int(periods_per_week),
                                'start_date': start_dt,
                                'end_date': end_dt,
                                'teacher_name': teacher_name,
                                'reference_source': reference_source,
                                'breaks': breaks,
                                'scheme_data': scheme_data,
                                'generated_by_ai': True,
                            }
                        )
                        saved_id = scheme_obj.id
            except Exception as save_err:
                print(f"Scheme save error (non-fatal): {save_err}")

            return JsonResponse({'success': True, 'data': scheme_data, 'saved_id': saved_id})

        except Exception as e:
            import traceback
            traceback.print_exc()
            raw = str(e)
            if 'PERMISSION_DENIED' in raw or 'suspended' in raw.lower() or '403' in raw:
                user_msg = (
                    "Huduma ya AI imesimamishwa kwa sababu ya ufunguo wa API uliosimamishwa. "
                    "Tafadhali wasiliana na msimamizi ili upate ufunguo mpya wa Google AI."
                )
            elif 'quota' in raw.lower() or '429' in raw:
                user_msg = "Kikomo cha matumizi ya AI kimefikiwa. Jaribu tena baadaye."
            elif 'API_KEY' in raw or 'api_key' in raw.lower():
                user_msg = "Ufunguo wa API ya AI haujawekwa. Wasiliana na msimamizi."
            else:
                user_msg = "Hitilafu ya ndani imetokea. Tafadhali jaribu tena."
            return JsonResponse({'success': False, 'error': user_msg}, status=500)

    return JsonResponse({'success': False}, status=400)

@login_required
def download_scheme_pdf(request):
    """Generate PDF ya Scheme of Work kwa mtindo wa KitabuSmart"""
    if request.method == 'POST':
        data = json.loads(request.body)
        scheme_data = data.get('scheme_data')
        subject = data.get('subject')
        class_name = data.get('class_name')
        term = data.get('term')
        year = data.get('year')
        syllabus = data.get('syllabus')
        teacher_name = data.get('teacher_name')
        school_name = data.get('school_name')
        total_weeks = data.get('total_weeks')
        
        buffer = BytesIO()
        # Use landscape orientation for wide table
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), 
                                rightMargin=20, leftMargin=20, topMargin=30, bottomMargin=30)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=14, spaceAfter=10, alignment=1)
        elements.append(Paragraph(f"{subject} - {class_name} - {term} {year} - {syllabus}", title_style))
        elements.append(Spacer(1, 6))
        
        # Teacher and School info
        info_style = ParagraphStyle('InfoStyle', parent=styles['Normal'], fontSize=9)
        elements.append(Paragraph(f"Teacher: {teacher_name} | School: {school_name} | Total Weeks: {total_weeks}", info_style))
        elements.append(Spacer(1, 12))
        
        if scheme_data:
            headers = list(scheme_data[0].keys())
            # Wrap long header text
            wrapped_headers = [h.replace(' & ', '&\n') for h in headers]
            table_data = [wrapped_headers]
            
            for row in scheme_data:
                row_data = []
                for h in headers:
                    val = row.get(h, '')
                    # Convert to string and handle long text
                    row_data.append(str(val) if val else '')
                table_data.append(row_data)
            
            # Calculate column widths based on content
            col_widths = [70, 80, 70, 80, 50, 40, 40, 60, 70, 70, 60, 60]  # approximate
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 6),
                ('WORDWRAP', (0, 0), (-1, -1), True),
            ]))
            elements.append(table)
        
        doc.build(elements)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Scheme_of_Work_{subject}_{class_name}.pdf"'
        return response

    return HttpResponse("Invalid request", status=400)


@login_required
def download_scheme_word(request):
    """Export Scheme of Work as Word (.docx)"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as _json

    data = _json.loads(request.body)
    scheme_data = data.get('scheme_data', [])
    subject = data.get('subject', '')
    class_name = data.get('class_name', '')
    term = data.get('term', '')
    year = data.get('year', '')
    teacher_name = data.get('teacher_name', '')
    school_name = data.get('school_name', '')

    doc = Document()
    doc.core_properties.title = f"Scheme of Work — {subject}"

    # Title
    title = doc.add_heading(f"{subject} — {class_name} — Term {term} {year}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph(f"Teacher: {teacher_name}    School: {school_name}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if scheme_data:
        headers = list(scheme_data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(8)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for row in scheme_data:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                row_cells[i].text = str(row.get(h, '') or '')
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"Scheme_{subject}_{class_name}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


@login_required
def download_lesson_plan_pdf(request):
    """Export Lesson Plan as PDF"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    import json as _json

    data = _json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form = data.get('form_data', {})

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    elements = []
    styles = getSampleStyleSheet()
    bold_style = ParagraphStyle('Bold', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10)
    normal_style = ParagraphStyle('Normal9', parent=styles['Normal'], fontSize=9, leading=14)
    heading_style = ParagraphStyle('Heading', parent=styles['Normal'], fontName='Helvetica-Bold',
                                   fontSize=11, textColor=colors.HexColor('#0A2B5E'),
                                   spaceAfter=4, spaceBefore=10)

    elements.append(Paragraph("LESSON PLAN", ParagraphStyle('Title', parent=styles['Heading1'],
                                                              fontSize=16, alignment=1,
                                                              textColor=colors.HexColor('#0A2B5E'))))
    elements.append(Spacer(1, 10))

    # Meta info table
    meta = [
        ['Teacher:', form.get('teacher_name', ''), 'Subject:', form.get('subject', '')],
        ['Class:', form.get('class_name', ''), 'Term/Year:', f"Term {form.get('term','')} {form.get('year','')}"],
        ['Topic:', form.get('topic', ''), 'Subtopic:', form.get('subtopic', '')],
        ['Duration:', f"{form.get('duration','')} min", 'Date:', str(timezone.now().date())],
        ['Students:', f"{form.get('total_students','')} total / {form.get('present_students','')} present", '', ''],
    ]
    meta_table = Table(meta, colWidths=[80, 180, 80, 175])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#EEF1F6')),
        ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#EEF1F6')),
    ]))
    elements.append(meta_table)
    elements.append(Spacer(1, 10))

    # Competences
    for label, key in [('Main Competence', 'main_competence'),
                       ('Specific Competence', 'specific_competence'),
                       ('Previous Knowledge', 'previous_knowledge')]:
        val = lesson.get(key, '')
        if val:
            elements.append(Paragraph(f"<b>{label}:</b> {val}", normal_style))

    # Lists
    for label, key in [('Learning Objectives', 'learning_objectives'),
                       ('Teaching Methods', 'teaching_methods'),
                       ('Teaching Resources', 'teaching_resources')]:
        items = lesson.get(key, [])
        if items:
            elements.append(Paragraph(label, heading_style))
            for item in items:
                elements.append(Paragraph(f"• {item}", normal_style))

    # Lesson Development table
    ld = lesson.get('lesson_development', [])
    if ld:
        elements.append(Paragraph("Lesson Development", heading_style))
        ld_headers = ['Time', 'Stage', 'Teacher Activities', 'Student Activities', 'Assessment']
        ld_data = [ld_headers]
        for stage in ld:
            ld_data.append([
                stage.get('time', ''),
                stage.get('stage', stage.get('phase', '')),
                stage.get('teacher_activities', ''),
                stage.get('student_activities', ''),
                stage.get('assessment_criteria', ''),
            ])
        ld_table = Table(ld_data, colWidths=[45, 70, 120, 120, 100], repeatRows=1)
        ld_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0A2B5E')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('WORDWRAP', (0, 0), (-1, -1), True),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F7FAFF')]),
        ]))
        elements.append(ld_table)

    remarks = lesson.get('remarks', '')
    if remarks:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(f"<b>Remarks:</b> {remarks}", normal_style))

    doc.build(elements)
    buffer.seek(0)
    safe_name = f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.pdf"'
    return response


@login_required
def download_lesson_plan_word(request):
    """Export Lesson Plan as Word (.docx)"""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as _json

    data = _json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form = data.get('form_data', {})

    doc = Document()
    doc.add_heading('LESSON PLAN', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta_table = doc.add_table(rows=5, cols=4)
    meta_table.style = 'Table Grid'
    rows = meta_table.rows
    def _set(r, c, text, bold=False):
        cell = rows[r].cells[c]
        cell.text = text
        if bold and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    _set(0,0,'Teacher:', True); _set(0,1, form.get('teacher_name',''))
    _set(0,2,'Subject:', True); _set(0,3, form.get('subject',''))
    _set(1,0,'Class:', True);   _set(1,1, form.get('class_name',''))
    _set(1,2,'Term/Year:', True);_set(1,3, f"Term {form.get('term','')} {form.get('year','')}")
    _set(2,0,'Topic:', True);   _set(2,1, form.get('topic',''))
    _set(2,2,'Subtopic:', True);_set(2,3, form.get('subtopic',''))
    _set(3,0,'Duration:', True);_set(3,1, f"{form.get('duration','')} min")
    _set(3,2,'Date:', True);    _set(3,3, str(timezone.now().date()))
    _set(4,0,'Students:', True);_set(4,1, f"{form.get('total_students','')} / {form.get('present_students','')}")

    doc.add_paragraph()

    for label, key in [('Main Competence', 'main_competence'),
                       ('Specific Competence', 'specific_competence'),
                       ('Previous Knowledge', 'previous_knowledge')]:
        val = lesson.get(key, '')
        if val:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val)

    for label, key in [('Learning Objectives', 'learning_objectives'),
                       ('Teaching Methods', 'teaching_methods'),
                       ('Teaching Resources', 'teaching_resources')]:
        items = lesson.get(key, [])
        if items:
            doc.add_heading(label, level=2)
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

    ld = lesson.get('lesson_development', [])
    if ld:
        doc.add_heading('Lesson Development', level=2)
        ld_table = doc.add_table(rows=1, cols=5)
        ld_table.style = 'Table Grid'
        hdr = ld_table.rows[0].cells
        for i, h in enumerate(['Time', 'Stage', 'Teacher Activities', 'Student Activities', 'Assessment']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for stage in ld:
            row = ld_table.add_row().cells
            row[0].text = stage.get('time', '')
            row[1].text = stage.get('stage', stage.get('phase', ''))
            row[2].text = stage.get('teacher_activities', '')
            row[3].text = stage.get('student_activities', '')
            row[4].text = stage.get('assessment_criteria', '')

    remarks = lesson.get('remarks', '')
    if remarks:
        p = doc.add_paragraph()
        p.add_run('Remarks: ').bold = True
        p.add_run(remarks)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


# views.py - Ongeza hizi

from django.http import JsonResponse
from .models import EducationLevel, ClassLevel, Subject, Textbook

@login_required
def get_classes_by_level(request):
    """AJAX endpoint to get classes based on education level"""
    level_id = request.GET.get('level_id')
    if level_id:
        classes = ClassLevel.objects.filter(education_level_id=level_id).values('id', 'name')
        return JsonResponse(list(classes), safe=False)
    return JsonResponse([], safe=False)

@login_required
def get_subjects_by_level(request):
    """AJAX endpoint to get subjects based on education level"""
    level_id = request.GET.get('level_id')
    if level_id:
        try:
            education_level = EducationLevel.objects.get(id=level_id)
            level_name = education_level.name.lower()
            
            # Filter subjects based on education level
            if 'primary' in level_name:
                subjects = Subject.objects.filter(level='primary')
            elif 'ordinary' in level_name or 'secondary' in level_name:
                subjects = Subject.objects.filter(level='secondary')
                subjects = Subject.objects.all()
            
            return JsonResponse(list(subjects.values('id', 'name')), safe=False)
        except:
            return JsonResponse([], safe=False)
    return JsonResponse([], safe=False)

@login_required
def get_textbooks_by_level(request):
    """AJAX endpoint to get textbooks based on education level"""
    level_id = request.GET.get('level_id')
    if level_id:
        try:
            education_level = EducationLevel.objects.get(id=level_id)
            textbooks = Textbook.objects.filter(
                education_level=education_level.name.lower(),
                is_active=True
            ).values('id', 'title')
            return JsonResponse(list(textbooks), safe=False)
        except:
            return JsonResponse([], safe=False)
    return JsonResponse([], safe=False)    
# =========================
# LESSON PLAN GENERATOR VIEWS
# =========================

# =========================
# LESSON PLAN GENERATOR VIEWS - CLEAN VERSION
# =========================

@login_required
@login_required
def lesson_plan_view(request):
    """Display lesson plan generator form"""
    from .models import EducationLevel, Subject

    education_levels = EducationLevel.objects.all().order_by('order')
    subjects = Subject.objects.all().order_by('name')

    return render(request, 'field_app/lesson_plan.html', {
        'education_levels': education_levels,
        'subjects': subjects,
    })

@login_required
def ajax_generate_lessonplan(request):
    """Generate lesson plan using AI"""
    if request.method == 'POST':
        try:
            import json as json_module
            import re
            
            data = json_module.loads(request.body)

            # Extract form data
            education_level = data.get('education_level', '')
            class_name = data.get('class_name', '')
            subject = data.get('subject', '')
            subject_id = data.get('subject_id', '')
            topic = data.get('topic', '')
            subtopic = data.get('subtopic', '')
            term = data.get('term', 'I')
            year = data.get('year', 2026)
            duration = data.get('duration', 40)
            total_students = data.get('total_students', '')
            present_students = data.get('present_students', '')
            learning_objectives = data.get('learning_objectives', '')
            teaching_methods = data.get('teaching_methods', '')
            reference_source = data.get('reference_source', '')
            teacher_name = data.get('teacher_name', '')
            
            # Build AI prompt
            prompt = f"""
You are an AI assistant for Tanzanian teachers. Generate a detailed LESSON PLAN following TIE Tanzania standards.

Input Details:
- Education Level: {education_level}
- Class: {class_name}
- Subject: {subject}
- Topic: {topic}
- Subtopic: {subtopic}
- Term: {term}
- Year: {year}
- Duration: {duration} minutes
- Total Students: {total_students}
- Present Students: {present_students}
- Learning Objectives: {learning_objectives}
- Teaching Methods: {teaching_methods}
- Reference Source: {reference_source}

Output MUST be ONLY valid JSON. Do not include any other text. Use this exact structure:
{{
    "lesson_title": "{subject} - {topic}",
    "date": "today's date",
    "main_competence": "Main competence here",
    "specific_competence": "Specific competence here",
    "previous_knowledge": "Previous knowledge here",
    "learning_objectives": ["Objective 1", "Objective 2"],
    "teaching_methods": ["Method 1", "Method 2"],
    "teaching_resources": ["Resource 1", "Resource 2"],
    "lesson_development": [
        {{"time": "5 min", "stage": "Introduction", "teacher_activities": "Activities", "student_activities": "Activities", "assessment_criteria": "Criteria"}},
        {{"time": "15 min", "stage": "Presentation", "teacher_activities": "Activities", "student_activities": "Activities", "assessment_criteria": "Criteria"}},
        {{"time": "15 min", "stage": "Practice", "teacher_activities": "Activities", "student_activities": "Activities", "assessment_criteria": "Criteria"}},
        {{"time": "5 min", "stage": "Conclusion", "teacher_activities": "Activities", "student_activities": "Activities", "assessment_criteria": "Criteria"}}
    ],
    "remarks": "Remarks here"
}}
"""
            from .ai_utils import client, model_name
            
            print("🤖 Calling Gemini AI for Lesson Plan...")
            
            response = client.models.generate_content(model=model_name, contents=prompt)
            response_text = response.text
            
            print(f"✅ Gemini response received ({len(response_text)} characters)")
            
            # Clean response
            cleaned_text = re.sub(r'```json\s*', '', response_text)
            cleaned_text = re.sub(r'```\s*', '', cleaned_text)
            cleaned_text = cleaned_text.strip()
            
            # Find JSON object
            start_idx = cleaned_text.find('{')
            end_idx = cleaned_text.rfind('}')
            
            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_data = cleaned_text[start_idx:end_idx + 1]
                lesson_data = json_module.loads(json_data)
            else:
                # Fallback response
                lesson_data = {
                    "lesson_title": f"{subject} - {topic}",
                    "date": timezone.now().strftime('%Y-%m-%d'),
                    "main_competence": f"Demonstrate understanding of {topic}",
                    "specific_competence": f"Explain key concepts of {topic}",
                    "previous_knowledge": "Basic knowledge from previous lessons",
                    "learning_objectives": [f"Define {topic}", f"Explain {topic}", f"Apply {topic}"],
                    "teaching_methods": ["Lecture", "Discussion", "Question and Answer"],
                    "teaching_resources": ["Chalkboard", "Textbook", "Handouts"],
                    "lesson_development": [
                        {"time": "5 min", "stage": "Introduction", "teacher_activities": f"Introduce {topic}", "student_activities": "Listen and respond", "assessment_criteria": "Participation"},
                        {"time": "15 min", "stage": "Presentation", "teacher_activities": "Explain key concepts", "student_activities": "Take notes and ask questions", "assessment_criteria": "Understanding demonstrated"},
                        {"time": "15 min", "stage": "Practice", "teacher_activities": "Guide through examples", "student_activities": "Work on exercises", "assessment_criteria": "Correct answers"},
                        {"time": "5 min", "stage": "Conclusion", "teacher_activities": "Summarize key points", "student_activities": "Review and ask questions", "assessment_criteria": "Recall of main points"}
                    ],
                    "remarks": "Students participated well"
                }

            # Save to database
            saved_id = None
            save_error_msg = None
            try:
                student = StudentTeacher.objects.get(user=request.user)
                school = student.selected_school
                if not school:
                    save_error_msg = 'Student has no selected school'
                else:
                    level_map = {
                        'primary school': 'primary',
                        'ordinary level': 'ordinary',
                        'advanced level': 'advanced',
                    }
                    edu_level = level_map.get((education_level or '').lower(), 'ordinary')

                    # Try subject lookup by ID first (reliable), then by name
                    subj_obj = None
                    if subject_id:
                        try:
                            subj_obj = Subject.objects.get(id=int(subject_id))
                        except (Subject.DoesNotExist, ValueError):
                            pass
                    if not subj_obj and subject:
                        subj_obj = Subject.objects.filter(name__iexact=subject).first()

                    if not subj_obj:
                        save_error_msg = f'Subject not found: id={subject_id} name={subject}'
                    else:
                        lp_date = timezone.now().date()
                        lo_list = lesson_data.get('learning_objectives', [])
                        if isinstance(lo_list, str):
                            lo_list = [x.strip() for x in lo_list.split('\n') if x.strip()]
                        tm_list = lesson_data.get('teaching_methods', [])
                        if isinstance(tm_list, str):
                            tm_list = [x.strip() for x in tm_list.split('\n') if x.strip()]
                        lp_obj = LessonPlan.objects.create(
                            student=student,
                            school=school,
                            subject=subj_obj,
                            education_level=edu_level,
                            class_name=class_name,
                            term=term,
                            year=int(year),
                            topic=topic,
                            subtopic=subtopic or '',
                            date=lp_date,
                            duration=int(duration) if duration else 40,
                            total_students=int(total_students) if total_students else 0,
                            present_students=int(present_students) if present_students else 0,
                            teacher_name=teacher_name or student.full_name,
                            reference_source=reference_source or '',
                            main_competence=lesson_data.get('main_competence', ''),
                            specific_competence=lesson_data.get('specific_competence', ''),
                            previous_knowledge=lesson_data.get('previous_knowledge', ''),
                            learning_objectives=lo_list,
                            teaching_methods=tm_list,
                            teaching_resources=lesson_data.get('teaching_resources', []),
                            lesson_development=lesson_data.get('lesson_development', []),
                            remarks=lesson_data.get('remarks', ''),
                            generated_by_ai=True,
                        )
                        saved_id = lp_obj.id
            except StudentTeacher.DoesNotExist:
                save_error_msg = f'No StudentTeacher for user {request.user}'
            except Exception as save_err:
                import traceback as _tb
                save_error_msg = str(save_err)
                print(f"Lesson Plan save error: {save_error_msg}")
                print(_tb.format_exc())

            if save_error_msg:
                print(f"[LessonPlan] NOT saved — {save_error_msg}")

            return JsonResponse({'success': True, 'data': lesson_data, 'saved_id': saved_id})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raw = str(e)
            if 'PERMISSION_DENIED' in raw or 'suspended' in raw.lower() or '403' in raw:
                user_msg = (
                    "Huduma ya AI imesimamishwa kwa sababu ya ufunguo wa API uliosimamishwa. "
                    "Tafadhali wasiliana na msimamizi ili upate ufunguo mpya wa Google AI."
                )
            elif 'quota' in raw.lower() or '429' in raw:
                user_msg = "Kikomo cha matumizi ya AI kimefikiwa. Jaribu tena baadaye."
            elif 'API_KEY' in raw or 'api_key' in raw.lower():
                user_msg = "Ufunguo wa API ya AI haujawekwa. Wasiliana na msimamizi."
            else:
                user_msg = "Hitilafu ya ndani imetokea. Tafadhali jaribu tena."
            return JsonResponse({'success': False, 'error': user_msg}, status=500)

    return JsonResponse({'success': False}, status=400)
@login_required
def api_get_schools(request):
    """API endpoint for AJAX school search"""
    district_id = request.GET.get('district_id')
    level = request.GET.get('level', 'Secondary')
    query = request.GET.get('q', '').strip()
    
    if not district_id:
        return JsonResponse({'success': False, 'error': 'District ID required'})
    
    district = get_object_or_404(District, id=district_id)
    current_year = get_current_academic_year()
    
    # Get pinned schools
    pinned_school_ids = []
    if current_year:
        pinned_school_ids = list(SchoolPin.objects.filter(
            academic_year=current_year,
            is_pinned=True
        ).values_list('school_id', flat=True))
    
    # Base queryset
    schools_qs = School.objects.filter(district=district, level=level)
    
    if query:
        schools_qs = schools_qs.filter(name__icontains=query)
    
    # Process schools
    schools_data = []
    available_count = 0
    pinned_count = 0
    full_count = 0
    
    for school in schools_qs:
        is_pinned = school.id in pinned_school_ids
        is_selectable = (not is_pinned) and (school.current_students < school.capacity)
        
        if is_pinned:
            pinned_count += 1
        elif not is_selectable:
            full_count += 1
            available_count += 1
        
        occupancy = round((school.current_students / school.capacity) * 100) if school.capacity > 0 else 0
        
        schools_data.append({
            'id': school.id,
            'name': school.name,
            'level_display': school.get_level_display(),
            'current_students': school.current_students,
            'capacity': school.capacity,
            'occupancy_percentage': occupancy,
            'is_pinned': is_pinned,
            'is_selectable': is_selectable,
        })
    
    return JsonResponse({
        'success': True,
        'schools': schools_data,
        'total_schools': schools_qs.count(),
        'available_schools': available_count,
        'pinned_schools': pinned_count,
        'full_schools': full_count,
    })
@login_required
def api_select_school_temp(request):
    """Temporarily store selected school in session"""
    if request.method == 'POST':
        data = json.loads(request.body)
        school_id = data.get('school_id')
        if school_id:
            request.session['temp_selected_school_id'] = school_id
            return JsonResponse({'success': True})
    return JsonResponse({'success': False})
@login_required
def api_districts_by_region(request):
    """Return districts for a given region - used in admin BoardMember form"""
    region_id = request.GET.get('region_id')
    if not region_id:
        return JsonResponse([], safe=False)
    districts = District.objects.filter(region_id=region_id).order_by('name').values('id', 'name')
    return JsonResponse(list(districts), safe=False)


def api_clear_selected_school(request):
    """Clear selected school from session"""
    if request.method == 'POST':
        request.session.pop('temp_selected_school', None)
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})    
@login_required
def api_filter_schools(request):
    """API endpoint for filtering schools - shows only selected school"""
    district_id = request.GET.get('district_id')
    selected_school_id = request.GET.get('selected_school_id')
    
    if not district_id:
        return JsonResponse({'success': False, 'error': 'District ID required'})
    
    district = get_object_or_404(District, id=district_id)
    current_year = get_current_academic_year()
    
    # Get pinned schools
    pinned_school_ids = []
    if current_year:
        pinned_school_ids = list(SchoolPin.objects.filter(
            academic_year=current_year,
            is_pinned=True
        ).values_list('school_id', flat=True))
    
    # Get all schools first
    schools_qs = School.objects.filter(district=district)
    
    # If selected_school_id is provided, show ONLY that school
    if selected_school_id:
        schools_qs = schools_qs.filter(id=selected_school_id)
    
    schools_data = []
    for school in schools_qs:
        is_pinned = school.id in pinned_school_ids
        is_selectable = (not is_pinned) and (school.current_students < school.capacity)
        occupancy = round((school.current_students / school.capacity) * 100) if school.capacity > 0 else 0
        
        schools_data.append({
            'id': school.id,
            'name': school.name,
            'level_display': school.get_level_display(),
            'current_students': school.current_students,
            'capacity': school.capacity,
            'occupancy_percentage': occupancy,
            'is_pinned': is_pinned,
            'is_selectable': is_selectable,
        })
    
    return JsonResponse({
        'success': True,
        'schools': schools_data,
        'total_schools': schools_qs.count(),
    })    
def set_language(request):
    """Switch UI language between English and Swahili. Stored in session."""
    lang = request.GET.get('lang', 'en')
    if lang not in ('en', 'sw'):
        lang = 'en'
    request.session['ui_lang'] = lang

    # Ruhusu redirect kwa URLs za ndani tu (zuia open redirect)
    next_url = request.GET.get('next') or request.META.get('HTTP_REFERER') or '/'
    if next_url.startswith('http') or not next_url.startswith('/'):
        next_url = '/'
    return redirect(next_url)


@staff_member_required
def create_admin(request):
    User = get_user_model()
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '').strip()
        if not email or not password:
            messages.error(request, "Email na password zinahitajika.")
        elif User.objects.filter(email=email).exists():
            messages.error(request, f"Mtumiaji mwenye email '{email}' tayari yupo.")
        elif len(password) < 6:
            messages.error(request, "Password iwe na herufi 6 au zaidi.")
        else:
            User.objects.create_superuser(email=email, password=password)
            messages.success(request, f"✅ Admin '{email}' ametengenezwa.")
            return redirect('admin_dashboard')
    return render(request, 'field_app/create_admin.html')


# =========================
# ADMIN REPORTS
# =========================

@staff_member_required
def admin_report_pdf(request):
    """Generate comprehensive admin report PDF"""
    try:
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors as rl_colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
    except ImportError as e:
        return HttpResponse(f"ReportLab library error: {e}", status=500)

    # Alias so rest of code uses rl_colors not module-level colors
    colors = rl_colors

    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        elements = []
        styles = getSampleStyleSheet()

        NAVY = colors.HexColor('#0A2B5E')
        GOLD = colors.HexColor('#C8900A')
        LIGHT = colors.HexColor('#EEF1F6')
        LIGHT2 = colors.HexColor('#FFF8ED')

        bold_style = ParagraphStyle('IMSBold', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, spaceAfter=6, textColor=NAVY)
        sub_style  = ParagraphStyle('IMSSub',  parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, spaceAfter=4, textColor=GOLD)
        norm_style = ParagraphStyle('IMSNorm', parent=styles['Normal'], fontSize=9, spaceAfter=4)

        current_year = get_current_academic_year()
        report_date  = timezone.now().strftime('%d %B %Y')

        elements.append(Paragraph('INTERNSHIP MANAGEMENT SYSTEM (IMS)', bold_style))
        elements.append(Paragraph('Teacher Colleges in Tanzania', sub_style))
        elements.append(Paragraph(f'Comprehensive Report — {current_year.year if current_year else "N/A"} | Generated: {report_date}', norm_style))
        elements.append(HRFlowable(width="100%", thickness=2, color=NAVY))
        elements.append(Spacer(1, 0.3*cm))

        # ── Summary stats ──
        total_students    = StudentTeacher.objects.count()
        approved_students = StudentTeacher.objects.filter(approval_status='approved').count()
        pending_students  = StudentTeacher.objects.filter(approval_status='pending').count()
        total_assessors   = Assessor.objects.count()
        total_schools     = School.objects.count()
        total_logbooks    = LogbookEntry.objects.count()
        verified_logbooks = LogbookEntry.objects.filter(is_location_verified=True).count()

        elements.append(Paragraph('SUMMARY STATISTICS', sub_style))
        summary_data = [
            ['Metric', 'Count'],
            ['Total Students Registered', str(total_students)],
            ['Approved Students', str(approved_students)],
            ['Pending Students', str(pending_students)],
            ['Total Assessors', str(total_assessors)],
            ['Partner Schools', str(total_schools)],
            ['Total Logbook Entries', str(total_logbooks)],
            ['GPS Verified Entries', str(verified_logbooks)],
        ]
        summary_table = Table(summary_data, colWidths=[10*cm, 5*cm])
        sum_style_cmds = [
            ('BACKGROUND', (0,0), (-1,0), NAVY),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E0')),
            ('ALIGN', (1,0), (1,-1), 'CENTER'),
            ('PADDING', (0,0), (-1,-1), 6),
        ]
        for row_i in range(1, len(summary_data)):
            bg = LIGHT if row_i % 2 == 0 else colors.white
            sum_style_cmds.append(('BACKGROUND', (0,row_i), (-1,row_i), bg))
        summary_table.setStyle(TableStyle(sum_style_cmds))
        elements.append(summary_table)
        elements.append(Spacer(1, 0.5*cm))

        # ── Students table ──
        elements.append(Paragraph('STUDENTS LIST', sub_style))
        students = StudentTeacher.objects.select_related('user', 'selected_school', 'selected_school__district').order_by('full_name')
        student_data = [['#', 'Full Name', 'Registration No.', 'School', 'District', 'Status', 'Logbook']]
        for i, st in enumerate(students, 1):
            logbook_count = LogbookEntry.objects.filter(student=st).count()
            student_data.append([
                str(i),
                (st.full_name or '-')[:40],
                st.registration_number or '-',
                (st.selected_school.name if st.selected_school else 'Not Selected')[:30],
                (st.selected_school.district.name if st.selected_school and st.selected_school.district else '-')[:20],
                st.approval_status.title(),
                str(logbook_count),
            ])
        st_table = Table(student_data, colWidths=[0.8*cm, 4.5*cm, 3*cm, 3.5*cm, 2.5*cm, 2*cm, 1.5*cm])
        st_cmds = [
            ('BACKGROUND', (0,0), (-1,0), NAVY),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 7),
            ('GRID', (0,0), (-1,-1), 0.3, colors.HexColor('#CBD5E0')),
            ('PADDING', (0,0), (-1,-1), 4),
            ('ALIGN', (0,0), (0,-1), 'CENTER'),
            ('ALIGN', (6,0), (6,-1), 'CENTER'),
        ]
        for row_i in range(1, len(student_data)):
            st_cmds.append(('BACKGROUND', (0,row_i), (-1,row_i), LIGHT if row_i % 2 == 0 else colors.white))
        st_table.setStyle(TableStyle(st_cmds))
        elements.append(st_table)
        elements.append(Spacer(1, 0.5*cm))

        # ── Assessors table ──
        elements.append(Paragraph('ASSESSORS LIST', sub_style))
        assessors = Assessor.objects.all().order_by('full_name')
        assessor_data = [['#', 'Full Name', 'Email', 'Phone', 'Academic Year', 'Schools']]
        for i, a in enumerate(assessors, 1):
            schools_count = SchoolAssessment.objects.filter(assessor=a, academic_year=current_year).count() if current_year else 0
            assessor_data.append([
                str(i),
                (a.full_name or '-')[:35],
                (a.email or '-')[:35],
                a.phone_number or '-',
                a.current_academic_year.year if a.current_academic_year else '-',
                str(schools_count),
            ])
        a_table = Table(assessor_data, colWidths=[0.8*cm, 4*cm, 5*cm, 3*cm, 3*cm, 2*cm])
        a_cmds = [
            ('BACKGROUND', (0,0), (-1,0), GOLD),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 7),
            ('GRID', (0,0), (-1,-1), 0.3, colors.HexColor('#CBD5E0')),
            ('PADDING', (0,0), (-1,-1), 4),
            ('ALIGN', (0,0), (0,-1), 'CENTER'),
            ('ALIGN', (5,0), (5,-1), 'CENTER'),
        ]
        for row_i in range(1, len(assessor_data)):
            a_cmds.append(('BACKGROUND', (0,row_i), (-1,row_i), LIGHT2 if row_i % 2 == 0 else colors.white))
        a_table.setStyle(TableStyle(a_cmds))
        elements.append(a_table)

        doc.build(elements)
        buffer.seek(0)
        filename = f"IMS_Report_{timezone.now().strftime('%Y%m%d')}.pdf"
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        import traceback
        return HttpResponse(
            f"<h3>PDF Generation Error</h3><pre>{traceback.format_exc()}</pre>",
            status=500
        )


# ============================================================
# BODI YA WALIMU — Teacher Board System
# ============================================================

def _get_board_member(request):
    """Return BoardMember for current user or None."""
    try:
        return request.user.board_member
    except Exception:
        return None


def _can_access_region(bm, region):
    """Check if board member can view/access a region."""
    if bm.role in ('chair', 'inspector', 'member'):
        return True
    if bm.role == 'reo':
        return bm.region_id == region.id
    if bm.role == 'deo':
        return bm.district and bm.district.region_id == region.id
    if bm.role == 'head_teacher':
        return bm.district and bm.district.region_id == region.id
    return False


def _can_access_district(bm, district):
    """Check if board member can view/access a district."""
    if bm.role in ('chair', 'inspector', 'member'):
        return True
    if bm.role == 'reo':
        return bm.region_id == district.region_id
    if bm.role == 'deo':
        return bm.district_id == district.id
    if bm.role == 'head_teacher':
        return bm.district_id == district.id
    return False


def _normalize_school_code(raw):
    raw = raw.strip().upper().replace(' ', '').replace('-', '').replace('.', '')
    if re.match(r'^PS\d+$', raw):
        return raw
    return re.sub(r'^([SP])(\d+)$', r'\1.\2', raw)


def board_login(request):
    if request.user.is_authenticated and _get_board_member(request):
        return redirect('board_home')

    # Pata GET params (kutoka /ombi/ form redirect)
    prefill_email = request.GET.get('email', '')
    prefill_school_id = request.GET.get('school_id', '')
    first_time = request.GET.get('first_time', '0') == '1'
    active_tab = request.GET.get('tab', 'staff')  # 'head' au 'staff'

    if request.method == 'POST':
        mode = request.POST.get('mode', 'staff')

        # ── Mkuu wa Shule: Weka Nywila Mara ya Kwanza ──
        if mode == 'head_set_password':
            email = request.POST.get('email', '').strip().lower()
            school_id = request.POST.get('school_id', '').strip()
            password1 = request.POST.get('password1', '')
            password2 = request.POST.get('password2', '')
            school = School.objects.filter(id=school_id).first()

            bm = BoardMember.objects.filter(
                user__email__iexact=email, role='head_teacher', is_active=True
            ).select_related('school', 'user').first()

            ctx = {'mode': 'head_set_password', 'prefill_email': email,
                   'prefill_school_id': school_id, 'school': school, 'active_tab': 'head'}

            if not bm:
                ctx['error'] = f'Email "{email}" haipo kwenye orodha ya wakuu wa shule. Wasiliana na DEO wako.'
                return render(request, 'field_app/board_login.html', ctx)
            if len(password1) < 6:
                ctx['error'] = 'Nywila iwe na herufi 6 au zaidi.'
                return render(request, 'field_app/board_login.html', ctx)
            if password1 != password2:
                ctx['error'] = 'Nywila mbili hazifanani. Jaribu tena.'
                return render(request, 'field_app/board_login.html', ctx)

            bm.user.set_password(password1)
            bm.user.save()
            login(request, bm.user, backend='field_app.backends.EmailBackend')
            messages.success(request, f'Karibu {bm.full_name}! Nywila yako imewekwa.')
            return redirect('board_head_teacher', school_id=bm.school.id)

        # ── Mkuu wa Shule: Ingia kwa Email + Nywila ──
        elif mode == 'head_login':
            email = request.POST.get('email', '').strip().lower()
            password = request.POST.get('password', '')
            school_id = request.POST.get('school_id', '').strip()
            school = School.objects.filter(id=school_id).first() if school_id else None

            bm = BoardMember.objects.filter(
                user__email__iexact=email, role='head_teacher', is_active=True
            ).select_related('school', 'user').first()

            ctx = {'mode': 'head_login', 'prefill_email': email,
                   'prefill_school_id': school_id, 'school': school, 'active_tab': 'head'}

            if not bm:
                ctx['error'] = f'Email "{email}" haipo kwenye orodha ya wakuu wa shule.'
                return render(request, 'field_app/board_login.html', ctx)

            user = authenticate(request, username=email, password=password,
                                backend='field_app.backends.EmailBackend')
            if not user:
                ctx['error'] = 'Nywila si sahihi. Jaribu tena.'
                return render(request, 'field_app/board_login.html', ctx)

            login(request, user, backend='field_app.backends.EmailBackend')
            return redirect('board_head_teacher', school_id=bm.school.id)

        # ── DEO / REO / Chair / Wengine: Email + Nywila (si mkuu wa shule) ──
        else:
            email = request.POST.get('email', '').strip()
            password = request.POST.get('password', '')
            user = authenticate(request, username=email, password=password,
                                backend='field_app.backends.EmailBackend')
            if user:
                try:
                    bm = user.board_member
                    if bm and bm.is_active:
                        if bm.role == 'head_teacher':
                            # Mkuu wa shule lazima atumie tab yake mwenyewe
                            messages.error(request, 'Wewe ni Mkuu wa Shule. Tumia tab ya "MKUU WA SHULE" upande wa kushoto.')
                        else:
                            login(request, user, backend='field_app.backends.EmailBackend')
                            return redirect('board_home')
                    else:
                        messages.error(request, 'Akaunti yako imezimwa. Wasiliana na msimamizi.')
                except Exception:
                    messages.error(request, 'Barua pepe au nywila si sahihi, au huna ruhusa ya Bodi ya Walimu.')
            else:
                messages.error(request, 'Barua pepe au nywila si sahihi.')

    # GET — onyesha form na pre-fill kutoka /ombi/ redirect
    school = School.objects.filter(id=prefill_school_id).first() if prefill_school_id else None
    return render(request, 'field_app/board_login.html', {
        'prefill_email': prefill_email,
        'prefill_school_id': prefill_school_id,
        'first_time': first_time,
        'active_tab': active_tab,
        'school': school,
    })


def board_logout(request):
    logout(request)
    return redirect('board_login')


@login_required
def board_home(request):
    bm = _get_board_member(request)
    if not bm:
        messages.error(request, 'Huna ruhusa ya Bodi ya Walimu.')
        return redirect('board_login')

    # DEO → redirect directly to their district (skip region selection)
    if bm.role == 'deo':
        if bm.district:
            return redirect('board_school_list', district_id=bm.district.id)
        messages.warning(request, 'Akaunti yako ya DEO haina wilaya iliyowekwa. Wasiliana na msimamizi.')

    # REO → redirect directly to their region
    if bm.role == 'reo':
        if bm.region:
            return redirect('board_district_list', region_id=bm.region.id)
        messages.warning(request, 'Akaunti yako ya REO haina mkoa iliyowekwa. Wasiliana na msimamizi.')

    # Head Teacher → redirect to their school dashboard
    if bm.role == 'head_teacher':
        if bm.school:
            return redirect('board_head_teacher', school_id=bm.school.id)
        messages.warning(request, 'Akaunti yako haina shule iliyowekwa. Wasiliana na msimamizi.')

    from datetime import date as _date
    today = _date.today()
    seven_days_ago = today - timedelta(days=7)

    # Chair/Inspector see all regions; others see only accessible ones
    if bm.role in ('chair', 'inspector', 'member'):
        regions_qs = Region.objects.all()
    else:
        regions_qs = Region.objects.none()

    regions = regions_qs.order_by('name').annotate(
        student_count=Count('district__school__studentteacher',
                            filter=Q(district__school__studentteacher__selected_school__isnull=False),
                            distinct=True),
        district_count=Count('district', distinct=True),
    )
    total_students = StudentTeacher.objects.filter(selected_school__isnull=False).count()
    total_schools = School.objects.filter(studentteacher__isnull=False).distinct().count()
    active_this_week = LogbookEntry.objects.filter(
        date__gte=seven_days_ago
    ).values('student').distinct().count()
    inactive_count = StudentTeacher.objects.filter(
        selected_school__isnull=False
    ).exclude(
        logbookentry__date__gte=seven_days_ago
    ).distinct().count()

    return render(request, 'field_app/board_home.html', {
        'bm': bm,
        'regions': regions,
        'total_students': total_students,
        'total_schools': total_schools,
        'active_this_week': active_this_week,
        'inactive_count': inactive_count,
        'total_regions': regions.count(),
    })


@board_login_required
def board_district_list(request, region_id):
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')
    region = get_object_or_404(Region, id=region_id)

    # DEO should not be browsing region pages — send to their district directly
    if bm.role == 'deo' and bm.district:
        return redirect('board_school_list', district_id=bm.district.id)

    # Access check — REO can only view their own region
    if not _can_access_region(bm, region):
        messages.error(request, f'Huna ruhusa ya kuona mkoa wa {region.name}.')
        return redirect('board_home')

    current_year = _cached_active_year()

    # REO sees only their region's districts; others see all districts in region
    districts = District.objects.filter(region=region).order_by('name').annotate(
        school_count=Count('school', distinct=True),
        student_count=Count('school__studentteacher',
                            filter=Q(school__studentteacher__selected_school__isnull=False),
                            distinct=True),
    )

    alloc_qs = DistrictAllocation.objects.filter(district__region=region)
    if current_year:
        alloc_qs = alloc_qs.filter(academic_year=current_year)
    alloc_map = {a.district_id: a for a in alloc_qs}
    for d in districts:
        d.allocation = alloc_map.get(d.id)

    # Maombi ya wakuu wa shule kwa mkoa huu (REO anaona, hawezi kubadilisha)
    region_requests = SchoolHeadRequest.objects.filter(
        district__region=region
    ).select_related('school', 'district', 'reviewed_by').order_by('-submitted_at')
    if current_year:
        region_requests = region_requests.filter(academic_year=current_year)

    return render(request, 'field_app/board_district_list.html', {
        'bm': bm,
        'region': region,
        'districts': districts,
        'region_requests': region_requests,
        'pending_count': region_requests.filter(status='pending').count(),
        'applied_count': region_requests.filter(status='applied').count(),
    })


@board_login_required
def board_school_list(request, district_id):
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')
    district = get_object_or_404(District, id=district_id)

    if not _can_access_district(bm, district):
        messages.error(request, f'Huna ruhusa ya kuona wilaya ya {district.name}.')
        return redirect('board_home')

    schools = School.objects.filter(
        district=district,
        studentteacher__isnull=False
    ).distinct().order_by('level', 'name').annotate(
        student_count=Count('studentteacher', distinct=True),
        logbook_count=Count('logbookentry', distinct=True),
    )
    def _lb_first_lesson(entry):
        if entry and entry.lessons_data:
            return entry.lessons_data[0]
        return {}

    from datetime import date as _date
    today = _date.today()

    students_data = []
    for school in schools:
        sts = StudentTeacher.objects.filter(selected_school=school).select_related('user')
        for st in sts:
            lb_count = LogbookEntry.objects.filter(student=st).count()
            latest_lb = LogbookEntry.objects.filter(student=st).order_by('-date').first()
            latest_lp = LessonPlan.objects.filter(student=st, school=school).first()
            first_lesson = _lb_first_lesson(latest_lb)
            display_subject = (
                first_lesson.get('subject') or
                (latest_lb.subject_taught.name if latest_lb and latest_lb.subject_taught else '') or
                (latest_lp.subject.name if latest_lp else '')
            )
            display_topic = (
                first_lesson.get('main_topic') or
                (latest_lb.topic_taught if latest_lb else '') or
                (latest_lp.topic if latest_lp else '')
            )
            display_class = first_lesson.get('class') or (latest_lb.class_taught if latest_lb else '')
            display_present = first_lesson.get('present') or (latest_lb.num_students_present if latest_lb else '')
            display_activity = first_lesson.get('activity_type') or (latest_lb.activity_type if latest_lb else '')
            last_lb_date = latest_lb.date if latest_lb else None
            days_inactive = (today - last_lb_date).days if last_lb_date else None
            inactive_alert = days_inactive is not None and days_inactive >= 7
            has_scheme = SchemeOfWork.objects.filter(student=st).exists()
            has_lesson_plan = LessonPlan.objects.filter(student=st).exists()
            students_data.append({
                'student': st,
                'school': school,
                'logbook_count': lb_count,
                'latest_logbook': latest_lb,
                'latest_lesson': latest_lp,
                'display_subject': display_subject,
                'display_topic': display_topic,
                'display_class': display_class,
                'display_present': display_present,
                'display_activity': display_activity,
                'days_inactive': days_inactive,
                'inactive_alert': inactive_alert,
                'has_scheme': has_scheme,
                'has_lesson_plan': has_lesson_plan,
                'last_lb_date': last_lb_date,
            })
    inactive_count = sum(1 for item in students_data if item.get('inactive_alert'))

    # Group students by school for template iteration (avoids O(n*m) nested filtering)
    schools_with_students = []
    for school in schools:
        school_students = [item for item in students_data if item['school'].id == school.id]
        schools_with_students.append({'school': school, 'students': school_students})

    current_year = _cached_active_year()
    _da_qs = DistrictAllocation.objects.filter(district=district)
    if current_year:
        _da_qs = _da_qs.filter(academic_year=current_year)
    district_alloc = _da_qs.first()
    school_alloc_map = {}
    if district_alloc:
        school_alloc_map = {
            sa.school_id: sa
            for sa in SchoolAllocation.objects.filter(district_allocation=district_alloc)
        }

    # Maombi ya wakuu wa shule (SchoolHeadRequest) — hizi ndizo zinaundwa na /ombi/ form
    head_req_qs = SchoolHeadRequest.objects.filter(
        district=district
    ).select_related('school').order_by('-submitted_at')
    if current_year:
        head_req_qs = head_req_qs.filter(academic_year=current_year)

    # Kwa kila shule, pata ombi la hivi karibuni
    latest_req_map = {}  # school_id → latest SchoolHeadRequest
    for req in head_req_qs:
        sid = req.school_id if req.school_id else req.school_name_submitted
        if sid not in latest_req_map:
            latest_req_map[sid] = req

    # Tengeneza orodha — shule zenye maombi tu + allocation zao
    schools_alloc_list = []
    seen_schools = set()
    for req in head_req_qs:
        if req.school_id and req.school_id not in seen_schools:
            seen_schools.add(req.school_id)
            sa = school_alloc_map.get(req.school_id)
            schools_alloc_list.append({
                'school': req.school,
                'alloc': sa,
                'requested': req.students_needed,
                'quota': sa.quota if sa else 0,
                'notes': req.notes,
                'status': req.status,
                'submitted_at': req.submitted_at,
                'head_name': req.head_name,
                'has_request': True,
            })
        elif not req.school_id and req.school_name_submitted not in seen_schools:
            # Ombi bila shule iliyolinganishwa
            seen_schools.add(req.school_name_submitted)
            schools_alloc_list.append({
                'school': None,
                'school_name': req.school_name_submitted,
                'alloc': None,
                'requested': req.students_needed,
                'quota': 0,
                'notes': req.notes,
                'status': req.status,
                'submitted_at': req.submitted_at,
                'head_name': req.head_name,
                'has_request': True,
            })

    return render(request, 'field_app/board_school_list.html', {
        'bm': bm,
        'district': district,
        'schools': schools,
        'schools_with_students': schools_with_students,
        'students_data': students_data,
        'today': today,
        'inactive_count': inactive_count,
        'district_alloc': district_alloc,
        'school_alloc_map': school_alloc_map,
        'schools_alloc_list': schools_alloc_list,
    })


@board_login_required
def board_deo_report(request, district_id):
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')
    district = get_object_or_404(District, id=district_id)

    if not _can_access_district(bm, district):
        messages.error(request, f'Huna ruhusa ya kuona wilaya ya {district.name}.')
        return redirect('board_home')

    HOLIDAY_MONTHS = {5, 8, 12}  # Mei, Agosti, Desemba
    now = timezone.now()
    selected_month = int(request.GET.get('month', now.month))
    selected_year = int(request.GET.get('year', now.year))

    existing = MonthlyReport.objects.filter(
        district=district, month=selected_month, year=selected_year
    ).first()

    if request.method == 'POST' and 'generate' in request.POST:
        if selected_month in HOLIDAY_MONTHS:
            messages.error(request, 'Ripoti haizalishwi kwa miezi ya likizo (Mei, Agosti, Desemba).')
            return redirect(request.path + f'?month={selected_month}&year={selected_year}')

        entries = LogbookEntry.objects.filter(
            date__month=selected_month,
            date__year=selected_year,
            school__district=district,
        ).select_related('student', 'subject_taught', 'school')

        if not entries.exists():
            messages.warning(request, 'Hakuna rekodi za logbook kwa mwezi huu katika wilaya hii.')
            return redirect(request.path + f'?month={selected_month}&year={selected_year}')

        # Build summary data per student
        from collections import defaultdict
        student_summary = defaultdict(lambda: {
            'name': '', 'school': '', 'level': '', 'subject': '',
            'topics': [], 'days_recorded': 0,
        })
        for entry in entries:
            sid = entry.student_id
            st = entry.student
            student_summary[sid]['name'] = st.full_name
            student_summary[sid]['school'] = entry.school.name if entry.school else ''
            student_summary[sid]['level'] = entry.school.level if entry.school else ''
            student_summary[sid]['days_recorded'] += 1
            if entry.lessons_data:
                for lesson in entry.lessons_data:
                    if lesson.get('main_topic'):
                        student_summary[sid]['topics'].append(lesson['main_topic'])
                    if not student_summary[sid]['subject'] and lesson.get('subject'):
                        student_summary[sid]['subject'] = lesson['subject']
            elif entry.topic_taught:
                student_summary[sid]['topics'].append(entry.topic_taught)
                if not student_summary[sid]['subject'] and entry.subject_taught:
                    student_summary[sid]['subject'] = entry.subject_taught.name

        # Months remaining in year
        months_remaining = 12 - selected_month

        summary_text = ""
        for sid, data in student_summary.items():
            topics_count = len(data['topics'])
            unique_topics = list(dict.fromkeys(data['topics']))[:10]
            summary_text += (
                f"- {data['name']} | Shule: {data['school']} ({data['level']}) | "
                f"Somo: {data['subject'] or 'Haijabainishwa'} | "
                f"Mada {topics_count} | Siku {data['days_recorded']} | "
                f"Mada za mwisho: {', '.join(unique_topics[:5]) or 'Hakuna'}\n"
            )

        month_name_map = dict(MonthlyReport.MONTH_CHOICES)
        prompt = f"""Wewe ni mshauri wa elimu kwa Tanzania. Toa ripoti ya kina kwa DEO (Afisa Elimu wa Wilaya) ya wilaya ya {district.name} kwa mwezi wa {month_name_map[selected_month]} {selected_year}.

Data ya walimu wanafunzi (student teachers) waliofanya kazi mwezi huu:
{summary_text}

Miezi iliyobaki mwaka huu: {months_remaining}

Toa ripoti ifuatayo kwa muundo wa JSON:
{{
  "muhtasari": "Muhtasari mfupi wa mwezi (2-3 sentensi)",
  "secondary": [
    {{
      "nafasi": 1,
      "jina": "Jina la mwalimu mwanafunzi",
      "shule": "Jina la shule",
      "somo": "Somo analofundisha",
      "mada_zote": 10,
      "siku": 20,
      "hali": "vizuri|wastani|nyuma",
      "mapendekezo": "Pendekezo fupi la kibinafsi kulingana na miezi iliyobaki"
    }}
  ],
  "primary": [
    {{
      "nafasi": 1,
      "jina": "Jina la mwalimu mwanafunzi",
      "shule": "Jina la shule",
      "somo": "Somo analofundisha",
      "mada_zote": 8,
      "siku": 18,
      "hali": "vizuri|wastani|nyuma",
      "mapendekezo": "Pendekezo fupi la kibinafsi"
    }}
  ],
  "hitimisho": "Hitimisho na mapendekezo ya jumla kwa DEO (3-4 sentensi)"
}}

Panga walimu kwa secondary na primary tofauti. Mpanga kutoka mwenye mada nyingi zaidi hadi chache. Mwalimu mwenye mada nyingi zaidi 'vizuri', wastani 'wastani', chini ya wastani 'nyuma'. Jibu kwa JSON tu bila maelezo mengine."""

        try:
            from .ai_utils import client, model_name as ai_model
            response = client.models.generate_content(model=ai_model, contents=prompt)
            raw = response.text.strip()
            raw = re.sub(r'^```(?:json)?\s*', '', raw)
            raw = re.sub(r'\s*```$', '', raw)
            ai_data = json.loads(raw)
        except Exception as e:
            messages.error(request, f'Hitilafu ya AI: {e}')
            return redirect(request.path + f'?month={selected_month}&year={selected_year}')

        report, _ = MonthlyReport.objects.update_or_create(
            district=district, month=selected_month, year=selected_year,
            defaults={'ai_content': ai_data, 'generated_by': bm},
        )
        existing = report
        messages.success(request, 'Ripoti imezalishwa kwa mafanikio.')

    available_months = [
        (m, n) for m, n in MonthlyReport.MONTH_CHOICES if m not in HOLIDAY_MONTHS
    ]
    past_reports = MonthlyReport.objects.filter(district=district).order_by('-year', '-month')

    available_years = list(range(2024, 2046))

    return render(request, 'field_app/board_deo_report.html', {
        'bm': bm,
        'district': district,
        'selected_month': selected_month,
        'selected_year': selected_year,
        'existing': existing,
        'available_months': available_months,
        'available_years': available_years,
        'past_reports': past_reports,
        'holiday_months': HOLIDAY_MONTHS,
        'is_holiday': selected_month in HOLIDAY_MONTHS,
        'month_name': dict(MonthlyReport.MONTH_CHOICES).get(selected_month, ''),
    })


@board_login_required
def board_student_progress(request, student_id):
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')
    student = get_object_or_404(StudentTeacher, id=student_id)
    school = student.selected_school

    # Access check — must be in accessible district
    if school and not _can_access_district(bm, school.district):
        messages.error(request, 'Huna ruhusa ya kuona mwanafunzi huyu.')
        return redirect('board_home')

    logbook_entries = LogbookEntry.objects.filter(student=student).select_related('subject_taught').order_by('-date')[:30]
    lesson_plans = LessonPlan.objects.filter(student=student).order_by('-date')[:10]
    schemes = SchemeOfWork.objects.filter(student=student).order_by('-created_at')[:5]
    applications = StudentApplication.objects.filter(student=student, status='approved').select_related('subject', 'school')
    board_comments = BoardComment.objects.filter(student=student).select_related('board_member').order_by('-created_at')

    from datetime import date as _date
    today = _date.today()
    month_entries = LogbookEntry.objects.filter(
        student=student, date__year=today.year, date__month=today.month,
    )
    total_entries = LogbookEntry.objects.filter(student=student).count()

    # Build enriched logbook list — extract lessons_data for display
    logbook_display = []
    for lb in logbook_entries:
        lessons = lb.lessons_data or []
        if lessons:
            first = lessons[0]
            disp_subject = first.get('subject', '') or (lb.subject_taught.name if lb.subject_taught else '—')
            disp_topic   = first.get('main_topic', '') or lb.topic_taught or '—'
            disp_class   = first.get('class', '') or lb.class_taught or '—'
            disp_present = first.get('present', '') or lb.num_students_present or '—'
            disp_activity = first.get('activity_type', '') or lb.activity_type or '—'
            num_periods  = len(lessons)
        else:
            disp_subject  = lb.subject_taught.name if lb.subject_taught else '—'
            disp_topic    = lb.topic_taught or '—'
            disp_class    = lb.class_taught or '—'
            disp_present  = lb.num_students_present or '—'
            disp_activity = lb.activity_type or '—'
            num_periods   = 0
        logbook_display.append({
            'entry': lb,
            'subject': disp_subject,
            'topic': disp_topic,
            'class': disp_class,
            'present': disp_present,
            'activity': disp_activity,
            'num_periods': num_periods,
            'lessons': lessons,
        })

    # Topics covered — from lessons_data first, fallback to old fields
    topics_covered = []
    seen = set()
    for lb in LogbookEntry.objects.filter(student=student).order_by('-date')[:50]:
        if lb.lessons_data:
            for lesson in lb.lessons_data:
                key = (lesson.get('subject',''), lesson.get('main_topic',''))
                if key not in seen and (key[0] or key[1]):
                    seen.add(key)
                    topics_covered.append({'subject': key[0], 'topic': key[1], 'date': lb.date})
        elif lb.topic_taught:
            key = (lb.subject_taught.name if lb.subject_taught else '', lb.topic_taught)
            if key not in seen:
                seen.add(key)
                topics_covered.append({'subject': key[0], 'topic': key[1], 'date': lb.date})
        if len(topics_covered) >= 20:
            break

    last_entry_for_inactive = LogbookEntry.objects.filter(student=student).order_by('-date').first()
    days_inactive = (today - last_entry_for_inactive.date).days if last_entry_for_inactive else None
    inactive_alert = days_inactive is not None and days_inactive >= 7

    lesson_plans_all = LessonPlan.objects.filter(student=student).order_by('-date')
    raw_schemes = SchemeOfWork.objects.filter(student=student).order_by('-year', '-term')

    # Normalize scheme_data rows — KitabuSmart uses spaced keys ('Main Competence',
    # 'Week', 'Periods' …). The template expects snake_case keys. Build a mapping
    # so any saved format is displayed correctly without raising VariableDoesNotExist.
    def _normalize_row(row):
        if not isinstance(row, dict):
            return {}
        get = lambda *keys: next((str(row[k]) for k in keys if k in row and row[k] not in (None, '')), '')
        return {
            'week':       get('week', 'Week'),
            'month':      get('month', 'Month'),
            'topic':      get('topic', 'main_topic', 'Main Competence', 'Learning Activities'),
            'subtopic':   get('subtopic', 'sub_topic', 'Specific Competence', 'Specific Learning Activities'),
            'activities': get('activities', 'learning_activities', 'Learning Activities', 'Specific Learning Activities'),
            'resources':  get('resources', 'teaching_resources', 'Teaching & Learning Resources'),
            'assessment': get('assessment', 'assessment_criteria', 'Assessment Tools'),
            'periods':    get('periods', 'Periods'),
            'remarks':    get('remarks', 'Remarks'),
            'reference':  get('reference', 'Reference'),
            'methods':    get('methods', 'Teaching & Learning Methods'),
        }

    schemes_all = []
    for scheme in raw_schemes:
        normalized_data = [_normalize_row(r) for r in (scheme.scheme_data or [])]
        schemes_all.append({
            'obj': scheme,
            'data': normalized_data,
        })

    return render(request, 'field_app/board_student_progress.html', {
        'bm': bm,
        'student': student,
        'school': school,
        'logbook_entries': logbook_entries,
        'logbook_display': logbook_display,
        'lesson_plans': lesson_plans,
        'lesson_plans_all': lesson_plans_all,
        'schemes': schemes,
        'schemes_all': schemes_all,
        'applications': applications,
        'board_comments': board_comments,
        'month_entry_count': month_entries.count(),
        'total_entries': total_entries,
        'topics_covered': topics_covered,
        'today': today,
        'days_inactive': days_inactive,
        'inactive_alert': inactive_alert,
    })


@board_login_required
def board_add_comment(request):
    if request.method != 'POST':
        return redirect('board_home')
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')

    student_id = request.POST.get('student_id', '').strip()
    comment_text = request.POST.get('comment', '').strip()
    status = request.POST.get('status', 'good')

    if not student_id or not comment_text:
        messages.error(request, 'Jaza maoni yote yanayohitajika.')
        return redirect(request.META.get('HTTP_REFERER', 'board_home'))

    student = get_object_or_404(StudentTeacher, id=student_id)
    current_year = get_current_academic_year()

    BoardComment.objects.create(
        board_member=bm,
        student=student,
        school=student.selected_school,
        comment=comment_text,
        status=status,
        academic_year=current_year,
    )
    messages.success(request, f'Maoni yametumwa kwa {student.full_name}.')
    return redirect('board_student_progress', student_id=student.id)


@board_login_required
def board_head_teacher(request, school_id):
    """Dashboard ya Mkuu wa Shule - management kamili ya walimu wanafunzi."""
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')

    school = get_object_or_404(School, id=school_id)

    if bm.role == 'head_teacher' and (not bm.school or bm.school_id != school.id):
        messages.error(request, 'Huna ruhusa ya kuona shule hii.')
        return redirect('board_home')

    current_year = _cached_active_year()
    today = timezone.now().date()

    # POST: ongeza maoni ya mwanafunzi
    if request.method == 'POST' and request.POST.get('action') == 'add_comment':
        student_id = request.POST.get('student_id')
        comment_text = request.POST.get('comment', '').strip()
        status_val = request.POST.get('status', 'good')
        if student_id and comment_text:
            student_obj = StudentTeacher.objects.filter(id=student_id, selected_school=school).first()
            if student_obj:
                BoardComment.objects.create(
                    board_member=bm, student=student_obj, school=school,
                    comment=comment_text, status=status_val, academic_year=current_year
                )
                messages.success(request, f'Maoni kwa {student_obj.full_name} yamehifadhiwa.')
        return redirect('board_head_teacher', school_id=school.id)

    # Wanafunzi wote wa shule hii
    students = StudentTeacher.objects.filter(
        selected_school=school
    ).select_related('user').order_by('full_name')

    total = students.count()
    approved_count = students.filter(approval_status='approved').count()
    pending_count  = students.filter(approval_status='pending').count()

    # Waliofika leo (GPS check-in)
    today_checkins = set(
        LogbookEntry.objects.filter(school=school, date=today, morning_check_in__isnull=False)
        .values_list('student_id', flat=True)
    )
    present_today = len(today_checkins)

    # Logbook compliance ya wiki hii
    week_ago = today - timedelta(days=7)
    recent_entries = LogbookEntry.objects.filter(
        school=school, date__gte=week_ago
    ).values('student_id', 'date').distinct().count()
    working_days = sum(1 for i in range(7) if (today - timedelta(days=i)).weekday() < 5)
    expected = max(1, total * working_days)
    compliance_pct = min(100, int(recent_entries / expected * 100))

    # Data ya kina kwa kila mwanafunzi
    latest_logs = {
        e['student_id']: e['date']
        for e in LogbookEntry.objects.filter(school=school)
        .values('student_id').annotate(date=models.Max('date')).values('student_id', 'date')
    }
    logbook_counts = {
        e['student_id']: e['cnt']
        for e in LogbookEntry.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    scheme_counts = {
        e['student_id']: e['cnt']
        for e in SchemeOfWork.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    lp_counts = {
        e['student_id']: e['cnt']
        for e in LessonPlan.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    last_comments = {}
    if BoardComment.objects.filter(school=school).exists():
        last_comments = {
            c.student_id: c
            for c in BoardComment.objects.filter(school=school)
            .select_related('board_member').order_by('student_id', '-created_at')
            .distinct('student_id')
        }
    # Approved applications per student
    approved_apps = {}
    for app in StudentApplication.objects.filter(school=school, status='approved').select_related('subject'):
        approved_apps.setdefault(app.student_id, []).append(app.subject.name)

    # Internship length: start of current academic year to today
    year_start = current_year.start_date if current_year and hasattr(current_year, 'start_date') else None
    total_working_days = 0
    if year_start:
        d = year_start
        while d <= today:
            if d.weekday() < 5:
                total_working_days += 1
            d += timedelta(days=1)
    total_working_days = max(1, total_working_days)

    students_data = []
    for st in students:
        last_log  = latest_logs.get(st.id)
        days_since = (today - last_log).days if last_log else None
        lb_count   = logbook_counts.get(st.id, 0)
        pct = min(100, int(lb_count / total_working_days * 100)) if total_working_days > 0 else 0
        students_data.append({
            'obj': st,
            'present_today': st.id in today_checkins,
            'last_log': last_log,
            'days_since': days_since,
            'log_status': 'ok' if days_since is not None and days_since <= 1
                         else ('warn' if days_since is not None and days_since <= 3 else 'late'),
            'last_comment': last_comments.get(st.id),
            'logbook_count': lb_count,
            'scheme_count': scheme_counts.get(st.id, 0),
            'lp_count': lp_counts.get(st.id, 0),
            'progress_pct': pct,
            'subjects': approved_apps.get(st.id, []),
        })

    # DEO Allocation
    alloc = None
    if current_year:
        try:
            dist_alloc = DistrictAllocation.objects.filter(
                district=school.district, academic_year=current_year
            ).first()
            if dist_alloc:
                alloc = SchoolAllocation.objects.filter(
                    district_allocation=dist_alloc, school=school
                ).first()
        except Exception:
            pass

    head_requests = SchoolHeadRequest.objects.filter(school=school).order_by('-submitted_at')[:5]
    current_month = today.strftime('%B %Y')

    return render(request, 'field_app/board_head_teacher.html', {
        'bm': bm,
        'school': school,
        'students_data': students_data,
        'total': total,
        'approved_count': approved_count,
        'pending_count': pending_count,
        'present_today': present_today,
        'compliance_pct': compliance_pct,
        'alloc': alloc,
        'head_requests': head_requests,
        'today': today,
        'current_month': current_month,
        'can_edit': bm.role in ('head_teacher', 'deo', 'chair'),
    })


def head_teacher_monthly_report(request, school_id):
    """Ripoti ya mwezi — mkuu wa shule anaweza kuona na kuchapisha."""
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')

    school = get_object_or_404(School, id=school_id)
    if bm.role == 'head_teacher' and (not bm.school or bm.school_id != school.id):
        messages.error(request, 'Huna ruhusa ya kuona shule hii.')
        return redirect('board_home')

    today = timezone.now().date()
    current_year = _cached_active_year()
    month_name = request.GET.get('month', today.strftime('%B %Y'))

    students = StudentTeacher.objects.filter(
        selected_school=school, approval_status='approved'
    ).select_related('user').order_by('full_name')

    year_start = current_year.start_date if current_year and hasattr(current_year, 'start_date') else None
    total_working_days = 0
    if year_start:
        d = year_start
        while d <= today:
            if d.weekday() < 5:
                total_working_days += 1
            d += timedelta(days=1)
    total_working_days = max(1, total_working_days)

    approved_apps = {}
    for app in StudentApplication.objects.filter(school=school, status='approved').select_related('subject'):
        approved_apps.setdefault(app.student_id, []).append(app.subject.name)

    logbook_counts = {
        e['student_id']: e['cnt']
        for e in LogbookEntry.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    scheme_counts = {
        e['student_id']: e['cnt']
        for e in SchemeOfWork.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    lp_counts = {
        e['student_id']: e['cnt']
        for e in LessonPlan.objects.filter(school=school)
        .values('student_id').annotate(cnt=models.Count('id'))
    }
    last_comments = {}
    if BoardComment.objects.filter(school=school).exists():
        last_comments = {
            c.student_id: c
            for c in BoardComment.objects.filter(school=school)
            .select_related('board_member').order_by('student_id', '-created_at')
            .distinct('student_id')
        }

    report_data = []
    for st in students:
        lb = logbook_counts.get(st.id, 0)
        pct = min(100, int(lb / total_working_days * 100))
        comment = last_comments.get(st.id)
        report_data.append({
            'name': st.full_name,
            'email': st.user.email if st.user else '—',
            'subjects': ', '.join(approved_apps.get(st.id, [])) or '—',
            'logbook_days': lb,
            'scheme_count': scheme_counts.get(st.id, 0),
            'lp_count': lp_counts.get(st.id, 0),
            'progress_pct': pct,
            'status': comment.get_status_display() if comment else '—',
            'comment': comment.comment if comment else '',
        })

    return render(request, 'field_app/head_teacher_monthly_report.html', {
        'bm': bm,
        'school': school,
        'report_data': report_data,
        'month_name': month_name,
        'today': today,
        'total_students': len(report_data),
        'avg_pct': int(sum(r['progress_pct'] for r in report_data) / max(1, len(report_data))),
    })


@staff_member_required
@login_required
def import_head_teachers(request):
    """Admin: import wakuu wa shule kwa CSV - school_identifier,email,jina
    school_identifier inaweza kuwa: school_code (S.0306) AU jina kamili la shule.
    """
    if not request.user.is_staff:
        return redirect('dashboard')

    User = get_user_model()
    results = []

    if request.method == 'POST':
        import csv, io
        mode = request.POST.get('mode', 'csv_text')

        if mode == 'csv_file' and request.FILES.get('csv_file'):
            f = request.FILES['csv_file']
            content = f.read().decode('utf-8', errors='ignore')
        else:
            content = request.POST.get('csv_text', '')

        reader = csv.reader(io.StringIO(content.strip()))
        for i, row in enumerate(reader, 1):
            if not row or row[0].strip().startswith('#'):
                continue
            parts = [p.strip() for p in row]
            if len(parts) < 2:
                results.append({'row': i, 'status': 'error', 'msg': f'Muundo mbaya: {",".join(parts)}'})
                continue

            school_identifier = parts[0]
            email = parts[1].lower()
            full_name = parts[2] if len(parts) > 2 else email.split('@')[0].replace('.', ' ').title()

            # Jaribu kupata shule kwa njia 4:
            # 1) school_code (S.0306, S0306)
            raw = school_identifier.upper().replace(' ', '').replace('-', '').replace('.', '')
            code = raw if re.match(r'^PS\d+$', raw) else re.sub(r'^([SP])(\d+)$', r'\1.\2', raw)
            school = School.objects.filter(school_code__iexact=code).first()
            # 2) jina kamili kama ilivyo
            if not school:
                school = School.objects.filter(name__iexact=school_identifier).first()
            # 3) jina bila suffix ya level (Secondary/Primary/Sekondari/Msingi)
            if not school:
                clean = re.sub(r'\s+(secondary|primary|sekondari|msingi|shule ya sekondari|shule ya msingi)$',
                               '', school_identifier, flags=re.IGNORECASE).strip()
                if clean != school_identifier:
                    school = School.objects.filter(name__iexact=clean).first()
            # 4) icontains — angalau sehemu ya jina
            if not school:
                school = School.objects.filter(name__icontains=school_identifier.split()[0]).first() \
                    if school_identifier.strip() else None
            if not school:
                results.append({'row': i, 'status': 'error', 'msg': f'Shule haikupatikana: "{school_identifier}"'})
                continue

            # Angalia kama tayari ipo
            if BoardMember.objects.filter(user__email__iexact=email, school=school, role='head_teacher').exists():
                results.append({'row': i, 'status': 'skip', 'msg': f'{school.name} — {email} tayari ipo'})
                continue

            # Unda user na BoardMember kwa bulk-efficient way
            user, created = User.objects.get_or_create(email=email)
            if created:
                user.set_unusable_password()
                user.save(update_fields=['password'])

            BoardMember.objects.create(
                user=user, full_name=full_name, role='head_teacher',
                school=school, district=school.district,
                region=school.district.region if school.district else None,
                is_active=True,
            )
            results.append({'row': i, 'status': 'ok', 'msg': f'{school.name} — {full_name} ({email})'})

    return render(request, 'field_app/import_head_teachers.html', {'results': results})


def create_board_member(request):
    """Admin: create board member - moja au bulk kwa wakuu wa shule."""
    User = get_user_model()
    regions = Region.objects.all().order_by('name')
    districts = District.objects.select_related('region').order_by('name')
    schools = School.objects.select_related('district__region').order_by('name')

    if request.method == 'POST':
        mode = request.POST.get('mode', 'single')

        if mode == 'bulk':
            # Bulk create: kila mstari = school_code,email,jina
            lines = request.POST.get('bulk_data', '').strip().splitlines()
            created, skipped = 0, []
            for line in lines:
                parts = [p.strip() for p in line.split(',')]
                if len(parts) < 2:
                    skipped.append(f"Mstari mbaya: {line}")
                    continue
                school_code_raw = parts[0]
                email = parts[1].lower()
                full_name = parts[2] if len(parts) > 2 else email.split('@')[0]

                raw = school_code_raw.upper().replace(' ', '').replace('-', '').replace('.', '')
                code = raw if re.match(r'^PS\d+$', raw) else re.sub(r'^([SP])(\d+)$', r'\1.\2', raw)
                school = School.objects.filter(school_code__iexact=code).first()
                if not school:
                    skipped.append(f"Shule haikupatikana: {school_code_raw}")
                    continue
                if User.objects.filter(email=email).exists():
                    skipped.append(f"Email tayari ipo: {email}")
                    continue

                user = User.objects.create_user(email=email, password=None)
                user.set_unusable_password()
                user.save()
                BoardMember.objects.create(
                    user=user, full_name=full_name, role='head_teacher',
                    school=school, district=school.district,
                    region=school.district.region if school.district else None,
                )
                created += 1

            msg = f'Akaunti {created} zimeundwa.'
            if skipped:
                msg += f' Zilizoshindwa ({len(skipped)}): ' + '; '.join(skipped[:5])
            messages.success(request, msg)
            return redirect('admin_dashboard')

        else:
            # Single create
            full_name = request.POST.get('full_name', '').strip()
            email = request.POST.get('email', '').strip().lower()
            phone = request.POST.get('phone_number', '').strip()
            role = request.POST.get('role', 'member')
            password = request.POST.get('password', '').strip()
            region_id = request.POST.get('region') or None
            district_id = request.POST.get('district') or None
            school_id = request.POST.get('school') or None

            if not full_name or not email:
                messages.error(request, 'Jaza jina na barua pepe.')
            elif not password or len(password) < 6:
                messages.error(request, 'Weka nywila ya angalau herufi 6.')
            elif User.objects.filter(email=email).exists():
                messages.error(request, f"Email '{email}' tayari ipo.")
            else:
                user = User.objects.create_user(email=email, password=password)
                region = Region.objects.filter(id=region_id).first() if region_id else None
                district = District.objects.filter(id=district_id).first() if district_id else None
                school = School.objects.filter(id=school_id).first() if school_id else None
                bm_new = BoardMember.objects.create(
                    user=user, full_name=full_name, phone_number=phone,
                    role=role, region=region, district=district, school=school,
                )
                # Onyesha credentials kwa admin azitume
                return render(request, 'field_app/create_board_member.html', {
                    'regions': regions, 'districts': districts, 'schools': schools,
                    'role_choices': BoardMember.ROLE_CHOICES,
                    'created_credentials': {
                        'full_name': full_name,
                        'email': email,
                        'password': password,
                        'role': bm_new.get_role_display(),
                        'region': region.name if region else '—',
                        'district': district.name if district else '—',
                    },
                })

    return render(request, 'field_app/create_board_member.html', {
        'regions': regions,
        'districts': districts,
        'schools': schools,
        'role_choices': BoardMember.ROLE_CHOICES,
    })


# =============================================================
# DEO ALLOCATION VIEWS
# =============================================================

@board_login_required
def deo_allocation(request):
    """Redirect DEO to their district's allocation page."""
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')
    if bm.district:
        return redirect('deo_district_allocation', district_id=bm.district.id)
    messages.error(request, 'Akaunti yako haina wilaya. Wasiliana na msimamizi.')
    return redirect('board_home')


@board_login_required
def deo_district_allocation(request, district_id):
    """DEO/Chair anaweka idadi ya walimu wanafunzi kwa wilaya maalumu."""
    bm = _get_board_member(request)
    if not bm:
        messages.error(request, 'Huna ruhusa ya Bodi ya Walimu.')
        return redirect('board_login')

    district = get_object_or_404(District, id=district_id)

    # Chair can edit any district; DEO can edit only their own district; others view-only
    can_edit = (
        bm.role == 'chair' or
        (bm.role == 'deo' and bm.district_id == district.id)
    )

    current_year = _cached_active_year()

    allocation, _ = DistrictAllocation.objects.get_or_create(
        district=district,
        academic_year=current_year,
        defaults={'uploaded_by': bm}
    )

    schools = School.objects.filter(district=district).order_by('level', 'name')
    school_alloc_map = {
        sa.school_id: sa
        for sa in SchoolAllocation.objects.filter(district_allocation=allocation).select_related('school')
    }

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'set_totals':
            allocation.primary_needed = int(request.POST.get('primary_needed', 0) or 0)
            allocation.secondary_needed = int(request.POST.get('secondary_needed', 0) or 0)
            allocation.notes = request.POST.get('notes', '').strip()
            if 'document' in request.FILES:
                allocation.document = request.FILES['document']
            allocation.uploaded_by = bm
            allocation.save()
            messages.success(request, 'Mahitaji ya wilaya yamehifadhiwa.')
            return redirect('deo_district_allocation', district_id=district.id)

        elif action == 'set_school_quotas':
            for school in schools:
                key = f'quota_{school.id}'
                quota_val = int(request.POST.get(key, 0) or 0)
                SchoolAllocation.objects.update_or_create(
                    district_allocation=allocation,
                    school=school,
                    defaults={'quota': quota_val}
                )
                # Sasisha School.capacity ili iwe sawa na quota
                if quota_val > 0:
                    School.objects.filter(id=school.id).update(capacity=quota_val)
            messages.success(request, 'Mgawanyo wa shule zote umehifadhiwa na capacity zimesasishwa.')
            return redirect('deo_district_allocation', district_id=district.id)

        elif action == 'ai_parse':
            if not allocation.document:
                messages.error(request, 'Pakia hati kwanza kabla ya kutumia AI.')
                return redirect('deo_district_allocation', district_id=district.id)
            result = _ai_parse_allocation_document(allocation)
            if result.get('success'):
                allocation.primary_needed = result.get('primary_needed', allocation.primary_needed)
                allocation.secondary_needed = result.get('secondary_needed', allocation.secondary_needed)
                allocation.ai_parsed = True
                allocation.save()
                for school_data in result.get('schools', []):
                    school_name = school_data.get('name', '').strip().lower()
                    quota = int(school_data.get('quota', 0) or 0)
                    matched = next(
                        (s for s in schools if school_name in s.name.lower() or s.name.lower() in school_name),
                        None
                    )
                    if matched and quota > 0:
                        SchoolAllocation.objects.update_or_create(
                            district_allocation=allocation,
                            school=matched,
                            defaults={'quota': quota}
                        )
                messages.success(
                    request,
                    f'AI imechambua hati. Msingi: {result.get("primary_needed", 0)}, '
                    f'Sekondari: {result.get("secondary_needed", 0)}'
                )
            else:
                messages.warning(
                    request,
                    f'AI haikuweza kuchambua hati: {result.get("error", "")}. Weka idadi wewe mwenyewe.'
                )
            return redirect('deo_district_allocation', district_id=district.id)

    # Rebuild after possible POST changes
    school_alloc_map = {
        sa.school_id: sa
        for sa in SchoolAllocation.objects.filter(district_allocation=allocation).select_related('school')
    }

    schools_with_alloc = []
    for school in schools:
        sa = school_alloc_map.get(school.id)
        filled = sa.filled if sa else StudentApplication.objects.filter(
            school=school, status='approved'
        ).values('student').distinct().count()
        quota = sa.quota if sa else 0
        schools_with_alloc.append({
            'school': school,
            'quota': quota,
            'filled': filled,
            'remaining': max(0, quota - filled),
            'pct': min(100, round(filled / quota * 100)) if quota > 0 else 0,
            'ht_requested': sa.head_teacher_requested if sa else 0,
            'ht_notes': sa.head_teacher_notes if sa else '',
        })

    return render(request, 'field_app/deo_allocation.html', {
        'allocation': allocation,
        'district': district,
        'schools_with_alloc': schools_with_alloc,
        'current_year': current_year,
        'can_edit': can_edit,
        'bm': bm,
        'primary_schools': [s for s in schools_with_alloc if s['school'].level == 'Primary'],
        'secondary_schools': [s for s in schools_with_alloc if s['school'].level == 'Secondary'],
    })


def _ai_parse_allocation_document(allocation):
    """Tumia AI (Gemini) kuchambua hati ya mahitaji ya walimu wanafunzi."""
    import os
    try:
        doc_path = allocation.document.path
        text = ''

        if doc_path.lower().endswith('.pdf'):
            import pdfplumber
            with pdfplumber.open(doc_path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or '') + '\n'
        elif doc_path.lower().endswith(('.docx', '.doc')):
            from docx import Document as DocxDocument
            doc = DocxDocument(doc_path)
            text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            return {'success': False, 'error': 'Aina ya faili haijulikani. Tumia PDF au Word.'}

        if not text.strip():
            return {'success': False, 'error': 'Hati haina maandishi yanayoweza kusomwa.'}

        from ai_utils import client, model_name
        prompt = f"""
Soma hati hii ya mahitaji ya walimu wanafunzi na utoe majibu kwa JSON tu.

HATI:
{text[:4000]}

Toa JSON katika muundo huu (numbers tu, bila maelezo):
{{
  "primary_needed": <jumla ya walimu wanafunzi wa shule za msingi>,
  "secondary_needed": <jumla ya walimu wanafunzi wa shule za sekondari>,
  "schools": [
    {{"name": "<jina la shule>", "level": "Primary|Secondary", "quota": <idadi>}},
    ...
  ]
}}

Kama hauwezi kupata idadi, weka 0. Jibu kwa JSON tu, bila maelezo mengine.
"""
        response = client.models.generate_content(
            model=model_name,
            contents=prompt,
        )
        raw = response.text.strip()
        # Extract JSON from response
        import re, json as _json
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            data = _json.loads(match.group())
            data['success'] = True
            return data
        return {'success': False, 'error': 'AI haikutoa JSON sahihi.'}

    except Exception as e:
        return {'success': False, 'error': str(e)}


# =============================================================
# SCHOOL HEAD REQUESTS
# =============================================================

def public_school_head_form(request):
    """Universal public form — school head verifies via school code then submits."""
    current_year = _cached_active_year()

    # Stage 2: full form submission
    if request.method == 'POST' and request.POST.get('stage') == 'submit':
        school_id = request.POST.get('school_id', '').strip()
        school = School.objects.filter(id=school_id).first()
        if not school:
            return render(request, 'field_app/public_school_head_form.html', {
                'error': 'Shule haikutambuliwa. Jaribu tena.',
                'current_year': current_year,
            })

        head_name   = request.POST.get('head_name', '').strip()
        head_email  = request.POST.get('head_email', '').strip().lower()
        head_phone  = request.POST.get('head_phone', '').strip()
        level       = school.level
        notes       = request.POST.get('notes', '').strip()
        try:
            students_needed = int(request.POST.get('students_needed', 0) or 0)
        except ValueError:
            students_needed = 0

        if not head_name or not head_email or students_needed < 1:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'form', 'school': school, 'current_year': current_year,
                'error': 'Tafadhali jaza sehemu zote zinazohitajika (jina, email, idadi).',
            })

        req = SchoolHeadRequest.objects.create(
            district=school.district,
            academic_year=current_year,
            school_name_submitted=school.name,
            school=school,
            head_name=head_name,
            head_phone=head_phone,
            submitter_email=head_email,
            level=level,
            students_needed=students_needed,
            notes=notes,
        )
        request.session['ombi_request_id'] = req.id
        request.session['ombi_school_id'] = school.id

        User = get_user_model()
        # Unda au pata BoardMember kwa mkuu huyu
        bm = BoardMember.objects.filter(school=school, role='head_teacher', is_active=True).first()
        needs_password = False

        if not bm:
            # Angalia kama email ipo kwenye pre-registered list ya shule hii
            pre_registered = BoardMember.objects.filter(
                user__email__iexact=head_email,
                school=school,
                role='head_teacher',
            ).first()
            if not pre_registered:
                return render(request, 'field_app/public_school_head_form.html', {
                    'stage': 'form', 'school': school, 'current_year': current_year,
                    'error': f'Email "{head_email}" haipo kwenye orodha ya wakuu wa shule ya {school.name}. Wasiliana na DEO wa {school.district.name} ili aongeze email yako.',
                })
            bm = pre_registered
            needs_password = not bm.user.has_usable_password()
        else:
            needs_password = not bm.user.has_usable_password()

        # Peleka moja kwa moja kwenye board login page na email tayari imejazwa
        from urllib.parse import urlencode
        params = {'email': head_email, 'school_id': school.id, 'tab': 'head'}
        if needs_password:
            params['first_time'] = '1'
        return redirect(reverse('board_login') + '?' + urlencode(params))

    # Stage: mkuu anaweka nywila yake mwenyewe (mara ya kwanza)
    if request.method == 'POST' and request.POST.get('stage') == 'set_password':
        school_id = request.session.get('ombi_school_id')
        school = School.objects.filter(id=school_id).first()
        email = request.POST.get('email', '').strip().lower()
        password1 = request.POST.get('password1', '')
        password2 = request.POST.get('password2', '')

        bm = BoardMember.objects.filter(
            user__email__iexact=email, school=school, role='head_teacher', is_active=True
        ).first() if school else None

        if not bm:
            # Jaribu kupata kwa school tu (email inaweza kutofautiana kidogo)
            bm = BoardMember.objects.filter(school=school, role='head_teacher', is_active=True).first()
            if bm and bm.user.email.lower() != email:
                return render(request, 'field_app/public_school_head_form.html', {
                    'stage': 'set_password',
                    'school': school,
                    'error': 'Email hii hailingani na iliyosajiliwa. Tumia email uliyoiweka kwenye form.',
                    'current_year': current_year,
                })
        if not bm:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'set_password',
                'school': school,
                'error': 'Akaunti haikupatikana. Rudi na jaza form upya.',
                'current_year': current_year,
            })
        if len(password1) < 6:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'set_password', 'school': school,
                'error': 'Nywila iwe na herufi 6 au zaidi.', 'current_year': current_year,
            })
        if password1 != password2:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'set_password', 'school': school,
                'error': 'Nywila hazifanani. Jaribu tena.', 'current_year': current_year,
            })

        bm.user.set_password(password1)
        bm.user.save()
        req_id = request.session.get('ombi_request_id')
        if req_id:
            SchoolHeadRequest.objects.filter(id=req_id).update(submitter_email=email)
        login(request, bm.user, backend='field_app.backends.EmailBackend')
        request.session.pop('ombi_request_id', None)
        request.session.pop('ombi_school_id', None)
        return redirect('board_head_teacher', school_id=school.id)

    # Stage: tuma OTP baada ya mkuu kuweka email
    if request.method == 'POST' and request.POST.get('stage') == 'send_otp':
        email = request.POST.get('email', '').strip().lower()
        school_id = request.session.get('ombi_school_id')
        school = School.objects.filter(id=school_id).first()

        if not school:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'success_email',
                'error': 'Taarifa za shule zimepotea. Jaribu tena.',
                'current_year': current_year,
            })

        if not email:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'success_email',
                'school': school,
                'error': 'Tafadhali weka barua pepe.',
                'current_year': current_year,
            })

        # Sasisha request na email (kwa DEO aione)
        req_id = request.session.get('ombi_request_id')
        if req_id:
            SchoolHeadRequest.objects.filter(id=req_id).update(submitter_email=email)

        # Hifadhi email kwenye session ili itumike baadaye
        request.session['ombi_email'] = email

        import random
        otp = str(random.randint(100000, 999999))
        cache.set(f'ombi_otp_{school.id}', otp, timeout=600)

        # Jina la mtumiaji - tumia BoardMember kama apo, vinginevyo tumia "Mkuu wa Shule"
        bm = BoardMember.objects.filter(
            user__email__iexact=email, school=school, role='head_teacher', is_active=True
        ).first()
        display_name = bm.full_name if bm else 'Mkuu wa Shule'

        try:
            from django.core.mail import send_mail
            send_mail(
                subject='Msimbo wa Kuthibitisha — IMS',
                message=(
                    f'Habari {display_name},\n\n'
                    f'Ombi lako la shule ya {school.name} limepokelewa.\n\n'
                    f'Msimbo wako wa kuthibitisha ni:\n\n  {otp}\n\n'
                    f'Msimbo huu utaisha baada ya dakika 10.\n\n— IMS'
                ),
                from_email=None,
                recipient_list=[email],
                fail_silently=False,
            )
        except Exception as e:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'success_email',
                'school': school,
                'error': f'Imeshindwa kutuma email: {e}',
                'current_year': current_year,
            })

        masked = email[:2] + '***' + email[email.index('@'):]
        return render(request, 'field_app/public_school_head_form.html', {
            'stage': 'verify_otp',
            'school': school,
            'masked_email': masked,
            'current_year': current_year,
        })

    # Stage: verify OTP na login
    if request.method == 'POST' and request.POST.get('stage') == 'verify_otp':
        entered = request.POST.get('otp', '').strip()
        school_id = request.session.get('ombi_school_id')
        school = School.objects.filter(id=school_id).first()

        if not school:
            return render(request, 'field_app/public_school_head_form.html', {'current_year': current_year})

        stored = cache.get(f'ombi_otp_{school.id}')
        if not stored:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'success_email',
                'school': school,
                'error': 'Msimbo umeisha muda wake. Weka email tena.',
                'current_year': current_year,
            })

        if entered != stored:
            masked_email = request.POST.get('masked_email', '')
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'verify_otp',
                'school': school,
                'masked_email': masked_email,
                'error': 'Msimbo si sahihi. Jaribu tena.',
                'current_year': current_year,
            })

        cache.delete(f'ombi_otp_{school.id}')
        email = request.session.pop('ombi_email', '')
        req_id = request.session.pop('ombi_request_id', None)
        request.session.pop('ombi_school_id', None)

        # Kama ana BoardMember → ingia kwenye full dashboard
        bm = BoardMember.objects.filter(
            user__email__iexact=email, school=school, role='head_teacher', is_active=True
        ).first()
        if bm:
            login(request, bm.user, backend='field_app.backends.EmailBackend')
            return redirect('board_head_teacher', school_id=school.id)

        # Kama hana BoardMember → onyesha status ya ombi lake tu
        req = SchoolHeadRequest.objects.filter(id=req_id).first() if req_id else \
              SchoolHeadRequest.objects.filter(school=school, submitter_email=email).order_by('-submitted_at').first()
        return render(request, 'field_app/public_school_head_form.html', {
            'stage': 'ombi_status',
            'school': school,
            'req': req,
            'current_year': current_year,
        })

    # Stage 1b: user selected school from name-search list
    if request.method == 'POST' and request.POST.get('stage') == 'select':
        school_id = request.POST.get('school_id', '').strip()
        school = School.objects.filter(id=school_id).first()
        if not school:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'verify',
                'error': 'Shule haikutambuliwa. Jaribu tena.',
                'current_year': current_year,
            })
        already = SchoolHeadRequest.objects.filter(
            school=school, academic_year=current_year,
        ).exclude(status='rejected').exists()
        if already:
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'verify',
                'error': f'Shule ya {school.name} tayari imetuma ombi kwa mwaka huu.',
                'current_year': current_year,
            })
        return render(request, 'field_app/public_school_head_form.html', {
            'stage': 'form', 'school': school, 'current_year': current_year,
        })

    # Stage 1: verify by code or name
    if request.method == 'POST' and request.POST.get('stage') == 'verify':
        query = request.POST.get('school_code', '').strip()
        raw = query.upper().replace(' ', '').replace('-', '').replace('.', '')
        # PS0101114 / ps0101114 → PS0101114 (primary)
        # S2895 / S.2895 → S.2895 (secondary)
        if re.match(r'^PS\d+$', raw):
            code = raw  # PS0101114 — tumia kama ilivyo
        else:
            code = re.sub(r'^([SP])(\d+)$', r'\1.\2', raw)

        # Jaribu kutafuta kwa code kwanza (secondary au primary)
        school = School.objects.filter(school_code__iexact=code).first()

        if not school:
            # Gundua kama user anataka primary au secondary
            q_lower = query.lower()
            name_query = re.sub(r'\b(primary|secondary|msingi|sekondari)\b', '', q_lower, flags=re.IGNORECASE).strip()
            qs = School.objects.select_related('district')
            if any(w in q_lower for w in ['primary', 'msingi']):
                qs = qs.filter(level='Primary')
            elif any(w in q_lower for w in ['secondary', 'sekondari']):
                qs = qs.filter(level='Secondary')
            matches = qs.filter(name__icontains=name_query or query).order_by('level', 'name')[:10]

            if not matches:
                return render(request, 'field_app/public_school_head_form.html', {
                    'stage': 'verify',
                    'error': f'"{query}" haikupatikana. Jaribu namba ya usajili (S.2895) au sehemu ya jina la shule.',
                    'current_year': current_year,
                })
            if len(matches) == 1:
                school = matches[0]
            else:
                return render(request, 'field_app/public_school_head_form.html', {
                    'stage': 'select',
                    'matches': matches,
                    'query': query,
                    'current_year': current_year,
                })

        already = SchoolHeadRequest.objects.filter(
            school=school, academic_year=current_year,
        ).exclude(status='rejected').exists()
        if already:
            # Ombi lipo → peleka login moja kwa moja
            bm = BoardMember.objects.filter(school=school, role='head_teacher', is_active=True).first()
            request.session['ombi_school_id'] = school.id
            return render(request, 'field_app/public_school_head_form.html', {
                'stage': 'goto_login',
                'school': school,
                'has_account': bm is not None,
                'needs_password': bm is not None and not bm.user.has_usable_password(),
                'current_year': current_year,
            })
        return render(request, 'field_app/public_school_head_form.html', {
            'stage': 'form', 'school': school, 'current_year': current_year,
        })

    # Stage 0: initial code entry
    return render(request, 'field_app/public_school_head_form.html', {
        'stage': 'verify',
        'current_year': current_year,
    })


def school_head_submit(request, district_id):
    """Public form — school head submits required student/teacher numbers."""
    district = get_object_or_404(District, id=district_id)
    current_year = _cached_active_year()
    schools = School.objects.filter(district=district).order_by('level', 'name')

    if request.method == 'POST':
        school_name = request.POST.get('school_name', '').strip()
        head_name = request.POST.get('head_name', '').strip()
        head_phone = request.POST.get('head_phone', '').strip()
        level = request.POST.get('level', '').strip()
        try:
            students_needed = int(request.POST.get('students_needed', 0) or 0)
        except ValueError:
            students_needed = 0
        notes = request.POST.get('notes', '').strip()

        if not school_name or not head_name or not level or students_needed < 1:
            messages.error(request, 'Tafadhali jaza sehemu zote zinazohitajika.')
        else:
            # Try to auto-match school from DB
            matched_school = None
            name_lower = school_name.lower()
            for s in schools:
                if name_lower in s.name.lower() or s.name.lower() in name_lower:
                    matched_school = s
                    break

            SchoolHeadRequest.objects.create(
                district=district,
                academic_year=current_year,
                school_name_submitted=school_name,
                school=matched_school,
                head_name=head_name,
                head_phone=head_phone,
                level=level,
                students_needed=students_needed,
                notes=notes,
            )
            messages.success(request, f'Asante {head_name}! Ombi lako limetumwa kwa DEO wa {district.name}.')
            return redirect('school_head_submit', district_id=district.id)

    return render(request, 'field_app/school_head_submit.html', {
        'district': district,
        'schools': schools,
        'current_year': current_year,
    })


@board_login_required
def deo_review_requests(request, district_id):
    """DEO anaona na kusimamia maombi ya wakuu wa shule."""
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')

    district = get_object_or_404(District, id=district_id)
    current_year = _cached_active_year()
    _rqs = SchoolHeadRequest.objects.filter(district=district)
    if current_year:
        _rqs = _rqs.filter(academic_year=current_year)
    requests_qs = _rqs.select_related('school', 'reviewed_by')

    # Chair or DEO of this district can edit; others view-only
    can_edit = bm.role == 'chair' or (bm.role == 'deo' and bm.district_id == district.id)

    if request.method == 'POST' and can_edit:
        action = request.POST.get('action')

        if action == 'reject':
            req_id = request.POST.get('request_id')
            SchoolHeadRequest.objects.filter(id=req_id, district=district).update(
                status='rejected', reviewed_by=bm, reviewed_at=timezone.now()
            )
            # AJAX reject returns JSON
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': True})
            messages.success(request, 'Ombi limekataliwa.')

        elif action == 'apply_selected':
            selected_ids = request.POST.getlist('selected_requests')
            if not selected_ids:
                messages.error(request, 'Chagua maombi kwanza.')
            else:
                allocation, _ = DistrictAllocation.objects.get_or_create(
                    district=district, academic_year=current_year,
                    defaults={'uploaded_by': bm}
                )
                applied = 0
                school_capacity_map = {}  # {school_id: capacity} — bulk update baadaye

                reqs = list(SchoolHeadRequest.objects.filter(
                    id__in=selected_ids, district=district
                ).select_related('school'))

                for req in reqs:
                    if req.school:
                        SchoolAllocation.objects.update_or_create(
                            district_allocation=allocation,
                            school=req.school,
                            defaults={'quota': req.students_needed}
                        )
                        school_capacity_map[req.school_id] = req.students_needed
                    req.status = 'applied'
                    req.reviewed_by = bm
                    req.reviewed_at = timezone.now()
                    applied += 1

                # Bulk save requests (moja kwa moja — si save() ndani ya loop)
                SchoolHeadRequest.objects.bulk_update(reqs, ['status', 'reviewed_by', 'reviewed_at'])

                # Sasisha School.capacity kwa mara moja nje ya loop
                if school_capacity_map:
                    from django.db.models import Case, When, IntegerField
                    School.objects.filter(id__in=school_capacity_map.keys()).update(
                        capacity=Case(
                            *[When(id=sid, then=cap) for sid, cap in school_capacity_map.items()],
                            output_field=IntegerField(),
                        )
                    )

                # Recalculate district totals
                all_applied = SchoolHeadRequest.objects.filter(district=district, status='applied')
                if current_year:
                    all_applied = all_applied.filter(academic_year=current_year)
                allocation.primary_needed = sum(r.students_needed for r in all_applied if r.level == 'Primary')
                allocation.secondary_needed = sum(r.students_needed for r in all_applied if r.level == 'Secondary')
                allocation.uploaded_by = bm
                allocation.save(update_fields=['primary_needed', 'secondary_needed', 'uploaded_by', 'updated_at'])
                messages.success(request, f'Maombi {applied} yamewekwa. Quota ya shule zimesasishwa.')

        elif action == 'ai_format_pdf':
            return _generate_requests_pdf(requests_qs, district, current_year)

        return redirect('deo_review_requests', district_id=district.id)

    pending = requests_qs.filter(status='pending')
    applied_list = requests_qs.filter(status='applied')
    rejected_list = requests_qs.filter(status='rejected')

    # Summary counts
    primary_pending = sum(r.students_needed for r in pending if r.level == 'Primary')
    secondary_pending = sum(r.students_needed for r in pending if r.level == 'Secondary')
    primary_applied = sum(r.students_needed for r in applied_list if r.level == 'Primary')
    secondary_applied = sum(r.students_needed for r in applied_list if r.level == 'Secondary')

    return render(request, 'field_app/deo_review_requests.html', {
        'bm': bm,
        'district': district,
        'pending': pending,
        'applied_list': applied_list,
        'rejected_list': rejected_list,
        'can_edit': can_edit,
        'current_year': current_year,
        'total_pending': pending.count(),
        'total_applied': applied_list.count(),
        'primary_pending': primary_pending,
        'secondary_pending': secondary_pending,
        'primary_applied': primary_applied,
        'secondary_applied': secondary_applied,
    })


def _generate_requests_pdf(requests_qs, district, current_year):
    """Generate a formatted PDF of school head requests using AI summary + ReportLab."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # AI summary
    ai_summary = _ai_summarise_requests(requests_qs)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    navy = colors.HexColor('#0A2B5E')
    gold = colors.HexColor('#C8900A')
    light = colors.HexColor('#EEF1F6')

    title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                 fontSize=16, textColor=navy, spaceAfter=4,
                                 alignment=TA_CENTER, fontName='Helvetica-Bold')
    sub_style = ParagraphStyle('Sub', parent=styles['Normal'],
                               fontSize=10, textColor=gold, spaceAfter=2, alignment=TA_CENTER)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, spaceAfter=4)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'],
                                   fontSize=11, textColor=navy, spaceBefore=12, spaceAfter=4,
                                   fontName='Helvetica-Bold')

    story = []
    story.append(Paragraph('JAMHURI YA MUUNGANO WA TANZANIA', sub_style))
    story.append(Paragraph('Wizara ya Elimu, Sayansi na Teknolojia', sub_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width='100%', thickness=3, color=gold))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(f'MAHITAJI YA WALIMU WANAFUNZI', title_style))
    story.append(Paragraph(f'Wilaya ya {district.name} — Mwaka {current_year.year if current_year else "—"}', sub_style))
    story.append(Spacer(1, 0.2*cm))
    story.append(HRFlowable(width='100%', thickness=1, color=navy))
    story.append(Spacer(1, 0.4*cm))

    # AI Summary section
    if ai_summary:
        story.append(Paragraph('Muhtasari wa AI', section_style))
        story.append(Paragraph(ai_summary, body_style))
        story.append(Spacer(1, 0.4*cm))

    # Primary schools table
    primary_reqs = [r for r in requests_qs if r.level == 'Primary' and r.status != 'rejected']
    secondary_reqs = [r for r in requests_qs if r.level == 'Secondary' and r.status != 'rejected']

    def make_table(reqs, level_label):
        if not reqs:
            return
        story.append(Paragraph(f'Shule za {level_label}', section_style))
        headers = ['#', 'Jina la Shule', 'Mkuu wa Shule', 'Simu', 'Idadi', 'Hali']
        data = [headers]
        total = 0
        for i, r in enumerate(reqs, 1):
            school_name = r.school.name if r.school else r.school_name_submitted
            status_map = {'pending': 'Inasubiri', 'reviewed': 'Imepitiwa',
                          'applied': 'Imewekwa', 'rejected': 'Imekataliwa'}
            data.append([
                str(i), school_name, r.head_name,
                r.head_phone or '—', str(r.students_needed), status_map.get(r.status, r.status)
            ])
            total += r.students_needed
        data.append(['', 'JUMLA', '', '', str(total), ''])

        col_widths = [1*cm, 5*cm, 4*cm, 3*cm, 1.5*cm, 2.5*cm]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), navy),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (4, 0), (4, -1), 'CENTER'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, light]),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#FEF3C7')),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E0')),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    make_table(primary_reqs, 'Msingi')
    make_table(secondary_reqs, 'Sekondari')

    # Grand total
    all_reqs = primary_reqs + secondary_reqs
    grand_total = sum(r.students_needed for r in all_reqs)
    story.append(HRFlowable(width='100%', thickness=1, color=navy))
    story.append(Spacer(1, 0.2*cm))
    total_data = [
        ['Jumla — Msingi', str(sum(r.students_needed for r in primary_reqs))],
        ['Jumla — Sekondari', str(sum(r.students_needed for r in secondary_reqs))],
        ['JUMLA KUU', str(grand_total)],
    ]
    t2 = Table(total_data, colWidths=[12*cm, 5*cm])
    t2.setStyle(TableStyle([
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, -1), (-1, -1), navy),
        ('TEXTCOLOR', (0, -1), (-1, -1), colors.white),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(t2)
    story.append(Spacer(1, 1*cm))
    from django.utils import timezone as tz
    story.append(Paragraph(f'Imetengenezwa: {tz.now().strftime("%d/%m/%Y %H:%M")} | Mfumo wa IMS — AI Generated', body_style))

    doc.build(story)
    buffer.seek(0)

    from django.http import HttpResponse
    filename = f'mahitaji_{district.name.replace(" ", "_")}_{(current_year.year if current_year else "")}.pdf'
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _send_sms_africastalking(phone, message):
    """Tuma SMS moja kupitia Africa's Talking API. Rudisha (True, '') au (False, error)."""
    import urllib.request
    import urllib.parse
    username = os.environ.get('AT_USERNAME', '')
    api_key = os.environ.get('AT_API_KEY', '')
    if not username or not api_key:
        return False, 'AT_USERNAME au AT_API_KEY hazijawekwa kwenye .env'

    # Normalize phone to +255XXXXXXXXX
    phone = phone.strip().replace(' ', '').replace('-', '')
    if phone.startswith('0') and len(phone) == 10:
        phone = '+255' + phone[1:]
    elif phone.startswith('255') and len(phone) == 12:
        phone = '+' + phone
    elif not phone.startswith('+'):
        phone = '+255' + phone

    data = urllib.parse.urlencode({
        'username': username,
        'to': phone,
        'message': message,
    }).encode()
    req = urllib.request.Request(
        'https://api.africastalking.com/version1/messaging',
        data=data,
        headers={
            'apiKey': api_key,
            'Accept': 'application/json',
            'Content-Type': 'application/x-www-form-urlencoded',
        },
        method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = resp.read().decode()
            if '"status":"Success"' in body or '"Status":"Success"' in body:
                return True, ''
            return False, body
    except Exception as e:
        return False, str(e)


@board_login_required
def send_links_to_heads(request, district_id):
    """Tuma SMS ya link ya kujaza form kwa wakuu wa shule zote kwenye wilaya."""
    bm = _get_board_member(request)
    if not bm:
        return redirect('board_login')

    district = get_object_or_404(District, id=district_id)
    can_edit = bm.role == 'chair' or (bm.role == 'deo' and bm.district_id == district.id)
    if not can_edit:
        messages.error(request, 'Huna ruhusa ya kutuma ujumbe.')
        return redirect('deo_review_requests', district_id=district.id)

    if request.method != 'POST':
        return redirect('deo_review_requests', district_id=district.id)

    # Build the submission link
    scheme = request.scheme
    host = request.get_host()
    submit_path = reverse('school_head_submit', args=[district.id])
    link = f'{scheme}://{host}{submit_path}'

    current_year = _cached_active_year()
    year_str = current_year.year if current_year else ''

    schools_with_phone = School.objects.filter(
        district=district
    ).exclude(head_phone='').order_by('level', 'name')

    schools_without_phone = School.objects.filter(
        district=district, head_phone=''
    ).count()

    sent, failed, skipped = [], [], []

    for school in schools_with_phone:
        # Filter by level if requested
        level_filter = request.POST.get('level', 'both')
        if level_filter != 'both' and school.level != level_filter:
            skipped.append(school.name)
            continue

        message = (
            f"Habari{' ' + school.head_name if school.head_name else ''},\n"
            f"Tafadhali jaza fomu ya mahitaji ya walimu wanafunzi "
            f"kwa mwaka {year_str} kwa Wilaya ya {district.name}:\n"
            f"{link}\n"
            f"- Mfumo wa IMS"
        )
        ok, err = _send_sms_africastalking(school.head_phone, message)
        if ok:
            sent.append(school.name)
        else:
            failed.append({'school': school.name, 'phone': school.head_phone, 'error': err})

    return render(request, 'field_app/send_links_result.html', {
        'district': district,
        'link': link,
        'sent': sent,
        'failed': failed,
        'skipped_count': skipped,
        'no_phone_count': schools_without_phone,
        'total': schools_with_phone.count(),
        'bm': bm,
    })


def _ai_summarise_requests(requests_qs):
    """Tumia AI kutoa muhtasari wa maombi."""
    try:
        from ai_utils import client, model_name
        lines = []
        for r in requests_qs:
            if r.status != 'rejected':
                sn = r.school.name if r.school else r.school_name_submitted
                lines.append(f"- {sn} ({r.level}): {r.students_needed} walimu wanafunzi — Mkuu: {r.head_name}")
        if not lines:
            return ''
        data_text = '\n'.join(lines)
        prompt = f"""Toa muhtasari mfupi (mistari 3-4) wa maombi haya ya walimu wanafunzi kwa ajili ya ripoti rasmi:

{data_text}

Andika kwa Kiswahili, muhtasari wa kitaalamu unaofaa ripoti rasmi ya serikali. Taja jumla na usambazaji kwa shule za msingi na sekondari."""
        resp = client.models.generate_content(model=model_name, contents=prompt)
        return resp.text.strip()
    except Exception:
        return ''
