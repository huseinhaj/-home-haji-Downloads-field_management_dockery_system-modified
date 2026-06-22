"""
Generate a plain black-and-white version of the CIVE DAY 2026 report.
- No colors (black text, white/light-grey backgrounds only)
- No images (logo replaced with text)
- Proper software development methodology added (Agile Iterative Model)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette (plain black & white only) ─────────────────────────────────
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK  = RGBColor(0x22, 0x22, 0x22)
MID   = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='000000', sz='6', val='single'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   val)
        tag.set(qn('w:sz'),    sz)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading_para(doc, text, size=13, bold=True, center=False, space_before=10, space_after=4, underline=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold           = bold
    run.underline      = underline
    run.font.size      = Pt(size)
    run.font.color.rgb = BLACK
    return p

def body_para(doc, text, size=10.5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = DARK
    return p

def divider(doc):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr= OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE  (logo + university header + title + members — all in one place)
# ══════════════════════════════════════════════════════════════════════════════
LOGO_PATH = '/home/haji/New Folder 2/udom_logo.jpeg'

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(10)
p_logo.paragraph_format.space_after  = Pt(6)
run_logo = p_logo.add_run()
run_logo.add_picture(LOGO_PATH, width=Inches(1.4))

for line, sz, bold in [
    ('THE UNIVERSITY OF DODOMA', 16, True),
    ('COLLEGE OF INFORMATICS AND VIRTUAL EDUCATION', 12, True),
    ('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', 11, True),
    ('CIVE DAY 2026 — PROJECT PRESENTATION', 11, False),
    ('ACADEMIC YEAR: 2025/2026', 10, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(line)
    run.bold           = bold
    run.font.size      = Pt(sz)
    run.font.color.rgb = BLACK

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

heading_para(doc, 'PROJECT TITLE:', size=11, bold=False, center=True, space_before=8, space_after=2)
heading_para(doc,
    'A DIGITAL PLATFORM FOR INTERNSHIP STUDENT TEACHER MANAGEMENT SYSTEM',
    size=13, bold=True, center=True, space_before=2, space_after=10)

doc.add_paragraph()
heading_para(doc, 'MEMBERS', size=11, bold=True, center=True, space_before=6, space_after=4)

members_tbl = doc.add_table(rows=3, cols=4)
members_tbl.style = 'Table Grid'
members_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for col, h in enumerate(['NO.', 'NAME', 'REG. NUMBER', 'PROGRAM']):
    cell = members_tbl.cell(0, col)
    set_cell_bg(cell, 'DDDDDD')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK

rows_data = [
    ('1', 'Haji Hamis Huseni',    'T22-03-11756', 'BSC. SE4'),
    ('2', 'Hamis Selemani Hamis', 'T22-03-11748', 'BSC. SE4'),
]
for ri, cols in enumerate(rows_data):
    for ci, val in enumerate(cols):
        cell = members_tbl.cell(ri + 1, ci)
        set_cell_bg(cell, 'FFFFFF')
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

p_sup = doc.add_paragraph()
p_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sup.paragraph_format.space_before = Pt(10)
r_sup = p_sup.add_run('SUPERVISOR:  Mr. Barongo')
r_sup.font.size = Pt(10)
r_sup.bold = True
r_sup.font.color.rgb = BLACK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION  (no duplication of cover page info)
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading_para(doc, '1. INTRODUCTION', size=12, space_before=8)
divider(doc)

heading_para(doc, 'a. Background', size=10.5, bold=True, space_before=6, space_after=2)
body_para(doc,
    'In Tanzania, student teachers from universities and colleges are required to complete a '
    'practical internship in secondary and primary schools across the country. '
    'This involves placing students in schools, assigning supervisors (assessors), monitoring '
    'teaching activities through daily logbooks, and issuing completion certificates at the end '
    'of the internship period. Coordination spans multiple stakeholders: students, school assessors, '
    'head teachers, District Education Officers (DEOs), and university administration. '
    'Traditionally, this entire process has been handled manually — through paper logbooks, '
    'physical assessment forms, and in-person coordination — leading to significant inefficiencies.'
)

heading_para(doc, 'b. Problem Statement', size=10.5, bold=True, space_before=6, space_after=2)
body_para(doc, 'The manual management of student teacher internships creates the following critical challenges:')
problems = [
    'No centralized system to track hundreds of student teachers across regions and districts.',
    'Paper logbooks are prone to loss, forgery, and inconsistency in record-keeping.',
    'Assessors have no efficient digital tool to record evaluations and observations.',
    'DEOs and head teachers lack real-time visibility into student progress.',
    'Certificate issuance is delayed due to multi-level manual approval processes.',
    'Communication between university admin, schools, and students is slow and unreliable.',
]
for prob in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(prob)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

heading_para(doc, 'c. Main Objective', size=10.5, bold=True, space_before=6, space_after=2)
body_para(doc,
    'To design and develop a web-based system that digitizes and automates the management '
    'of student teacher internship in Tanzania.'
)

heading_para(doc, 'Specific Objectives:', size=10, bold=True, space_before=4, space_after=2)
objectives = [
    'Automate student registration, school selection, and subject assignment.',
    'Enable digital logbook submission with assessor feedback and remarks.',
    'Provide assessors with tools to evaluate and grade student performance.',
    'Give DEOs and head teachers a dashboard for district-wide progress monitoring.',
    'Generate tamper-proof, digitally-signed completion certificates and official letters.',
    'Integrate AI tools to assist students in creating Schemes of Work and Lesson Plans.',
]
for obj in objectives:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(obj)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE CORE INNOVATION
# ══════════════════════════════════════════════════════════════════════════════
heading_para(doc, '2. THE CORE INNOVATION', size=12, space_before=10)
divider(doc)

# ── System Architecture (plain text table, no colors) ─────────────────────────
heading_para(doc, 'a. System Architecture Overview', size=10.5, bold=True, space_before=6, space_after=4)

arch = doc.add_table(rows=7, cols=5)
arch.style = 'Table Grid'
arch.alignment = WD_TABLE_ALIGNMENT.CENTER

def arch_cell(cell, text, header=False, sz=8):
    set_cell_bg(cell, 'EEEEEE' if header else 'FFFFFF')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    r = p.add_run(text)
    r.bold = header
    r.font.size = Pt(sz)
    r.font.color.rgb = BLACK

# Row 0 — User layer
arch.cell(0, 0).merge(arch.cell(0, 4))
arch_cell(arch.cell(0, 0),
    'ACCESS LAYER: Web Browser (HTTPS)\nStudent Teacher | Assessor | Admin | DEO | Head Teacher',
    header=True, sz=8.5)

# Row 1 — application tier
arch.cell(1, 0).merge(arch.cell(1, 4))
arch_cell(arch.cell(1, 0),
    'APPLICATION TIER: Django Web Application (Python 3.12 / Django 5.2 — MVT Architecture)',
    header=True, sz=8.5)

# Row 2 — core modules
for col, txt in enumerate([
    'AUTH MODULE\nLogin / Register\nPassword Reset',
    'STUDENT MODULE\nSchool Selection\nLogbook / Reports',
    'ASSESSOR MODULE\nEvaluations\nRemarks / Grades',
    'BOARD / DEO\nDashboard\nApprovals',
    'ADMIN MODULE\nUser Management\nAssignments',
]):
    arch_cell(arch.cell(2, col), txt, sz=7.5)

# Row 3 — support layers
for col, txt in enumerate([
    'AI TOOLS\nScheme of Work\nLesson Plans',
    'PDF ENGINE\nCertificates\nLogbook PDFs',
    'SMS / EMAIL\nAfrica\'s Talking\nGmail SMTP',
    'CACHING\nRedis\nSession Store',
    'REST APIs\nAJAX Endpoints\nJSON Responses',
]):
    arch_cell(arch.cell(3, col), txt, sz=7.5)

# Row 4 — data label
arch.cell(4, 0).merge(arch.cell(4, 4))
arch_cell(arch.cell(4, 0), 'DATA LAYER', header=True, sz=8)

# Row 5 — databases & external
arch.cell(5, 0).merge(arch.cell(5, 1))
arch_cell(arch.cell(5, 0),
    'PostgreSQL 16 Database\n(Schools / Students /\nAssessments / Logbooks)',
    sz=7.5)
arch_cell(arch.cell(5, 2),
    'Docker Container\n(Web + DB\nIsolated)',
    sz=7.5)
arch.cell(5, 3).merge(arch.cell(5, 4))
arch_cell(arch.cell(5, 3),
    'External Services\n- OpenAI / Google Gemini (AI)\n- Gmail SMTP (Email)\n- Africa\'s Talking (SMS)',
    sz=7.5)

# Row 6 — legend
arch.cell(6, 0).merge(arch.cell(6, 4))
arch_cell(arch.cell(6, 0),
    'LEGEND: Grey header rows = Tiers/Labels  |  White rows = Modules/Components',
    header=False, sz=7.5)

doc.add_paragraph()

# ── Methodology ───────────────────────────────────────────────────────────────
heading_para(doc, 'b. Methodology', size=10.5, bold=True, space_before=6, space_after=4)

# -- Development Methodology --------------------------------------------------
heading_para(doc, 'Development Methodology: Agile Iterative Model',
    size=10, bold=True, space_before=4, space_after=2)
body_para(doc,
    'This project used the Agile Iterative Development Methodology. '
    'The system was built in short iterations, each producing a working part of the system. '
    'After each iteration, the team reviewed progress, gathered feedback from stakeholders, '
    'and improved the next iteration. This approach allowed the team to add new features '
    'gradually, fix issues early, and deliver a complete and tested internship management system at the end.'
)
body_para(doc,
    'Each iteration followed this cycle:  Plan  →  Design  →  Implement  →  Test  →  Review  →  Next Iteration.'
)

doc.add_paragraph()

# -- Technical Stack table ---------------------------------------------------
heading_para(doc, 'Technical Tools and Technologies Used', size=10, bold=True, space_before=4, space_after=4)

meth_rows = [
    ('Programming Language', 'Python 3.12'),
    ('Web Framework',        'Django 5.2 (MVT Architecture)'),
    ('Database',             'PostgreSQL 16 (via Docker)'),
    ('AI Integration',       'OpenAI GPT-4 & Google Gemini — AI-powered Scheme of Work & Lesson Plan generation'),
    ('PDF Generation',       'ReportLab — Certificates, Logbooks, Official Letters with watermarks & signatures'),
    ('Containerization',     'Docker & Docker Compose — isolated, reproducible deployment'),
    ('Frontend',             'HTML5, CSS3, JavaScript — Responsive multi-role web interface'),
    ('Communication',        'Gmail SMTP (email) · Africa\'s Talking API (SMS notifications to head teachers)'),
    ('Version Control',      'Git — source code management and collaboration'),
    ('Testing',              'Django TestCase (unit tests) · Manual functional testing per iteration'),
]

meth_tbl = doc.add_table(rows=len(meth_rows), cols=2)
meth_tbl.style = 'Table Grid'
meth_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(meth_rows):
    row = meth_tbl.rows[i]
    lc = row.cells[0]
    vc = row.cells[1]

    set_cell_bg(lc, 'DDDDDD')
    set_cell_border(lc)
    lc.width = Cm(4.5)
    lp = lc.paragraphs[0]
    lrun = lp.add_run(label)
    lrun.bold = True
    lrun.font.size = Pt(9)
    lrun.font.color.rgb = BLACK
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    set_cell_bg(vc, 'FFFFFF')
    set_cell_border(vc)
    vp = vc.paragraphs[0]
    vrun = vp.add_run(value)
    vrun.font.size = Pt(9.5)
    vrun.font.color.rgb = DARK

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
divider(doc)
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(4)
fr = fp.add_run('CIVE DAY 2026  |  UDOM  |  CSE Department  |  NI ELIMU TU  |  MUNGU IBARIKI UDOM')
fr.font.size = Pt(8.5)
fr.font.color.rgb = BLACK
fr.bold = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/haji/Downloads/field_management_dockery_system-modifiedied/CIVE_DAY_2026_IMS_Plain.docx'
doc.save(out)
print(f'SAVED: {out}')
