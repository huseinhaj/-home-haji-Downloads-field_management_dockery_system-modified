"""
Curriculum app views — standalone Scheme of Work, Lesson Plan & Logbook.
Uses models from field_app but provides its own views, templates, and URLs.
"""
import json
import os
import re
import threading
from io import BytesIO
from datetime import datetime, timedelta

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.utils import IntegrityError
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.cache import never_cache
from django.contrib.auth import logout as django_logout


from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable, PageBreak

import logging

# Optional PIL for watermark support
try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


from .ai_utils import client, model_name, OPENROUTER_API_KEY, GROQ_API_KEY, GOOGLE_API_KEY
from .forms import SchemeOfWorkForm, LogbookForm

logger = logging.getLogger(__name__)

from field_app.models import (
    StudentTeacher, Subject, SchemeOfWork, LessonPlan,
    LogbookEntry, EducationLevel, ClassLevel, School,
    SchoolSubjectCapacity, StudentApplication, SchoolAssessment,
    AcademicYear, Textbook, BoardMember, BoardComment, Region, District,
)
from field_app.views.utils import (
    _cached_active_year, _cached_subjects, _cached_today_logbook,
    _invalidate_today_logbook, get_or_create_student_profile,
    get_current_academic_year, invalidate_student_cache,
)

from .models import (
    TLMTeacher, Testimonial, LessonNote, SubjectTopic, TopicSubtopic,
    TLMLogbookEntry, GeneratedExam,
)


# =============================================================================
# PWA: Progressive Web App — manifest & service worker
# =============================================================================

def pwa_manifest(request):
    """Serve Web App Manifest for PWA install on mobile devices."""
    manifest = {
        "name": "TLM Tanzania — Teaching & Learning Materials",
        "short_name": "TLM Tanzania",
        "description": "Mfumo wa Nyenzo za Kufundishia na Kujifunza — Scheme of Work, Lesson Plan & Logbook kwa Walimu wa Tanzania",
        "start_url": "/curriculum/",
        "scope": "/",
        "display": "standalone",
        "orientation": "portrait-primary",
        "background_color": "#0D4F2B",
        "theme_color": "#0D4F2B",
        "categories": ["education", "productivity", "books"],
        "lang": "sw",
        "dir": "ltr",
        "icons": [
            {
                "src": "/static/curriculum/pwa-icon-192.png",
                "sizes": "192x192",
                "type": "image/png",
                "purpose": "any"
            },
            {
                "src": "/static/curriculum/pwa-icon-512.png",
                "sizes": "512x512",
                "type": "image/png",
                "purpose": "any maskable"
            },
            {
                "src": "/static/curriculum/pwa-icon.svg",
                "sizes": "any",
                "type": "image/svg+xml",
                "purpose": "any"
            }
        ],
        "screenshots": [
            {
                "src": "/static/curriculum/pwa-splash.png",
                "sizes": "1280x640",
                "type": "image/png",
                "form_factor": "wide",
                "label": "TLM Tanzania — Scheme of Work, Lesson Plan & Logbook"
            },
            {
                "src": "/static/curriculum/pwa-splash.png",
                "sizes": "1280x640",
                "type": "image/png",
                "form_factor": "narrow",
                "label": "TLM Tanzania — Nyenzo za Kufundishia"
            }
        ],
        "shortcuts": [
            {
                "name": "Scheme of Work",
                "short_name": "Scheme",
                "description": "Tengeneza Scheme of Work kwa AI",
                "url": "/curriculum/scheme/",
                "icons": [{"src": "/static/curriculum/pwa-icon-192.png", "sizes": "192x192"}]
            },
            {
                "name": "Lesson Plan",
                "short_name": "Lesson",
                "description": "Tengeneza Lesson Plan kwa AI",
                "url": "/curriculum/lesson-plan/",
                "icons": [{"src": "/static/curriculum/pwa-icon-192.png", "sizes": "192x192"}]
            },
            {
                "name": "Library",
                "short_name": "Library",
                "description": "Tazama mifano ya walimu wengine",
                "url": "/curriculum/library/",
                "icons": [{"src": "/static/curriculum/pwa-icon-192.png", "sizes": "192x192"}]
            },
            {
                "name": "Lesson Notes",
                "short_name": "Notes",
                "description": "Andika na hifadhi maelezo ya somo",
                "url": "/curriculum/lesson-notes/",
                "icons": [{"src": "/static/curriculum/pwa-icon-192.png", "sizes": "192x192"}]
            }
        ]
    }
    return JsonResponse(manifest)


def pwa_service_worker(request):
    """Serve the Service Worker JS for offline caching & PWA install."""
    sw_code = '''const CACHE_NAME = "tlm-tanzania-v1";
const STATIC_ASSETS = [
  "/static/curriculum/pwa-icon-192.png",
  "/static/curriculum/pwa-icon-512.png",
  "/static/curriculum/pwa-icon.svg",
];

// ── Install: cache static assets ──
self.addEventListener("install", function(event) {
  event.waitUntil(
    caches.open(CACHE_NAME).then(function(cache) {
      return cache.addAll(STATIC_ASSETS);
    })
  );
  self.skipWaiting();
});

// ── Activate: clean old caches ──
self.addEventListener("activate", function(event) {
  event.waitUntil(
    caches.keys().then(function(keys) {
      return Promise.all(
        keys.filter(function(k) { return k !== CACHE_NAME; })
            .map(function(k) { return caches.delete(k); })
      );
    })
  );
  self.clients.claim();
});

// ── Fetch: network-first, fallback to cache ──
self.addEventListener("fetch", function(event) {
  // Only handle GET requests
  if (event.request.method !== "GET") return;
  
  // Skip non-http(s) URLs
  if (!event.request.url.startsWith("http")) return;
  
  event.respondWith(
    fetch(event.request)
      .then(function(response) {
        // Cache successful responses
        if (response.status === 200) {
          const clone = response.clone();
          caches.open(CACHE_NAME).then(function(cache) {
            cache.put(event.request, clone);
          });
        }
        return response;
      })
      .catch(function() {
        // Offline: serve from cache
        return caches.match(event.request).then(function(cached) {
          return cached || new Response("Offline", { status: 503 });
        });
      })
  );
});
'''
    return HttpResponse(sw_code, content_type="application/javascript; charset=utf-8")


def _sanitize_json_control_chars(text):
    """Replace control characters inside JSON strings with their escaped forms.
    This handles AI responses that include literal newlines/tabs in string values,
    which would otherwise cause json.loads() to fail with 'Invalid control character'."""
    result = []
    in_str = False
    prev = ''
    for ch in text:
        if ch == '"' and prev != '\\':
            in_str = not in_str
        if in_str and ch in ('\n', '\r', '\t'):
            result.append({'\n': '\\n', '\r': '\\r', '\t': '\\t'}[ch])
        else:
            result.append(ch)
        prev = ch
    return ''.join(result)


# =============================================================================
# HELPER: Generate one batch of Scheme of Work data
# =============================================================================

def _generate_scheme_batch(prompt_text):
    """Make a single AI call for one batch of scheme data and parse the JSON response.
    Returns (scheme_data_list, response_text) or (None, response_text) on failure."""
    response = client.models.generate_content(model=model_name, contents=prompt_text)
    response_text = response.text
    logger.info(f"[Scheme Batch] AI response length: {len(response_text)} chars")

    # Strip markdown code blocks
    cleaned = re.sub(r'```(?:json)?\s*', '', response_text)
    cleaned = re.sub(r'```\s*', '', cleaned).strip()

    # Helper: parse first valid JSON array from text using bracket matching
    def _extract(text):
        start = text.find('[')
        if start == -1:
            return None
        depth = 0
        in_str = False
        esc = False
        for i in range(start, len(text)):
            ch = text[i]
            if esc:
                esc = False
                continue
            if ch == '\\' and in_str:
                esc = True
                continue
            if ch == '"' and not esc:
                in_str = not in_str
                continue
            if in_str:
                continue
            if ch == '[':
                depth += 1
            elif ch == ']':
                depth -= 1
                if depth == 0:
                    candidate = text[start:i+1]
                    try:
                        return json.loads(candidate)
                    except json.JSONDecodeError:
                        try:
                            return json.loads(_sanitize_json_control_chars(candidate))
                        except json.JSONDecodeError:
                            continue
        # No closing bracket - extract complete objects
        candidate = text[start:]
        complete_objs = []
        i = 0
        obj_start = None
        obj_depth = 0
        in_str2 = False
        esc2 = False
        while i < len(candidate):
            ch = candidate[i]
            if esc2:
                esc2 = False
                i += 1
                continue
            if ch == '\\' and in_str2:
                esc2 = True
                i += 1
                continue
            if ch == '"' and not esc2:
                in_str2 = not in_str2
                i += 1
                continue
            if in_str2:
                i += 1
                continue
            if ch == '{':
                if obj_depth == 0:
                    obj_start = i
                obj_depth += 1
            elif ch == '}':
                obj_depth -= 1
                if obj_depth == 0 and obj_start is not None:
                    complete_objs.append(candidate[obj_start:i+1])
                    obj_start = None
            i += 1
        if complete_objs:
            rebuilt = '[' + ','.join(complete_objs) + ']'
            try:
                return json.loads(rebuilt)
            except json.JSONDecodeError:
                try:
                    return json.loads(_sanitize_json_control_chars(rebuilt))
                except json.JSONDecodeError:
                    pass
        # Bracket-closing fallback
        if candidate.count('"') % 2 == 1:
            candidate += '"'
        open_b = candidate.count('{')
        close_b = candidate.count('}')
        if open_b > close_b:
            candidate += '}' * (open_b - close_b)
        open_arr = candidate.count('[')
        close_arr = candidate.count(']')
        if open_arr > close_arr:
            candidate += ']' * (open_arr - close_arr)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            try:
                return json.loads(_sanitize_json_control_chars(candidate))
            except json.JSONDecodeError:
                return None

    scheme_data = _extract(cleaned)
    
    # Fallbacks
    if scheme_data is None:
        try:
            parsed = json.loads(cleaned)
            if isinstance(parsed, dict):
                for v in parsed.values():
                    if isinstance(v, list):
                        scheme_data = v
                        break
            elif isinstance(parsed, list):
                scheme_data = parsed
        except (json.JSONDecodeError, TypeError):
            try:
                parsed = json.loads(_sanitize_json_control_chars(cleaned))
                if isinstance(parsed, dict):
                    for v in parsed.values():
                        if isinstance(v, list):
                            scheme_data = v
                            break
                elif isinstance(parsed, list):
                        scheme_data = parsed
            except (json.JSONDecodeError, TypeError):
                pass

    if scheme_data is None:
        for attempt in [cleaned, _sanitize_json_control_chars(cleaned)]:
            try:
                m = re.search(r'\[.*?\]', attempt, re.DOTALL)
                candidate = m.group() if m else None
                if not candidate:
                    start = attempt.find('[')
                    if start != -1:
                        candidate = attempt[start:]
                if candidate:
                    open_b = candidate.count('{')
                    close_b = candidate.count('}')
                    if open_b > close_b:
                        candidate += '}' * (open_b - close_b)
                    open_arr = candidate.count('[')
                    close_arr = candidate.count(']')
                    if open_arr > close_arr:
                        candidate += ']' * (open_arr - close_arr)
                    try:
                        scheme_data = json.loads(candidate)
                    except json.JSONDecodeError:
                        try:
                            scheme_data = json.loads(_sanitize_json_control_chars(candidate))
                        except json.JSONDecodeError:
                            continue
                    if scheme_data:
                        break
            except (json.JSONDecodeError, TypeError):
                continue

    # Post-process: convert arrays to strings
    if scheme_data and isinstance(scheme_data, list):
        for row in scheme_data:
            for key in row:
                if isinstance(row[key], list):
                    row[key] = ', '.join(str(v) for v in row[key] if v)
                elif not isinstance(row[key], str):
                    row[key] = str(row[key] or '')

    return scheme_data, response_text


# =============================================================================
# HELPER: Check/Get TLM teacher from session
# =============================================================================

def get_tlm_teacher(request):
    """Get the registered TLM teacher from session, or None."""
    teacher_id = request.session.get('tlm_teacher_id')
    if teacher_id:
        try:
            return TLMTeacher.objects.get(id=teacher_id)
        except TLMTeacher.DoesNotExist:
            del request.session['tlm_teacher_id']
    return None


# =============================================================================
# LANDING PAGE (public) — never_cache to show LIVE statistics
# =============================================================================

@never_cache
def landing(request):
    """Public landing page — shows tools. If teacher is registered, greet them."""
    teacher = get_tlm_teacher(request)
    
    # Real DB statistics for the landing page
    from django.db.models import Count
    total_teachers = TLMTeacher.objects.count()
    total_schemes = SchemeOfWork.objects.count()
    total_lesson_plans = LessonPlan.objects.count()
    total_logbooks = LogbookEntry.objects.count()
    total_notes = LessonNote.objects.count()
    
    # Real testimonials from teachers
    testimonials = Testimonial.objects.filter(is_approved=True).select_related('teacher__school')[:6]
    
    return render(request, 'curriculum/landing.html', {
        'teacher': teacher,
        'total_teachers': total_teachers,
        'total_schemes': total_schemes,
        'total_lesson_plans': total_lesson_plans,
        'total_logbooks': total_logbooks,
        'total_notes': total_notes,
        'testimonials': testimonials,
    })


# =============================================================================
# TEMPLATE LIBRARY — browse saved schemes & lesson plans
# =============================================================================

@never_cache
def template_library(request):
    """Browse ALL saved schemes and lesson plans from all teachers."""
    teacher = get_tlm_teacher(request)
    
    # Show ALL schemes and lesson plans from ALL teachers — not just the logged-in user
    # This creates a shared library where everyone can see what others have created
    all_schemes = SchemeOfWork.objects.all().select_related(
        'subject', 'school'
    ).order_by('-updated_at')[:50]
    
    all_lessons = LessonPlan.objects.all().select_related(
        'subject'
    ).order_by('-created_at')[:50]
    
    # Counts for stats
    total_schemes_count = SchemeOfWork.objects.count()
    total_lessons_count = LessonPlan.objects.count()
    
    return render(request, 'curriculum/library.html', {
        'teacher': teacher,
        'schemes': all_schemes,
        'lesson_plans': all_lessons,
        'total_schemes_count': total_schemes_count,
        'total_lessons_count': total_lessons_count,
    })


# =============================================================================
# TLM LOGOUT — clear session for TLM teachers
# =============================================================================

def tlm_logout(request):
    """Logout from TLM system — clears the TLM teacher session and Django auth."""
    # Clear the TLM teacher session
    if 'tlm_teacher_id' in request.session:
        del request.session['tlm_teacher_id']
    # Also clear Django auth session (in case user is both TLM + Django logged in)
    django_logout(request)
    return redirect('curriculum:landing')


# =============================================================================
# SUBMIT TESTIMONIAL / FEEDBACK
# =============================================================================

@require_POST
def ajax_submit_testimonial(request):
    """Submit a testimonial/feedback from a teacher."""
    try:
        data = json.loads(request.body)
        message = data.get('message', '').strip()
        teacher_name = data.get('teacher_name', '').strip()
        school_name = data.get('school_name', '').strip()
        
        if not message or not teacher_name:
            return JsonResponse({'success': False, 'error': 'Tafadhali jaza jina na ujumbe.'}, status=400)
        
        teacher = get_tlm_teacher(request)
        
        testimonial = Testimonial.objects.create(
            teacher=teacher,
            teacher_name=teacher_name,
            school_name=school_name,
            message=message,
            is_approved=True,  # Auto-approve for now
        )
        
        return JsonResponse({'success': True, 'id': testimonial.id})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


# =============================================================================
# TEACHER REGISTRATION (simple — no login needed)
# =============================================================================

def teacher_register(request):
    """
    Registration page: teacher selects Region → District → School → 
    (auto Education Level) → Class → Stream → Subject → Registered Students.
    If already registered (session), redirect to landing.
    """
    # Check if already registered
    teacher = get_tlm_teacher(request)
    if teacher:
        # Already registered — redirect where they were going
        next_url = request.GET.get('next', 'curriculum:landing')
        return redirect(next_url)

    regions = Region.objects.all().order_by('name')
    subjects = Subject.objects.all().order_by('name')
    
    import json as _json
    education_levels = EducationLevel.objects.all().order_by('order')
    classes_by_level = {}
    for cl in ClassLevel.objects.select_related('education_level').order_by('education_level', 'order'):
        classes_by_level.setdefault(cl.education_level_id, []).append({'id': cl.id, 'name': cl.name})

    # Masomo kwa kila ngazi ya elimu (Primary→primary, Technical/VETA→technical, n.k.)
    def _subjects_for_level(level_name):
        ln = (level_name or '').lower()
        if 'primary' in ln:
            return list(Subject.objects.filter(level='primary').order_by('name').values('id', 'name'))
        if 'technical' in ln or 'veta' in ln:
            return list(Subject.objects.filter(level='technical').order_by('name').values('id', 'name'))
        if 'advanced' in ln:
            return list(Subject.objects.filter(level='advanced').order_by('name').values('id', 'name'))
        return list(Subject.objects.filter(level='secondary').order_by('name').values('id', 'name'))

    subjects_by_level = {
        lvl.id: _subjects_for_level(lvl.name)
        for lvl in education_levels
    }
    
    return render(request, 'curriculum/teacher_register.html', {
        'regions': regions,
        'subjects': subjects,
        'education_levels': education_levels,
        'classes_by_level_json': _json.dumps(classes_by_level),
        'subjects_by_level_json': _json.dumps(subjects_by_level),
    })


def ajax_get_districts(request):
    """AJAX: Get districts for a given region."""
    region_id = request.GET.get('region_id')
    if not region_id:
        return JsonResponse([], safe=False)
    districts = District.objects.filter(region_id=region_id).order_by('name')
    return JsonResponse([{'id': d.id, 'name': d.name} for d in districts], safe=False)


def ajax_get_schools(request):
    """AJAX: Get schools for a given district."""
    district_id = request.GET.get('district_id')
    if not district_id:
        return JsonResponse([], safe=False)
    schools = School.objects.filter(district_id=district_id).order_by('name')
    return JsonResponse([{'id': s.id, 'name': s.name, 'level': s.level} for s in schools], safe=False)


def ajax_save_teacher(request):
    """AJAX: Save teacher registration."""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required'}, status=400)
    
    data = json.loads(request.body)
    
    full_name = data.get('full_name', '').strip()
    phone_number = data.get('phone_number', '').strip()
    region_id = data.get('region_id')
    district_id = data.get('district_id')
    school_id = data.get('school_id')
    class_name = data.get('class_name', '').strip()
    stream = data.get('stream', '').strip()
    subject_id = data.get('subject_id')
    total_boys = data.get('total_boys', 0)
    total_girls = data.get('total_girls', 0)
    
    if not all([full_name, phone_number, region_id, district_id, school_id, subject_id]):
        return JsonResponse({'success': False, 'error': 'Tafadhali jaza sehemu zote.'}, status=400)
    
    # ENFORCE unique phone number: reject if already registered to another user
    existing = TLMTeacher.objects.filter(phone_number=phone_number).first()
    if existing:
        return JsonResponse({
            'success': False,
            'error': 'Namba hii ya simu tayari imesajiliwa kwa mwalimu mwingine. Tafadhali tumia sehemu ya "Tayari umesajiliwa?" juu ya ukurasa au wasiliana na msimamizi.',
            'phone_exists': True,
        }, status=409)
    
    # Wrap in try-except for security: prevent duplicate phone numbers (race condition)
    try:
        teacher = TLMTeacher.objects.create(
            full_name=full_name,
            phone_number=phone_number,
            region_id=region_id,
            district_id=district_id,
            school_id=school_id,
            class_name=class_name,
            stream=stream,
            subject_id=subject_id,
            total_boys=int(total_boys) if total_boys else 0,
            total_girls=int(total_girls) if total_girls else 0,
        )
        request.session['tlm_teacher_id'] = teacher.id
        return JsonResponse({'success': True, 'is_new': True})
    except IntegrityError:
        # Race condition: another request got there first
        return JsonResponse({
            'success': False,
            'error': 'Namba hii ya simu tayari imesajiliwa. Tafadhali tumia sehemu ya utafutaji juu ya ukurasa.',
            'phone_exists': True,
        }, status=409)


def ajax_lookup_teacher(request):
    """AJAX: Lookup returning teacher by phone number. Accepts ?remember=1 for 30-day session."""
    phone = request.GET.get('phone', '').strip()
    remember = request.GET.get('remember', '0') == '1'
    if not phone:
        return JsonResponse({'found': False})
    
    teacher = TLMTeacher.objects.filter(phone_number=phone).first()
    if teacher:
        request.session['tlm_teacher_id'] = teacher.id
        # Set session to expire in 30 days if remember me is checked
        if remember:
            request.session.set_expiry(60 * 60 * 24 * 30)  # 30 days
        else:
            request.session.set_expiry(0)  # Browser session
        return JsonResponse({
            'found': True,
            'full_name': teacher.full_name,
            'school_name': teacher.school.name if teacher.school else '',
        })
    return JsonResponse({'found': False})


# =============================================================================
# DASHBOARD (login required)
# =============================================================================

@login_required
def dashboard(request):
    """Standalone curriculum dashboard — overview of schemes, lesson plans & logbooks."""
    try:
        student = StudentTeacher.objects.select_related('selected_school').get(user=request.user)
    except StudentTeacher.DoesNotExist:
        messages.error(request, "Tafadhali jaza wasifu wako kwanza.")
        return redirect(reverse('curriculum:dashboard'))

    school = student.selected_school
    today = timezone.now().date()

    # Stats
    schemes_count = SchemeOfWork.objects.filter(student=student).count()
    lesson_plans_count = LessonPlan.objects.filter(student=student).count()
    logbook_today = LogbookEntry.objects.filter(student=student, date=today).first()
    logbook_count = LogbookEntry.objects.filter(student=student).count()
    this_month_logbooks = LogbookEntry.objects.filter(
        student=student, date__year=today.year, date__month=today.month
    ).count()
    # Lesson Notes — scoped to the student's school (LessonNote links to TLM teachers at a school)
    notes_count = LessonNote.objects.filter(school=school).count() if school else 0

    # Recent items
    recent_schemes = SchemeOfWork.objects.filter(student=student).select_related('subject').order_by('-updated_at')[:3]
    recent_lessons = LessonPlan.objects.filter(student=student).select_related('subject').order_by('-created_at')[:3]
    recent_logbooks = LogbookEntry.objects.filter(student=student).select_related('subject_taught').order_by('-date')[:5]

    return render(request, 'curriculum/dashboard.html', {
        'student': student,
        'school': school,
        'schemes_count': schemes_count,
        'lesson_plans_count': lesson_plans_count,
        'logbook_today': logbook_today,
        'logbook_count': logbook_count,
        'this_month_logbooks': this_month_logbooks,
        'notes_count': notes_count,
        'recent_schemes': recent_schemes,
        'recent_lessons': recent_lessons,
        'recent_logbooks': recent_logbooks,
        'today': today,
    })


# =============================================================================
# AZIMIO LA KAZI (KISWAHILI) — Column name mappings & helpers
# =============================================================================

# English → Swahili column names for Scheme of Work / Azimio la Kazi
SWAHILI_SCHEME_KEYS = {
    'Main Competence': 'UMAHIRI MKUU',
    'Specific Competences': 'UMAHIRI MAHUSUSI',
    'Main Learning Activities': 'SHUGHULI KUU ZA UJIFUNZAJI',
    'Specific Learning Activities': 'SHUGHULI NDOGO ZA UJIFUNZAJI',
    'Month': 'MWEZI',
    'Week': 'WIKI',
    'Number of Periods': 'VIPINDI',
    'Teaching and Learning Methods': 'MBINU ZA UJIFUNZAJI NA UFUNDISHAJI',
    'Teaching and Learning Resources': 'ZANA ZA UJIFUNZAJI NA UFUNDISHAJI',
    'Assessment Tools': 'ZANA ZA UPIMAJI',
    'References': 'REJEA',
    'Remarks': 'MAONI',
}

# Reverse map: Swahili → English
ENGLISH_SCHEME_KEYS = {v: k for k, v in SWAHILI_SCHEME_KEYS.items()}

# English → Swahili month names
SWAHILI_MONTHS = {
    'JANUARY': 'JANUARI', 'FEBRUARY': 'FEBRUARI', 'MARCH': 'MACHI',
    'APRIL': 'APRILI', 'MAY': 'MEI', 'JUNE': 'JUNI',
    'JULY': 'JULAI', 'AUGUST': 'AGOSTI', 'SEPTEMBER': 'SEPTEMBA',
    'OCTOBER': 'OKTOBA', 'NOVEMBER': 'NOVEMBA', 'DECEMBER': 'DESEMBA',
}

# =============================================================================
# SWAHILI PDF LABELS — English → Kiswahili translations for ALL PDF headers
# =============================================================================

# Comprehensive translation dictionary for ALL PDF labels
SWAHILI_ALL_LABELS = {
    # Ministry headers
    "PRIME MINISTER'S OFFICE": "OFISI YA WAZIRI MKUU",
    "REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT": "TAWALA ZA MIKOA NA SERIKALI ZA MITAA",
    "TEACHER'S LESSON PLAN": "MPANGO WA SOMO LA MWALIMU",
    "SCHEME OF WORK": "AZIMIO LA KAZI",
    # Info table labels (SOW & LP)
    "Teacher's Name": 'Jina la Mwalimu',
    'School Name': 'Jina la Shule',
    'Subject': 'Somo',
    'Class': 'Darasa',
    'Term': 'Muhula',
    'Year': 'Mwaka',
    'Total Weeks': 'Jumla ya Wiki',
    'Syllabus': 'Silabasi',
    'Duration': 'Muda',
    'Topic': 'Mada',
    'Subtopic': 'Mada Ndogo',
    'Date': 'Tarehe',
    'Students': 'Wanafunzi',
    'Teacher': 'Mwalimu',
    'Term/Year': 'Muhula/Mwaka',
    # Lesson Development
    'Main Competence': 'Umahiri Mkuu',
    'Specific Competence': 'Umahiri Mahususi',
    'Main Activity': 'Shughuli Kuu',
    'Specific Activity': 'Shughuli Mahususi',
    'Teaching & Learning Resources': 'Zana za Ufundishaji na Ujifunzaji',
    'References': 'Marejeo',
    'Lesson Development (IDDR Model)': 'Ukuzaji wa Somo (Muundo wa IDDR)',
    # IDDR Table headers
    'Stage': 'Hatua',
    'Time': 'Muda',
    'Teaching Activities': 'Shughuli za Ufundishaji',
    'Learning Activities': 'Shughuli za Ujifunzaji',
    'Assessment Criteria': 'Vigezo vya Upimaji',
    'Remarks': 'Maoni',
    # Student statistics
    'Registered': 'Waliojiandikisha',
    'Present': 'Waliohudhuria',
    'Absent': 'Wasiohudhuria',
    'Boys': 'Wavulana',
    'Girls': 'Wasichana',
    'Total': 'Jumla',
    # Remarks sub-labels
    '1. Strength': '1. Nguvu',
    '2. Weakness': '2. Udhaifu',
    '3. Way Forward': '3. Hatua Zifuatazo',
}


def _tl(label, mode):
    """Get translated label. If mode is Swahili (truthy), return Kiswahili translation."""
    if mode:
        return SWAHILI_ALL_LABELS.get(label, label)
    return label


def _has_swahili_keys(rows):
    """Check if scheme data uses Swahili column keys."""
    if not rows:
        return False
    first = rows[0] if isinstance(rows, list) else rows
    if not isinstance(first, dict):
        return False
    sw_keys = set(SWAHILI_SCHEME_KEYS.values())
    return bool(sw_keys & set(first.keys()))


def _normalize_scheme_keys(rows):
    """Convert Swahili keys to English keys for internal processing."""
    if not rows or not _has_swahili_keys(rows):
        return rows
    result = []
    for row in rows:
        new_row = {}
        for k, v in row.items():
            en_k = ENGLISH_SCHEME_KEYS.get(k, k)
            new_row[en_k] = v
        result.append(new_row)
    return result


def _to_swahili_keys(rows):
    """Convert English keys to Swahili keys for Azimio la Kazi output."""
    if not rows:
        return rows
    result = []
    for row in rows:
        new_row = {}
        for k, v in row.items():
            sw_k = SWAHILI_SCHEME_KEYS.get(k, k)
            sw_v = v
            # Convert month values to Swahili too
            if k in ('Month', 'MWEZI') and v:
                mth_upper = v.strip().upper()
                sw_v = SWAHILI_MONTHS.get(mth_upper, v)
            # Convert week format: keep as numbers for Kiswahili
            elif k in ('Week', 'WIKI') and v:
                week_str = str(v).strip()
                # Remove 'st', 'nd', 'rd', 'th' suffixes for Kiswahili format
                week_str = re.sub(r'(\d+)(st|nd|rd|th)', r'\1', week_str)
                sw_v = week_str
            else:
                sw_v = v
            new_row[sw_k] = sw_v
        result.append(new_row)
    return result


def _is_kiswahili_mode(language_param, subject, school_level):
    """
    Return True if the output should be in Kiswahili (Azimio la Kazi format).
    # NOTE: Keep sync'ed with _get_language_instruction() — same detection logic.
    
    Rules:
    - If user explicitly chose 'english' → False (English)
    - If user explicitly chose 'kiswahili' → True (Kiswahili)
    - Otherwise (auto) → auto-detect based on subject & school level
    """
    if language_param == 'kiswahili':
        return True
    if language_param == 'english':
        return False
    subj_lower = subject.lower()
    lvl_lower = (school_level or '').lower()
    if 'primary' in lvl_lower:
        if subj_lower in ('english', 'english language'):
            return False
        return True
    if 'secondary' in lvl_lower or 'ordinary' in lvl_lower or 'advanced' in lvl_lower:
        if subj_lower in ('kiswahili', 'swahili'):
            return True
        return False
    return False


# =============================================================================
# HELPER: Expand compact AI rows into a bigger document
# =============================================================================

def _expand_scheme_rows(rows, target_multiplier=3):
    """
    Take compact AI rows and expand them by splitting multi-week rows
    into individual week rows. This creates a bigger document without
    requiring the AI to generate more tokens.
    
    Example: "Week: 1st-4th" becomes 4 rows: "1st", "2nd", "3rd", "4th"
    Each expanded row keeps all other fields from the original.
    Handles both English and Swahili (Azimio la Kazi) column keys.
    """
    import re as _re
    
    # Normalize Swahili keys to English for internal processing
    was_swahili = _has_swahili_keys(rows)
    if was_swahili:
        rows = _normalize_scheme_keys(rows)
        # Swahili months -> English months for month_order sorting
        reverse_sw_months = {v: k for k, v in SWAHILI_MONTHS.items()}
        for row in rows:
            mth = (row.get('Month') or '').strip().upper()
            if mth in reverse_sw_months:
                row['Month'] = reverse_sw_months[mth]
    
    expanded = []
    
    for row in rows:
        week = (row.get('Week') or '').strip()
        
        # Try to split week ranges like "1st-4th", "1st & 2nd", "1st - 4th"
        weeks = _parse_week_range(week) if week else None
        
        if weeks and len(weeks) > 1:
            # Split into individual weeks, adjust periods proportionally
            total_periods = _parse_periods(row.get('Number of Periods', ''))
            periods_per = max(1, total_periods // len(weeks)) if total_periods > 0 else 0
            
            for w in weeks:
                new_row = dict(row)
                new_row['Week'] = w
                if periods_per > 0:
                    new_row['Number of Periods'] = str(periods_per)
                expanded.append(new_row)
        else:
            expanded.append(row)
    
    # If still too small, duplicate with varied weeks
    if len(expanded) < 30 and target_multiplier > 1:
        extra_rows = []
        for i, row in enumerate(expanded):
            if len(expanded) + len(extra_rows) >= 30 * target_multiplier:
                break
            if i % 2 == 0:  # Duplicate every other row with "2nd & 3rd" style
                new_row = dict(row)
                w = row.get('Week', '')
                if w and (w in ('1st', '2nd', '3rd', '4th') or w.isdigit()):
                    # For Kiswahili week format (just numbers), use "1&2" style
                    if w.isdigit():
                        next_w = str(int(w) + 1)
                        new_row['Week'] = f"{w}&{next_w}"
                    else:
                        next_w = {'1st': '2nd', '2nd': '3rd', '3rd': '4th', '4th': '5th'}.get(w, w)
                        new_row['Week'] = f"{w} & {next_w}"
                    extra_rows.append(new_row)
        expanded.extend(extra_rows)
    
    # Sort by month order for proper presentation
    month_order = {
        'JANUARY': 1, 'FEBRUARY': 2, 'MARCH': 3, 'APRIL': 4,
        'MAY': 5, 'JUNE': 6, 'JULY': 7, 'AUGUST': 8,
        'SEPTEMBER': 9, 'OCTOBER': 10, 'NOVEMBER': 11, 'DECEMBER': 12
    }
    expanded.sort(key=lambda r: (
        month_order.get((r.get('Month') or '').strip().upper(), 99),
        r.get('Week', '')
    ))
    
    return expanded


def _parse_week_range(week_str):
    """Parse week string like '1st-4th', '1st & 2nd', '1st,2nd,3rd' into list.
    Handles ranges (1st-4th → all 4 weeks) and comma/ampersand lists."""
    import re as _re
    original = week_str.strip()
    # Detect range: Nst-Mth or Nst - Mth
    range_m = _re.match(r'(\d+)(?:st|nd|rd|th)\s*[-–]\s*(\d+)(?:st|nd|rd|th)', original, _re.IGNORECASE)
    if range_m:
        start = int(range_m.group(1))
        end = int(range_m.group(2))
        if 1 <= start <= end <= 52:
            result = []
            for num in range(start, end + 1):
                suffix = 'st' if num == 1 else 'nd' if num == 2 else 'rd' if num == 3 else 'th'
                result.append(f"{num}{suffix}")
            return result
    # Fallback: comma/ampersand separated list
    cleaned = original.lower()
    cleaned = cleaned.replace('&', ',').replace('-', ',').replace('to', ',')
    cleaned = _re.sub(r'\s+', '', cleaned)
    parts = cleaned.split(',')
    result = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        m = _re.match(r'(\d+)(?:st|nd|rd|th)?', p)
        if m:
            num = int(m.group(1))
            if 1 <= num <= 52:
                suffix = 'st' if num == 1 else 'nd' if num == 2 else 'rd' if num == 3 else 'th'
                result.append(f"{num}{suffix}")
    return result if len(result) > 1 else None


def _parse_periods(period_str):
    """Parse period string like '6' or '3-6' into total number."""
    import re as _re
    if not period_str:
        return 0
    m = _re.match(r'(\d+)', str(period_str).strip())
    if m:
        return int(m.group(1))
    return 0


# =============================================================================
# LANGUAGE HELPER — Used by both Scheme & Lesson Plan generation
# =============================================================================

def _get_language_instruction(language_param, subject, school_level):
    """
    STRONG language instruction for Scheme of Work generation.
    Lists EVERY field explicitly so the AI follows it 100%.
    """
    subject_lower = subject.lower()
    school_level_lower = (school_level or '').lower()

    # ── ENGLISH: force everything to English ──
    if language_param == 'english':
        return (
("\n"
            "============================================\n"
            "⚠️  LANGUAGE RULE — ALL FIELDS IN ENGLISH ⚠️\n"
            "============================================\n"
            "EVERY field value below MUST be in ENGLISH ONLY:\n"
            "  - Main Competence → ENGLISH | Specific Competences → ENGLISH\n"
            "  - Main Learning Activities → ENGLISH | Specific Learning Activities → ENGLISH\n"
            "  - Month → ENGLISH (JANUARY, FEBRUARY etc)\n"
            "  - Week → ENGLISH (1st, 2nd, 3rd...) | Periods → numbers\n"
            "  - Methods, Resources, Assessment, References, Remarks → ALL ENGLISH\n"
            "FAIL: If any value is Swahili, the output is WRONG.\n"
            "============================================\n")
        )

    # ── KISWAHILI: AZIMIO LA KAZI format - ULTRA-STRONG ----
    elif language_param == 'kiswahili':
        return (
"\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"⚠️  KANUNI KAKAWA YA LUGHA - LAZIMA UFUATE KAMILI KAMILI ⚠️\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Unazalisha AZIMIO LA KAZI kwa KISWAHILI - KILA KITU (column names na values) LAZIMA kiwe KISWAHILI TU:\n" +
"\n" +
"COLUMN NAMES za JSON LAZIMA ziwe hizi kwa KISWAHILI (sio Kiingereza):\n" +
"  UMAHIRI MKUU      (sio 'Main Competence')\n" +
"  UMAHIRI MAHUSUSI   (sio 'Specific Competences')\n" +
"  SHUGHULI KUU ZA UJIFUNZAJI  (sio 'Main Learning Activities')\n" +
"  SHUGHULI NDOGO ZA UJIFUNZAJI  (sio 'Specific Learning Activities')\n" +
"  MWEZI              (sio 'Month' - tumia JANUARI, FEBRUARI, MACHI, APRILI, MEI, JUNI, JULAI, AGOSTI, SEPTEMBA, OKTOBA, NOVEMBA, DESEMBA)\n" +
"  WIKI               (sio 'Week' - tumia namba tu: 1, 2, 3, 4... au 1&2, 2&3)\n" +
"  VIPINDI            (sio 'Number of Periods' - tumia namba tu)\n" +
"  MBINU ZA UJIFUNZAJI NA UFUNDISHAJI  (sio 'Teaching and Learning Methods')\n" +
"  ZANA ZA UJIFUNZAJI NA UFUNDISHAJI  (sio 'Teaching and Learning Resources')\n" +
"  ZANA ZA UPIMAJI    (sio 'Assessment Tools')\n" +
"  REJEA              (sio 'References')\n" +
"  MAONI              (sio 'Remarks')\n" +
"\n" +
"VALUES zote LAZIMA ziwe KISWAHILI TU:\n" +
"  UMAHIRI MKUU -> Mfano: Kumudu misingi ya awali ya kihisabati\n" +
"  UMAHIRI MAHUSUSI -> Mfano: Kuonesha uelewa wa dhana ya namba\n" +
"  SHUGHULI KUU -> Mfano: Kutambua namba za kirumi (hadi M)\n" +
"  SHUGHULI NDOGO -> Mfano: Mwanafunzi kutambua namba za kirumi (hadi M)\n" +
"  MWEZI -> JANUARI, FEBRUARI, MACHI, APRILI, MEI, JUNI, JULAI, AGOSTI, SEPTEMBA, OKTOBA, NOVEMBA, DESEMBA\n" +
"  WIKI -> namba tu: 1, 2, 3, 4, 1&2, 3&4\n" +
"  VIPINDI -> namba tu: 6, 4, 8\n" +
"  MBINU -> KISWAHILI: Nyimbo, Michezo, Maswali na majibu, Onesho mbinu\n" +
"  ZANA -> KISWAHILI: Kadi, Picha, Chati\n" +
"  ZANA ZA UPIMAJI -> KISWAHILI: Orodha hakiki, Dodoso, Hojaji\n" +
"  REJEA -> KISWAHILI: TIE (2023), Somo, darasa...\n" +
"  MAONI -> KISWAHILI\n" +
"\n" +
"KOSA KUBWA: Kama COLUMN NAME au VALUE yoyote iko kwa Kiingereza -> MAKOSA MAKUBWA!\n" +
"JSON NAMES na VALUES zote LAZIMA ziwe KISWAHILI KAMILI.\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
        )

    # ── AUTO-DETECT ──
    if 'primary' in school_level_lower:
        if subject_lower in ('english', 'english language'):
            return "LANGUAGE: Write ALL 12 fields in ENGLISH — this is an English subject for Primary school."
        else:
            return (
"\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"⚠️  KANUNI KAKAWA - AZIMIO LA KAZI KISWAHILI KAMILI ⚠️\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Hii ni shule ya MSINGI (Primary) - OUTPUT LAZIMA iwe AZIMIO LA KAZI kwa KISWAHILI KAMILI:\n" +
"\n" +
"COLUMN NAMES LAZIMA ziwe KISWAHILI (sio Kiingereza):\n" +
"  UMAHIRI MKUU | UMAHIRI MAHUSUSI |\n" +
"  SHUGHULI KUU ZA UJIFUNZAJI | SHUGHULI NDOGO ZA UJIFUNZAJI |\n" +
"  MWEZI | WIKI | VIPINDI |\n" +
"  MBINU ZA UJIFUNZAJI NA UFUNDISHAJI | ZANA ZA UJIFUNZAJI NA UFUNDISHAJI |\n" +
"  ZANA ZA UPIMAJI | REJEA | MAONI\n" +
"\n" +
"VALUES zote LAZIMA ziwe KISWAHILI TU:\n" +
"  MWEZI -> JANUARI, FEBRUARI, MACHI, APRILI, MEI, JUNI, JULAI, AGOSTI, SEPTEMBA, OKTOBA, NOVEMBA, DESEMBA\n" +
"  WIKI -> namba tu: 1, 2, 3, 4, 1&2 (sio 1st, 2nd, 3rd)\n" +
"  VIPINDI -> namba tu: 6, 4, 8\n" +
"  MBINU -> KISWAHILI: Nyimbo, Michezo, Maswali na majibu\n" +
"  ZANA -> KISWAHILI: Kadi, Picha, Chati\n" +
"KOSA: Kama COLUMN NAME au VALUE yoyote iko Kiingereza -> MAKOSA!\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
            )
    elif 'secondary' in school_level_lower or 'ordinary' in school_level_lower or 'advanced' in school_level_lower:
        if subject_lower in ('kiswahili', 'swahili'):
            return (
"\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"⚠️  KANUNI KAKAWA - AZIMIO LA KAZI KISWAHILI KAMILI ⚠️\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Hili ni somo la Kiswahili -> COLUMN NAMES na VALUES zote LAZIMA ziwe KISWAHILI KAMILI:\n" +
"\n" +
"COLUMN NAMES (sio Kiingereza):\n" +
"  UMAHIRI MKUU | UMAHIRI MAHUSUSI | SHUGHULI KUU ZA UJIFUNZAJI |\n" +
"  SHUGHULI NDOGO ZA UJIFUNZAJI | MWEZI | WIKI | VIPINDI |\n" +
"  MBINU ZA UJIFUNZAJI NA UFUNDISHAJI | ZANA ZA UJIFUNZAJI NA UFUNDISHAJI |\n" +
"  ZANA ZA UPIMAJI | REJEA | MAONI\n" +
"\n" +
"VALUES zote LAZIMA ziwe KISWAHILI TU - hakuna Kiingereza kinachokubalika.\n" +
"MWEZI -> JANUARI, FEBRUARI, MACHI, APRILI, MEI, JUNI, JULAI, AGOSTI, SEPTEMBA, OKTOBA, NOVEMBA, DESEMBA\n" +
"WIKI -> namba tu (1, 2, 3 au 1&2)\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
            )
        else:
            return "LANGUAGE: Write ALL 12 fields in ENGLISH — this is a Secondary school subject taught in English."
    return ""


def _get_lp_language_instruction(language_param, subject, school_level):
    """
    ULTRA-STRONG language instruction for Lesson Plan generation.
    Lists EVERY JSON field explicitly so the AI follows it 100%.
    """
    subject_lower = subject.lower()
    school_level_lower = (school_level or '').lower()

    # ── ENGLISH: force EVERY field to English ──
    if language_param == 'english':
        return (
("\n"
            "===============================================================\n"
            "🚨 LANGUAGE RULE — ALL JSON FIELDS IN ENGLISH 🚨\n"
            "===============================================================\n"
            "ALL fields in the JSON output below → ENGLISH ONLY:\n"
            "  main_competence, specific_competence, main_activity, specific_activity → ENGLISH\n"
            "  teaching_resources, references → ENGLISH\n"
            "  lesson_development[].stage → ENGLISH (Introduction/Competence Development/Design/Realisation)\n"
            "  lesson_development[].teaching_activities → ENGLISH\n"
            "  lesson_development[].learning_activities → ENGLISH\n"
            "  lesson_development[].assessment_criteria → ENGLISH\n"
            "  remarks → ENGLISH paragraph\n"
            "FAIL: If any field value is Swahili → WRONG output.\n"
            "===============================================================\n")
        )

    # ── KISWAHILI: ULTRA-STRONG - EVERY JSON field in Kiswahili ──
    elif language_param == 'kiswahili':
        return (
"\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"🚨 KANUNI KAKAWA YA LUGHA - LESSON PLAN JSON FIELDS ZOTE KWA KISWAHILI 🚨\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Unazalisha LESSON PLAN kwa KISWAHILI - KILA field VALUE LAZIMA iwe KISWAHILI TU:\n" +
"\n" +
"JSON field names (keys) zinaweza kubaki English, lakini VALUES zote LAZIMA ziwe KISWAHILI:\n" +
"\n" +
"  main_competence -> KISWAHILI (mf. 1.0 Mada Kuu - Maelezo ya kompetensia)\n" +
"  specific_competence -> KISWAHILI (mf. Kufikia mwisho wa mada, mwanafunzi aweze...)\n" +
"  main_activity -> KISWAHILI (mf. Kwa muda wa kipindi kimoja mwanafunzi aweze...)\n" +
"  specific_activity -> KISWAHILI (mf. Kufikia mwisho wa somo, mwanafunzi aweze...)\n" +
"  teaching_resources -> KISWAHILI (mf. Vitabu vya TIE, Chati, Kadi)\n" +
"  references -> KISWAHILI (mf. TET (2023). Hisabati darasa la nne. Tanzania: TET.)\n" +
"\n" +
"  lesson_development array -> STAGES ZOTE kwa KISWAHILI:\n" +
"    stage names: Utangulizi | Ukuzaji Kompetensia | Muundo | Utekelezaji\n" +
"      (sio 'Introduction', 'Competence Development', 'Design', 'Realisation')\n" +
"    teaching_activities -> KISWAHILI (shughuli za mwalimu)\n" +
"    learning_activities -> KISWAHILI (shughuli za wanafunzi)\n" +
"    assessment_criteria -> KISWAHILI (vigezo vya tathmini)\n" +
"\n" +
"  student_statistics -> VALUES KISWAHILI\n" +
"  remarks -> KISWAHILI (tafakari na mapendekezo kwa somo lijalo)\n" +
"\n" +
"KOSA KUBWA: Field VALUE yoyote iko Kiingereza -> MAKOSA MAKUBWA!\n" +
"JSON keys zinaweza kubaki English, lakini VALUES zote LAZIMA ziwe KISWAHILI KAMILI.\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
        )

    # ── AUTO-DETECT ──
    if 'primary' in school_level_lower:
        if subject_lower in ('english', 'english language'):
            return "LANGUAGE: Write ALL fields in ENGLISH — this is an English subject for Primary school."
        else:
            return (
"\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"🚨 KANUNI KAKAWA YA LUGHA - LESSON PLAN KWA KISWAHILI KAMILI 🚨\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Hii ni shule ya MSINGI (Primary) na somo si English -> LESSON PLAN kwa KISWAHILI KAMILI:\n" +
"\n" +
"JSON field VALUES zote LAZIMA ziwe KISWAHILI TU:\n" +
"  main_competence, specific_competence -> KISWAHILI\n" +
"  main_activity, specific_activity -> KISWAHILI\n" +
"  teaching_resources, references -> KISWAHILI\n" +
"  lesson_development[].stage -> Utangulizi | Ukuzaji Kompetensia | Muundo | Utekelezaji\n" +
"    teaching_activities -> KISWAHILI\n" +
"    learning_activities -> KISWAHILI\n" +
"    assessment_criteria -> KISWAHILI\n" +
"  remarks -> KISWAHILI\n" +
"\n" +
"KOSA: Field value yoyote iko Kiingereza -> MAKOSA!\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
            )
    elif 'secondary' in school_level_lower or 'ordinary' in school_level_lower or 'advanced' in school_level_lower:
        if subject_lower in ('kiswahili', 'swahili'):
            return (
                "\n"
                "\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"🚨 KANUNI KAKAWA YA LUGHA - LESSON PLAN KWA KISWAHILI KAMILI 🚨\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
"Hili ni somo la Kiswahili kwa shule ya SECONDARY -> LESSON PLAN values zote KISWAHILI:\n" +
"\n" +
"JSON field VALUES zote LAZIMA ziwe KISWAHILI TU (keys zinaweza kubaki English):\n" +
"  main_competence, specific_competence -> KISWAHILI\n" +
"  lesson_development[].stage -> Utangulizi | Ukuzaji Kompetensia | Muundo | Utekelezaji\n" +
"  teaching_activities, learning_activities, assessment_criteria -> KISWAHILI\n" +
"  remarks -> KISWAHILI\n" +
"\n" +
"KOSA: Value yoyote iko Kiingereza -> MAKOSA!\n" +
"!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n"
            )
        else:
            return "LANGUAGE: Write ALL fields in ENGLISH — this is a Secondary school subject taught in English."
    return ""


# =============================================================================
# SCHEME OF WORK
# =============================================================================

def generate_scheme_view(request):
    """Display scheme of work generator form — requires TLM registration."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        # Need to register first
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:generate_scheme')}")
    
    form = SchemeOfWorkForm()
    education_levels = EducationLevel.objects.all().order_by('order')

    import json as _json
    classes_by_level = {}
    for cl in ClassLevel.objects.select_related('education_level').order_by('education_level', 'order'):
        classes_by_level.setdefault(cl.education_level_id, []).append({'id': cl.id, 'name': cl.name})

    # Filter subjects by education level (Primary → primary, Advanced → advanced, Ordinary → secondary)
    def _subjects_for_level(level_name):
        name_lower = level_name.lower()
        if 'primary' in name_lower:
            return list(Subject.objects.filter(level='primary').order_by('name').values('id', 'name'))
        elif 'advanced' in name_lower:
            return list(Subject.objects.filter(level='advanced').order_by('name').values('id', 'name'))
        elif 'technical' in name_lower or 'veta' in name_lower:
            # Technical / VETA → technical subjects
            return list(Subject.objects.filter(level='technical').order_by('name').values('id', 'name'))
        else:
            # Ordinary Level → secondary subjects
            return list(Subject.objects.filter(level='secondary').order_by('name').values('id', 'name'))

    subjects_by_level = {
        lvl.id: _subjects_for_level(lvl.name)
        for lvl in education_levels
    }

    student = None
    school = None
    if request.user.is_authenticated:
        try:
            student = StudentTeacher.objects.select_related('selected_school').get(user=request.user)
            school = student.selected_school
            if not school:
                app = StudentApplication.objects.filter(
                    student=student, status='approved'
                ).select_related('school').first()
                if app:
                    school = app.school
        except StudentTeacher.DoesNotExist:
            pass

    # Get teacher info for auto-fill
    teacher_name = teacher.full_name if teacher else ''
    teacher_school_name = teacher.school.name if teacher and teacher.school else ''
    teacher_school_level = teacher.school.level if teacher and teacher.school else ''
    teacher_subject_id = teacher.subject.id if teacher and teacher.subject else ''
    teacher_subject_name = teacher.subject.name if teacher and teacher.subject else ''
    teacher_subject_level = teacher.subject.level if teacher and teacher.subject else ''
    teacher_class_name = teacher.class_name if teacher else ''
    teacher_stream = teacher.stream if teacher else ''
    teacher_total_boys = teacher.total_boys if teacher else 0
    teacher_total_girls = teacher.total_girls if teacher else 0

    return render(request, 'curriculum/generate_scheme.html', {
        'form': form,
        'education_levels': education_levels,
        'classes_by_level_json': _json.dumps(classes_by_level),
        'subjects_by_level_json': _json.dumps(subjects_by_level),
        'student': student,
        'school': school,
        'teacher': teacher,
        'teacher_name': teacher_name,
        'teacher_school_name': teacher_school_name,
        'teacher_school_level': teacher_school_level,
        'teacher_subject_id': teacher_subject_id,
        'teacher_subject_name': teacher_subject_name,
        'teacher_subject_level': teacher_subject_level,
        'teacher_class_name': teacher_class_name,
        'teacher_stream': teacher_stream,
        'teacher_total_boys': teacher_total_boys,
        'teacher_total_girls': teacher_total_girls,
    })


@csrf_exempt
def ajax_generate_scheme(request):
    """
    Generate a Scheme of Work using AI with PARALLEL AI calls.
    All month groups are generated simultaneously via ThreadPoolExecutor,
    cutting total wait time from ~2-3 minutes to ~40-60 seconds.
    Returns the complete data synchronously (no task_id polling needed).
    """
    if client is None:
        return JsonResponse({
            'success': False,
            'error': 'Huduma ya AI haitumiki. Ufunguo wa API haujawekwa.'
        }, status=503)

    if request.method != 'POST':
        return JsonResponse({'success': False}, status=400)

    try:
        data = json.loads(request.body)

        education_level = data.get('education_level', '')
        class_name = data.get('class_name', '')
        stream = data.get('stream', '')
        subject = data.get('subject', '')
        term = data.get('term', 'I')
        year = data.get('year', '2026')
        syllabus = data.get('syllabus', 'New Syllabus')
        total_weeks = int(data.get('total_weeks', 12))
        periods_per_week = int(data.get('periods_per_week', 8))
        start_date = data.get('start_date', '')
        end_date = data.get('end_date', '')
        teacher_name = data.get('teacher_name', '')
        school_name = data.get('school_name', '')
        reference_source = data.get('reference_source', '')
        breaks = data.get('breaks', [])
        # ── Language: manual selection (english/kiswahili) or auto-detect ──
        _lang_tlm = get_tlm_teacher(request)
        _scheme_language = data.get('language', getattr(_lang_tlm, 'preferred_language', 'auto') if _lang_tlm else 'auto')
        _school_level = _lang_tlm.school.level if _lang_tlm and _lang_tlm.school else ''
        language_instruction = _get_language_instruction(_scheme_language, subject, _school_level)

        full_class_name = f"{class_name}{stream}" if stream else class_name
        ref_text = f"\nReference source: {reference_source}" if reference_source else ''

        # ── school_id and user_id for saving ──
        school_id = None
        user_id = None
        if request.user.is_authenticated:
            user_id = request.user.id
            try:
                student = StudentTeacher.objects.get(user=request.user)
                if student.selected_school:
                    school_id = student.selected_school.id
            except StudentTeacher.DoesNotExist:
                pass
        elif _lang_tlm:
            if _lang_tlm.school:
                school_id = _lang_tlm.school.id

        # ── Compute month groups ──
        import calendar as _cal
        term_groups_map = {
            'Full Year': [['JANUARY', 'FEBRUARY', 'MARCH'], ['APRIL', 'MAY', 'JUNE'], ['JULY', 'AUGUST', 'SEPTEMBER'], ['OCTOBER']],
            'I': [['JANUARY', 'FEBRUARY', 'MARCH'], ['APRIL', 'MAY', 'JUNE']],
            'II': [['JULY', 'AUGUST', 'SEPTEMBER'], ['OCTOBER']],
            'III': [['AUGUST', 'SEPTEMBER'], ['OCTOBER', 'NOVEMBER']],
        }
        month_groups = term_groups_map.get(term, [['JANUARY', 'FEBRUARY', 'MARCH']])
        total_groups = len(month_groups)

        # ── Breaks by month ──
        _all_month_names = []
        for grp in month_groups:
            _all_month_names.extend(grp)
        breaks_by_month = {m: [] for m in _all_month_names}
        if breaks:
            for b in breaks:
                start_str = b.get('start', '')
                if start_str:
                    try:
                        start_dt_brk = datetime.strptime(start_str, '%Y-%m-%d')
                        brk_month_name = _cal.month_name[start_dt_brk.month].upper()
                        if brk_month_name in _all_month_names:
                            breaks_by_month[brk_month_name].append(b)
                    except (ValueError, IndexError):
                        pass

        # ── Build a SINGLE prompt covering ALL months ──
        all_months_flat = []
        for grp in month_groups:
            all_months_flat.extend(grp)
        all_months_str = ', '.join(all_months_flat)
        # If Kiswahili mode, use Swahili month names
        if _is_kiswahili_mode(_scheme_language, subject, _school_level):
            all_months_str_sw = ', '.join(SWAHILI_MONTHS.get(m, m) for m in all_months_flat)
        else:
            all_months_str_sw = all_months_str
        rows_per_month = max(2, 10 // max(1, len(all_months_flat)))  # Compact AI output, expanded by Python

        scope_lines = [
            f"MONTHS TO COVER: {all_months_str_sw}",
            f"TOTAL WEEKS: {total_weeks}",
            f"START FROM the FIRST topics in the syllabus for {subject} {full_class_name}",
            f"Progress through the syllabus TOPIC BY TOPIC across all months.",
            f"Generate {rows_per_month}+ rows for EACH month listed above.",
            f"Do NOT skip any month.",
        ]

        # Add breaks
        for m in all_months_flat:
            month_breaks = breaks_by_month.get(m, [])
            for b in month_breaks:
                b_name = b.get('name', 'Break')
                b_start = b.get('start', '')
                b_end = b.get('end', '')
                scope_lines.append(
                    f"  INCLUDE a BREAK ROW for '{b_name.upper()} ({b_start} - {b_end})' in {m}."
                )

        scope_text = '\n'.join(scope_lines)

        prompt = f"""You are a Tanzanian TIE/SEQUIP curriculum expert. Generate a Scheme of Work in TAMISEMI format.

TEACHER: {teacher_name or '____________________'}
SCHOOL: {school_name or '____________________'}
SUBJECT: {subject}
CLASS: {full_class_name}
TERM: {term} | YEAR: {year}
SYLLABUS: {syllabus} | WEEKS: {total_weeks} | PERIODS/WEEK: {periods_per_week}{ref_text}
{language_instruction}

MONTHS: {all_months_str_sw}
Generate {rows_per_month}+ rows PER month covering the full syllabus from start to end.
{scope_text}

OUTPUT: JSON array only. Each object must use the EXACT column names specified by the LANGUAGE instruction above. If Kiswahili -> use KISWAHILI column names. If English -> use English column names.

RULES:
- Real TIE competences, numbered (1.0, 2.1 etc)
- (a)/(b)/(c) each on its OWN separate row (do NOT combine)
- ALL 12 field values MUST follow the LANGUAGE instruction above
- ALL text values → MUST be in the language specified above by LANGUAGE instruction
- ALL 12 column names MUST be in the LANGUAGE specified above
- ALL 12 field values MUST be in the LANGUAGE specified above
- NO arrays inside values, only strings
- Return ONLY the JSON array. No other text.

!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
⚠️  LANGUAGE CHECK: EVERY field above must be in the language specified by LANGUAGE instruction.
    If Kiswahili -> ALL column names AND values = KISWAHILI
    If English -> ALL column names AND values = ENGLISH
    If ANY field is in the wrong language -> WRONG OUTPUT!
!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!"""

        # ── Make a SINGLE AI call (no parallelism) ──
        logger.info(f"[Scheme] Starting single AI call for {len(all_months_flat)} months...")
        all_scheme_data, response_text = _generate_scheme_batch(prompt)

        if not all_scheme_data:
            raise RuntimeError("AI failed to generate scheme data")

        logger.info(f"[Scheme] AI returned {len(all_scheme_data)} rows")

        # ── Expand rows: AI generates compact, Python expands for big document ──
        expanded_count = len(all_scheme_data)
        all_scheme_data = _expand_scheme_rows(all_scheme_data)
        logger.info(f"[Scheme] Expanded from {expanded_count} to {len(all_scheme_data)} rows")

        # ── If Kiswahili mode, convert keys back to Swahili ──
        if _is_kiswahili_mode(_scheme_language, subject, _school_level):
            all_scheme_data = _to_swahili_keys(all_scheme_data)

        # ── Validate months ──
        expected_months = set(all_months_flat)
        present_months = set()
        for row in all_scheme_data:
            mth = (row.get('Month') or '').strip().upper()
            if mth:
                present_months.add(mth)
        missing_months = expected_months - present_months
        if missing_months:
            _is_sw = _is_kiswahili_mode(_scheme_language, subject, _school_level)
            for mm in sorted(missing_months):
                mth_sw = SWAHILI_MONTHS.get(mm, mm) if _is_sw else mm
                if _is_sw:
                    all_scheme_data.append({
                        "UMAHIRI MKUU": "Endelea na mada ya silabasi",
                        "UMAHIRI MAHUSUSI": "Endelea na mada ndogo",
                        "SHUGHULI KUU ZA UJIFUNZAJI": f"Shughuli za {mth_sw}",
                        "SHUGHULI NDOGO ZA UJIFUNZAJI": "Endelea na ujifunzaji",
                        "MWEZI": mth_sw,
                        "WIKI": "1-4",
                        "VIPINDI": str(max(3, periods_per_week // 2)),
                        "MBINU ZA UJIFUNZAJI NA UFUNDISHAJI": "Mbinu mbalimbali za CBC",
                        "ZANA ZA UJIFUNZAJI NA UFUNDISHAJI": "Vitabu vya TIE, Chati",
                        "ZANA ZA UPIMAJI": "Mazoezi, Maswali",
                        "REJEA": "Vitabu vya TIE",
                        "MAONI": "Endelea na silabasi"
                    })
                else:
                    all_scheme_data.append({
                        "Main Competence": "Continue with syllabus topics",
                        "Specific Competences": "Continue with subtopics",
                        "Main Learning Activities": f"Learning activities for {mm}",
                        "Specific Learning Activities": "Continue with learning",
                        "Month": mm,
                        "Week": "1st - 4th",
                        "Number of Periods": str(max(3, periods_per_week // 2)),
                        "Teaching and Learning Methods": "Various CBC methods",
                        "Teaching and Learning Resources": "TIE textbook, Charts",
                        "Assessment Tools": "Exercises, Questions",
                        "References": "TIE textbooks",
                        "Remarks": "Proceed with syllabus"
                    })

        # ── Save to DB ──
        saved_id = None
        try:
            from field_app.models import School, StudentTeacher, Subject, SchemeOfWork
            school = None
            student = None
            if school_id:
                try:
                    school = School.objects.get(id=school_id)
                except School.DoesNotExist:
                    pass
            if user_id:
                try:
                    student = StudentTeacher.objects.get(user_id=user_id)
                    if student.selected_school:
                        school = student.selected_school
                except StudentTeacher.DoesNotExist:
                    pass
            if school:
                level_map = {'primary school': 'primary', 'ordinary level': 'ordinary', 'advanced level': 'advanced'}
                edu_level = level_map.get((education_level or '').lower(), 'ordinary')
                subj_obj = Subject.objects.filter(name__iexact=subject).first()
                if subj_obj:
                    start_dt = None
                    end_dt_obj = None
                    if start_date:
                        try:
                            from datetime import date as _dt
                            start_dt = _dt.fromisoformat(start_date)
                        except Exception:
                            pass
                    if end_date:
                        try:
                            from datetime import date as _dt
                            end_dt_obj = _dt.fromisoformat(end_date)
                        except Exception:
                            pass
                    defaults = {
                        'school': school,
                        'education_level': edu_level,
                        'class_name': class_name,
                        'syllabus': syllabus,
                        'total_weeks': int(total_weeks),
                        'periods_per_week': int(periods_per_week),
                        'start_date': start_dt,
                        'end_date': end_dt_obj,
                        'teacher_name': teacher_name,
                        'reference_source': reference_source,
                        'breaks': breaks,
                        'scheme_data': all_scheme_data,
                        'generated_by_ai': True,
                    }
                    if student:
                        scheme_obj, _ = SchemeOfWork.objects.update_or_create(
                            student=student, subject=subj_obj, term=term, year=int(year),
                            defaults=defaults,
                        )
                    else:
                        scheme_obj = SchemeOfWork.objects.create(student=None, **defaults)
                    saved_id = scheme_obj.id
        except Exception as save_err:
            logger.warning(f"[Scheme] Save error (non-fatal): {save_err}")

        logger.info(f"[Scheme] DONE: {len(all_scheme_data)} rows from single AI call")

        return JsonResponse({
            'success': True,
            'data': all_scheme_data,
            'saved_id': saved_id,
            'groups': [','.join(grp) for grp in month_groups],
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        logger.error(f"[Scheme] Fatal: {type(e).__name__}: {str(e)[:300]}")
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


# =============================================================================
# AJAX: Get Scheme generation progress (legacy stub)
# =============================================================================

def ajax_get_scheme_progress(request):
    """Legacy stub - no longer needed since parallel calls return synchronously."""
    return JsonResponse({'status': 'done', 'current': 1, 'total': 1})


def ajax_get_scheme_result(request):
    """Legacy stub - no longer needed."""
    return JsonResponse({'success': False, 'error': 'No task'}, status=404)

def download_scheme_pdf(request):
    """Generate PDF ya Scheme of Work."""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)

    data = json.loads(request.body)
    scheme_data = data.get('scheme_data') or []
    subject = data.get('subject', '')
    class_name = data.get('class_name', '')
    term = data.get('term', '')
    year = data.get('year', '')
    syllabus = data.get('syllabus', '')
    teacher_name = data.get('teacher_name', '')
    school_name = data.get('school_name', '')
    total_weeks = data.get('total_weeks', '')

    # Normalize Swahili keys to English for PDF processing
    was_swahili = _has_swahili_keys(scheme_data)
    if was_swahili:
        scheme_data = _normalize_scheme_keys(scheme_data)

    # ── Get teacher theme for PDF ──
    teacher = get_tlm_teacher(request)
    theme_name = getattr(teacher, 'theme', 'classic') if teacher else 'classic'
    
    THEMES = {
        'classic': {
            'primary': colors.HexColor('#0A2B5E'),
            'accent': colors.HexColor('#C8900A'),
            'accent_dark': colors.HexColor('#A67B07'),
            'stripe': colors.HexColor('#EBF0FB'),
            'border': colors.HexColor('#9BAAC4'),
            'light_bg': colors.HexColor('#FAFBFD'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'tanzania': {
            'primary': colors.HexColor('#1E7A34'),
            'accent': colors.HexColor('#F5C518'),
            'accent_dark': colors.HexColor('#D4A000'),
            'stripe': colors.HexColor('#EBF8F0'),
            'border': colors.HexColor('#6BA87B'),
            'light_bg': colors.HexColor('#F5FBF7'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'ocean': {
            'primary': colors.HexColor('#1A365D'),
            'accent': colors.HexColor('#0EA5E9'),
            'accent_dark': colors.HexColor('#0284C7'),
            'stripe': colors.HexColor('#E8F4FD'),
            'border': colors.HexColor('#7FB8D9'),
            'light_bg': colors.HexColor('#F4F9FD'),
            'font_size': 8.0,
            'border_width': 2.5,
        },
        'royal': {
            'primary': colors.HexColor('#4C1D95'),
            'accent': colors.HexColor('#EC4899'),
            'accent_dark': colors.HexColor('#DB2775'),
            'stripe': colors.HexColor('#F3EEFB'),
            'border': colors.HexColor('#B794D4'),
            'light_bg': colors.HexColor('#FAF5FF'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'executive': {
            'primary': colors.HexColor('#1F2937'),
            'accent': colors.HexColor('#9CA3AF'),
            'accent_dark': colors.HexColor('#6B7280'),
            'stripe': colors.HexColor('#F0F1F3'),
            'border': colors.HexColor('#B0B5BD'),
            'light_bg': colors.HexColor('#F8F9FA'),
            'font_size': 8.5,
            'border_width': 2.0,
        },
        'sunset': {
            'primary': colors.HexColor('#C2410C'),
            'accent': colors.HexColor('#FB923C'),
            'accent_dark': colors.HexColor('#EA580C'),
            'stripe': colors.HexColor('#FFF7ED'),
            'border': colors.HexColor('#FDBA74'),
            'light_bg': colors.HexColor('#FFF7ED'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'forest': {
            'primary': colors.HexColor('#065F46'),
            'accent': colors.HexColor('#34D399'),
            'accent_dark': colors.HexColor('#059669'),
            'stripe': colors.HexColor('#ECFDF5'),
            'border': colors.HexColor('#6EE7B7'),
            'light_bg': colors.HexColor('#F0FDF4'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'midnight': {
            'primary': colors.HexColor('#1E1B4B'),
            'accent': colors.HexColor('#6366F1'),
            'accent_dark': colors.HexColor('#4F46E5'),
            'stripe': colors.HexColor('#EEF2FF'),
            'border': colors.HexColor('#A5B4FC'),
            'light_bg': colors.HexColor('#F8FAFC'),
            'font_size': 8.0,
            'border_width': 2.5,
        },
        'cherry': {
            'primary': colors.HexColor('#9B1C1C'),
            'accent': colors.HexColor('#F43F5E'),
            'accent_dark': colors.HexColor('#E11D48'),
            'stripe': colors.HexColor('#FFF1F2'),
            'border': colors.HexColor('#FDA4AF'),
            'light_bg': colors.HexColor('#FFF5F6'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'safari': {
            'primary': colors.HexColor('#78350F'),
            'accent': colors.HexColor('#D97706'),
            'accent_dark': colors.HexColor('#B45309'),
            'stripe': colors.HexColor('#FFFBEB'),
            'border': colors.HexColor('#FCD34D'),
            'light_bg': colors.HexColor('#FFFBEB'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'dawn': {
            'primary': colors.HexColor('#7E22CE'),
            'accent': colors.HexColor('#D946EF'),
            'accent_dark': colors.HexColor('#C026D3'),
            'stripe': colors.HexColor('#FAF5FF'),
            'border': colors.HexColor('#D8B4FE'),
            'light_bg': colors.HexColor('#FDF4FF'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
    }
    T = THEMES.get(theme_name, THEMES['classic'])
    NAVY = T['primary']
    GOLD = T['accent']
    DARK_GOLD = T['accent_dark']
    STRIPE = T['stripe']
    BORDER = T['border']
    _fs = T.get('font_size', 8.5)

    # ── Create Tanzania flag watermark ──
    _tz_watermark = None
    if PILImage:
        try:
            _flag_path = os.path.join(os.path.dirname(__file__), 'static', 'curriculum', 'tz_flag.png')
            if os.path.exists(_flag_path):
                _pil_img = PILImage.open(_flag_path).convert('RGBA')
                _r, _g, _b, _a = _pil_img.split()
                _new_a = _a.point(lambda x: int(x * 0.10))
                _faded = PILImage.merge('RGBA', (_r, _g, _b, _new_a))
                _buf = BytesIO()
                _faded.save(_buf, format='PNG')
                _buf.seek(0)
                _tz_watermark = _buf
        except Exception:
            pass

    # ── Cover page: professional border drawing ──
    def _scheme_cover(can, doc_obj):
        """Draw professional double-border frame with flag watermark."""
        can.saveState()
        pw = doc_obj.pagesize[0]
        ph = doc_obj.pagesize[1]
        # Outer gold border (thick)
        can.setStrokeColor(GOLD)
        can.setLineWidth(3.5)
        can.rect(12, 12, pw - 24, ph - 24)
        # Inner navy border (thin)
        can.setStrokeColor(NAVY)
        can.setLineWidth(1.5)
        can.rect(18, 18, pw - 36, ph - 36)
        # Top gold accent bar
        can.setStrokeColor(GOLD)
        can.setLineWidth(4)
        can.line(18, ph - 55, pw - 18, ph - 55)
        # Bottom gold accent bar
        can.line(18, 55, pw - 18, 55)
        # Top-left corner square
        can.setFillColor(GOLD)
        can.rect(18, ph - 26, 14, 8, fill=1, stroke=0)
        # Top-right corner square
        can.rect(pw - 32, ph - 26, 14, 8, fill=1, stroke=0)
        # Bottom-left corner square
        can.rect(18, 18, 14, 8, fill=1, stroke=0)
        # Bottom-right corner square
        can.rect(pw - 32, 18, 14, 8, fill=1, stroke=0)
        # ── Draw Tanzania flag watermark ──
        if _tz_watermark:
            try:
                _tz_watermark.seek(0)
                fw = pw * 0.50
                fh = fw * 2.0 / 3.0
                can.drawImage(_tz_watermark, pw/2 - fw/2, ph/2 - fh/2,
                              width=fw, height=fh, mask='auto')
            except Exception:
                pass
        can.restoreState()

    # ── Page header (appears from page 2 onwards) ──
    def _scheme_header(can, doc_obj):
        """Draw NAME OF SCHOOL / TEACHER header bar on pages 2+."""
        can.saveState()
        can.setFont('Helvetica-Bold', 7)
        can.setFillColor(NAVY)
        pw = doc_obj.pagesize[0]
        ph = doc_obj.pagesize[1]
        # Row 1: School | Subject | Class
        can.drawString(18, ph - 18, f"NAME OF SCHOOL: {school_name or '____________________'}")
        can.drawCentredString(pw / 2, ph - 18, f"SUBJECT: {subject}")
        can.drawRightString(pw - 18, ph - 18, f"CLASS: {class_name}")
        # Row 2: Teacher | Term | Year
        can.drawString(18, ph - 30, f"TEACHER'S NAME: {teacher_name or '____________________'}")
        can.drawCentredString(pw / 2, ph - 30, f"TERM: {term}")
        can.drawRightString(pw - 18, ph - 30, f"YEAR: {year}")
        # Gold underline
        can.setStrokeColor(GOLD)
        can.setLineWidth(0.8)
        can.line(18, ph - 34, pw - 18, ph - 34)
        can.restoreState()

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4),
                            rightMargin=18, leftMargin=18, topMargin=40, bottomMargin=25)
    elements = []

    # ── Cover page: decorative top lines ──
    elements.append(HRFlowable(width="100%", thickness=3, color=GOLD, spaceAfter=2))
    elements.append(HRFlowable(width="100%", thickness=1.5, color=NAVY, spaceAfter=12))

    cell_style = ParagraphStyle('SchCell', fontName='Helvetica', fontSize=_fs,
                                leading=_fs + 3.5, wordWrap='LTR')
    hdr_style = ParagraphStyle('SchHdr', fontName='Helvetica-Bold', fontSize=_fs,
                               leading=_fs + 3.5, textColor=colors.white, wordWrap='LTR', alignment=1)

    # ── TAMISEMI Header (centered, bold) — bilingual ──
    _pm = _tl("PRIME MINISTER'S OFFICE", was_swahili)
    _ralg = _tl("REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT", was_swahili)
    _sow_title = _tl("SCHEME OF WORK", was_swahili)
    
    # Visual style: alternate between styles for variety based on hash of subject
    _style_seed = sum(ord(c) for c in subject + (class_name or '')) % 5
    if _style_seed == 0:
        # Classic centered
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1', fontName='Helvetica-Bold', fontSize=13, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2', fontName='Helvetica-Bold', fontSize=10, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
        elements.append(HRFlowable(width="50%", thickness=1, color=GOLD, spaceAfter=8))
    elif _style_seed == 1:
        # Gold background bar style
        elements.append(Spacer(1, 2))
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1b', fontName='Helvetica-Bold', fontSize=14, alignment=1,
                           textColor=GOLD, spaceAfter=1)))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2b', fontName='Helvetica-Bold', fontSize=11, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
        elements.append(HRFlowable(width="100%", thickness=3, color=NAVY, spaceAfter=6))
    elif _style_seed == 2:
        # Modern compact
        elements.append(HRFlowable(width="100%", thickness=2, color=GOLD, spaceAfter=2))
        elements.append(Paragraph(f"{_pm}  |  {_ralg}",
            ParagraphStyle('MH1c', fontName='Helvetica-Bold', fontSize=9, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
        elements.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=6))
    elif _style_seed == 3:
        # Navy bg with gold text
        elements.append(Spacer(1, 4))
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1d', fontName='Helvetica-Bold', fontSize=13, alignment=1,
                           textColor=NAVY, spaceAfter=0, backColor=colors.HexColor('#F0F4FF'))))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2d', fontName='Helvetica-Bold', fontSize=10, alignment=1,
                           textColor=GOLD, spaceAfter=2)))
        elements.append(HRFlowable(width="40%", thickness=0.5, color=GOLD, spaceAfter=6))
    else:
        # Gold underline style
        elements.append(Spacer(1, 2))
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1e', fontName='Helvetica-Bold', fontSize=12, alignment=1,
                           textColor=NAVY, spaceAfter=1,
                           borderWidth=0, borderPadding=2)))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2e', fontName='Helvetica-Bold', fontSize=9, alignment=1,
                           textColor=colors.HexColor('#555555'), spaceAfter=4)))

    # ── Title ──
    elements.append(Paragraph(_sow_title,
        ParagraphStyle('ST', fontName='Helvetica-Bold', fontSize=18, alignment=1,
                       textColor=NAVY, spaceAfter=10)))

    # ── Cover info: styled table (bilingual labels) ──
    lbl = ParagraphStyle('lbl', fontName='Helvetica-Bold', fontSize=10, textColor=colors.white, leading=14)
    val = ParagraphStyle('val', fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#222222'), leading=14)
    info_rows = [
        [Paragraph(_tl("Teacher's Name", was_swahili), lbl), Paragraph(teacher_name, val)],
        [Paragraph(_tl('School Name', was_swahili), lbl), Paragraph(school_name or '____________________', val)],
        [Paragraph(_tl('Subject', was_swahili), lbl), Paragraph(subject, val)],
        [Paragraph(_tl('Class', was_swahili), lbl), Paragraph(class_name, val)],
        [Paragraph(_tl('Term', was_swahili), lbl), Paragraph(term, val)],
        [Paragraph(_tl('Year', was_swahili), lbl), Paragraph(str(year), val)],
        [Paragraph(_tl('Total Weeks', was_swahili), lbl), Paragraph(str(total_weeks), val)],
        [Paragraph(_tl('Syllabus', was_swahili), lbl), Paragraph(syllabus, val)],
    ]
    info_table = Table(info_rows, colWidths=[180, 280])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), NAVY),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#FAFBFD')),
        ('ROWBACKGROUNDS', (1, 0), (1, -1), [colors.HexColor('#FAFBFD'), STRIPE]),
        ('GRID', (0, 0), (-1, -1), 0.5, BORDER),
        ('BOX', (0, 0), (-1, -1), 2, GOLD),
        ('LINEBELOW', (0, 0), (-1, 0), 1.2, GOLD),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
    ]))
    elements.append(info_table)
    # ── PageBreak: Cover page ends here, data table starts on page 2 ──
    elements.append(PageBreak())

    if scheme_data:
        headers = list(scheme_data[0].keys())
        WIDTH_MAP = {
            'Main Competence': 72, 'Specific Competences': 72,
            'Main Learning Activities': 68, 'Specific Learning Activities': 76,
            'Month': 38, 'Week': 34, 'Number of Periods': 32,
            'Teaching and Learning Methods': 64,
            'Teaching and Learning Resources': 64,
            'Assessment Tools': 54, 'References': 64, 'Remarks': 68,
        }
        TOTAL = 806
        col_widths = []
        for h in headers:
            w = WIDTH_MAP.get(h)
            if w is None:
                for k, v in WIDTH_MAP.items():
                    if k.lower() in h.lower() or h.lower() in k.lower():
                        w = v
                        break
            col_widths.append(w if w else TOTAL // len(headers))
        scale = TOTAL / sum(col_widths)
        col_widths = [w * scale for w in col_widths]

        table_data = [[Paragraph(h, hdr_style) for h in headers]]
        for row in scheme_data:
            table_data.append([Paragraph(str(row.get(h, '') or ''), cell_style) for h in headers])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        ts = [
            ('BACKGROUND', (0, 0), (-1, 0), NAVY),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.35, BORDER),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('LINEBELOW', (0, 0), (-1, 0), 1.2, GOLD),
        ]
        for i in range(1, len(table_data)):
            bg = STRIPE if i % 2 == 0 else colors.white
            ts.append(('BACKGROUND', (0, i), (-1, i), bg))
        tbl.setStyle(TableStyle(ts))
        elements.append(tbl)

    doc.build(elements, onFirstPage=_scheme_cover, onLaterPages=_scheme_header)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Scheme_of_Work_{subject}_{class_name}.pdf"'
    return response


def download_scheme_word(request):
    """Export Scheme of Work as Word (.docx)."""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)

    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = json.loads(request.body)
    scheme_data = data.get('scheme_data', [])
    subject = data.get('subject', '')
    class_name = data.get('class_name', '')
    term = data.get('term', '')
    year = data.get('year', '')
    teacher_name = data.get('teacher_name', '')
    school_name = data.get('school_name', '')

    # Normalize Swahili keys to English for Word processing
    was_swahili = _has_swahili_keys(scheme_data)
    if was_swahili:
        scheme_data = _normalize_scheme_keys(scheme_data)

    doc = Document()
    # Kiswahili Azimio la Kazi title
    if was_swahili:
        doc.core_properties.title = f"AZIMIO LA KAZI — {subject}"
        title = doc.add_heading(f"AZIMIO LA SOMO LA {subject.upper()} — MUHULA WA {term} {year}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph(f"JINA LA MWALIMU: {teacher_name}    SHULE: {school_name}    DARASA: {class_name}")
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.core_properties.title = f"Scheme of Work — {subject}"
        title = doc.add_heading(f"{subject} — {class_name} — Term {term} {year}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_paragraph(f"Teacher: {teacher_name}    School: {school_name}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if scheme_data:
        headers = list(scheme_data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        WIDTH_CM = {
            'Main Competence': 2.6, 'Specific Competences': 2.6,
            'Main Learning Activities': 2.4, 'Specific Learning Activities': 2.8,
            'Month': 1.3, 'Week': 1.2, 'Number of Periods': 1.1,
            'Teaching and Learning Methods': 2.2,
            'Teaching and Learning Resources': 2.2,
            'Assessment Tools': 1.9, 'References': 2.2, 'Remarks': 2.5,
        }
        TOTAL_CM = 24.0
        col_cms = []
        for h in headers:
            w = WIDTH_CM.get(h)
            if w is None:
                for k, v in WIDTH_CM.items():
                    if k.lower() in h.lower() or h.lower() in k.lower():
                        w = v
                        break
            col_cms.append(w if w else TOTAL_CM / len(headers))
        scale_cm = TOTAL_CM / sum(col_cms)
        col_cms = [w * scale_cm for w in col_cms]

        section = doc.sections[0]
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = section.right_margin = Cm(1.2)
        section.top_margin = section.bottom_margin = Cm(1.2)

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].width = Cm(col_cms[i])
            p = hdr_cells[i].paragraphs[0]
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0A2B5E')
            tcPr.append(shd)

        for ri, row in enumerate(scheme_data):
            row_cells = table.add_row().cells
            fill_hex = 'EBF0FB' if ri % 2 == 0 else 'FFFFFF'
            for i, h in enumerate(headers):
                row_cells[i].width = Cm(col_cms[i])
                row_cells[i].text = str(row.get(h, '') or '')
                p = row_cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(8)
                tc = row_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"Scheme_{subject}_{class_name}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


def ajax_load_saved_scheme(request):
    """Load most recent saved SchemeOfWork from DB (works for both Django users & TLM teachers)."""
    try:
        tlm_teacher = get_tlm_teacher(request)
        scheme = None
        if request.user.is_authenticated:
            try:
                student = StudentTeacher.objects.get(user=request.user)
                scheme = (SchemeOfWork.objects
                          .filter(student=student)
                          .select_related('subject')
                          .order_by('-updated_at')
                          .first())
            except StudentTeacher.DoesNotExist:
                pass
        if not scheme and tlm_teacher and tlm_teacher.school:
            # TLM teacher: load by school and teacher_name
            scheme = (SchemeOfWork.objects
                      .filter(school=tlm_teacher.school, teacher_name=tlm_teacher.full_name)
                      .select_related('subject')
                      .order_by('-updated_at')
                      .first())
        if not scheme or not scheme.scheme_data:
            return JsonResponse({'success': False, 'error': 'Hakuna mpango uliohifadhiwa. Tengeneza kwanza.'}, status=404)
        return JsonResponse({
            'success': True,
            'data': scheme.scheme_data,
            'saved_id': scheme.id,
            'meta': {
                'subject': scheme.subject.name,
                'class_name': scheme.class_name,
                'term': scheme.term,
                'year': scheme.year,
                'teacher_name': scheme.teacher_name,
                'school_name': scheme.school.name if scheme.school else '',
                'total_weeks': scheme.total_weeks,
                'syllabus': scheme.syllabus,
                'updated_at': scheme.updated_at.strftime('%d %b %Y, %H:%M'),
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': 'Hitilafu: ' + str(e)[:100]}, status=500)


def ajax_load_scheme_by_id(request, scheme_id):
    """Load a specific SchemeOfWork by ID for viewing/editing from the library."""
    try:
        scheme = SchemeOfWork.objects.select_related('subject', 'school').get(id=scheme_id)
        if not scheme.scheme_data:
            return JsonResponse({'success': False, 'error': 'Scheme haina data.'}, status=404)
        return JsonResponse({
            'success': True,
            'data': scheme.scheme_data,
            'saved_id': scheme.id,
            'meta': {
                'subject': scheme.subject.name if scheme.subject else '',
                'class_name': scheme.class_name,
                'term': scheme.term,
                'year': scheme.year,
                'teacher_name': scheme.teacher_name or '',
                'school_name': scheme.school.name if scheme.school else '',
                'total_weeks': scheme.total_weeks,
                'syllabus': scheme.syllabus,
                'updated_at': scheme.updated_at.strftime('%d %b %Y, %H:%M'),
            }
        })
    except SchemeOfWork.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Scheme haipatikani.'}, status=404)


# =============================================================================
# LESSON PLAN
# =============================================================================

def lesson_plan_view(request):
    """Display lesson plan generator form — requires TLM registration."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:lesson_plan')}")
    
    education_levels = EducationLevel.objects.all().order_by('order')

    import json as _json
    classes_by_level = {}
    for cl in ClassLevel.objects.select_related('education_level').order_by('education_level', 'order'):
        classes_by_level.setdefault(cl.education_level_id, []).append({'id': cl.id, 'name': cl.name})

    # Filter subjects by education level (Primary → primary, Advanced → advanced, Ordinary → secondary)
    def _subjects_for_level(level_name):
        name_lower = level_name.lower()
        if 'primary' in name_lower:
            return list(Subject.objects.filter(level='primary').order_by('name').values('id', 'name'))
        elif 'advanced' in name_lower:
            return list(Subject.objects.filter(level='advanced').order_by('name').values('id', 'name'))
        elif 'technical' in name_lower or 'veta' in name_lower:
            # Technical / VETA → technical subjects
            return list(Subject.objects.filter(level='technical').order_by('name').values('id', 'name'))
        else:
            # Ordinary Level → secondary subjects
            return list(Subject.objects.filter(level='secondary').order_by('name').values('id', 'name'))

    subjects_by_level = {
        lvl.id: _subjects_for_level(lvl.name)
        for lvl in education_levels
    }

    student = None
    school = None
    if request.user.is_authenticated:
        try:
            student = StudentTeacher.objects.select_related('selected_school').get(user=request.user)
            school = student.selected_school
            if not school:
                app = StudentApplication.objects.filter(
                    student=student, status='approved'
                ).select_related('school').first()
                if app:
                    school = app.school
        except StudentTeacher.DoesNotExist:
            pass

    # Get teacher info for auto-fill
    teacher_name = teacher.full_name if teacher else ''
    teacher_school_name = teacher.school.name if teacher and teacher.school else ''
    teacher_school_level = teacher.school.level if teacher and teacher.school else ''
    teacher_subject_id = teacher.subject.id if teacher and teacher.subject else ''
    teacher_subject_name = teacher.subject.name if teacher and teacher.subject else ''
    teacher_subject_level = teacher.subject.level if teacher and teacher.subject else ''
    teacher_class_name = teacher.class_name if teacher else ''
    teacher_stream = teacher.stream if teacher else ''
    teacher_total_boys = teacher.total_boys if teacher else 0
    teacher_total_girls = teacher.total_girls if teacher else 0

    # ── Statistics: lesson plans generated by this teacher ──
    _now = timezone.now()
    _today = _now.date()
    _week_start = _today - timedelta(days=_today.weekday())  # Jumatatu
    _month_start = _today.replace(day=1)

    if teacher and teacher.school:
        _my_lessons = LessonPlan.objects.filter(
            teacher_name=teacher.full_name,
            school=teacher.school,
        )
        lesson_plans_count = _my_lessons.count()
        lesson_plans_today = _my_lessons.filter(created_at__date=_today).count()
        lesson_plans_week = _my_lessons.filter(created_at__date__gte=_week_start).count()
        lesson_plans_month = _my_lessons.filter(created_at__date__gte=_month_start).count()
        # Lesson plans generated TODAY (for the 'Imetengenezwa Leo' section)
        todays_lesson_plans = list(
            _my_lessons.filter(created_at__date=_today)
            .select_related('subject')
            .order_by('-created_at')[:6]
        )
        # Most recent plans (fallback display + 'Load Saved' quick list)
        recent_lesson_plans = list(
            _my_lessons.select_related('subject')
            .order_by('-created_at')[:6]
        )
    else:
        lesson_plans_count = lesson_plans_today = lesson_plans_week = lesson_plans_month = 0
        todays_lesson_plans = []
        recent_lesson_plans = []

    return render(request, 'curriculum/lesson_plan.html', {
        'lesson_plans_count': lesson_plans_count,
        'lesson_plans_today': lesson_plans_today,
        'lesson_plans_week': lesson_plans_week,
        'lesson_plans_month': lesson_plans_month,
        'todays_lesson_plans': todays_lesson_plans,
        'recent_lesson_plans': recent_lesson_plans,
        'today': _today,
        'education_levels': education_levels,
        'classes_by_level_json': _json.dumps(classes_by_level),
        'subjects_by_level_json': _json.dumps(subjects_by_level),
        'student': student,
        'school': school,
        'teacher': teacher,
        'teacher_name': teacher_name,
        'teacher_school_name': teacher_school_name,
        'teacher_school_level': teacher_school_level,
        'teacher_subject_id': teacher_subject_id,
        'teacher_subject_name': teacher_subject_name,
        'teacher_subject_level': teacher_subject_level,
        'teacher_class_name': teacher_class_name,
        'teacher_stream': teacher_stream,
        'teacher_total_boys': teacher_total_boys,
        'teacher_total_girls': teacher_total_girls,
    })


@csrf_exempt
def ajax_generate_lessonplan(request):
    """Generate lesson plan using AI — follows Tanzanian Teacher's Lesson Plan format."""
    if client is None:
        return JsonResponse({'success': False, 'error': 'Huduma ya AI haitumiki.'}, status=503)

    if request.method != 'POST':
        return JsonResponse({'success': False}, status=400)

    try:
        data = json.loads(request.body)

        education_level = data.get('education_level', '')
        class_name = data.get('class_name', '')
        stream = data.get('stream', '')
        subject = data.get('subject', '')
        subject_id = data.get('subject_id', '')
        topic = data.get('topic', '')
        subtopic = data.get('subtopic', '')
        term = data.get('term', 'I')
        year = data.get('year', 2026)
        duration = int(data.get('duration', 40))
        total_boys = data.get('total_boys', '')
        total_girls = data.get('total_girls', '')
        total_students = data.get('total_students', '')
        present_boys = data.get('present_boys', '')
        present_girls = data.get('present_girls', '')
        present_students = data.get('present_students', '')
        teacher_name = data.get('teacher_name', '')
        school_name = data.get('school_name', '')

        full_class = f"{class_name}{stream}" if stream else class_name

        # Calculate IDDR time allocation
        intro_time = max(5, int(duration * 0.15))
        dev_time = max(10, int(duration * 0.40))
        design_time = max(8, int(duration * 0.30))
        real_time = max(5, int(duration * 0.15))

        # ── Language: manual selection (english/kiswahili) or auto-detect ──
        _lp_tlm = get_tlm_teacher(request)
        _lp_language = data.get('language', getattr(_lp_tlm, 'preferred_language', 'auto') if _lp_tlm else 'auto')
        _lp_school_level = _lp_tlm.school.level if _lp_tlm and _lp_tlm.school else ''
        lp_language_instruction = _get_lp_language_instruction(_lp_language, subject, _lp_school_level)

        prompt = f"""Generate a TEACHER'S LESSON PLAN for a Tanzanian {education_level} classroom following the OFFICIAL TAMISEMI (PRIME MINISTER'S OFFICE - REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT) format.

============================================
PRIME MINISTER'S OFFICE
REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT
TEACHER'S LESSON PLAN
============================================

School: {school_name or '[School Name]'}
Teacher's Name: {teacher_name}
Subject: {subject}
Form/Class: {full_class}
Date: {datetime.now().strftime("%d/%m/%Y")}

Main Topic: {topic}
Sub-topic: {subtopic or 'N/A'}

Main Competence: Numbered format "1.0 Topic Name" - the REAL numbered competence from the Tanzanian syllabus for this subject/class.
Specific Competence: The REAL specific competence from the syllabus for this subtopic.

Term: {term}, Year: {year}
Duration: {duration} minutes

{lp_language_instruction}

-- STUDENT STATISTICS --
Registered Boys: {total_boys or 'N/A'}, Registered Girls: {total_girls or 'N/A'}, Total: {total_students or 'N/A'}
Present Boys: {present_boys or 'N/A'}, Present Girls: {present_girls or 'N/A'}, Present Total: {present_students or 'N/A'}

-- TEACHING PROCESS (IDDR Model) --

CRITICAL REQUIREMENTS - MUST FOLLOW EXACTLY:

!!! CONTENT MATCHING (CRITICAL): ALL content below MUST be specifically about:
   - Subject: {subject} (do NOT use a different subject)
   - Class: {full_class} (do NOT use a different form level)
   - Main Topic: "{topic}" (ALL content must relate DIRECTLY to this topic)
   - Sub-topic: "{subtopic or 'N/A'}" (ALL content must relate DIRECTLY to this subtopic)
   FAILURE: If you use content for a different topic, subject, or class, the output is WRONG.

1. main_competence: Use numbered format like "1.0 Topic Name" - the REAL Tanzanian syllabus competence for {subject} {full_class}
2. specific_competence: REAL subtopic-specific competence from the {subject} syllabus for this subtopic
3. specific_activity: Start with "By the end of this lesson, students should be able to:" then list 2-3 measurable outcomes SPECIFIC to "{topic}" and "{subtopic}"
4. references: Use the LATEST TIE textbook format - "Tanzania Institute of Education. (2024). {subject} for Secondary Schools Student's Book. Tanzania Institute of Education." If primary: "Tanzania Institute of Education. (2024). {subject} for Primary Schools Pupil's Book. Tanzania Institute of Education."
5. teaching_resources: MUST include specific, real resources for THIS topic (not generic) - e.g. TIE textbook pages, charts, real specimens, manila sheets, markers
6. remarks: Must be a detailed paragraph evaluating student achievement, specific challenges faced, and concrete way forward
7. student_statistics: Provide registered/present counts for girls and boys

!!! ALL lesson_development stages MUST use the EXACT topic "{topic}" and subtopic "{subtopic}":
   - Introduction: relate directly to "{topic}"
   - Competence Development: explore "{topic}" / "{subtopic}" specifically
   - Design: apply "{topic}" in real-life contexts
   - Realisation: assess understanding of "{topic}" / "{subtopic}"

Lesson Development uses the IDDR Model with these 5 columns per stage: stage, time, teaching_activities, learning_activities, assessment_criteria. Each stage MUST have detailed, Tanzania-specific content using CBC methodologies.

Output ONLY valid JSON with this EXACT structure.
⚠️ LANGUAGE REMINDER: ALL text values MUST be in the language specified by the LANGUAGE instruction above. If Kiswahili -> ALL values KISWAHILI. If English -> ALL values ENGLISH. CHECK every field before outputting!
{{
    "main_competence": "1.0 [Competence description in the language specified above]",
    "specific_competence": "[Specific competence in the language specified above]",
    "main_activity": "[Main learning activity in the language specified above]",
    "specific_activity": "[Specific learning activity in the language specified above]",
    "teaching_resources": "[Teaching resources in the language specified above]",
    "references": "[References in the language specified above]",
    "student_statistics": {{"registered_girls": "", "registered_boys": "", "present_girls": "", "present_boys": ""}},
    "lesson_development": [
        {{"stage": "[Stage 1 name in the language specified]", "time": "{intro_time:02d}", "teaching_activities": "[Description in the language specified]", "learning_activities": "[Description in the language specified]", "assessment_criteria": "[Criteria in the language specified]"}},
        {{"stage": "[Stage 2 name in the language specified]", "time": "{dev_time:02d}", "teaching_activities": "[Description in the language specified]", "learning_activities": "[Description in the language specified]", "assessment_criteria": "[Criteria in the language specified]"}},
        {{"stage": "[Stage 3 name in the language specified]", "time": "{design_time:02d}", "teaching_activities": "[Description in the language specified]", "learning_activities": "[Description in the language specified]", "assessment_criteria": "[Criteria in the language specified]"}},
        {{"stage": "[Stage 4 name in the language specified]", "time": "{real_time:02d}", "teaching_activities": "[Description in the language specified]", "learning_activities": "[Description in the language specified]", "assessment_criteria": "[Criteria in the language specified]"}}
    ],
    "remarks": "[Remarks in the language specified above]"
}}

All text values must be plain strings. Use REAL Tanzanian content. Return ONLY the JSON object, no extra text."""

        response = client.models.generate_content(model=model_name, contents=prompt)
        response_text = response.text

        cleaned = re.sub(r'```json\s*', '', response_text)
        cleaned = re.sub(r'```\s*', '', cleaned).strip()

        start_idx = cleaned.find('{')
        end_idx = cleaned.rfind('}')
        if start_idx != -1 and end_idx != -1:
            json_str = cleaned[start_idx:end_idx + 1]
            try:
                lesson_data = json.loads(json_str)
            except json.JSONDecodeError:
                logger.warning(f"[Lesson] JSON parse error - sanitizing control chars...")
                try:
                    lesson_data = json.loads(_sanitize_json_control_chars(json_str))
                except json.JSONDecodeError:
                    lesson_data = None
        else:
            lesson_data = None
        
        if lesson_data is None:
            # Could not parse JSON, use fallback default data
            # Fallback: generate sensible defaults
            lesson_data = {
                "main_competence": f"1.0 {topic} - Demonstrate understanding of {topic}",
                "specific_competence": f"Explain key concepts of {topic} based on the Tanzanian syllabus",
                "main_activity": f"Within 1 period students should be able to describe and apply {topic}",
                "specific_activity": f"By the end of this lesson, students should be able to:\n- Define {topic}\n- Explain key concepts of {topic}\n- Apply {topic} in real-life situations",
                "teaching_resources": "TIE textbook, Chalkboard/Whiteboard, Charts, Real objects, Manila sheets, Markers",
                "references": f"Tanzania Institute of Education. (2024). {subject} for Secondary Schools Student's Book. Tanzania Institute of Education.",
                "lesson_development": [
                    {"stage": "Introduction", "time": f"{intro_time:02d}", "teaching_activities": f"Display pictures/video about {topic}. Ask students questions to activate prior knowledge.", "learning_activities": "Observe pictures and respond to questions.", "assessment_criteria": "Questions about the lesson are answered."},
                    {"stage": "Competence Development", "time": f"{dev_time:02d}", "teaching_activities": f"Guide students in groups to explore {topic}. Provide guiding questions and resources.", "learning_activities": "Discuss in groups and share findings.", "assessment_criteria": "Concepts taught are clearly explained."},
                    {"stage": "Design", "time": f"{design_time:02d}", "teaching_activities": f"Ask students to apply knowledge of {topic} in real-life contexts through exercises.", "learning_activities": "Complete exercises and present findings.", "assessment_criteria": "Correct application of concepts."},
                    {"stage": "Realisation", "time": f"{real_time:02d}", "teaching_activities": "Assess student understanding through oral questions or short quiz. Provide feedback.", "learning_activities": "Respond to assessment questions and reflect.", "assessment_criteria": "Achievement of lesson objectives."}
                ],
                "student_statistics": {"registered_girls": "", "registered_boys": "", "present_girls": "", "present_boys": ""},
                "remarks": f"The students were able to explain {topic} due to the use of interactive teaching and learning methods. However, some students need more clarification. I will address this in the next lesson through remedial activities."
            }

        saved_id = None
        try:
            tlm_teacher = get_tlm_teacher(request)
            student = None
            school = None
            if request.user.is_authenticated:
                try:
                    student = StudentTeacher.objects.get(user=request.user)
                    school = student.selected_school
                except StudentTeacher.DoesNotExist:
                    pass
            elif tlm_teacher and tlm_teacher.school:
                school = tlm_teacher.school
            
            if school:
                level_map = {'primary school': 'primary', 'ordinary level': 'ordinary', 'advanced level': 'advanced'}
                edu_level = level_map.get((education_level or '').lower(), 'ordinary')
                subj_obj = None
                if subject_id:
                    try:
                        subj_obj = Subject.objects.get(id=int(subject_id))
                    except (Subject.DoesNotExist, ValueError):
                        pass
                if not subj_obj and subject:
                    subj_obj = Subject.objects.filter(name__iexact=subject).first()
                if subj_obj:
                    # Build lesson_development for DB (map new keys to old keys)
                    lp_development = []
                    for stage in lesson_data.get('lesson_development', []):
                        lp_development.append({
                            'stage': stage.get('stage', ''),
                            'time': stage.get('time', '') + ' min',
                            'teacher_activities': stage.get('teaching_activities', ''),
                            'student_activities': stage.get('learning_activities', ''),
                            'assessment_criteria': stage.get('assessment_criteria', ''),
                        })
                    lp_obj = LessonPlan.objects.create(
                        student=student, school=school, subject=subj_obj,
                        education_level=edu_level, class_name=class_name,
                        term=term, year=int(year), topic=topic, subtopic=subtopic or '',
                        date=timezone.now().date(), duration=duration,
                        total_boys=int(total_boys) if total_boys else 0,
                        total_girls=int(total_girls) if total_girls else 0,
                        total_students=int(total_students) if total_students else 0,
                        present_boys=int(present_boys) if present_boys else 0,
                        present_girls=int(present_girls) if present_girls else 0,
                        present_students=int(present_students) if present_students else 0,
                        teacher_name=teacher_name or (student.full_name if student else ''),
                        main_competence=lesson_data.get('main_competence', ''),
                        specific_competence=lesson_data.get('specific_competence', ''),
                        previous_knowledge=lesson_data.get('specific_activity', ''),
                        learning_objectives=[lesson_data.get('specific_activity', '')],
                        teaching_methods=[],
                        teaching_resources=[lesson_data.get('teaching_resources', '')],
                        lesson_development=lp_development,
                        remarks=lesson_data.get('remarks', ''),
                        generated_by_ai=True,
                    )
                    saved_id = lp_obj.id
        except Exception as save_err:
            logger.warning(f"[Curriculum] LessonPlan save error: {save_err}")

        return JsonResponse({'success': True, 'data': lesson_data, 'saved_id': saved_id})

    except Exception as e:
        import traceback
        traceback.print_exc()
        raw = str(e)
        logger.error(f"[Lesson Gen Error] {type(e).__name__}: {raw[:300]}")
        if '413' in raw or 'request too large' in raw.lower():
            msg = f"Ombi ni kubwa mno kwa AI. Jaribu kupunguza maelezo au kubadili somo. Hitilafu: {raw[:200]}"
        elif 'quota' in raw.lower() or '429' in raw or 'rate' in raw.lower():
            msg = f"Kikomo cha matumizi: {raw[:200]}"
        else:
            msg = f"Hitilafu ya AI: {raw[:200]}"
        return JsonResponse({'success': False, 'error': msg}, status=500)


def download_lesson_plan_pdf(request):
    """Export Lesson Plan as PDF — supports new Tanzanian format + old format."""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)

    data = json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form = data.get('form_data', {})

    # Helper: get value from new key or old key
    def _get(new_key, old_key, default=''):
        return lesson.get(new_key, lesson.get(old_key, default))

    # ── Create Tanzania flag watermark ──
    _tz_watermark = None
    if PILImage:
        try:
            _flag_path = os.path.join(os.path.dirname(__file__), 'static', 'curriculum', 'tz_flag.png')
            if os.path.exists(_flag_path):
                _pil_img = PILImage.open(_flag_path).convert('RGBA')
                _r, _g, _b, _a = _pil_img.split()
                _new_a = _a.point(lambda x: int(x * 0.10))
                _faded = PILImage.merge('RGBA', (_r, _g, _b, _new_a))
                _buf = BytesIO()
                _faded.save(_buf, format='PNG')
                _buf.seek(0)
                _tz_watermark = _buf
        except Exception:
            pass

    # ── Cover page: professional border drawing ──
    def _lp_cover(can, doc_obj):
        """Draw professional double-border frame with flag watermark."""
        can.saveState()
        pw = doc_obj.pagesize[0]
        ph = doc_obj.pagesize[1]
        # Outer gold border (thick)
        can.setStrokeColor(GOLD)
        can.setLineWidth(3.5)
        can.rect(30, 30, pw - 60, ph - 60)
        # Inner navy border (thin)
        can.setStrokeColor(NAVY)
        can.setLineWidth(1.5)
        can.rect(36, 36, pw - 72, ph - 72)
        # Top gold accent bar
        can.setStrokeColor(GOLD)
        can.setLineWidth(4)
        can.line(36, ph - 55, pw - 36, ph - 55)
        # Bottom gold accent bar
        can.line(36, 55, pw - 36, 55)
        # Top-left corner square
        can.setFillColor(GOLD)
        can.rect(36, ph - 26, 14, 8, fill=1, stroke=0)
        # Top-right corner square
        can.rect(pw - 50, ph - 26, 14, 8, fill=1, stroke=0)
        # Bottom-left corner square
        can.rect(36, 36, 14, 8, fill=1, stroke=0)
        # Bottom-right corner square
        can.rect(pw - 50, 36, 14, 8, fill=1, stroke=0)
        # ── Draw Tanzania flag watermark ──
        if _tz_watermark:
            try:
                _tz_watermark.seek(0)
                fw = pw * 0.45
                fh = fw * 2.0 / 3.0
                can.drawImage(_tz_watermark, pw/2 - fw/2, ph/2 - fh/2,
                              width=fw, height=fh, mask='auto')
            except Exception:
                pass
        can.restoreState()

    # ── Page header (appears from page 2 onwards) ──
    def _lp_header(can, doc_obj):
        """Draw NAME OF SCHOOL / TEACHER header bar on pages 2+."""
        can.saveState()
        can.setFont('Helvetica-Bold', 7)
        can.setFillColor(NAVY)
        pw = doc_obj.pagesize[0]
        ph = doc_obj.pagesize[1]
        sn = form.get('school_name', '____________________')
        tn = form.get('teacher_name', '____________________')
        sbj = form.get('subject', '')
        cls = form.get('class_name', '')
        trm = form.get('term', '')
        yr = form.get('year', '')
        can.drawString(36, ph - 18, f"NAME OF SCHOOL: {sn}")
        can.drawCentredString(pw / 2, ph - 18, f"SUBJECT: {sbj}")
        can.drawRightString(pw - 36, ph - 18, f"CLASS: {cls}")
        can.drawString(36, ph - 30, f"TEACHER'S NAME: {tn}")
        can.drawCentredString(pw / 2, ph - 30, f"TERM: {trm}")
        can.drawRightString(pw - 36, ph - 30, f"YEAR: {yr}")
        can.setStrokeColor(GOLD)
        can.setLineWidth(0.8)
        can.line(36, ph - 34, pw - 36, ph - 34)
        can.restoreState()

    # ── Get teacher theme for PDF ──
    teacher = get_tlm_teacher(request)
    theme_name = getattr(teacher, 'theme', 'classic') if teacher else 'classic'
    
    THEMES = {
        'classic': {
            'primary': colors.HexColor('#0A2B5E'),
            'accent': colors.HexColor('#C8900A'),
            'accent_dark': colors.HexColor('#A67B07'),
            'light': colors.HexColor('#EEF1F6'),
            'stripe': colors.HexColor('#F4F7FF'),
            'border': colors.HexColor('#9BAAC4'),
            'light_bg': colors.HexColor('#FAFBFD'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'tanzania': {
            'primary': colors.HexColor('#1E7A34'),
            'accent': colors.HexColor('#F5C518'),
            'accent_dark': colors.HexColor('#D4A000'),
            'light': colors.HexColor('#E8F5E9'),
            'stripe': colors.HexColor('#EBF8F0'),
            'border': colors.HexColor('#6BA87B'),
            'light_bg': colors.HexColor('#F5FBF7'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'ocean': {
            'primary': colors.HexColor('#1A365D'),
            'accent': colors.HexColor('#0EA5E9'),
            'accent_dark': colors.HexColor('#0284C7'),
            'light': colors.HexColor('#E0F2FE'),
            'stripe': colors.HexColor('#E8F4FD'),
            'border': colors.HexColor('#7FB8D9'),
            'light_bg': colors.HexColor('#F4F9FD'),
            'font_size': 8.0,
            'border_width': 2.5,
        },
        'royal': {
            'primary': colors.HexColor('#4C1D95'),
            'accent': colors.HexColor('#EC4899'),
            'accent_dark': colors.HexColor('#DB2775'),
            'light': colors.HexColor('#EDE9FE'),
            'stripe': colors.HexColor('#F3EEFB'),
            'border': colors.HexColor('#B794D4'),
            'light_bg': colors.HexColor('#FAF5FF'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'executive': {
            'primary': colors.HexColor('#1F2937'),
            'accent': colors.HexColor('#9CA3AF'),
            'accent_dark': colors.HexColor('#6B7280'),
            'light': colors.HexColor('#E5E7EB'),
            'stripe': colors.HexColor('#F0F1F3'),
            'border': colors.HexColor('#B0B5BD'),
            'light_bg': colors.HexColor('#F8F9FA'),
            'font_size': 8.5,
            'border_width': 2.0,
        },
        'sunset': {
            'primary': colors.HexColor('#C2410C'),
            'accent': colors.HexColor('#FB923C'),
            'accent_dark': colors.HexColor('#EA580C'),
            'light': colors.HexColor('#FFEDD5'),
            'stripe': colors.HexColor('#FFF7ED'),
            'border': colors.HexColor('#FDBA74'),
            'light_bg': colors.HexColor('#FFF7ED'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'forest': {
            'primary': colors.HexColor('#065F46'),
            'accent': colors.HexColor('#34D399'),
            'accent_dark': colors.HexColor('#059669'),
            'light': colors.HexColor('#D1FAE5'),
            'stripe': colors.HexColor('#ECFDF5'),
            'border': colors.HexColor('#6EE7B7'),
            'light_bg': colors.HexColor('#F0FDF4'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'midnight': {
            'primary': colors.HexColor('#1E1B4B'),
            'accent': colors.HexColor('#6366F1'),
            'accent_dark': colors.HexColor('#4F46E5'),
            'light': colors.HexColor('#E0E7FF'),
            'stripe': colors.HexColor('#EEF2FF'),
            'border': colors.HexColor('#A5B4FC'),
            'light_bg': colors.HexColor('#F8FAFC'),
            'font_size': 8.0,
            'border_width': 2.5,
        },
        'cherry': {
            'primary': colors.HexColor('#9B1C1C'),
            'accent': colors.HexColor('#F43F5E'),
            'accent_dark': colors.HexColor('#E11D48'),
            'light': colors.HexColor('#FFE4E6'),
            'stripe': colors.HexColor('#FFF1F2'),
            'border': colors.HexColor('#FDA4AF'),
            'light_bg': colors.HexColor('#FFF5F6'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'safari': {
            'primary': colors.HexColor('#78350F'),
            'accent': colors.HexColor('#D97706'),
            'accent_dark': colors.HexColor('#B45309'),
            'light': colors.HexColor('#FEF3C7'),
            'stripe': colors.HexColor('#FFFBEB'),
            'border': colors.HexColor('#FCD34D'),
            'light_bg': colors.HexColor('#FFFBEB'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
        'dawn': {
            'primary': colors.HexColor('#7E22CE'),
            'accent': colors.HexColor('#D946EF'),
            'accent_dark': colors.HexColor('#C026D3'),
            'light': colors.HexColor('#F3E8FF'),
            'stripe': colors.HexColor('#FAF5FF'),
            'border': colors.HexColor('#D8B4FE'),
            'light_bg': colors.HexColor('#FDF4FF'),
            'font_size': 8.5,
            'border_width': 2.5,
        },
    }
    T = THEMES.get(theme_name, THEMES['classic'])
    NAVY = T['primary']
    GOLD = T['accent']
    DARK_GOLD = T['accent_dark']
    LIGHT = T['light']
    STRIPE = T['stripe']
    BORDER = T['border']
    _lp_fs = T.get('font_size', 8.5)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=36, leftMargin=36, topMargin=40, bottomMargin=30)
    elements = []

    normal = ParagraphStyle('LP_N', fontName='Helvetica', fontSize=_lp_fs + 0.5, leading=_lp_fs + 4.5, wordWrap='LTR', spaceAfter=3)
    section_hdr = ParagraphStyle('LP_H', fontName='Helvetica-Bold', fontSize=_lp_fs + 1.5, textColor=NAVY, spaceBefore=10, spaceAfter=4)
    cell_s = ParagraphStyle('LP_C', fontName='Helvetica', fontSize=_lp_fs, leading=_lp_fs + 3, wordWrap='LTR')
    hdr_s = ParagraphStyle('LP_CH', fontName='Helvetica-Bold', fontSize=_lp_fs, leading=_lp_fs + 3, textColor=colors.white, wordWrap='LTR', alignment=1)
    label_s = ParagraphStyle('LP_L', fontName='Helvetica-Bold', fontSize=_lp_fs, leading=_lp_fs + 3, textColor=NAVY)
    title_s = ParagraphStyle('LP_TITLE', fontName='Helvetica-Bold', fontSize=16, alignment=1, textColor=NAVY, spaceAfter=1)
    subtitle_s = ParagraphStyle('LP_SUBTITLE', fontName='Helvetica-Bold', fontSize=_lp_fs - 0.5, alignment=1, textColor=DARK_GOLD, spaceAfter=1)

    # ── Cover page: decorative top lines ──
    elements.append(HRFlowable(width="100%", thickness=2.5, color=GOLD, spaceAfter=2))
    elements.append(HRFlowable(width="100%", thickness=1.2, color=NAVY, spaceAfter=10))

    # ── Determine language mode for LP (check form data for language hint) ──
    _lp_sw = form.get('language', '') == 'kiswahili' or form.get('subject', '').lower() in ('kiswahili', 'swahili')
    if not _lp_sw:
        # Also check via teacher's school level
        _lp_teacher_lp = get_tlm_teacher(request)
        if _lp_teacher_lp and _lp_teacher_lp.school:
            _lp_sw = _is_kiswahili_mode(
                form.get('language', getattr(_lp_teacher_lp, 'preferred_language', 'auto')),
                form.get('subject', ''),
                _lp_teacher_lp.school.level
            )

    # ── TAMISEMI Header (centered, bold) — bilingual ──
    _pm = _tl("PRIME MINISTER'S OFFICE", _lp_sw)
    _ralg = _tl("REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT", _lp_sw)
    _lp_title = _tl("TEACHER'S LESSON PLAN", _lp_sw)

    # Visual style based on subject hash for variety
    _lp_style_seed = sum(ord(c) for c in form.get('subject','') + (form.get('class_name','') or '')) % 5
    if _lp_style_seed == 0:
        # Classic centered navy
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1', fontName='Helvetica-Bold', fontSize=13, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2', fontName='Helvetica-Bold', fontSize=10, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
    elif _lp_style_seed == 1:
        # Gold text on light background
        elements.append(Spacer(1, 3))
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1b', fontName='Helvetica-Bold', fontSize=14, alignment=1,
                           textColor=GOLD, spaceAfter=0)))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2b', fontName='Helvetica-Bold', fontSize=11, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
    elif _lp_style_seed == 2:
        # Single line combined
        elements.append(Paragraph(f"{_pm} ❖ {_ralg}",
            ParagraphStyle('MH1c', fontName='Helvetica-Bold', fontSize=9, alignment=1,
                           textColor=NAVY, spaceAfter=2)))
    elif _lp_style_seed == 3:
        # Navy with decorative border
        elements.append(Paragraph(_pm,
            ParagraphStyle('MH1d', fontName='Helvetica-Bold', fontSize=13, alignment=1,
                           textColor=NAVY, spaceAfter=0, backColor=colors.HexColor('#F0F4FF'))))
        elements.append(Paragraph(_ralg,
            ParagraphStyle('MH2d', fontName='Helvetica-Bold', fontSize=10, alignment=1,
                           textColor=GOLD, spaceAfter=2)))
    else:
        # Inline style
        elements.append(Paragraph(f"{_pm}  |  {_ralg}",
            ParagraphStyle('MH1e', fontName='Helvetica-Bold', fontSize=8, alignment=1,
                           textColor=colors.HexColor('#555555'), spaceAfter=2)))

    elements.append(HRFlowable(width="50%", thickness=1, color=GOLD, spaceAfter=6))

    # ── Title ──
    elements.append(Paragraph(_lp_title,
        ParagraphStyle('LP_T', fontName='Helvetica-Bold', fontSize=14, alignment=1,
                       textColor=NAVY, spaceAfter=2)))
    _term_label = _tl('Term', _lp_sw)
    _year_label = _tl('Year', _lp_sw)
    elements.append(Paragraph(
        f"{form.get('subject','')}  |  {form.get('class_name','')}  |  {_term_label} {form.get('term','')} {_year_label} {form.get('year','')}",
        ParagraphStyle('LP_S', fontSize=9, alignment=1, textColor=colors.grey, spaceAfter=6)))

    # ── Cover info: styled table (bilingual labels) ──
    lbl = ParagraphStyle('lbl', fontName='Helvetica-Bold', fontSize=10, textColor=colors.white, leading=14)
    val = ParagraphStyle('val', fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#222222'), leading=14)
    info_rows = [
        [Paragraph(_tl("Teacher's Name", _lp_sw), lbl), Paragraph(form.get('teacher_name',''), val)],
        [Paragraph(_tl('School Name', _lp_sw), lbl), Paragraph(form.get('school_name','____________________'), val)],
        [Paragraph(_tl('Subject', _lp_sw), lbl), Paragraph(form.get('subject',''), val)],
        [Paragraph(_tl('Class', _lp_sw), lbl), Paragraph(form.get('class_name',''), val)],
        [Paragraph(_tl('Term', _lp_sw), lbl), Paragraph(form.get('term',''), val)],
        [Paragraph(_tl('Year', _lp_sw), lbl), Paragraph(str(form.get('year','')), val)],
        [Paragraph(_tl('Duration', _lp_sw), lbl), Paragraph(str(form.get('duration', '')) + ' min', val)],
        [Paragraph(_tl('Topic', _lp_sw), lbl), Paragraph(form.get('topic',''), val)],
    ]
    info_table = Table(info_rows, colWidths=[180, 220])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), NAVY),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#FAFBFD')),
        ('ROWBACKGROUNDS', (1, 0), (1, -1), [colors.HexColor('#FAFBFD'), STRIPE]),
        ('GRID', (0, 0), (-1, -1), 0.5, BORDER),
        ('BOX', (0, 0), (-1, -1), 2, GOLD),
        ('LINEBELOW', (0, 0), (-1, 0), 1.2, GOLD),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
    ]))
    elements.append(info_table)
    # ── PageBreak: Cover page ends here, lesson content starts on page 2 ──
    elements.append(PageBreak())

    def P(txt, st=normal):
        return Paragraph(str(txt or ''), st)

    # ── Student statistics data (embedded inside the Students cell) ──
    tb = str(form.get('total_boys', '') or '—')
    tg = str(form.get('total_girls', '') or '—')
    ts = str(form.get('total_students', '') or '—')
    pb = str(form.get('present_boys', '') or '—')
    pg = str(form.get('present_girls', '') or '—')
    ps = str(form.get('present_students', '') or '—')
    ab = str(form.get('absent_boys', '') or '—')
    ag = str(form.get('absent_girls', '') or '—')
    a_stud = str(form.get('absent_students', '') or '—')
    
    # Calculate absent if not provided
    if ab == '—' and tb != '—' and pb != '—':
        try: ab = str(int(float(tb)) - int(float(pb)))
        except: pass
    if ag == '—' and tg != '—' and pg != '—':
        try: ag = str(int(float(tg)) - int(float(pg)))
        except: pass
    if a_stud == '—' and ts != '—' and ps != '—':
        try: a_stud = str(int(float(ts)) - int(float(ps)))
        except: pass
    
    reg_label = _tl('Registered', _lp_sw)
    pres_label = _tl('Present', _lp_sw)
    abs_label = _tl('Absent', _lp_sw)
    boys_label = _tl('Boys', _lp_sw)
    girls_label = _tl('Girls', _lp_sw)
    total_label = _tl('Total', _lp_sw)
    
    # Styles for embedded students stats table
    _s_hdr = ParagraphStyle('_sh', fontName='Helvetica-Bold', fontSize=7.5, textColor=colors.white, alignment=1, leading=10)
    _s_cel = ParagraphStyle('_sc', fontName='Helvetica', fontSize=7.5, alignment=1, leading=10)
    _s_lbl = ParagraphStyle('_sl', fontName='Helvetica-Bold', fontSize=7.5, textColor=NAVY, leading=10)
    _s_cb  = ParagraphStyle('_cb', fontName='Helvetica-Bold', fontSize=7.5, alignment=1, leading=10, textColor=NAVY)
    
    stat_rows = [
        [Paragraph('', _s_hdr), Paragraph(f'<b>{reg_label}</b>', _s_hdr), Paragraph(f'<b>{pres_label}</b>', _s_hdr), Paragraph(f'<b>{abs_label}</b>', _s_hdr)],
        [Paragraph(f'{boys_label}', _s_lbl),        Paragraph(tb, _s_cel),                 Paragraph(pb, _s_cel),                 Paragraph(ab, _s_cel)],
        [Paragraph(f'{girls_label}', _s_lbl),       Paragraph(tg, _s_cel),                 Paragraph(pg, _s_cel),                 Paragraph(ag, _s_cel)],
        [Paragraph(f'{total_label}', _s_cb),         Paragraph(ts, _s_cb),                  Paragraph(ps, _s_cb),                  Paragraph(a_stud, _s_cb)],
    ]
    stat_inner = Table(stat_rows, colWidths=[48, 48, 48, 44])
    stat_inner.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), NAVY),
        ('BACKGROUND', (0, 1), (0, -1), LIGHT),
        ('GRID', (0, 0), (-1, -1), 0.3, BORDER),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
    ]))

    meta_rows = [
        [P(_tl('Teacher', _lp_sw), label_s), P(form.get('teacher_name', '')),
         P(_tl('Subject', _lp_sw), label_s), P(form.get('subject', ''))],
        [P(_tl('Class', _lp_sw), label_s), P(form.get('class_name', '')),
         P(_tl('Term/Year', _lp_sw), label_s), P(f"{_tl('Term', _lp_sw)} {form.get('term','')} {form.get('year','')}")],
        [P(_tl('Topic', _lp_sw), label_s), P(form.get('topic', '')),
         P(_tl('Subtopic', _lp_sw), label_s), P(form.get('subtopic', ''))],
        [P(_tl('Duration', _lp_sw), label_s), P(f"{form.get('duration','')} min"),
         P(_tl('Date', _lp_sw), label_s), P(str(timezone.now().date()))],
        [P(_tl('Students', _lp_sw), label_s),
         stat_inner, P(''), P('')],
    ]
    meta_tbl = Table(meta_rows, colWidths=[72, 188, 72, 191])
    meta_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), LIGHT),
        ('BACKGROUND', (2, 0), (2, -1), LIGHT),
        ('GRID', (0, 0), (-1, -1), 0.4, BORDER),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('LINEBELOW', (0, -1), (-1, -1), 1.2, GOLD),
    ]))
    elements.append(meta_tbl)
    elements.append(Spacer(1, 10))

    for label, key, fallback in [('Main Competence', 'main_competence', ''),
                       ('Specific Competence', 'specific_competence', ''),
                       ('Main Activity', 'main_activity', ''),
                       ('Specific Activity', 'specific_activity', 'previous_knowledge')]:
        val = lesson.get(key, lesson.get(fallback, '')) if fallback else lesson.get(key, '')
        if val:
            elements.append(Paragraph(f"<b>{_tl(label, _lp_sw)}:</b>  {val}", normal))

    # Teaching & Learning Resources
    tlr = lesson.get('teaching_resources', '')
    if tlr:
        elements.append(Paragraph(_tl('Teaching & Learning Resources', _lp_sw), section_hdr))
        if isinstance(tlr, list):
            for item in tlr:
                elements.append(Paragraph(f"<bullet>•</bullet> {item}",
                    ParagraphStyle('LP_B', fontName='Helvetica', fontSize=9, leading=13, leftIndent=14, wordWrap='LTR')))
        else:
            elements.append(Paragraph(f"{tlr}", normal))

    # References
    ref = lesson.get('references', '')
    if ref:
        elements.append(Paragraph(_tl('References', _lp_sw), section_hdr))
        elements.append(Paragraph(f"{ref}", normal))

    ld = lesson.get('lesson_development', [])
    if ld:
        elements.append(Paragraph(_tl('Lesson Development (IDDR Model)', _lp_sw), section_hdr))
        ld_headers = [_tl('Stage', _lp_sw), _tl('Time', _lp_sw), _tl('Teaching Activities', _lp_sw), _tl('Learning Activities', _lp_sw), _tl('Assessment Criteria', _lp_sw)]
        ld_data = [[Paragraph(h, hdr_s) for h in ld_headers]]
        for i, stage in enumerate(ld):
            bg = colors.white if i % 2 == 0 else STRIPE
            ld_data.append([
                Paragraph(str(stage.get('stage', stage.get('phase', '')) or ''), cell_s),
                Paragraph(str(stage.get('time', '') or ''), cell_s),
                Paragraph(str(stage.get('teaching_activities', stage.get('teacher_activities', '')) or ''), cell_s),
                Paragraph(str(stage.get('learning_activities', stage.get('student_activities', '')) or ''), cell_s),
                Paragraph(str(stage.get('assessment_criteria', '') or ''), cell_s),
            ])
        ld_tbl = Table(ld_data, colWidths=[72, 38, 118, 118, 105], repeatRows=1)
        ld_ts = [
            ('BACKGROUND', (0, 0), (-1, 0), NAVY),
            ('LINEBELOW', (0, 0), (-1, 0), 1.2, GOLD),
            ('GRID', (0, 0), (-1, -1), 0.4, BORDER),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]
        for i in range(1, len(ld_data)):
            ld_ts.append(('BACKGROUND', (0, i), (-1, i), STRIPE if i % 2 == 0 else colors.white))
        ld_tbl.setStyle(TableStyle(ld_ts))
        elements.append(ld_tbl)

    remarks = lesson.get('remarks', '')
    if remarks:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(_tl('Remarks', _lp_sw), section_hdr))
        if isinstance(remarks, dict):
            for label, key in [('1. Strength', 'strength'), ('2. Weakness', 'weakness'), ('3. Way Forward', 'way_forward')]:
                val = remarks.get(key, '') or '...............................................'
                elements.append(Paragraph(f"<b>{_tl(label, _lp_sw)}:</b>  {val}", normal))
        else:
            elements.append(Paragraph(str(remarks), normal))

    doc.build(elements, onFirstPage=_lp_cover, onLaterPages=_lp_header)
    buffer.seek(0)
    safe_name = f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.pdf"'
    return response


def download_lesson_plan_word(request):
    """Export Lesson Plan as Word (.docx)."""
    if request.method != 'POST':
        return HttpResponse("Invalid request", status=400)

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = json.loads(request.body)
    lesson = data.get('lesson_data', {})
    form = data.get('form_data', {})

    doc = Document()
    doc.add_heading('LESSON PLAN', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_table = doc.add_table(rows=5, cols=4)
    meta_table.style = 'Table Grid'
    rows = meta_table.rows

    def _set(r, c, text, bold=False):
        cell = rows[r].cells[c]
        cell.text = text
        if bold and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    # Determine language mode
    _lp_sw_word = form.get('language', '') == 'kiswahili' or form.get('subject', '').lower() in ('kiswahili', 'swahili')

    _set(0, 0, _tl('Teacher', _lp_sw_word) + ':', True)
    _set(0, 1, form.get('teacher_name', ''))
    _set(0, 2, _tl('Subject', _lp_sw_word) + ':', True)
    _set(0, 3, form.get('subject', ''))
    _set(1, 0, _tl('Class', _lp_sw_word) + ':', True)
    _set(1, 1, form.get('class_name', ''))
    _set(1, 2, _tl('Term/Year', _lp_sw_word) + ':', True)
    _set(1, 3, f"{_tl('Term', _lp_sw_word)} {form.get('term','')} {form.get('year','')}")
    _set(2, 0, _tl('Topic', _lp_sw_word) + ':', True)
    _set(2, 1, form.get('topic', ''))
    _set(2, 2, _tl('Subtopic', _lp_sw_word) + ':', True)
    _set(2, 3, form.get('subtopic', ''))
    _set(3, 0, _tl('Duration', _lp_sw_word) + ':', True)
    _set(3, 1, f"{form.get('duration','')} min")
    _set(3, 2, _tl('Date', _lp_sw_word) + ':', True)
    _set(3, 3, str(timezone.now().date()))
    _set(4, 0, _tl('Students', _lp_sw_word) + ':', True)
    tb = form.get('total_boys', '')
    tg = form.get('total_girls', '')
    ts = form.get('total_students', '')
    pb = form.get('present_boys', '')
    pg = form.get('present_girls', '')
    ps = form.get('present_students', '')
    _set(4, 1, f"{_tl('Students', _lp_sw_word)} — B:{tb} G:{tg} T:{ts}")

    doc.add_paragraph()

    for label, key in [('Main Competence', 'main_competence'),
                       ('Specific Competence', 'specific_competence'),
                       ('Main Activity', 'main_activity'),
                       ('Specific Activity', 'specific_activity')]:
        val = lesson.get(key, '')
        if not val:
            val = lesson.get({'Main Activity': None, 'Specific Activity': 'previous_knowledge'}.get(key, ''), '')
        if val:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(val))

    # Teaching & Learning Resources
    tlr = lesson.get('teaching_resources', '')
    if tlr:
        doc.add_heading('Teaching & Learning Resources', level=2)
        if isinstance(tlr, list):
            for item in tlr:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(tlr))

    # References
    ref = lesson.get('references', '')
    if ref:
        p = doc.add_paragraph()
        p.add_run('References: ').bold = True
        p.add_run(str(ref))

    ld = lesson.get('lesson_development', [])
    if ld:
        doc.add_heading('Lesson Development (IDDR Model)', level=2)
        ld_table = doc.add_table(rows=1, cols=5)
        ld_table.style = 'Table Grid'
        hdr = ld_table.rows[0].cells
        for i, h in enumerate(['Stage', 'Time', 'Teaching Activities', 'Learning Activities', 'Assessment Criteria']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for stage in ld:
            row = ld_table.add_row().cells
            row[0].text = stage.get('stage', stage.get('phase', ''))
            row[1].text = stage.get('time', '')
            row[2].text = stage.get('teaching_activities', stage.get('teacher_activities', ''))
            row[3].text = stage.get('learning_activities', stage.get('student_activities', ''))
            row[4].text = stage.get('assessment_criteria', '')
            row[4].text = stage.get('assessment_criteria', '')

    remarks = lesson.get('remarks', '')
    if remarks:
        doc.add_heading('Remarks', level=2)
        if isinstance(remarks, dict):
            for label, key in [('1. Strength', 'strength'), ('2. Weakness', 'weakness'), ('3. Way Forward', 'way_forward')]:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(remarks.get(key, '') or '...............................................')
        else:
            p = doc.add_paragraph()
            p.add_run('Remarks: ').bold = True
            p.add_run(str(remarks))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_name = f"LessonPlan_{form.get('subject','')}_{form.get('topic','')}".replace(' ', '_')
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response


# =============================================================================
# SAVE EDITED SCHEME / LESSON PLAN
# =============================================================================

@require_POST
def ajax_save_scheme_edits(request):
    """Save edited scheme data back to DB."""
    try:
        data = json.loads(request.body)
        scheme_id = data.get('saved_id')
        scheme_data = data.get('scheme_data', [])
        
        if not scheme_data:
            return JsonResponse({'success': False, 'error': 'Hakuna data'}, status=400)
        
        if scheme_id:
            try:
                scheme = SchemeOfWork.objects.get(id=scheme_id)
                scheme.scheme_data = scheme_data
                scheme.save(update_fields=['scheme_data'])
            except SchemeOfWork.DoesNotExist:
                pass
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


@require_POST
def ajax_save_lesson_edits(request):
    """Save edited lesson plan data back to DB (supports new + old key formats)."""
    try:
        data = json.loads(request.body)
        lesson_id = data.get('saved_id')
        lesson_data = data.get('lesson_data', {})
        form_data = data.get('form_data', {})
        
        if lesson_id:
            try:
                lp = LessonPlan.objects.get(id=lesson_id)
                # Map new keys AND old keys to DB fields
                # Main competence
                if lesson_data.get('main_competence'):
                    lp.main_competence = lesson_data['main_competence']
                # Specific competence
                if lesson_data.get('specific_competence'):
                    lp.specific_competence = lesson_data['specific_competence']
                # Previous knowledge (new: specific_activity, old: previous_knowledge)
                spec_act = lesson_data.get('specific_activity', lesson_data.get('previous_knowledge', ''))
                if spec_act:
                    lp.previous_knowledge = spec_act
                # Learning objectives (new data may not have this)
                if lesson_data.get('learning_objectives'):
                    lp.learning_objectives = lesson_data['learning_objectives']
                elif spec_act:
                    lp.learning_objectives = [spec_act]
                # Teaching methods (new data may not have this)
                if lesson_data.get('teaching_methods'):
                    lp.teaching_methods = lesson_data['teaching_methods']
                # Teaching resources (new: string, old: array)
                tlr = lesson_data.get('teaching_resources', '')
                if tlr:
                    if isinstance(tlr, str):
                        lp.teaching_resources = [tlr]
                    else:
                        lp.teaching_resources = tlr
                # Lesson development - map new keys to old keys for DB storage
                if lesson_data.get('lesson_development'):
                    mapped_dev = []
                    for stage in lesson_data['lesson_development']:
                        mapped_dev.append({
                            'stage': stage.get('stage', stage.get('phase', '')),
                            'time': stage.get('time', ''),
                            'teacher_activities': stage.get('teaching_activities', stage.get('teacher_activities', '')),
                            'student_activities': stage.get('learning_activities', stage.get('student_activities', '')),
                            'methods': stage.get('methods', ''),
                            'assessment_criteria': stage.get('assessment_criteria', ''),
                        })
                    lp.lesson_development = mapped_dev
                # Remarks (new: string, old: dict with strength/weakness/way_forward)
                if lesson_data.get('remarks'):
                    lp.remarks = lesson_data['remarks']
                # Teacher name
                if form_data.get('teacher_name'):
                    lp.teacher_name = form_data['teacher_name']
                # Topic
                if form_data.get('topic'):
                    lp.topic = form_data['topic']
                # Subtopic
                if form_data.get('subtopic') is not None:
                    lp.subtopic = form_data['subtopic']
                # Student statistics breakdown
                if form_data.get('total_boys'):
                    lp.total_boys = int(form_data['total_boys'])
                if form_data.get('total_girls'):
                    lp.total_girls = int(form_data['total_girls'])
                if form_data.get('present_boys'):
                    lp.present_boys = int(form_data['present_boys'])
                if form_data.get('present_girls'):
                    lp.present_girls = int(form_data['present_girls'])
                if form_data.get('total_students'):
                    lp.total_students = int(form_data['total_students'])
                if form_data.get('present_students'):
                    lp.present_students = int(form_data['present_students'])
                lp.save()
            except LessonPlan.DoesNotExist:
                pass
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


def ajax_load_saved_lessonplan(request):
    """Load most recent saved LessonPlan from DB (works for both Django users & TLM teachers)."""
    try:
        tlm_teacher = get_tlm_teacher(request)
        lp = None
        if request.user.is_authenticated:
            try:
                student = StudentTeacher.objects.get(user=request.user)
                lp = (LessonPlan.objects
                      .filter(student=student)
                      .select_related('subject')
                      .order_by('-created_at')
                      .first())
            except StudentTeacher.DoesNotExist:
                pass
        if not lp and tlm_teacher and tlm_teacher.school:
            lp = (LessonPlan.objects
                  .filter(school=tlm_teacher.school, teacher_name=tlm_teacher.full_name)
                  .select_related('subject')
                  .order_by('-created_at')
                  .first())
        if not lp or not lp.lesson_development:
            return JsonResponse({'success': False, 'error': 'Hakuna mpango wa somo uliohifadhiwa.'}, status=404)
        # Build lesson_data with backward-compatible keys
        lesson_data = {
            'lesson_title': f"{lp.subject.name} - {lp.topic}",
            'main_competence': lp.main_competence,
            'specific_competence': lp.specific_competence,
            'main_activity': None,  # Not stored in old model
            'specific_activity': lp.previous_knowledge or '',
            'previous_knowledge': lp.previous_knowledge,
            'learning_objectives': lp.learning_objectives,
            'teaching_methods': lp.teaching_methods,
            'teaching_resources': lp.teaching_resources,
            'references': None,  # Not stored in old model
            'lesson_development': lp.lesson_development,
            'remarks': lp.remarks,
        }
        form_data = {
            'subject': lp.subject.name, 'class_name': lp.class_name,
            'term': lp.term, 'year': lp.year, 'topic': lp.topic,
            'subtopic': lp.subtopic or '', 'teacher_name': lp.teacher_name,
            'duration': str(lp.duration),  # Must be string for PDF generator
            'total_students': lp.total_students or '',
            'present_students': lp.present_students or '',
            'total_boys': lp.total_boys or '',
            'total_girls': lp.total_girls or '',
            'present_boys': lp.present_boys or '',
            'present_girls': lp.present_girls or '',
        }
        return JsonResponse({'success': True, 'data': lesson_data, 'form_data': form_data, 'saved_id': lp.id})
    except StudentTeacher.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Wasifu haupatikani.'}, status=404)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': 'Hitilafu: ' + str(e)[:200]}, status=500)


def ajax_load_lesson_by_id(request, lesson_id):
    """Load a specific LessonPlan by ID for viewing/editing from the library."""
    try:
        lp = LessonPlan.objects.select_related('subject', 'school').get(id=lesson_id)
        if not lp.lesson_development:
            return JsonResponse({'success': False, 'error': 'Lesson Plan haina data.'}, status=404)
        # Build lesson_data with backward-compatible keys
        lesson_data = {
            'lesson_title': f"{lp.subject.name} - {lp.topic}",
            'main_competence': lp.main_competence,
            'specific_competence': lp.specific_competence,
            'main_activity': None,  # Not stored in old model
            'specific_activity': lp.previous_knowledge or '',
            'previous_knowledge': lp.previous_knowledge,
            'learning_objectives': lp.learning_objectives,
            'teaching_methods': lp.teaching_methods,
            'teaching_resources': lp.teaching_resources,
            'references': None,  # Not stored in old model
            'lesson_development': lp.lesson_development,
            'remarks': lp.remarks,
        }
        form_data = {
            'subject': lp.subject.name, 'class_name': lp.class_name,
            'term': lp.term, 'year': lp.year, 'topic': lp.topic,
            'subtopic': lp.subtopic or '', 'teacher_name': lp.teacher_name,
            'duration': str(lp.duration),  # Must be string for PDF generator
            'total_students': lp.total_students or '',
            'present_students': lp.present_students or '',
            'total_boys': lp.total_boys or '',
            'total_girls': lp.total_girls or '',
            'present_boys': lp.present_boys or '',
            'present_girls': lp.present_girls or '',
        }
        return JsonResponse({'success': True, 'data': lesson_data, 'form_data': form_data, 'saved_id': lp.id})
    except LessonPlan.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Lesson Plan haipatikani.'}, status=404)


# =============================================================================
# LOGBOOK — works for BOTH TLM teachers (session-based) and authenticated IMS users
# =============================================================================

def _tlm_subjects(teacher):
    """Subjects suitable for the TLM teacher's school level."""
    school_level = teacher.school.level if teacher and teacher.school else ''
    if school_level:
        sl = school_level.lower()
        if 'primary' in sl:
            return Subject.objects.filter(level='primary').order_by('name')
        if 'advanced' in sl or 'a level' in sl:
            return Subject.objects.filter(level='advanced').order_by('name')
        if 'technical' in sl or 'veta' in sl:
            return Subject.objects.filter(level='technical').order_by('name')
        if 'secondary' in sl:
            return Subject.objects.filter(level='secondary').order_by('name')
    return Subject.objects.all().order_by('name')


def _tlm_logbook_context(teacher, logbook_entry, today):
    """Context for the logbook template when the user is a TLM teacher."""
    days_swahili = {0: 'Jumatatu', 1: 'Jumanne', 2: 'Jumatano', 3: 'Alhamisi', 4: 'Ijumaa'}
    days_english = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday'}
    school = teacher.school
    return {
        'form': None,
        'student': None,
        'logbook_entry': logbook_entry,
        'today': today,
        'today_name': days_swahili.get(today.weekday(), 'Leo'),
        'today_name_en': days_english.get(today.weekday(), 'Today'),
        'school': school,
        'subjects': _tlm_subjects(teacher),
        'teacher': teacher,
        'preferred_language': getattr(teacher, 'preferred_language', 'auto'),
        'tlm_teacher_name': teacher.full_name,
        'tlm_school_name': school.name if school else '',
        'tlm_subject_name': teacher.subject.name if teacher.subject else '',
        'tlm_class_name': teacher.class_name or '',
    }


def submit_logbook(request):
    """
    Submit daily logbook entry.
    FIX: TLM teachers (session-based, no Django login) no longer get redirected
    to the student dashboard — they use TLMLogbookEntry directly.
    """
    today = timezone.now().date()
    teacher = get_tlm_teacher(request)

    # ── TLM TEACHER FLOW (session-based, NO Django login needed) ──
    if teacher:
        if today.weekday() >= 5:
            messages.info(request, "Hakuna kazi ya uwanjani wikendi. Rudi tena Jumatatu.")
            return redirect(reverse('curriculum:logbook_history'))

        school = teacher.school
        logbook_entry = TLMLogbookEntry.objects.filter(teacher=teacher, date=today).first()

        if request.method == 'POST':
            if logbook_entry is None:
                logbook_entry = TLMLogbookEntry(
                    teacher=teacher, date=today, school=school,
                    day_of_week=['monday', 'tuesday', 'wednesday', 'thursday', 'friday'][today.weekday()],
                )
            logbook_entry.school = school
            logbook_entry.other_activities = request.POST.get('other_activities', '')
            logbook_entry.challenges_faced = request.POST.get('challenges_faced', '')
            logbook_entry.lessons_learned = request.POST.get('lessons_learned', '')
            try:
                logbook_entry.lessons_data = json.loads(request.POST.get('lessons_data', '[]'))
            except (ValueError, TypeError):
                logbook_entry.lessons_data = []
            logbook_entry.is_location_verified = True
            logbook_entry.save()
            messages.success(request, "✅ Logbook imesajiliwa kikamilifu!")
            return redirect(reverse('curriculum:logbook_history'))

        return render(request, 'curriculum/logbook.html',
                      _tlm_logbook_context(teacher, logbook_entry, today))

    # ── AUTHENTICATED IMS STUDENT FLOW (fallback) ──
    if not request.user.is_authenticated:
        return redirect(f"{reverse('login')}?next={reverse('curriculum:submit_logbook')}")

    student = get_or_create_student_profile(request.user)

    if today.weekday() >= 5:
        messages.info(request, "Hakuna kazi ya uwanjani wikendi. Rudi tena Jumatatu.")
        return redirect(reverse('curriculum:logbook_history'))

    if not student.selected_school:
        messages.error(request, "Lazima uchague shule kabla ya kujaza logbook.")
        return redirect(reverse('curriculum:logbook_history'))

    school = student.selected_school
    logbook_entry = _cached_today_logbook(student, school, today)

    if request.method == 'POST':
        if logbook_entry is None:
            logbook_entry, _ = LogbookEntry.objects.get_or_create(
                student=student, date=today,
                defaults={'school': school, 'morning_check_in': timezone.now()}
            )
            _invalidate_today_logbook(student, today)

        form = LogbookForm(request.POST, instance=logbook_entry)

        if form.is_valid():
            entry = form.save(commit=False)
            try:
                entry.lessons_data = json.loads(request.POST.get('lessons_data', '[]'))
            except (ValueError, TypeError):
                entry.lessons_data = []

            entry.is_location_verified = True
            entry.save()
            _invalidate_today_logbook(student, today)
            messages.success(request, "✅ Logbook imesajiliwa kikamilifu!")

            return redirect(reverse('curriculum:logbook_history'))
        else:
            messages.error(request, "Tafadhali kagua makosa yaliyomo kwenye fomu.")
    else:
        form = LogbookForm(instance=logbook_entry)

    subjects = _cached_subjects(student)
    days_swahili = {0: 'Jumatatu', 1: 'Jumanne', 2: 'Jumatano', 3: 'Alhamisi', 4: 'Ijumaa'}
    days_english = {0: 'Monday', 1: 'Tuesday', 2: 'Wednesday', 3: 'Thursday', 4: 'Friday'}

    return render(request, 'curriculum/logbook.html', {
        'form': form, 'student': student, 'logbook_entry': logbook_entry,
        'today': today, 'today_name': days_swahili.get(today.weekday(), 'Leo'),
        'today_name_en': days_english.get(today.weekday(), 'Today'),
        'school': school, 'subjects': subjects,
        'teacher': None,
        'preferred_language': 'auto',
        'tlm_teacher_name': student.full_name,
        'tlm_school_name': school.name if school else '',
        'tlm_subject_name': '',
        'tlm_class_name': '',
    })


def logbook_history(request):
    """View logbook history — TLM teachers see their own TLM logbooks."""
    teacher = get_tlm_teacher(request)

    week_filter = request.GET.get('week')
    month_filter = request.GET.get('month')

    # ── TLM TEACHER FLOW ──
    if teacher:
        entries = TLMLogbookEntry.objects.filter(teacher=teacher).select_related('school')
        owner_name = teacher.full_name
        is_tlm = True
    else:
        if not request.user.is_authenticated:
            return redirect(f"{reverse('login')}?next={reverse('curriculum:logbook_history')}")
        student = get_or_create_student_profile(request.user)
        entries = LogbookEntry.objects.filter(student=student).select_related('subject_taught', 'school')
        owner_name = student.full_name
        is_tlm = False

    if week_filter:
        try:
            year, week = map(int, week_filter.split('-W'))
            start_date = datetime.strptime(f'{year}-W{week}-1', "%Y-W%W-%w").date()
            entries = entries.filter(date__range=[start_date, start_date + timedelta(days=6)])
        except ValueError:
            messages.error(request, "Tarehe ya wiki si sahihi.")

    elif month_filter:
        try:
            year, month = map(int, month_filter.split('-'))
            entries = entries.filter(date__year=year, date__month=month)
        except ValueError:
            messages.error(request, "Tarehe ya mwezi si sahihi.")

    else:
        today = timezone.now().date()
        start_of_week = today - timedelta(days=today.weekday())
        entries = entries.filter(date__range=[start_of_week, start_of_week + timedelta(days=4)])

    return render(request, 'curriculum/logbook_history.html', {
        'entries': entries.order_by('-date'), 'student': None,
        'teacher': teacher, 'owner_name': owner_name, 'is_tlm': is_tlm,
    })


def download_logbook_pdf(request, period_type=None):
    """Download logbook as PDF — works for TLM teachers AND authenticated IMS users."""
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet

    today = timezone.now().date()
    teacher = get_tlm_teacher(request)
    period_value = period_type or 'week'

    # ── Decide data source: TLM teacher vs authenticated IMS student ──
    if teacher:
        entries = TLMLogbookEntry.objects.filter(teacher=teacher)
        owner_name = teacher.full_name
        owner_label = 'Jina la Mwalimu'
        school = teacher.school
    else:
        if not request.user.is_authenticated:
            return redirect(f"{reverse('login')}?next={reverse('curriculum:logbook_download_options')}")
        student = get_or_create_student_profile(request.user)
        entries = LogbookEntry.objects.filter(student=student)
        owner_name = student.full_name
        owner_label = 'Jina la Mwalimu'
        school = student.selected_school

    if period_value == 'today':
        entries = entries.filter(date=today)
        filename = f"logbook_{today}.pdf"
        title = f"Logbook — {today}"
    elif period_value == 'week':
        start_of_week = today - timedelta(days=today.weekday())
        end_of_week = start_of_week + timedelta(days=4)
        entries = entries.filter(date__range=[start_of_week, end_of_week])
        filename = f"logbook_wiki_{start_of_week}.pdf"
        title = f"Logbook — Wiki {start_of_week} hadi {end_of_week}"
    elif period_value == 'month':
        start_of_month = today.replace(day=1)
        next_month = today.replace(day=28) + timedelta(days=4)
        end_of_month = next_month - timedelta(days=next_month.day)
        entries = entries.filter(date__range=[start_of_month, end_of_month])
        filename = f"logbook_mwezi_{today.year}_{today.month:02d}.pdf"
        title = f"Logbook — Mwezi {today.month}/{today.year}"
    else:
        filename = f"logbook_zote_{today}.pdf"
        title = "Logbook Zote"

    entries = entries.order_by('date')

    NAVY = colors.HexColor('#0A2B5E')
    GOLD = colors.HexColor('#C8900A')
    LIGHT = colors.HexColor('#EEF1F6')
    WHITE = colors.white

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm)

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle('title', fontName='Helvetica-Bold', fontSize=14, textColor=NAVY, spaceAfter=4)
    s_sub = ParagraphStyle('sub', fontName='Helvetica', fontSize=9, textColor=colors.HexColor('#4A5568'), spaceAfter=2)
    s_head = ParagraphStyle('head', fontName='Helvetica-Bold', fontSize=10, textColor=WHITE)
    s_label = ParagraphStyle('label', fontName='Helvetica-Bold', fontSize=8, textColor=NAVY)
    s_body = ParagraphStyle('body', fontName='Helvetica', fontSize=8, textColor=colors.HexColor('#1A1A2E'), leading=11)
    s_small = ParagraphStyle('small', fontName='Helvetica', fontSize=7.5, textColor=colors.HexColor('#4A5568'), leading=10)

    school_name = school.name if school else '—'
    district_name = (school.district.name if school and school.district else '—')
    current_year = _cached_active_year()
    year_label = str(current_year) if current_year else '—'

    story = []
    story.append(Paragraph("WIZARA YA ELIMU, SAYANSI NA TEKNOLOJIA",
                           ParagraphStyle('gov', fontName='Helvetica-Bold', fontSize=11, textColor=NAVY, alignment=1)))
    story.append(Paragraph("Mfumo wa Ufuatiliaji wa Walimu (TLM — Teaching & Learning Materials)",
                           ParagraphStyle('gov2', fontName='Helvetica', fontSize=9, textColor=colors.HexColor('#4A5568'), alignment=1, spaceAfter=6)))
    story.append(HRFlowable(width="100%", thickness=2, color=GOLD, spaceAfter=6))
    story.append(Paragraph(title, s_title))

    info_data = [
        [Paragraph(f'<b>{owner_label}:</b>', s_label), Paragraph(owner_name, s_body),
         Paragraph('<b>Shule:</b>', s_label), Paragraph(school_name, s_body)],
        [Paragraph('<b>Wilaya:</b>', s_label), Paragraph(district_name, s_body),
         Paragraph('<b>Mwaka wa Masomo:</b>', s_label), Paragraph(year_label, s_body)],
        [Paragraph('<b>Tarehe ya Chapisha:</b>', s_label), Paragraph(str(today), s_body),
         Paragraph('<b>Idadi ya Siku:</b>', s_label), Paragraph(str(entries.count()), s_body)],
    ]
    info_tbl = Table(info_data, colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 5.5 * cm])
    info_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), LIGHT),
        ('BOX', (0, 0), (-1, -1), 0.5, NAVY),
        ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#CBD5E0')),
        ('PADDING', (0, 0), (-1, -1), 4),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 10))

    day_names = {'Monday': 'Jumatatu', 'Tuesday': 'Jumanne', 'Wednesday': 'Jumatano',
                 'Thursday': 'Alhamisi', 'Friday': 'Ijumaa', 'Saturday': 'Jumamosi', 'Sunday': 'Jumapili'}

    for entry in entries:
        day_sw = day_names.get(entry.date.strftime('%A'), entry.date.strftime('%A'))
        gps_status = "✓ Imehakikiwa" if entry.is_location_verified else "✗ Haijahakikiwa"
        gps_color = colors.HexColor('#10B981') if entry.is_location_verified else colors.HexColor('#F59E0B')

        day_hdr = Table(
            [[Paragraph(f"{day_sw} — {entry.date}", s_head),
              Paragraph(f"GPS: {gps_status}", ParagraphStyle('gps', fontName='Helvetica-Bold', fontSize=8, textColor=gps_color))]],
            colWidths=[13 * cm, 5 * cm]
        )
        day_hdr.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), NAVY),
            ('PADDING', (0, 0), (-1, -1), 6),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ]))
        story.append(day_hdr)

        lessons = entry.lessons_data if entry.lessons_data else []
        if not lessons:
            if entry.other_activities or entry.challenges_faced:
                extra_rows = []
                if entry.other_activities:
                    extra_rows.append([Paragraph('<b>Shughuli Nyingine</b>', s_label), Paragraph(entry.other_activities, s_body)])
                if entry.challenges_faced:
                    extra_rows.append([Paragraph('<b>Changamoto</b>', s_label), Paragraph(entry.challenges_faced, s_body)])
                if extra_rows:
                    extra_tbl = Table(extra_rows, colWidths=[4 * cm, 14 * cm])
                    extra_tbl.setStyle(TableStyle([
                        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E0')),
                        ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#CBD5E0')),
                        ('BACKGROUND', (0, 0), (0, -1), LIGHT),
                        ('PADDING', (0, 0), (-1, -1), 4),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    story.append(extra_tbl)
            else:
                story.append(Paragraph("Hakuna data ya masomo.", s_small))
        else:
            for lesson in lessons:
                period_num = lesson.get('period', '?')
                subj = lesson.get('subject', '—')
                cls = lesson.get('class', '—')
                main_topic = lesson.get('main_topic', '—')
                activity = lesson.get('activity_type', '—')
                
                # Boys/Girls breakdown
                enrolled_b = lesson.get('enrolled_boys', lesson.get('enrolled', '—'))
                enrolled_g = lesson.get('enrolled_girls', '—')
                present_b = lesson.get('present_boys', lesson.get('present', '—'))
                present_g = lesson.get('present_girls', '—')
                methods = lesson.get('methods', '—')
                aids = lesson.get('teaching_aids', '—')
                achievements = lesson.get('achievements', '')
                challenges_p = lesson.get('challenges', '')
                subtopic = lesson.get('subtopic', '—')

                period_hdr = Table(
                    [[Paragraph(
                        f"Kipindi {period_num}  |  {subj}  |  Darasa: {cls}  |  Aina: {activity}",
                        ParagraphStyle('ph', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE))]],
                    colWidths=[18 * cm]
                )
                period_hdr.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), GOLD),
                    ('PADDING', (0, 0), (-1, -1), 5),
                ]))
                story.append(period_hdr)

                lesson_rows = [
                    [Paragraph('<b>Mada Kuu</b>', s_label), Paragraph(main_topic, s_body),
                     Paragraph('<b>Mada Ndogo</b>', s_label), Paragraph(subtopic, s_body)],
                    [Paragraph('<b>Mbinu</b>', s_label), Paragraph(methods, s_body),
                     Paragraph('<b>Vifaa</b>', s_label), Paragraph(aids, s_body)],
                ]
                
                # Add Boys/Girls row
                enroll_boys_str = str(enrolled_b) if enrolled_b and enrolled_b != '—' else '0'
                enroll_girls_str = str(enrolled_g) if enrolled_g and enrolled_g != '—' else '0'
                pres_boys_str = str(present_b) if present_b and present_b != '—' else '0'
                pres_girls_str = str(present_g) if present_g and present_g != '—' else '0'
                
                lesson_rows.append([
                    Paragraph('<b>Waliojisajili</b>', s_label),
                    Paragraph(f"Wavulana: {enroll_boys_str}  |  Wasichana: {enroll_girls_str}", s_body),
                    Paragraph('<b>Waliohudhuria</b>', s_label),
                    Paragraph(f"Wavulana: {pres_boys_str}  |  Wasichana: {pres_girls_str}", s_body),
                ])
                
                # Add Achievements & Challenges row if present
                if achievements or challenges_p:
                    ach_text = achievements if achievements else '—'
                    ch_text = challenges_p if challenges_p else '—'
                    lesson_rows.append([
                        Paragraph('<b>Mafanikio</b>', s_label), Paragraph(ach_text, s_body),
                        Paragraph('<b>Changamoto</b>', s_label), Paragraph(ch_text, s_body),
                    ])
                
                lesson_tbl = Table(lesson_rows, colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 5.5 * cm])
                lesson_tbl.setStyle(TableStyle([
                    ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E0')),
                    ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#CBD5E0')),
                    ('BACKGROUND', (0, 0), (0, -1), LIGHT),
                    ('BACKGROUND', (2, 0), (2, -1), LIGHT),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(lesson_tbl)

        if entry.supervisor_remarks:
            sr_tbl = Table(
                [[Paragraph('<b>Maoni ya Msimamizi</b>', s_label), Paragraph(entry.supervisor_remarks, s_body)]],
                colWidths=[4 * cm, 14 * cm]
            )
            sr_tbl.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E0')),
                ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#CBD5E0')),
                ('BACKGROUND', (0, 0), (0, -1), LIGHT),
                ('PADDING', (0, 0), (-1, -1), 4),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(sr_tbl)

        story.append(Spacer(1, 10))

    if not entries:
        story.append(Spacer(1, 20))
        story.append(Paragraph("Hakuna rekodi za logbook kwa kipindi kilichochaguliwa.", s_sub))

    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceBefore=10, spaceAfter=4))
    story.append(Paragraph("© Wizara ya Elimu, Sayansi na Teknolojia — IMS v2.1.0",
                           ParagraphStyle('footer', fontName='Helvetica', fontSize=7, textColor=colors.HexColor('#4A5568'), alignment=1)))

    doc.build(story)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def logbook_download_options(request):
    """Page for choosing download options — works for TLM teachers too."""
    teacher = get_tlm_teacher(request)
    if teacher:
        total_entries = TLMLogbookEntry.objects.filter(teacher=teacher).count()
        this_week_entries = TLMLogbookEntry.objects.filter(
            teacher=teacher, date__gte=timezone.now().date() - timedelta(days=7)
        ).count()
        return render(request, 'curriculum/logbook_download.html', {
            'student': None, 'teacher': teacher, 'owner_name': teacher.full_name,
            'total_entries': total_entries, 'this_week_entries': this_week_entries,
        })

    if not request.user.is_authenticated:
        return redirect(f"{reverse('login')}?next={reverse('curriculum:logbook_download_options')}")

    student = get_or_create_student_profile(request.user)
    total_entries = LogbookEntry.objects.filter(student=student).count()
    this_week_entries = LogbookEntry.objects.filter(
        student=student, date__gte=timezone.now().date() - timedelta(days=7)
    ).count()

    return render(request, 'curriculum/logbook_download.html', {
        'student': student, 'teacher': None, 'owner_name': student.full_name,
        'total_entries': total_entries, 'this_week_entries': this_week_entries,
    })


# =============================================================================
# API HELPERS
# =============================================================================

def get_topics_by_subject(request):
    """AJAX: Get topics (from textbooks) for a given subject and class level."""
    subject_id = request.GET.get('subject_id')
    class_level_id = request.GET.get('class_level_id')
    
    if not subject_id:
        return JsonResponse([], safe=False)
    
    subject = Subject.objects.filter(id=subject_id).first()
    if not subject:
        return JsonResponse([], safe=False)
    
    # Map subject level to textbook education_level
    if subject.level == 'primary':
        edu_level = 'primary'
    else:
        # For secondary, check class level to determine ordinary vs advanced
        if class_level_id:
            cl = ClassLevel.objects.filter(id=class_level_id).first()
            if cl:
                cl_name = cl.name.lower()
                if 'form 5' in cl_name or 'form 6' in cl_name or 'advanced' in cl_name:
                    edu_level = 'advanced'
                else:
                    edu_level = 'ordinary'
            else:
                edu_level = 'ordinary'
        else:
            edu_level = 'ordinary'
    
    textbooks = Textbook.objects.filter(
        subject=subject,
        education_level=edu_level,
        is_active=True
    ).order_by('title').values('id', 'title')
    
    topics = list(textbooks)
    
    # If no textbooks found, return empty list
    return JsonResponse(topics, safe=False)


def get_classes_by_level(request):
    """AJAX: Get classes for a given education level."""
    level_id = request.GET.get('level_id')
    if not level_id:
        return JsonResponse([], safe=False)
    classes = ClassLevel.objects.filter(education_level_id=level_id).order_by('order')
    return JsonResponse([{'id': c.id, 'name': c.name} for c in classes], safe=False)


def get_subjects_by_level(request):
    """AJAX: Get subjects for a given education level.
    Primary School → primary subjects
    Ordinary/Advanced Level → secondary subjects
    """
    level_id = request.GET.get('level_id')
    if not level_id:
        return JsonResponse([], safe=False)
    level = EducationLevel.objects.filter(id=level_id).first()
    if not level:
        return JsonResponse([], safe=False)
    if 'primary' in level.name.lower():
        subjects = Subject.objects.filter(level='primary').order_by('name')
    elif 'advanced' in level.name.lower():
        subjects = Subject.objects.filter(level='advanced').order_by('name')
    elif 'technical' in level.name.lower() or 'veta' in level.name.lower():
        # Technical / VETA → technical subjects
        subjects = Subject.objects.filter(level='technical').order_by('name')
    else:
        # Ordinary Level → secondary subjects
        subjects = Subject.objects.filter(level='secondary').order_by('name')
    return JsonResponse([{'id': s.id, 'name': s.name} for s in subjects], safe=False)


def get_topics_ai(request):
    """
    AJAX: Use AI to generate a list of topics for a given subject and class level.
    Returns a JSON array of topic strings.
    """
    if client is None:
        return JsonResponse({'success': False, 'error': 'AI haitumiki'}, status=503)

    subject_id = request.GET.get('subject_id')
    class_name = request.GET.get('class_name', '')
    education_level = request.GET.get('education_level', '')

    if not subject_id:
        return JsonResponse({'success': False, 'error': 'subject_id inahitajika'}, status=400)

    subject = Subject.objects.filter(id=subject_id).first()
    if not subject:
        return JsonResponse({'success': False, 'error': 'Somo halipatikani'}, status=404)

    # ── Determine language based on school level (Primary vs Secondary) ──
    _is_primary = ('primary' in education_level.lower())
    _subject_lower = subject.name.lower()

    if _is_primary:
        if _subject_lower in ('english', 'english language'):
            language_rule = (
                f"LUGHA: Hii ni SHULE YA MSINGI na somo ni ENGLISH. "
                f"Topics ZOTE zirejeshwe kwa KIINGEREZA (English). "
                f"Kwa mfano: 'Parts of Speech', 'Tenses', 'Comprehension', 'Nouns', 'Verbs', 'Reading', 'Writing' n.k."
            )
        else:
            language_rule = (
                f"LUGHA: Hii ni SHULE YA MSINGI. Somo la {subject.name} linafundishwa kwa KISWAHILI. "
                f"Topics ZOTE lazima zirejeshwe kwa KISWAHILI. "
                f"Hata kama jina la somo ni Kiingereza (kama Science, Geography, civics), "
                f"topics zake kwa shule ya msingi lazima ziwe kwa KISWAHILI. "
                f"Kwa mfano: "
                f"- Sayansi (Science) → 'Viumbe Hai', 'Mwili na Afya', 'Mimea', 'Wanyama', 'Hali ya Hewa' n.k. "
                f"- Hisabati → 'Namba', 'Jiometri', 'Aljebra', 'Vipimo', 'Takwimu' n.k. "
                f"- Maarifa ya Jamii → 'Familia', 'Jamii', 'Uchumi', 'Utamaduni' n.k. "
                f"- Uraia na Maadili → 'Haki za Binadamu', 'Majukumu', 'Maadili' n.k. "
                f"- Mwili na Afya → 'Mwili wa Binadamu', 'Lishe', 'Afya na Usafi' n.k. "
                f"USIANDIKE topics kwa Kiingereza isipokuwa kwa somo la English pekee."
            )
    else:
        # Secondary / Advanced level
        if _subject_lower in ('kiswahili', 'swahili'):
            language_rule = (
                f"LUGHA: Hii ni SHULE YA SEKONDARI na somo ni Kiswahili. "
                f"Topics ZOTE zirejeshwe kwa KISWAHILI. "
                f"Kwa mfano: 'Fasihi Simulizi', 'Uchambuzi wa Riwaya', 'Matumizi ya Lugha' n.k."
            )
        else:
            language_rule = (
                f"LUGHA: Hii ni SHULE YA SEKONDARI. "
                f"Somo la {subject.name} linafundishwa kwa KIINGEREZA (isipokuwa Kiswahili). "
                f"Topics ZOTE zirejeshwe kwa KIINGEREZA. "
                f"Kwa mfano: 'Living Things', 'Matter', 'Energy', 'Algebra', 'Trigonometry', 'Cell Biology' n.k."
            )

    prompt = (
            f"Wewe ni mtaalamu wa mtaala wa Tanzania (TIE/SEQUIP). "
            f"Orodhesha mada kuu (topics) za somo la {subject.name} kwa darasa la {class_name} "
            f"katika ngazi ya {education_level} kwa mujibu wa mtaala wa TIE Tanzania. "
            f"\n\n"
            f"!!! HII NI MUHIMU SANA: TOPICS LAZIMA ZIWE ZA DARASA LA {class_name} PEKEE. "
            f"USIORODHESHE topics za darasa jingine! Kwa mfano: "
            f"- Kama ni Form 2 → orodhesha topics za Form 2 PEKEE (sio Form 1, sio Form 3, sio Form 4) "
            f"- Kama ni Form 1 → orodhesha topics za Form 1 PEKEE "
            f"- Kama ni Form 3 → orodhesha topics za Form 3 PEKEE "
            f"- Kama ni Form 4 → orodhesha topics za Form 4 PEKEE \n\n"
            f"MAKOSI: Topics lazima zilingane hasa na zile za vitabu vya TIE kwa darasa la {class_name}. "
            f"KAGUA: Hakikisha kila topic uliyoorodhesha ni ya {class_name} kulingana na silabasi ya TIE. "
            f"USIORODHESHE topic zozote za Form 1, Form 3, au Form 4 kama unataka topics za Form 2. \n"
            f"{language_rule}"
            f"\n"
            f"Rudisha TU JSON array ya string: [\"Topic 1\", \"Topic 2\", ...]. "
            f"USIANDIKE chochote kingine — JSON pekee."
        )

    try:
        response = client.models.generate_content(model=model_name, contents=prompt)
        text = response.text.strip()
        # Strip markdown code fences if present
        text = re.sub(r'```(?:json)?', '', text).strip()
        topics = json.loads(text)
        if isinstance(topics, list):
            return JsonResponse({'success': True, 'topics': topics})
        return JsonResponse({'success': True, 'topics': []})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


def get_subtopics_ai(request):
    """
    AJAX: Use AI to generate subtopics for a given subject, class level, and topic.
    Returns a JSON array of subtopic strings.
    """
    if client is None:
        return JsonResponse({'success': False, 'error': 'AI haitumiki'}, status=503)

    subject_id = request.GET.get('subject_id')
    topic = request.GET.get('topic', '')
    class_name = request.GET.get('class_name', '')
    education_level = request.GET.get('education_level', '')

    if not subject_id or not topic:
        return JsonResponse({'success': False, 'error': 'subject_id na topic vinahitajika'}, status=400)

    subject = Subject.objects.filter(id=subject_id).first()
    if not subject:
        return JsonResponse({'success': False, 'error': 'Somo halipatikani'}, status=404)

    # ── Determine language based on school level (Primary vs Secondary) ──
    _is_primary = ('primary' in education_level.lower())
    _subject_lower = subject.name.lower()

    if _is_primary:
        if _subject_lower in ('english', 'english language'):
            language_rule = (
                f"LUGHA: Hii ni SHULE YA MSINGI na somo ni ENGLISH. "
                f"Subtopics ZOTE zirejeshwe kwa KIINGEREZA (English). "
                f"Kwa mfano: 'Nouns', 'Verbs', 'Pronouns', 'Present Tense', 'Past Tense', 'Comprehension' n.k."
            )
        else:
            language_rule = (
                f"LUGHA: Hii ni SHULE YA MSINGI. Somo la {subject.name} linafundishwa kwa KISWAHILI. "
                f"Subtopics ZOTE lazima zirejeshwe kwa KISWAHILI. "
                f"Hata kama jina la somo ni Kiingereza (kama Science, Geography, Civics), "
                f"subtopics zake kwa shule ya msingi lazima ziwe kwa KISWAHILI. "
                f"Kwa mfano: "
                f"- Sayansi (Science) → 'Viumbe Hai', 'Mimea', 'Wanyama', 'Mfumo wa Mwili' n.k. "
                f"- Hisabati → 'Namba Asilia', 'Sehemu', 'Desimali', 'Mizani', 'Ulinganifu' n.k. "
                f"- Maarifa ya Jamii → 'Familia yangu', 'Jamii yetu', 'Shughuli za Kiuchumi' n.k. "
                f"- Uraia na Maadili → 'Haki na Wajibu', 'Maadili ya Kiislamu/Kikristo', 'Usalama Barabarani' n.k. "
                f"USIANDIKE subtopics kwa Kiingereza isipokuwa kwa somo la English pekee."
            )
    else:
        # Secondary / Advanced level
        if _subject_lower in ('kiswahili', 'swahili'):
            language_rule = (
                f"LUGHA: Hii ni SHULE YA SEKONDARI na somo ni Kiswahili. "
                f"Subtopics ZOTE zirejeshwe kwa KISWAHILI. "
                f"Kwa mfano: 'Tamthilia', 'Ushairi', 'Insha', 'Sarufi', 'Matumizi ya Lugha' n.k."
            )
        else:
            language_rule = (
                f"LUGHA: Hii ni SHULE YA SEKONDARI. "
                f"Somo la {subject.name} linafundishwa kwa KIINGEREZA (isipokuwa Kiswahili). "
                f"Subtopics ZOTE zirejeshwe kwa KIINGEREZA. "
                f"Kwa mfano: 'Cell Division', 'Quadratic Equations', 'Chemical Bonding', 'Market Structure' n.k."
            )

    prompt = (
            f"Wewe ni mtaalamu wa mtaala wa Tanzania (TIE/SEQUIP). "
            f"Orodhesha mada ndogo (subtopics) za somo la {subject.name}, kwa mada kuu '{topic}', "
            f"kwa darasa la {class_name} ({education_level}) kwa mujibu wa mtaala wa TIE Tanzania. "
            f"\n\n"
            f"!!! HII NI MUHIMU SANA: SUBTOPICS LAZIMA ZIWE ZA DARASA LA {class_name} PEKEE. "
            f"USIORODHESHE subtopics za darasa jingine! Kama mada kuu '{topic}' inapatikana kwa darasa zaidi ya moja, "
            f"chagua subtopics za {class_name} PEKEE, sio za darasa lingine. \n\n"
            f"MAKOSI: Subtopics lazima zilingane na zile za vitabu vya TIE kwa darasa la {class_name}. "
            f"KAGUA mara mbili: Hakikisha subtopics zote ni za {class_name} kulingana na silabasi ya TIE. "
            f"{language_rule}"
            f"\n"
            f"Rudisha TU JSON array ya string: [\"Subtopic 1\", \"Subtopic 2\", ...]. "
            f"Kama hakuna subtopics, rudisha [] (array tupu). "
            f"USIANDIKE chochote kingine — JSON pekee."
        )

    try:
        response = client.models.generate_content(model=model_name, contents=prompt)
        text = response.text.strip()
        text = re.sub(r'```(?:json)?', '', text).strip()
        subtopics = json.loads(text)
        if isinstance(subtopics, list):
            return JsonResponse({'success': True, 'subtopics': subtopics})
        return JsonResponse({'success': True, 'subtopics': []})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


def get_textbooks_by_level(request):
    """AJAX: Get textbooks/references for a given education level."""
    level_id = request.GET.get('level_id')
    if not level_id:
        return JsonResponse([], safe=False)
    level = EducationLevel.objects.filter(id=level_id).first()
    if not level:
        return JsonResponse([], safe=False)
    level_name = level.name.lower()
    level_map = {'primary': 'primary', 'ordinary': 'ordinary', 'advanced': 'advanced'}
    mapped = 'primary'
    for k, v in level_map.items():
        if k in level_name:
            mapped = v
            break
    textbooks = Textbook.objects.filter(education_level=mapped, is_active=True).order_by('title')
    return JsonResponse([{'id': t.id, 'title': t.title, 'publisher': t.publisher} for t in textbooks], safe=False)


# =============================================================================
# AJAX: Search schools (autocomplete)
# =============================================================================

def ajax_search_schools(request):
    """
    AJAX: Search schools by name (for autocomplete on registration).
    Returns JSON array of matching schools with id, name, level, district__name.
    """
    q = request.GET.get('q', '').strip()
    district_id = request.GET.get('district_id', '')
    if len(q) < 2 and not district_id:
        return JsonResponse([], safe=False)
    
    schools = School.objects.all()
    if q and len(q) >= 2:
        schools = schools.filter(name__icontains=q)
    if district_id:
        schools = schools.filter(district_id=district_id)
    
    schools = schools.select_related('district').order_by('name')[:30]
    return JsonResponse([{
        'id': s.id,
        'name': s.name,
        'level': s.level,
        'district_name': s.district.name,
    } for s in schools], safe=False)


# =============================================================================
# AJAX: Submit update comment/feedback
# =============================================================================

@require_POST
def ajax_submit_update_comment(request):
    """Submit a comment/feedback about system updates from teachers."""
    try:
        data = json.loads(request.body)
        message = data.get('message', '').strip()
        teacher_name = data.get('teacher_name', '').strip()
        
        if not message:
            return JsonResponse({'success': False, 'error': 'Tafadhali andika maoni yako.'}, status=400)
        if not teacher_name:
            return JsonResponse({'success': False, 'error': 'Tafadhali ingiza jina lako.'}, status=400)
        
        teacher = get_tlm_teacher(request)
        teacher_school_name = teacher.school.name if teacher and teacher.school else ''
        
        testimonial = Testimonial.objects.create(
            teacher=teacher,
            teacher_name=teacher_name,
            school_name=teacher_school_name,
            message=message,
            is_approved=True,
        )
        
        return JsonResponse({'success': True, 'id': testimonial.id})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


# =============================================================================
# LESSON NOTES — standalone teacher notes
# =============================================================================

def lesson_notes_view(request):
    """
    Lesson Notes page — teachers write their own reflections, methods, challenges.
    Notes are filtered by the teacher's school education level.
    Teachers only see notes from their own level (primary/ordinary/advanced).
    Shows live statistics + "Imetengenezwa Leo" badges.
    """
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:lesson_notes')}")
    
    school_level = teacher.school.level if teacher and teacher.school else ''
    
    # Determine the teacher's education level for filtering
    teacher_edu_level = 'ordinary'
    if school_level:
        sl = school_level.lower()
        if 'primary' in sl:
            teacher_edu_level = 'primary'
        elif 'advanced' in sl or 'a level' in sl or 'secondary' in sl:
            # For advanced level, we still show both ordinary and advanced
            # But check if specifically advanced
            if 'advanced' in sl or 'a level' in sl:
                teacher_edu_level = 'advanced'
        # Default stays 'ordinary'
    
    # Get notes for this teacher, filtered by education level
    notes = LessonNote.objects.filter(
        teacher=teacher,
        education_level=teacher_edu_level,
    ).order_by('-created_at')[:50]

    # ── LIVE STATISTICS ──
    today = timezone.now().date()
    week_ago = today - timedelta(days=7)
    month_start = today.replace(day=1)

    notes_qs = LessonNote.objects.filter(teacher=teacher, education_level=teacher_edu_level)
    total_notes = notes_qs.count()
    notes_today = notes_qs.filter(created_at__date=today).count()
    notes_this_week = notes_qs.filter(created_at__date__gte=week_ago).count()
    notes_this_month = notes_qs.filter(created_at__date__gte=month_start).count()

    from django.db.models import Count as _Count
    notes_by_subject = list(
        notes_qs.exclude(subject='').values('subject').annotate(count=_Count('id'))
        .order_by('-count')[:6]
    )
    notes_by_class = list(
        notes_qs.exclude(class_name='').values('class_name').annotate(count=_Count('id'))
        .order_by('-count')[:6]
    )
    
    # Subjects for dropdown
    subjects = Subject.objects.all().order_by('name')
    education_levels = EducationLevel.objects.all().order_by('order')
    
    # Get saved lesson plans for quick reference
    saved_lessons = LessonPlan.objects.filter(
        teacher_name=teacher.full_name,
        school=teacher.school,
    ).order_by('-created_at')[:10]
    
    return render(request, 'curriculum/lesson_notes.html', {
        'teacher': teacher,
        'school_level': school_level,
        'teacher_edu_level': teacher_edu_level,
        'notes': notes,
        'subjects': subjects,
        'education_levels': education_levels,
        'saved_lessons': saved_lessons,
        'teacher_name': teacher.full_name if teacher else '',
        'teacher_school_name': teacher.school.name if teacher and teacher.school else '',
        'teacher_subject_name': teacher.subject.name if teacher and teacher.subject else '',
        'teacher_class_name': teacher.class_name if teacher else '',
        # Statistics
        'today': today,
        'total_notes': total_notes,
        'notes_today': notes_today,
        'notes_this_week': notes_this_week,
        'notes_this_month': notes_this_month,
        'notes_by_subject': notes_by_subject,
        'notes_by_class': notes_by_class,
    })


@require_POST
def ajax_save_lesson_note(request):
    """AJAX: Save a new lesson note or update existing one."""
    try:
        data = json.loads(request.body)
        teacher = get_tlm_teacher(request)
        if not teacher:
            return JsonResponse({'success': False, 'error': 'Tafadhali jisajili kwanza.'}, status=401)
        
        note_id = data.get('note_id')
        content = data.get('content', '').strip()
        subject = data.get('subject', '').strip()
        class_name = data.get('class_name', '').strip()
        topic = data.get('topic', '').strip()
        education_level = data.get('education_level', '').strip()
        
        if not content:
            return JsonResponse({'success': False, 'error': 'Tafadhali andika maelezo ya somo.'}, status=400)
        
        school = teacher.school if teacher.school else None
        
        if note_id:
            # Update existing note
            note = get_object_or_404(LessonNote, id=note_id, teacher=teacher)
            note.content = content
            note.subject = subject
            note.class_name = class_name
            note.topic = topic
            note.education_level = education_level
            note.save()
        else:
            # Create new note
            note = LessonNote.objects.create(
                teacher=teacher,
                teacher_name=teacher.full_name,
                school=school,
                school_name=school.name if school else '',
                education_level=education_level or 'ordinary',
                class_name=class_name,
                subject=subject,
                topic=topic,
                content=content,
            )
        
        return JsonResponse({'success': True, 'note_id': note.id, 'created': note.created_at.isoformat()})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


def ajax_get_lesson_note(request, note_id):
    """AJAX: Get a single lesson note for editing."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Una hitaji kujiandikisha.'}, status=401)
    try:
        note = LessonNote.objects.get(id=note_id, teacher=teacher)
        return JsonResponse({
            'success': True,
            'note': {
                'id': note.id,
                'content': note.content,
                'subject': note.subject,
                'class_name': note.class_name,
                'topic': note.topic,
                'education_level': note.education_level,
                'created_at': note.created_at.strftime('%d %b %Y, %H:%M'),
            }
        })
    except LessonNote.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Note haipatikani.'}, status=404)


# =============================================================================
# AJAX: Generate Lesson Notes from a saved Lesson Plan (AI-powered)
# =============================================================================

@require_POST
def ajax_generate_lesson_note_from_lp(request):
    """Generate lesson notes (reflection + quiz) from a saved lesson plan."""
    if client is None:
        return JsonResponse({'success': False, 'error': 'AI haitumiki'}, status=503)
    
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Jisajili kwanza'}, status=401)
    
    try:
        data = json.loads(request.body)
        lp_id = data.get('lesson_plan_id')
        
        if not lp_id:
            return JsonResponse({'success': False, 'error': 'Lesson Plan ID inahitajika'}, status=400)
        
        lp = LessonPlan.objects.get(id=lp_id)
        
        # Determine language based on school level + subject
        school_name = teacher.school.name if teacher.school else ''
        school_level = teacher.school.level if teacher.school else ''
        subject_name = lp.subject.name if lp.subject else lp.topic
        
        # Build appropriate language instruction
        school_level_lower = (school_level or '').lower()
        subject_lower = subject_name.lower()
        
        if 'primary' in school_level_lower:
            if subject_lower in ('english', 'english language'):
                lang_instruction = "Write ALL content in ENGLISH."
            else:
                lang_instruction = "Write ALL content in KISWAHILI (Swahili) except column headers."
        else:
            if subject_lower in ('kiswahili', 'swahili'):
                lang_instruction = "Write ALL content in KISWAHILI (Swahili)."
            else:
                lang_instruction = "Write ALL content in ENGLISH."
        
        prompt = f"""You are an expert Tanzanian teacher writing DETAILED LESSON NOTES in NOTEBOOK format.
Your notes must be EXCEPTIONALLY THOROUGH — like a teacher's personal notebook that another teacher could use to teach the same lesson.

IMPORTANT LESSON PLAN DETAILS:
Subject: {subject_name}
Class: {lp.class_name}
Topic: {lp.topic}
Subtopic: {lp.subtopic or 'N/A'}
Date: {lp.date}
Duration: {lp.duration} minutes

Main Competence: {lp.main_competence or 'N/A'}
Specific Competence: {lp.specific_competence or 'N/A'}

{lang_instruction}

Based on the above lesson plan, create EXTREMELY DETAILED lesson notes in NOTEBOOK FORMAT.
These notes should be so complete that another teacher could pick them up and teach the lesson confidently.

Write in the following STRUCTURE - every section is required:

---

📖 [TOPIC NAME] — LESSON NOTES
═══════════════════════════════════════

1. MUHTASARI / SUMMARY (3-5 long paragraphs)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Write a VERY DETAILED summary covering:
- Ufafanuzi kamili wa dhana kuu (comprehensive definition of main concepts)
- Misingi ya kinadharia (theoretical foundations)
- Maelezo ya kina kwa kila hatua (detailed step-by-step explanations)
- Mifano halisi kutoka kwenye mada (real examples from the topic)
- Viungo na mada nyingine (connections to other topics)
- Matumizi ya kivitendo (practical applications)
- Maneno muhimu na istilahi (key vocabulary and terminology)

Each paragraph should be 5-8 SENTENCES long. Be thorough and educational.

2. NUKTA MUHIMU / KEY POINTS (at least 10 bullet points)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
List the MOST important concepts the student must remember, ordered from basic to advanced.
Each point should include a SHORT EXPLANATION, not just a heading.

3. MBINU ZA UFUNDISHAJI / TEACHING METHODS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- Describe the methods that worked BEST for this topic
- Nyimbo / Songs that can help students remember
- Michezo / Games for concept reinforcement
- Maswali ya haraka / Quick oral questions to ask during class
- Kazi za vikundi / Group work ideas
- Kazi za nyumbani / Homework assignments

4. TATHMINI / ASSESSMENT (5 MAJIBU SWALI / Q&A)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Generate 5 CHALLENGING questions with COMPLETE answers.
Questions should test: understanding, application, analysis and evaluation (not just recall).
Each answer should be 2-4 sentences explaining the concept clearly.

Question 1: [Question]
Jibu / Answer: [Detailed answer]

Question 2: [Question]
Jibu / Answer: [Detailed answer]

Question 3: [Question]
Jibu / Answer: [Detailed answer]

Question 4: [Question]
Jibu / Answer: [Detailed answer]

Question 5: [Question]
Jibu / Answer: [Detailed answer]

5. MWONGOZO / CONCLUSION
━━━━━━━━━━━━━━━━━━━━━━━━━━
- Things that went well during the lesson
- Challenges faced
- Recommendations for next lesson

OUTPUT as JSON with this exact structure:
{{
    "title": "Topic name - Lesson Notes",
    "summary_paragraphs": ["Long paragraph 1...", "Long paragraph 2...", "Long paragraph 3...", "Long paragraph 4...", "Long paragraph 5..."],
    "key_points": ["1. Concept - Explanation", "2. Concept - Explanation", ...],
    "teaching_methods": "Detailed description of teaching methods, songs, games, activities...",
    "quiz": [
        {{"question": "Question 1?", "answer": "Detailed answer 1..."}},
        {{"question": "Question 2?", "answer": "Detailed answer 2..."}},
        {{"question": "Question 3?", "answer": "Detailed answer 3..."}},
        {{"question": "Question 4?", "answer": "Detailed answer 4..."}},
        {{"question": "Question 5?", "answer": "Detailed answer 5..."}}
    ],
    "conclusion": "What went well, challenges, recommendations..."
}}

LENGTH REQUIREMENTS:
- summary_paragraphs: 5 paragraphs, each 5-8 sentences long (VERY DETAILED)
- key_points: at least 10 items
- quiz: exactly 5 questions with detailed answers (2-4 sentences each)
- teaching_methods: at least 3-4 sentences
- conclusion: 3-5 sentences

Return ONLY valid JSON. No other text."""

        response = client.models.generate_content(model=model_name, contents=prompt)
        response_text = response.text
        
        cleaned = re.sub(r'```json\\s*', '', response_text)
        cleaned = re.sub(r'```\\s*', '', cleaned).strip()
        
        start_idx = cleaned.find('{')
        end_idx = cleaned.rfind('}')
        note_data = None
        if start_idx != -1 and end_idx != -1:
            json_str = cleaned[start_idx:end_idx + 1]
            try:
                note_data = json.loads(json_str)
            except json.JSONDecodeError:
                try:
                    note_data = json.loads(_sanitize_json_control_chars(json_str))
                except json.JSONDecodeError:
                    pass
        
        if note_data:
            # Save the generated content as a LessonNote
            ed_level = ({'primary school': 'primary', 'ordinary level': 'ordinary', 'advanced level': 'advanced'}).get(
                (lp.education_level or '').lower(), 'ordinary')
            
            # Format the note content - NOTEBOOK STYLE with proper formatting
            content_parts = []
            
            # Title
            title = note_data.get('title', f"Lesson Notes - {lp.topic}")
            content_parts.append(f"📖 {title}")
            content_parts.append("")
            
            # Summary paragraphs
            content_parts.append("=" * 60)
            content_parts.append("1. MUHTASARI / SUMMARY")
            content_parts.append("=" * 60)
            paragraphs = note_data.get('summary_paragraphs', [])
            if not paragraphs and note_data.get('summary'):
                paragraphs = [note_data['summary']]
            for para in paragraphs:
                content_parts.append(para)
                content_parts.append("")
            
            # Key points
            key_points = note_data.get('key_points', [])
            if key_points:
                content_parts.append("=" * 60)
                content_parts.append("2. NUKTA MUHIMU / KEY POINTS")
                content_parts.append("=" * 60)
                for kp in key_points:
                    content_parts.append(f"  • {kp}")
                content_parts.append("")
            
            # Teaching methods
            tm = note_data.get('teaching_methods', '')
            if tm:
                content_parts.append("=" * 60)
                content_parts.append("3. MBINU ZA UFUNDISHAJI / TEACHING METHODS")
                content_parts.append("=" * 60)
                content_parts.append(tm)
                content_parts.append("")
            
            # Quiz
            quiz = note_data.get('quiz', [])
            if quiz:
                content_parts.append("=" * 60)
                content_parts.append("4. TATHMINI / ASSESSMENT — MASWALI NA MAJIBU")
                content_parts.append("=" * 60)
                for i, q in enumerate(quiz, 1):
                    content_parts.append(f"\nSwali {i}: {q.get('question', '')}")
                    content_parts.append(f"\nJibu: {q.get('answer', '')}")
                content_parts.append("")
            
            # Conclusion
            conclusion = note_data.get('conclusion', '')
            if conclusion:
                content_parts.append("=" * 60)
                content_parts.append("5. MWONGOZO / CONCLUSION & RECOMMENDATIONS")
                content_parts.append("=" * 60)
                content_parts.append(conclusion)
                content_parts.append("")
            
            note_content = '\n'.join(content_parts)
            
            note = LessonNote.objects.create(
                teacher=teacher,
                teacher_name=teacher.full_name,
                school=teacher.school if teacher.school else None,
                school_name=school_name,
                education_level=ed_level,
                class_name=lp.class_name,
                subject=subject_name,
                topic=lp.topic,
                content=note_content,
            )
            
            return JsonResponse({
                'success': True,
                'note_id': note.id,
                'note_data': note_data,
                'note_html': note_content,
                'created': note.created_at.isoformat(),
            })
        else:
            return JsonResponse({'success': False, 'error': 'AI ilishindwa kuzalisha notes. Jaribu tena.'}, status=422)
            
    except LessonPlan.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Lesson Plan haipatikani'}, status=404)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


@require_POST
def ajax_delete_lesson_note(request):
    """AJAX: Delete a lesson note."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False}, status=401)
    try:
        data = json.loads(request.body)
        note_id = data.get('note_id')
        note = LessonNote.objects.get(id=note_id, teacher=teacher)
        note.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)

# =============================================================================
# TEACHER PROFILE UPDATE — auto-save teacher's class/stream/subject/students
# =============================================================================

@require_POST
def ajax_update_teacher_profile(request):
    """
    Update the TLM teacher's profile with data from scheme/lesson plan forms.
    Ensures OLD users who registered before auto-fill fields were added
    get their class_name, stream, subject, total_boys, total_girls saved.
    """
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Haujasajiliwa. Tafadhali jisajili kwanza.'}, status=401)
    
    try:
        data = json.loads(request.body)
        changed = False
        
        class_name = data.get('class_name', '').strip()
        if class_name and teacher.class_name != class_name:
            teacher.class_name = class_name
            changed = True
        
        stream = data.get('stream', '').strip()
        if stream and teacher.stream != stream:
            teacher.stream = stream
            changed = True
        
        subject_name = data.get('subject_name', '').strip()
        if subject_name:
            subj = Subject.objects.filter(name__iexact=subject_name).first()
            if subj and teacher.subject != subj:
                teacher.subject = subj
                changed = True
        
        total_boys = data.get('total_boys')
        if total_boys:
            try:
                val = int(total_boys)
                if val > 0 and teacher.total_boys != val:
                    teacher.total_boys = val
                    changed = True
            except (ValueError, TypeError):
                pass
        
        total_girls = data.get('total_girls')
        if total_girls:
            try:
                val = int(total_girls)
                if val > 0 and teacher.total_girls != val:
                    teacher.total_girls = val
                    changed = True
            except (ValueError, TypeError):
                pass
        
        # ── Handle theme ──
        theme = data.get('theme', '').strip()
        if theme and teacher.theme != theme:
            teacher.theme = theme
            changed = True
        
        # ── Handle preferred_language ──
        preferred_language = data.get('preferred_language', '').strip()
        if preferred_language and preferred_language in ('auto', 'english', 'kiswahili') and teacher.preferred_language != preferred_language:
            teacher.preferred_language = preferred_language
            changed = True
        
        if changed:
            teacher.save(update_fields=['class_name', 'stream', 'subject', 'total_boys', 'total_girls', 'theme', 'preferred_language'])
            logger.info(f"[Profile] Updated teacher {teacher.id} ({teacher.full_name})")
            return JsonResponse({'success': True, 'updated': True, 'message': 'Profile imehifadhiwa!'})
        
        return JsonResponse({'success': True, 'updated': False, 'message': 'Hakuna mabadiliko.'})
        
    except Exception as e:
        logger.error(f"[Profile Update Error] {e}")
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


# =============================================================================
# SYLLABUS TOPICS & SUBTOPICS — Database-powered (TIE Syllabus)
# =============================================================================

def ajax_get_topics_db(request):
    """
    AJAX: Get topics from DB for a given subject and class name.
    Uses the TIE syllabus data seeded via the seed_tie_syllabus management command.
    Returns a JSON array of topic objects with id and name.
    """
    subject_id = request.GET.get('subject_id')
    class_name = request.GET.get('class_name', '').strip()
    
    if not subject_id or not class_name:
        return JsonResponse({'success': False, 'error': 'subject_id na class_name vinahitajika'}, status=400)
    
    # Map common class names (e.g., "Form 1" matches "Form 1" in DB)
    topics = SubjectTopic.objects.filter(
        subject_id=subject_id,
        class_name__iexact=class_name
    ).order_by('order').values('id', 'name')
    
    # If no topics found, try with broader matching
    if not topics.exists():
        # Try just with the form number (e.g., "1" from "Form 1")
        words = class_name.split()
        for w in words:
            if w.isdigit():
                topics = SubjectTopic.objects.filter(
                    subject_id=subject_id,
                    class_name__icontains=w
                ).order_by('order').values('id', 'name')
                break
    
    return JsonResponse({
        'success': True,
        'topics': list(topics),
        'count': topics.count(),
        'source': 'database'
    })


def ajax_get_subtopics_db(request):
    """
    AJAX: Get subtopics from DB for a given topic.
    Uses the TIE syllabus data seeded via the seed_tie_syllabus management command.
    Returns a JSON array of subtopic strings.
    """
    topic_id = request.GET.get('topic_id')
    
    if not topic_id:
        return JsonResponse({'success': False, 'error': 'topic_id inahitajika'}, status=400)
    
    subtopics = TopicSubtopic.objects.filter(
        topic_id=topic_id
    ).order_by('order').values('id', 'name')
    
    return JsonResponse({
        'success': True,
        'subtopics': list(subtopics),
        'count': subtopics.count(),
        'source': 'database'
    })


# =============================================================================
# DIAGNOSTIC — test AI API keys directly on Railway
# =============================================================================

def ajax_ai_diagnostic(request):
    """Test all configured AI API keys and return status. No auth required."""
    import requests as _req
    results = {}
    
    # 1. Test OpenRouter key — via BOTH OpenAI library (old) AND direct HTTP (new)
    or_key = os.environ.get("OPENROUTER_API_KEY", "") or OPENROUTER_API_KEY or ""
    if or_key:
        or_status = {}
        # 1a. Via OpenAI library (old method)
        try:
            from openai import OpenAI
            client_test = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=or_key, timeout=10)
            resp = client_test.chat.completions.create(
                model="deepseek/deepseek-chat",
                messages=[{"role": "user", "content": "Say OK"}],
                max_tokens=5,
            )
            or_status["via_openai_lib"] = "✅ OK"
        except Exception as e:
            or_status["via_openai_lib"] = f"❌ {str(e)[:100]}"
        # 1b. Via direct HTTP (new method, same as ai_utils.py)
        try:
            resp = _req.post(
                "https://openrouter.ai/api/v1/chat/completions",
                json={
                    "model": "deepseek/deepseek-chat",
                    "messages": [{"role": "user", "content": "Say OK"}],
                    "max_tokens": 5,
                },
                headers={
                    "Authorization": f"Bearer {or_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://tlm-tanzania.railway.app",
                    "X-Title": "TLM Tanzania",
                },
                timeout=15,
            )
            if resp.status_code == 200:
                or_status["via_direct_http"] = "✅ OK"
            else:
                or_status["via_direct_http"] = f"❌ HTTP {resp.status_code}: {resp.text[:100]}"
        except Exception as e:
            or_status["via_direct_http"] = f"❌ {str(e)[:100]}"
        results["openrouter"] = or_status
    else:
        results["openrouter"] = {"status": "⚠️ NOT SET"}
    
    # 2. Test Groq key
    groq_key = os.environ.get("GROQ_API_KEY", "") or GROQ_API_KEY or ""
    if groq_key:
        try:
            from groq import Groq
            client_test = Groq(api_key=groq_key)
            resp = client_test.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{"role": "user", "content": "Say OK"}],
                max_tokens=5,
            )
            results["groq"] = {"status": "✅ OK", "response": resp.choices[0].message.content[:50]}
        except Exception as e:
            results["groq"] = {"status": "❌ FAILED", "error": str(e)[:200]}
    else:
        results["groq"] = {"status": "⚠️ NOT SET"}
    
    # 3. Test Gemini key
    gem_key = os.environ.get("GOOGLE_API_KEY", "") or GOOGLE_API_KEY or ""
    if gem_key:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gem_key}"
            payload = {"contents": [{"parts": [{"text": "Say OK"}]}]}
            resp = _req.post(url, json=payload, timeout=10)
            if resp.status_code == 200:
                results["gemini"] = {"status": "✅ OK"}
            else:
                results["gemini"] = {"status": "❌ FAILED", "http": resp.status_code, "error": resp.text[:200]}
        except Exception as e:
            results["gemini"] = {"status": "❌ FAILED", "error": str(e)[:200]}
    else:
        results["gemini"] = {"status": "⚠️ NOT SET"}
    
    return JsonResponse({"diagnostic": results})


@csrf_exempt
def ajax_generate_all_lessons(request):
    """Generate lesson plans for ALL topics of a subject in one batch."""
    if client is None:
        return JsonResponse({'success': False, 'error': 'AI haitumiki'}, status=503)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required'}, status=400)

    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Jisajili kwanza'}, status=401)

    try:
        data = json.loads(request.body)
        subject_name = data.get('subject', '').strip()
        class_name = data.get('class_name', '').strip()
        education_level = data.get('education_level', '')
        stream = data.get('stream', '')
        term = data.get('term', 'I')
        year = data.get('year', 2026)
        duration = int(data.get('duration', 40))
        total_boys = data.get('total_boys', '')
        total_girls = data.get('total_girls', '')
        teacher_name = data.get('teacher_name', '')
        school_name = data.get('school_name', '')

        # Get all topics for this subject from syllabus
        subj_obj = Subject.objects.filter(name__iexact=subject_name).first()
        if not subj_obj:
            return JsonResponse({'success': False, 'error': f'Somo "{subject_name}" halipatikani'}, status=404)
        
        topics = SubjectTopic.objects.filter(subject=subj_obj, class_name=class_name).order_by('order')
        
        # ── Fallback: if no syllabus topics in DB, use AI to generate a topic list ──
        topic_names = []
        if not topics:
            try:
                # Use AI to get topic names for this subject
                ai_prompt = f"List 6-10 main topics for {subject_name} {class_name} (Tanzanian TIE syllabus). Return as a comma-separated list. No numbers. No extra text."
                ai_resp = client.models.generate_content(model=model_name, contents=ai_prompt)
                ai_text = ai_resp.text.strip()
                # Parse comma-separated list
                topic_names = [t.strip() for t in ai_text.replace('\n', ',').split(',') if t.strip() and len(t.strip()) > 3]
            except Exception:
                pass
            
            if not topic_names:
                # Hardcoded fallback for common subjects
                common_topics = {
                    'mathematics': ['Numbers', 'Algebra', 'Geometry', 'Trigonometry', 'Statistics', 'Probability'],
                    'biology': ['Cell Structure', 'Classification', 'Nutrition', 'Transport', 'Respiration', 'Reproduction'],
                    'chemistry': ['Matter', 'Atomic Structure', 'Bonding', 'Chemical Reactions', 'Acids and Bases', 'Organic Chemistry'],
                    'physics': ['Mechanics', 'Heat', 'Light', 'Sound', 'Electricity', 'Magnetism'],
                    'english': ['Parts of Speech', 'Tenses', 'Comprehension', 'Composition', 'Literature', 'Vocabulary'],
                    'kiswahili': ['Sarufi', 'Ufahamu', 'Insha', 'Fasihi', 'Msamiati', 'Matumizi ya Lugha'],
                    'geography': ['Location', 'Climate', 'Population', 'Agriculture', 'Mining', 'Tourism'],
                    'history': ['Early Man', 'Development of Agriculture', 'Trade', 'Colonialism', 'Independence', 'Modern Africa'],
                }
                subject_lower = subject_name.lower()
                for key, topics_list in common_topics.items():
                    if key in subject_lower or subject_lower in key:
                        topic_names = topics_list
                        break
                if not topic_names:
                    topic_names = [f"Topic {i+1}" for i in range(6)]
            
            # Create mock topic objects with .name attribute
            class MockTopic:
                def __init__(self, name):
                    self.name = name
            topics = [MockTopic(t) for t in topic_names]
        
        full_class = f"{class_name}{stream}" if stream else class_name
        intro_time = max(5, int(duration * 0.15))
        dev_time = max(10, int(duration * 0.40))
        design_time = max(8, int(duration * 0.30))
        real_time = max(5, int(duration * 0.15))

        # Handle missing school
        school_obj = teacher.school
        if not school_obj and school_name:
            school_obj = School.objects.filter(name__iexact=school_name).first()
        if not school_obj:
            return JsonResponse({'success': False, 'error': 'Shule yako haijapatiKANA kwenye mfumo. Sasisha wasifu wako.'}, status=400)

        results = []
        errors = []

        for idx, topic in enumerate(topics):
            try:
                # Get first subtopic (only for real DB topics, skip for fallback MockTopic)
                subtopic_name = ''
                if hasattr(topic, '_meta'):
                    subtopic = TopicSubtopic.objects.filter(topic=topic).order_by('order').first()
                    subtopic_name = subtopic.name if subtopic else ''

                prompt = f"""Generate a TEACHER'S LESSON PLAN for a Tanzanian {education_level} classroom.

============================================
PRIME MINISTER'S OFFICE
REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT
TEACHER'S LESSON PLAN
============================================

School: {school_name or '[School Name]'}
Teacher's Name: {teacher_name}
Subject: {subject_name}
Form/Class: {full_class}
Date: {datetime.now().strftime("%d/%m/%Y")}

Main Topic: {topic.name}
Sub-topic: {subtopic_name or 'N/A'}

Main Competence: Numbered format from Tanzanian syllabus for {subject_name} {full_class}.
Specific Competence: The specific competence for {topic.name}.

Term: {term}, Year: {year}
Duration: {duration} minutes

Content MUST relate to Subject: {subject_name}, Class: {full_class}, Topic: "{topic.name}".

Lesson Development uses the IDDR Model (Introduction, Competence Development, Design, Realisation).
Use CBC methodologies: Brainstorming, Group discussion, Jigsaw, Q&A, Demonstration.

Output ONLY valid JSON with this EXACT structure.
⚠️ LANGUAGE REMINDER: ALL text values MUST be in the language specified by the LANGUAGE instruction above. If Kiswahili -> ALL values KISWAHILI. If English -> ALL values ENGLISH. CHECK every field before outputting!
{{
    "main_competence": "[Competence in the language specified above]",
    "specific_competence": "[Specific competence in the language specified above]",
    "main_activity": "[Main activity in the language specified above]",
    "specific_activity": "[Specific activity in the language specified above]",
    "teaching_resources": "[Resources in the language specified above]",
    "references": "[References in the language specified above]",
    "lesson_development": [
        {{"stage": "[Stage 1 - language specified]", "time": "{intro_time:02d}", "teaching_activities": "[language specified]", "learning_activities": "[language specified]", "assessment_criteria": "[language specified]"}},
        {{"stage": "[Stage 2 - language specified]", "time": "{dev_time:02d}", "teaching_activities": "[language specified]", "learning_activities": "[language specified]", "assessment_criteria": "[language specified]"}},
        {{"stage": "[Stage 3 - language specified]", "time": "{design_time:02d}", "teaching_activities": "[language specified]", "learning_activities": "[language specified]", "assessment_criteria": "[language specified]"}},
        {{"stage": "[Stage 4 - language specified]", "time": "{real_time:02d}", "teaching_activities": "[language specified]", "learning_activities": "[language specified]", "assessment_criteria": "[language specified]"}}
    ],
    "remarks": "[Remarks in the language specified above]"
}}

Return ONLY the JSON object. ALL content must be specific to {topic.name}."""

                response = client.models.generate_content(model=model_name, contents=prompt)
                response_text = response.text

                cleaned = re.sub(r'```json\\s*', '', response_text)
                cleaned = re.sub(r'```\\s*', '', cleaned).strip()

                start_idx = cleaned.find('{')
                end_idx = cleaned.rfind('}')
                lesson_data = None
                if start_idx != -1 and end_idx != -1:
                    json_str = cleaned[start_idx:end_idx + 1]
                    try:
                        lesson_data = json.loads(json_str)
                    except json.JSONDecodeError:
                        try:
                            lesson_data = json.loads(_sanitize_json_control_chars(json_str))
                        except json.JSONDecodeError:
                            pass

                if lesson_data:
                    # Convert teaching_resources to list if it's a string
                    tr = lesson_data.get('teaching_resources', '')
                    if isinstance(tr, str):
                        tr = [x.strip() for x in tr.split(',') if x.strip()] or ['TIE textbook']
                    
                    # Save lesson plan to DB
                    lp = LessonPlan.objects.create(
                        student=None,
                        school=school_obj,
                        subject=subj_obj,
                        class_name=class_name,
                        term=term,
                        year=int(year),
                        teacher_name=teacher_name or teacher.full_name,
                        topic=topic.name,
                        subtopic=subtopic_name or '',
                        date=timezone.now().date(),
                        duration=duration,
                        education_level=({'primary school': 'primary', 'ordinary level': 'ordinary', 'advanced level': 'advanced'}).get((education_level or '').lower(), 'ordinary'),
                        main_competence=lesson_data.get('main_competence', ''),
                        specific_competence=lesson_data.get('specific_competence', ''),
                        previous_knowledge=lesson_data.get('main_activity', lesson_data.get('specific_activity', '')),
                        learning_objectives=[lesson_data.get('specific_activity', 'By the end students should be able to...')],
                        teaching_methods=[],
                        teaching_resources=tr,
                        lesson_development=lesson_data.get('lesson_development', []),
                        remarks=lesson_data.get('remarks', ''),
                        generated_by_ai=True,
                    )
                    results.append({
                        'id': lp.id,
                        'topic': topic.name,
                        'subtopic': subtopic_name,
                        'success': True,
                    })
                else:
                    errors.append(f"{topic.name}: AI parsing failed")
                    results.append({
                        'topic': topic.name,
                        'success': False,
                        'error': 'AI parsing failed',
                    })
            except Exception as topic_err:
                errors.append(f"{topic.name}: {str(topic_err)[:100]}")
                results.append({
                    'topic': topic.name,
                    'success': False,
                    'error': str(topic_err)[:100],
                })

        return JsonResponse({
            'success': True,
            'results': results,
            'total': len(topics),
            'generated': sum(1 for r in results if r.get('success')),
            'failed': sum(1 for r in results if not r.get('success')),
            'errors': errors[:5],
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


# =============================================================================
# AJAX: Generate ONE lesson plan (for topic-by-topic progress bar)
# =============================================================================

@require_POST
def ajax_generate_one_lesson(request):
    """Generate a single lesson plan for a given topic. Used by the frontend
    to iterate topic-by-topic with real-time progress."""
    if client is None:
        return JsonResponse({'success': False, 'error': 'AI haitumiki'}, status=503)
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST required'}, status=400)

    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Jisajili kwanza'}, status=401)

    try:
        data = json.loads(request.body)
        subject_id = data.get('subject_id', '')
        subject_name = data.get('subject', '').strip()
        class_name = data.get('class_name', '').strip()
        education_level = data.get('education_level', '')
        stream = data.get('stream', '')
        term = data.get('term', 'I')
        year = data.get('year', 2026)
        duration = int(data.get('duration', 40))
        total_boys = data.get('total_boys', '')
        total_girls = data.get('total_girls', '')
        teacher_name = data.get('teacher_name', '')
        school_name = data.get('school_name', '')
        topic_name = data.get('topic', '').strip()
        subtopic_name = data.get('subtopic', '').strip()
        topic_index = int(data.get('topic_index', 0))
        total_topics = int(data.get('total_topics', 1))

        if not topic_name:
            return JsonResponse({'success': False, 'error': 'Topic inahitajika'}, status=400)

        full_class = f"{class_name}{stream}" if stream else class_name
        intro_time = max(5, int(duration * 0.15))
        dev_time = max(10, int(duration * 0.40))
        design_time = max(8, int(duration * 0.30))
        real_time = max(5, int(duration * 0.15))

        # Handle missing school (match ajax_generate_lessonplan leniency)
        school_obj = teacher.school
        if not school_obj and school_name:
            school_obj = School.objects.filter(name__iexact=school_name).first()
        # Allow empty school - LP saves without school if not found

        # Get subject object
        subj_obj = Subject.objects.filter(id=subject_id).first() or Subject.objects.filter(name__iexact=subject_name).first()
        if not subj_obj:
            return JsonResponse({'success': False, 'error': f'Somo "{subject_name}" halipatikani'}, status=404)

        # ── Language: manual selection or auto-detect ──
        _lp_tlm = teacher  # Already fetched
        _lp_language = data.get('language', getattr(_lp_tlm, 'preferred_language', 'auto') if _lp_tlm else 'auto')
        _lp_school_level = _lp_tlm.school.level if _lp_tlm and _lp_tlm.school else ''
        lp_language_instruction = _get_lp_language_instruction(_lp_language, subject_name, _lp_school_level)

        # Build prompt
        prompt = f"""Generate a TEACHER'S LESSON PLAN for a Tanzanian {education_level} classroom.

============================================
PRIME MINISTER'S OFFICE
REGIONAL ADMINISTRATION AND LOCAL GOVERNMENT
TEACHER'S LESSON PLAN
============================================

School: {school_name or '[School Name]'}
Teacher's Name: {teacher_name}
Subject: {subject_name}
Form/Class: {full_class}
Date: {datetime.now().strftime('%d/%m/%Y')}

Main Topic: {topic_name}
Sub-topic: {subtopic_name or 'N/A'}

Main Competence: Numbered format from Tanzanian syllabus for {subject_name} {full_class}.
Specific Competence: The specific competence for {topic_name}.

Term: {term}, Year: {year}
Duration: {duration} minutes
{lp_language_instruction}

Content MUST relate to Subject: {subject_name}, Class: {full_class}, Topic: \"{topic_name}\".

Lesson Development uses the IDDR Model (Introduction, Competence Development, Design, Realisation).
Use CBC methodologies: Brainstorming, Group discussion, Jigsaw, Q&A, Demonstration.

Output ONLY valid JSON with this EXACT structure.
⚠️ LANGUAGE REMINDER: ALL text values MUST be in the language specified by the LANGUAGE instruction above. If Kiswahili -> ALL values KISWAHILI. If English -> ALL values ENGLISH. CHECK every field before outputting!
{{
    \"main_competence\": \"[Competence in the language specified above]\",
    \"specific_competence\": \"[Specific competence in the language specified above]\",
    \"main_activity\": \"[Main activity in the language specified above]\",
    \"specific_activity\": \"[Specific activity in the language specified above]\",
    \"teaching_resources\": \"[Resources in the language specified above]\",
    \"references\": \"[References in the language specified above]\",
    \"lesson_development\": [
        {{\"stage\": \"[Stage 1 - language specified]\", \"time\": \"{intro_time:02d}\", \"teaching_activities\": \"[language specified]\", \"learning_activities\": \"[language specified]\", \"assessment_criteria\": \"[language specified]\"}},
        {{\"stage\": \"[Stage 2 - language specified]\", \"time\": \"{dev_time:02d}\", \"teaching_activities\": \"[language specified]\", \"learning_activities\": \"[language specified]\", \"assessment_criteria\": \"[language specified]\"}},
        {{\"stage\": \"[Stage 3 - language specified]\", \"time\": \"{design_time:02d}\", \"teaching_activities\": \"[language specified]\", \"learning_activities\": \"[language specified]\", \"assessment_criteria\": \"[language specified]\"}},
        {{\"stage\": \"[Stage 4 - language specified]\", \"time\": \"{real_time:02d}\", \"teaching_activities\": \"[language specified]\", \"learning_activities\": \"[language specified]\", \"assessment_criteria\": \"[language specified]\"}}
    ],
    \"remarks\": \"[Remarks in the language specified above]\"
}}

Return ONLY the JSON object. ALL content must be specific to {topic_name}."""

        response = client.models.generate_content(model=model_name, contents=prompt)
        response_text = response.text

        cleaned = re.sub(r'```json\\s*', '', response_text)
        cleaned = re.sub(r'```\\s*', '', cleaned).strip()

        start_idx = cleaned.find('{')
        end_idx = cleaned.rfind('}')
        lesson_data = None
        if start_idx != -1 and end_idx != -1:
            json_str = cleaned[start_idx:end_idx + 1]
            try:
                lesson_data = json.loads(json_str)
            except json.JSONDecodeError:
                try:
                    lesson_data = json.loads(_sanitize_json_control_chars(json_str))
                except json.JSONDecodeError:
                    pass

        if lesson_data:
            tr = lesson_data.get('teaching_resources', '')
            if isinstance(tr, str):
                tr = [x.strip() for x in tr.split(',') if x.strip()] or ['TIE textbook']

            ed_level = ({'primary school': 'primary', 'ordinary level': 'ordinary', 'advanced level': 'advanced'}).get((education_level or '').lower(), 'ordinary')

            lp = LessonPlan.objects.create(
                student=None,
                school=school_obj,
                subject=subj_obj,
                class_name=class_name,
                term=term,
                year=int(year),
                teacher_name=teacher_name or teacher.full_name,
                topic=topic_name,
                subtopic=subtopic_name or '',
                date=timezone.now().date(),
                duration=duration,
                education_level=ed_level,
                main_competence=lesson_data.get('main_competence', ''),
                specific_competence=lesson_data.get('specific_competence', ''),
                previous_knowledge=lesson_data.get('main_activity', lesson_data.get('specific_activity', '')),
                learning_objectives=[lesson_data.get('specific_activity', 'By the end students should be able to...')],
                teaching_methods=[],
                teaching_resources=tr,
                lesson_development=lesson_data.get('lesson_development', []),
                remarks=lesson_data.get('remarks', ''),
                generated_by_ai=True,
            )

            return JsonResponse({
                'success': True,
                'lp_id': lp.id,
                'topic': topic_name,
                'topic_index': topic_index,
                'total_topics': total_topics,
                'data': lesson_data,
            })
        else:
            return JsonResponse({
                'success': False,
                'error': 'AI ilishindwa kuchakata response',
                'topic': topic_name,
                'topic_index': topic_index,
                'total_topics': total_topics,
            }, status=422)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'success': False,
            'error': str(e)[:200],
            'topic': data.get('topic', '') if 'data' in dir() else '',
        }, status=500)


def my_lessons(request):
    """Show all lesson plans generated by this teacher, grouped by subject."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:my_lessons')}")
    
    # Get all lesson plans for this teacher (by teacher_name + school)
    lessons = LessonPlan.objects.filter(
        teacher_name=teacher.full_name,
        school=teacher.school,
    ).select_related('subject').order_by('-created_at')[:100]
    
    # Group by subject
    from collections import defaultdict
    grouped = defaultdict(list)
    for lp in lessons:
        subj_name = lp.subject.name if lp.subject else 'General'
        grouped[subj_name].append(lp)
    
    return render(request, 'curriculum/my_lessons.html', {
        'teacher': teacher,
        'grouped_lessons': dict(grouped),
        'total_lessons': lessons.count(),
    })


# =============================================================================
# EXAM GENERATOR — NECTA FORMAT (Primary / O-Level / A-Level)
# =============================================================================

_EXAM_TYPE_LABELS = dict(GeneratedExam.EXAM_TYPE_CHOICES)
_EXAM_LEVEL_LABELS = dict(GeneratedExam.EDUCATION_LEVEL_CHOICES)


def _safe_marks(val):
    """Safely convert a marks value (int, float or numeric string) to int."""
    try:
        return int(float(val))
    except (ValueError, TypeError):
        return 0


def _exam_probe(education_level, class_name, language='english'):
    """Build an unsaved GeneratedExam so we reuse the model's NECTA helpers
    (assessment_name / necta_header_lines / duration_display) — single source of truth."""
    return GeneratedExam(
        education_level=education_level,
        class_name=class_name or '',
        language=language or 'english',
    )


def _exam_necta_header_lines(education_level, class_name, language):
    """Official NECTA paper header lines — via GeneratedExam.necta_header_lines."""
    return _exam_probe(education_level, class_name, language).necta_header_lines


def _parse_exam_json(response_text):
    """Parse AI exam response into a dict with 'title', 'instructions', 'sections'."""
    cleaned = re.sub(r'```(?:json)?\s*', '', response_text)
    cleaned = re.sub(r'```\s*', '', cleaned).strip()

    def _extract_obj(text):
        start = text.find('{')
        if start == -1:
            return None
        depth = 0
        in_str = False
        esc = False
        for i in range(start, len(text)):
            ch = text[i]
            if esc:
                esc = False
                continue
            if ch == '\\' and in_str:
                esc = True
                continue
            if ch == '"' and not esc:
                in_str = not in_str
                continue
            if in_str:
                continue
            if ch == '{':
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0:
                    candidate = text[start:i + 1]
                    try:
                        return json.loads(candidate)
                    except json.JSONDecodeError:
                        try:
                            return json.loads(_sanitize_json_control_chars(candidate))
                        except json.JSONDecodeError:
                            continue
        # Fallback: fix unbalanced braces
        candidate = text[start:]
        if candidate.count('"') % 2 == 1:
            candidate += '"'
        open_b = candidate.count('{')
        close_b = candidate.count('}')
        if open_b > close_b:
            candidate += '}' * (open_b - close_b)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            try:
                return json.loads(_sanitize_json_control_chars(candidate))
            except json.JSONDecodeError:
                return None

    data = _extract_obj(cleaned)
    if data is None:
        # Sometimes AI returns a bare array of sections
        arr_start = cleaned.find('[')
        if arr_start != -1:
            try:
                arr = json.loads(cleaned[arr_start:])
                if isinstance(arr, list):
                    data = {'sections': arr}
            except json.JSONDecodeError:
                pass

    if not isinstance(data, dict):
        data = {'sections': []}

    sections = data.get('sections') or data.get('questions') or []
    # Normalize: if sections is a flat list of questions, wrap into one section
    normalized = []
    if isinstance(sections, list):
        for sec in sections:
            if isinstance(sec, dict):
                qs = sec.get('questions') or []
                # If this dict looks like a single question, wrap it
                if not qs and (sec.get('text') or sec.get('question')):
                    qs = [sec]
                if qs or sec.get('section'):
                    normalized.append({
                        'section': sec.get('section') or sec.get('name') or 'A',
                        'instructions': sec.get('instructions') or '',
                        'questions': qs,
                    })
    if not normalized:
        # Flat list of questions
        normalized = [{'section': 'A', 'instructions': '', 'questions': sections}]

    # Clean questions
    for sec in normalized:
        clean_qs = []
        for i, q in enumerate(sec.get('questions', []) or []):
            if isinstance(q, dict):
                clean_qs.append({
                    'number': q.get('number') or (i + 1),
                    'text': q.get('text') or q.get('question') or '',
                    'marks': _safe_marks(q.get('marks')),
                    'answer': q.get('answer') or q.get('correct_answer') or '',
                    'topic': q.get('topic') or '',
                })
        sec['questions'] = clean_qs

    total = sum(
        q.get('marks', 0)
        for sec in normalized for q in sec.get('questions', [])
    )

    return {
        'title': data.get('title') or '',
        'instructions': data.get('instructions') or '',
        'sections': normalized,
        'total_marks': data.get('total_marks') or total,
    }


def _exam_language_instruction(language, education_level, subject_name):
    """Strong language rule for exam generation.

    Returns a tuple: (instruction_text, resolved_language)
    - 'auto' (default): Primary (except English subject) → KISWAHILI;
      Secondary / A-Level (except Kiswahili subject) → ENGLISH.
    - Explicit 'kiswahili' / 'english' choice wins.
    """
    subj_lower = (subject_name or '').lower()
    lang = (language or 'auto').lower()
    level_lower = (education_level or '').lower()

    is_english_subject = subj_lower in ('english', 'english language')
    is_kiswahili_subject = subj_lower in ('kiswahili', 'swahili')
    is_primary = level_lower == 'primary'

    # ── Resolve 'auto' ──
    if lang == 'auto':
        if is_kiswahili_subject:
            lang = 'kiswahili'
        elif is_primary and not is_english_subject:
            lang = 'kiswahili'
        else:
            lang = 'english'

    if lang == 'kiswahili':
        return (
            "\n" +
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
            "🚨 KANUNI YA LUGHA: KILA KITU (maswali, majibu, maelekezo ya sections, na\n" +
            "instructions za mtihani mzima) LAZIMA kiandikwe kwa KISWAHILI KAMILI.\n" +
            "Hakuna Kiingereza kinachokubalika katika question text au answers.\n" +
            "(Isipokuwa majina ya istilahi za kipekee kama 'cell', 'DNA' n.k.)\n" +
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n",
            'kiswahili',
        )
    # ENGLISH default (Secondary / A-Level / Primary English subject)
    return (
        "\n" +
        "!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n" +
        "🚨 LANGUAGE RULE: EVERYTHING (questions, answers, section instructions and\n" +
        "the whole paper instructions) MUST be written in ENGLISH ONLY. Do NOT mix\n" +
        "any Kiswahili into the question text or answers.\n" +
        "!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!\n",
        'english',
    )


def exam_generator_view(request):
    """Exam generator page — teacher picks level/class/subject/type then AI generates."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:exam_generator')}")

    education_levels = EducationLevel.objects.all().order_by('order')
    subjects = Subject.objects.all().order_by('name')

    from django.db.models import Count as _C
    my_exams = GeneratedExam.objects.filter(teacher=teacher).order_by('-created_at')[:30]
    exams_count = GeneratedExam.objects.filter(teacher=teacher).count()
    exams_by_type = [
        {'exam_type': t['exam_type'],
         'label': _EXAM_TYPE_LABELS.get(t['exam_type'], t['exam_type']),
         'count': t['count']}
        for t in GeneratedExam.objects.filter(teacher=teacher)
        .values('exam_type').annotate(count=_C('id')).order_by('-count')
    ]
    exams_by_subject = list(
        GeneratedExam.objects.filter(teacher=teacher).values('subject_name')
        .annotate(count=_C('id')).order_by('-count')
    )

    teacher_info = {
        'name': teacher.full_name,
        'school': teacher.school.name if teacher.school else '',
        'class_name': teacher.class_name or '',
        'subject_id': teacher.subject_id or '',
        'subject_name': teacher.subject.name if teacher.subject else '',
    }

    return render(request, 'curriculum/exam_generator.html', {
        'teacher': teacher,
        'education_levels': education_levels,
        'subjects': subjects,
        'my_exams': my_exams,
        'exams_count': exams_count,
        'exams_by_type': exams_by_type,
        'exams_by_subject': exams_by_subject,
        'teacher_info': teacher_info,
    })


@csrf_exempt
def ajax_generate_exam(request):
    """AI generates an exam in NECTA format and saves it."""
    if client is None:
        return JsonResponse({'success': False, 'error': 'Huduma ya AI haitumiki. Ufunguo wa API haujawekwa.'}, status=503)
    if request.method != 'POST':
        return JsonResponse({'success': False}, status=400)

    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Tafadhali jisajili kwanza.'}, status=401)

    try:
        data = json.loads(request.body)
        education_level = data.get('education_level', 'ordinary')
        class_name = (data.get('class_name') or '').strip()
        subject_id = data.get('subject_id')
        subject_name = (data.get('subject') or '').strip()
        exam_type = data.get('exam_type', 'TEST')
        duration = int(data.get('duration', 120))
        total_marks = int(data.get('total_marks', 100))
        question_count = int(data.get('question_count', 20))
        language = data.get('language', 'english')
        topics_input = (data.get('topics') or '').strip()
        term = (data.get('term') or '').strip()
        year = int(data.get('year', 2026))

        if not class_name or not subject_name:
            return JsonResponse({'success': False, 'error': 'Tafadhali jaza darasa (class) na somo (subject).'}, status=400)

        subject_obj = Subject.objects.filter(id=subject_id).first() if subject_id else None
        school_name = teacher.school.name if teacher.school else ''
        teacher_name = teacher.full_name

        level_label = _EXAM_LEVEL_LABELS.get(education_level, education_level)
        type_label = _EXAM_TYPE_LABELS.get(exam_type, exam_type)
        lang_rule, resolved_language = _exam_language_instruction(language, education_level, subject_name)

        topics_hint = f"COVER THESE TOPICS (spread questions evenly): {topics_input}" if topics_input else "COVER the official TIE syllabus topics for this subject/class."

        # NECTA-style time label, e.g. 150 → '2:30 Hours' / 'Saa 2:30'
        _dur_h = duration // 60
        _dur_m = duration % 60
        _time_label = f"{_dur_h}:{_dur_m:02d} Hours" if _dur_h else f"{_dur_m} minutes"
        if resolved_language == 'kiswahili':
            _time_label = f"Saa {_dur_h}:{_dur_m:02d}" if _dur_h else f"Dakika {_dur_m}"
        # Official NECTA paper header lines for the prompt
        _hdr = _exam_necta_header_lines(education_level, class_name, resolved_language)
        _header_text = '\n'.join(_hdr)

        # ── Level-specific NECTA format (Primary = PSLE 2024 / SFNA, Secondary = FTNA/CSEE CBA) ──
        is_primary = education_level == 'primary'
        cls_lower = (class_name or '').lower()
        is_lower_primary = bool(re.search(r'\b(std|standard|darasa(?:\s+la)?)\s*\.?\s*[1-4]\b', cls_lower))
        is_upper_primary = bool(re.search(r'\b(std|standard|darasa(?:\s+la)?)\s*\.?\s*[5-7]\b', cls_lower))

        if is_primary:
            if is_lower_primary or (not is_upper_primary):
                # SFNA (Std 1-4): Sections A + B, ~50 marks, 90 min
                level_format_rule = (
                    "PRIMARY SCHOOL FORMAT (SFNA — Standard Four National Assessment, NECTA exact format):\n"
                    "1. The paper consists of SECTIONS A and B ONLY (no Section C). Total marks = 50.\n"
                    "2. Section A (ALAMA 20): OBJECTIVE — multiple-choice items (chaguo A, B, C, D), matching items (linganisha), and true/false items. Each item carries 1 mark.\n"
                    "3. Section B (ALAMA 30): SHORT ANSWER / structured items — fill-in-the-blanks, short responses, simple problem solving with sub-items (a), (b), (c). Each item carries 2 marks.\n"
                    "4. Questions MUST be based on the official TIE new syllabus for {subject_name} {class_name} (Tanzania primary).\n"
                    "5. Use simple age-appropriate Kiswahili sentences (unless subject is English).\n"
                    "6. Time allowed: 1 hour 30 minutes.\n"
                    "7. Number questions continuously (1, 2, 3...).\n"
                    "8. Marks must sum to approximately {total_marks}."
                )
            else:
                # PSLE 2024 (Std 5-7): 6-8 questions, A=20/B=20/C=10, 50 marks, 100 min
                level_format_rule = (
                    "PRIMARY SCHOOL FORMAT (PSLE 2024 — Primary School Leaving Examination, improved NECTA format):\n"
                    "1. The paper consists of SECTIONS A, B and C with a TOTAL OF 6 TO 8 QUESTIONS only (NOT 30+). Total marks = 50.\n"
                    "2. Section A (ALAMA 20): OBJECTIVE — multiple-choice items (chaguo A, B, C, D or A-E), matching items (linganisha), and true/false items. Each item carries 1 mark.\n"
                    "3. Section B (ALAMA 20): SHORT ANSWER — fill-in-the-blanks, complete sentences, short responses based on passages/stories. Each item carries 2 marks.\n"
                    "4. Section C (ALAMA 10): ONE (1) structured question (composition / picture interpretation / map reading) divided into five (5) sub-items (a)-(e), each carrying 2 marks. NO long essay.\n"
                    "5. Follow the real subject question counts: Hisabati/Mathematics = 8 questions, Sayansi na Teknolojia = 8 questions, Kiswahili = 6 questions, English = 7 questions, Maarifa ya Jamii = 7 questions.\n"
                    "6. Questions MUST be based on the official TIE new syllabus for {subject_name} {class_name} (Tanzania primary).\n"
                    "7. Use simple age-appropriate Kiswahili sentences (unless subject is English).\n"
                    "8. Time allowed: 1 hour 40 minutes.\n"
                    "9. Number questions continuously (1, 2, 3...).\n"
                    "10. Marks must sum to approximately {total_marks}."
                )
        elif education_level == 'advanced':
            # ACSEE (Form 5-6): 3 hours, 100 marks — A=20 objective, B=40 (4×10), C=40 (2 essays of 20, answer one)
            level_format_rule = (
                "ADVANCED LEVEL FORMAT (ACSEE — Advanced Certificate of Secondary Education, NECTA exact style):\n"
                "1. The paper consists of sections A, B and C with a total of EIGHT (8) questions. Total marks = 100.\n"
                "2. Section A (20 marks): OBJECTIVE — multiple-choice items (chaguo A, B, C, D), matching items and true/false items. Each item carries 1 mark.\n"
                "3. Section B (40 marks): FOUR (4) short-answer / structured questions, each carrying 10 marks (4 x 10 = 40). Calculations, explanations, diagram interpretation, applying methods.\n"
                "4. Section C (40 marks): TWO (2) essay / long-answer questions, each carrying 20 marks. Candidates answer ONE (1) question from this section.\n"
                "5. Questions MUST be based on the official TIE advanced level syllabus for {subject_name} {class_name} in Tanzania.\n"
                "6. Number questions continuously across the whole paper (1 to 8).\n"
                "7. Difficulty increases from Section A to Section C.\n"
                "8. Total marks must sum to approximately {total_marks}. Time allowed: 3 hours."
            )
        else:
            level_format_rule = (
                "SECONDARY SCHOOL FORMAT (FTNA/CSEE — NECTA Competence-Based Assessment, exact format as LusaElimu):\n"
                "1. The paper consists of sections A, B and C with a total of TEN (10) questions. Answer ALL questions.\n"
                "2. Section A (15 marks): TWO (2) objective questions —\n"
                "   - Question 1: MULTIPLE CHOICE — ten (10) items, each carrying one (1) mark (total 10 marks). Each item has four options A, B, C, D on separate lines.\n"
                "   - Question 2: MATCHING ITEMS — five (5) items, each carrying one (1) mark (total 5 marks). Provide List A and List B.\n"
                "3. Section B (70 marks): SEVEN (7) short-answer questions, each carrying 10 marks (7 x 10 = 70 marks). Calculations, explanations, diagram/map interpretation, applying methods.\n"
                "4. Section C (15 marks): ONE (1) essay / long-answer question carrying 15 marks — analysis, problem-solving and real-life application.\n"
                "5. Questions MUST be based on the official TIE new syllabus (CBA) for {subject_name} {class_name} in Tanzania.\n"
                "6. Number questions continuously across the whole paper (1 to 10).\n"
                "7. Difficulty increases from Section A to Section C.\n"
                "8. Total marks must sum to approximately {total_marks}. Time allowed: 2 hours 30 minutes."
            )

        prompt = f"""You are a senior Tanzanian NECTA exam setter and curriculum expert. Generate a complete {type_label} examination paper for {subject_name} — {class_name} ({level_label}) in the official NECTA format.

The paper MUST carry this official NECTA header (print it at the top of the paper):
{_header_text}

EXAM DETAILS:
- Subject: {subject_name} | Class: {class_name} | Level: {level_label}
- Exam type: {type_label} | Year: {year} | Term: {term or 'N/A'}
- Time allowed: {_time_label} | Total marks: {total_marks}
- Number of questions: approximately {question_count}
- School: {school_name} | Teacher: {teacher_name}
- {topics_hint}
{lang_rule}

NECTA FORMAT RULES:
{level_format_rule}

OUTPUT: Return ONLY valid JSON (no markdown, no extra text) with this EXACT structure:
{{
  "title": "{subject_name} — {class_name} {type_label} Examination",
  "instructions": "General instructions for candidates, written EXACTLY like a real NECTA paper in the exam language. Example: 'This paper consists of sections A, B and C with a total of {total_marks} marks. Answer ALL questions in sections A and B and one (1) question from section C. Show all your working. Calculators may be used...' — for primary use: 'Karatasi hii ina sehemu A, B na C zenye jumla ya alama {total_marks}. Jibu maswali YOTE...'",
  "sections": [
    {{
      "section": "A",
      "instructions": "Official section instruction in the exam language, e.g. 'SECTION A (15 Marks): Answer ALL questions in this section.' / 'SEHEMU A (Alama 20): Jibu maswali YOTE katika sehemu hii.'",
      "questions": [
        {{"number": 1, "text": "...", "marks": 2, "answer": "...", "topic": "..."}}
      ]
    }}
  ],
  "total_marks": {total_marks}
}}

The "answer" field MUST contain the correct answer AND brief marking points (for the marking scheme). Question "text" must be the full question including options if multiple choice. Do NOT omit any field."""

        logger.info(f"[Exam Gen] Generating {type_label} for {subject_name} {class_name}")
        response = client.models.generate_content(model=model_name, contents=prompt)
        exam_data = _parse_exam_json(response.text)

        if not exam_data.get('sections'):
            return JsonResponse({'success': False, 'error': 'AI haikurejesha maswali. Jaribu tena.'}, status=500)

        total_qs = sum(len(s.get('questions', [])) for s in exam_data['sections'])
        if total_qs == 0:
            return JsonResponse({'success': False, 'error': 'Hakuna maswali yaliyopatikana. Jaribu tena.'}, status=500)

        # Collect topics covered
        topics_covered = []
        for sec in exam_data['sections']:
            for q in sec.get('questions', []):
                if q.get('topic') and q['topic'] not in topics_covered:
                    topics_covered.append(q['topic'])

        exam = GeneratedExam.objects.create(
            teacher=teacher,
            school=teacher.school,
            title=exam_data.get('title') or f"{subject_name} — {class_name} {type_label}",
            education_level=education_level,
            class_name=class_name,
            subject=subject_obj,
            subject_name=subject_name,
            exam_type=exam_type,
            term=term,
            year=year,
            duration_minutes=duration,
            total_marks=exam_data.get('total_marks') or total_marks,
            instructions=exam_data.get('instructions') or '',
            language=resolved_language,
            questions=exam_data['sections'],
            topics_covered=topics_covered,
        )

        return JsonResponse({
            'success': True,
            'exam_id': exam.id,
            'title': exam.title,
            'question_count': total_qs,
            'total_marks': exam.total_marks,
            'url': reverse('curriculum:exam_detail', args=[exam.id]),
        })
    except Exception as e:
        logger.exception("[Exam Gen] Error")
        return JsonResponse({'success': False, 'error': str(e)[:300]}, status=500)


def exam_detail_view(request, exam_id):
    """View a generated exam — NECTA-style paper + toggleable marking scheme."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:exam_generator')}")
    exam = get_object_or_404(GeneratedExam, id=exam_id, teacher=teacher)
    return render(request, 'curriculum/exam_detail.html', {
        'exam': exam,
        'teacher': teacher,
    })


def my_exams(request):
    """List all generated exams + progress statistics."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:my_exams')}")

    from django.db.models import Count as _C
    exams = GeneratedExam.objects.filter(teacher=teacher).order_by('-created_at')[:100]
    total = GeneratedExam.objects.filter(teacher=teacher).count()
    by_type = [
        {'exam_type': t['exam_type'],
         'label': _EXAM_TYPE_LABELS.get(t['exam_type'], t['exam_type']),
         'count': t['count']}
        for t in GeneratedExam.objects.filter(teacher=teacher)
        .values('exam_type').annotate(count=_C('id')).order_by('-count')
    ]
    by_subject = list(GeneratedExam.objects.filter(teacher=teacher).values('subject_name').annotate(count=_C('id')).order_by('-count'))
    by_class = list(GeneratedExam.objects.filter(teacher=teacher).values('class_name').annotate(count=_C('id')).order_by('-count'))

    return render(request, 'curriculum/my_exams.html', {
        'exams': exams,
        'total': total,
        'by_type': by_type,
        'by_subject': by_subject,
        'by_class': by_class,
        'teacher': teacher,
    })


@require_POST
def ajax_delete_exam(request):
    """Delete a generated exam."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False}, status=401)
    try:
        data = json.loads(request.body)
        exam = GeneratedExam.objects.get(id=data.get('exam_id'), teacher=teacher)
        exam.delete()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)[:200]}, status=500)


def download_exam_pdf(request, exam_id, mode='paper'):
    """Download exam as PDF — mode: 'paper' (questions only) or 'marking' (with answers)."""
    from reportlab.lib.units import cm

    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('login')}?next={request.path}")
    exam = get_object_or_404(GeneratedExam, id=exam_id, teacher=teacher)
    include_answers = mode == 'marking'

    NAVY = colors.HexColor('#0A2B5E')
    GOLD = colors.HexColor('#C8900A')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=1.6 * cm, rightMargin=1.6 * cm,
                            topMargin=1.4 * cm, bottomMargin=1.4 * cm)
    s_title = ParagraphStyle('t', fontName='Helvetica-Bold', fontSize=13, textColor=NAVY, alignment=1, spaceAfter=2)
    s_sub = ParagraphStyle('s', fontName='Helvetica', fontSize=8.5, textColor=colors.HexColor('#4A5568'), alignment=1, spaceAfter=1)
    s_inst = ParagraphStyle('i', fontName='Helvetica', fontSize=8.5, textColor=colors.HexColor('#1A1A2E'), leading=11, spaceAfter=6)
    s_sec = ParagraphStyle('sec', fontName='Helvetica-Bold', fontSize=10, textColor=colors.white, backColor=NAVY, spaceAfter=2, leading=13)
    s_q = ParagraphStyle('q', fontName='Helvetica', fontSize=9, textColor=colors.HexColor('#1A1A2E'), leading=12, spaceAfter=2)
    s_a = ParagraphStyle('a', fontName='Helvetica', fontSize=8.5, textColor=colors.HexColor('#0D4F2B'), leading=11, leftIndent=18, spaceAfter=6)

    story = []
    for line in exam.necta_header_lines:
        story.append(Paragraph(line, s_sub))
    story.append(Paragraph(exam.title, s_title))
    story.append(Paragraph(f"{exam.subject_name}  |  {exam.class_name}  |  {_EXAM_LEVEL_LABELS.get(exam.education_level, exam.education_level)}", s_sub))
    time_label = f"Time: {exam.duration_display}" if exam.language != 'kiswahili' else f"Muda: {exam.duration_display}"
    marks_label = f"Total Marks: {exam.total_marks}" if exam.language != 'kiswahili' else f"Alama Zote: {exam.total_marks}"
    story.append(Paragraph(f"{_EXAM_TYPE_LABELS.get(exam.exam_type, exam.exam_type)}  |  Muhula: {exam.term or '-'}  |  Mwaka: {exam.year}  |  {time_label}  |  {marks_label}", s_sub))
    story.append(HRFlowable(width="100%", thickness=1.5, color=GOLD, spaceBefore=4, spaceAfter=8))

    if exam.instructions:
        story.append(Paragraph(f"<b>INSTRUCTIONS:</b> {exam.instructions}", s_inst))

    sw = exam.language == 'kiswahili'
    for sec in exam.sections:
        sec_label = f"SEHEMU {sec.get('section', 'A')}" if sw else f"SECTION {sec.get('section', 'A')}"
        story.append(Paragraph(sec_label, s_sec))
        if sec.get('instructions'):
            story.append(Paragraph(sec['instructions'], s_inst))
        for q in sec.get('questions', []):
            num = q.get('number', '')
            marks = q.get('marks', 0)
            txt = q.get('text', '')
            marks_suffix = 'alama' if sw else 'marks'
            story.append(Paragraph(f"<b>{num}.</b> {txt} <b>({marks} {marks_suffix})</b>", s_q))
            if include_answers and q.get('answer'):
                ans_label = 'Jibu' if sw else 'Answer'
                story.append(Paragraph(f"<b>{ans_label}:</b> {q['answer']}", s_a))
        story.append(Spacer(1, 8))

    if include_answers:
        story.append(HRFlowable(width="100%", thickness=1.5, color=GOLD, spaceBefore=6))
        story.append(Paragraph("MARKING SCHEME / SCHEME YA KUSAHLIHIA",
                               ParagraphStyle('ms', fontName='Helvetica-Bold', fontSize=10, textColor=NAVY, alignment=1, spaceBefore=4)))

    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Prepared by: {teacher.full_name}", s_sub))
    story.append(Paragraph(f"School: {teacher.school.name if teacher.school else '-'}", s_sub))

    doc.build(story)
    buffer.seek(0)
    suffix = 'marking' if include_answers else 'paper'
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{exam.title[:40].replace(chr(32), "_")}_{suffix}.pdf"'
    return response


def download_exam_word(request, exam_id, mode='paper'):
    """Download exam as Word (.docx)."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('login')}?next={request.path}")
    exam = get_object_or_404(GeneratedExam, id=exam_id, teacher=teacher)
    include_answers = mode == 'marking'

    doc = Document()
    navy = RGBColor(0x0A, 0x2B, 0x5E)
    gold = RGBColor(0xC8, 0x90, 0x0A)
    green = RGBColor(0x0D, 0x4F, 0x2B)

    for line in exam.necta_header_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = navy

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(exam.title)
    tr.bold = True
    tr.font.size = Pt(15)
    tr.font.color.rgb = navy

    time_label = f"Time: {exam.duration_display}" if exam.language != 'kiswahili' else f"Muda: {exam.duration_display}"
    marks_label = f"Total Marks: {exam.total_marks}" if exam.language != 'kiswahili' else f"Alama Zote: {exam.total_marks}"
    for line in [
        f"{exam.subject_name} | {exam.class_name} | {_EXAM_LEVEL_LABELS.get(exam.education_level, exam.education_level)}",
        f"{_EXAM_TYPE_LABELS.get(exam.exam_type, exam.exam_type)} | Muhula: {exam.term or '-'} | Mwaka: {exam.year} | {time_label} | {marks_label}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    if exam.instructions:
        ip = doc.add_paragraph()
        ip.add_run('INSTRUCTIONS: ').bold = True
        ip.add_run(exam.instructions)

    sw = exam.language == 'kiswahili'
    for sec in exam.sections:
        hp = doc.add_paragraph()
        sec_label = f"SEHEMU {sec.get('section', 'A')}" if sw else f"SECTION {sec.get('section', 'A')}"
        hr = hp.add_run(sec_label)
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = navy
        if sec.get('instructions'):
            doc.add_paragraph(sec['instructions'])
        for q in sec.get('questions', []):
            qp = doc.add_paragraph()
            marks_suffix = 'alama' if sw else 'marks'
            qr = qp.add_run(f"{q.get('number', '')}. {q.get('text', '')} ({q.get('marks', 0)} {marks_suffix})")
            qr.font.size = Pt(10)
            if include_answers and q.get('answer'):
                ap = doc.add_paragraph()
                ans_label = 'Jibu' if sw else 'Answer'
                ar = ap.add_run(f"{ans_label}: {q['answer']}")
                ar.font.size = Pt(9.5)
                ar.font.color.rgb = green
                ap.paragraph_format.left_indent = Pt(18)

    if include_answers:
        mp = doc.add_paragraph()
        mr = mp.add_run("MARKING SCHEME / SCHEME YA KUSAHLIHIA")
        mr.bold = True
        mr.font.size = Pt(12)
        mr.font.color.rgb = gold

    fp = doc.add_paragraph()
    fp.add_run(f"Prepared by: {teacher.full_name}  |  School: {teacher.school.name if teacher.school else '-'}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    suffix = 'marking' if include_answers else 'paper'
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{exam.title[:40].replace(chr(32), "_")}_{suffix}.docx"'
    return response


# =============================================================================
# AI PACE ADVISOR — remaining months vs remaining topics
# =============================================================================

_MONTH_NUM = {'JANUARY': 1, 'FEBRUARY': 2, 'MARCH': 3, 'APRIL': 4, 'MAY': 5, 'JUNE': 6,
              'JULY': 7, 'AUGUST': 8, 'SEPTEMBER': 9, 'OCTOBER': 10, 'NOVEMBER': 11, 'DECEMBER': 12}
_SW_TO_EN_MONTH = {v: k for k, v in SWAHILI_MONTHS.items()}


def _scheme_progress(scheme):
    """Compute progress stats from a SchemeOfWork: months elapsed, rows done, topics remaining."""
    rows = scheme.scheme_data if isinstance(scheme.scheme_data, list) else []
    today = timezone.now().date()
    current_month_num = today.month

    months_in_scheme = []
    for row in rows:
        m = (row.get('Month') or '').strip().upper()
        m = _SW_TO_EN_MONTH.get(m, m)
        if m in _MONTH_NUM and m not in months_in_scheme:
            months_in_scheme.append(m)
    months_in_scheme.sort(key=lambda m: _MONTH_NUM[m])

    total_rows = len(rows)
    rows_done = 0
    for row in rows:
        m = (row.get('Month') or '').strip().upper()
        m = _SW_TO_EN_MONTH.get(m, m)
        if m in _MONTH_NUM and _MONTH_NUM[m] < current_month_num:
            rows_done += 1
        elif m in _MONTH_NUM and _MONTH_NUM[m] == current_month_num:
            rows_done += 0.5  # halfway through current month
    rows_remaining = max(0, total_rows - int(rows_done))

    # Unique topic-ish entries from scheme rows
    topic_keys = []
    for row in rows:
        for key in ('Main Competence', 'Specific Competences', 'UMAHIRI MKUU', 'UMAHIRI MAHUSUSI'):
            val = (row.get(key) or '').strip()
            if val and val not in topic_keys:
                topic_keys.append(val)
                break

    total_months = len(months_in_scheme)
    months_done = sum(1 for m in months_in_scheme if _MONTH_NUM[m] < current_month_num)
    months_remaining = max(0, total_months - months_done)

    # End of school year = December
    end_of_year = today.replace(month=12, day=31)
    days_left = (end_of_year - today).days
    weeks_left = max(0, days_left // 7)

    return {
        'scheme': scheme,
        'total_rows': total_rows,
        'rows_done': int(rows_done),
        'rows_remaining': rows_remaining,
        'total_months': total_months,
        'months_done': months_done,
        'months_remaining': months_remaining,
        'months_in_scheme': months_in_scheme,
        'current_month': today.strftime('%B').upper(),
        'weeks_left': weeks_left,
        'days_left': days_left,
        'topics': topic_keys[:30],
        'percent_done': round((int(rows_done) / total_rows) * 100) if total_rows else 0,
    }


def pace_advisor(request):
    """AI Pace Advisor page — teacher picks a scheme; AI recommends teaching pace."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:pace_advisor')}")

    schemes = SchemeOfWork.objects.filter(
        school=teacher.school, teacher_name=teacher.full_name
    ).select_related('subject').order_by('-created_at')[:30]

    # Also allow schemes from other schools by the same teacher name
    if not schemes:
        schemes = SchemeOfWork.objects.filter(
            teacher_name=teacher.full_name
        ).select_related('subject').order_by('-created_at')[:30]

    return render(request, 'curriculum/pace_advisor.html', {
        'teacher': teacher,
        'schemes': schemes,
    })


@csrf_exempt
def ajax_pace_advice(request):
    """Generate AI advice about teaching pace from a scheme."""
    if request.method != 'POST':
        return JsonResponse({'success': False}, status=400)

    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Jisajili kwanza.'}, status=401)

    try:
        data = json.loads(request.body)
        scheme_id = data.get('scheme_id')
        # Scope to this teacher's own schemes
        scheme = SchemeOfWork.objects.filter(id=scheme_id, teacher_name=teacher.full_name).first()
        if not scheme:
            return JsonResponse({'success': False, 'error': 'Scheme haipatikani.'}, status=404)
        progress = _scheme_progress(scheme)

        subject_name = scheme.subject.name if scheme.subject else 'N/A'
        class_name = scheme.class_name
        term = scheme.term
        year = scheme.year

        topics_text = '\n'.join(f"- {t}" for t in progress['topics']) or "(hakuna topics zilizopatikana)"

        # Build a manual (non-AI) summary too
        summary = {
            'subject': subject_name,
            'class_name': class_name,
            'term': term,
            'year': year,
            'total_rows': progress['total_rows'],
            'rows_done': progress['rows_done'],
            'rows_remaining': progress['rows_remaining'],
            'percent_done': progress['percent_done'],
            'months_in_scheme': progress['months_in_scheme'],
            'current_month': progress['current_month'],
            'weeks_left': progress['weeks_left'],
            'days_left': progress['days_left'],
            'topics_remaining_count': len(progress['topics']),
        }

        if client is None:
            return JsonResponse({'success': True, 'summary': summary, 'advice': None,
                                 'warning': 'Huduma ya AI haipo — unaona hesabu tu.'})

        prompt = f"""You are an experienced Tanzanian education advisor for teachers. A teacher needs help finishing the syllabus on time.

SUBJECT: {subject_name}
CLASS: {class_name} | TERM: {term} | YEAR: {year}

PROGRESS DATA:
- Total scheme entries: {progress['total_rows']}
- Completed: {int(progress['rows_done'])} | Remaining: {progress['rows_remaining']} ({progress['percent_done']}% done)
- Months covered in scheme: {', '.join(progress['months_in_scheme'])}
- Current month: {progress['current_month']}
- Weeks left until end of December: {progress['weeks_left']} (about {progress['days_left']} days)
- Remaining topics/subjects in the syllabus: {progress['topics_remaining_count']}

TOPICS IN THE SYLLABUS:
{topics_text}

Give practical, specific advice in a mix of Kiswahili and English (teachers understand both). Include:
1. Current pace assessment: is the teacher ON TRACK, SLIGHTLY BEHIND, or FAR BEHIND?
2. Recommended pace: how many topics to cover per week and per month to finish by December.
3. A suggested weekly schedule for the remaining months (which topics, in what order, how many periods).
4. Strategies to catch up if behind (e.g. combining similar topics, prioritising examinable topics, revision drills).
5. What to do if the teacher is ahead (e.g. deepen understanding, mocks, project work).

Be concrete and encouraging. Use bullet points and short paragraphs."""

        response = client.models.generate_content(model=model_name, contents=prompt)
        return JsonResponse({'success': True, 'summary': summary, 'advice': response.text})
    except Exception as e:
        logger.exception("[Pace] Error")
        return JsonResponse({'success': False, 'error': str(e)[:300]}, status=500)


# =============================================================================
# RESULTS UPLOAD + AI ANALYSIS
# =============================================================================

def results_analysis(request):
    """Upload subject results (CSV/Excel) → AI analyses → recommends topics & students to support."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return redirect(f"{reverse('curriculum:teacher_register')}?next={reverse('curriculum:results_analysis')}")
    return render(request, 'curriculum/results_analysis.html', {'teacher': teacher})


def _parse_results_file(uploaded_file):
    """Parse CSV/Excel upload into (headers, rows)."""
    import csv
    import io
    name = (uploaded_file.name or '').lower()
    if name.endswith('.csv'):
        content = uploaded_file.read().decode('utf-8-sig', errors='replace')
        reader = csv.reader(io.StringIO(content))
        rows = [r for r in reader if any(c.strip() for c in r)]
        if not rows:
            raise ValueError('Faili halina data')
        return rows[0], rows[1:]
    if name.endswith(('.xlsx', '.xlsm')):
        from openpyxl import load_workbook
        wb = load_workbook(uploaded_file, data_only=True)
        ws = wb.active
        all_rows = list(ws.iter_rows(values_only=True))
        rows = [[('' if v is None else str(v).strip()) for v in r] for r in all_rows]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            raise ValueError('Faili halina data')
        return rows[0], rows[1:]
    raise ValueError('Aina ya faili haitambuliki. Tumia CSV au Excel (.xlsx).')


def _extract_number(val):
    """Extract a float from a value or None."""
    try:
        return float(str(val).strip().replace('%', '').replace(',', ''))
    except (ValueError, TypeError):
        return None


@csrf_exempt
def ajax_analyze_results(request):
    """Analyze uploaded results: per-topic stats + AI recommendations."""
    teacher = get_tlm_teacher(request)
    if not teacher:
        return JsonResponse({'success': False, 'error': 'Jisajili kwanza.'}, status=401)

    if request.method != 'POST':
        return JsonResponse({'success': False}, status=400)

    try:
        uploaded = request.FILES.get('file')
        if not uploaded:
            return JsonResponse({'success': False, 'error': 'Hakuna faili iliyopakiwa.'}, status=400)

        subject_name = request.POST.get('subject', '').strip()
        class_name = request.POST.get('class_name', '').strip()
        pass_mark = float(request.POST.get('pass_mark', 50) or 50)

        headers, rows = _parse_results_file(uploaded)
        if not rows:
            return JsonResponse({'success': False, 'error': 'Faili halina rekodi za wanafunzi.'}, status=400)

        # Identify name column (first column usually) and score columns
        score_cols = []
        for i, h in enumerate(headers[1:], start=1):
            hh = (h or '').strip()
            if hh and not _extract_number(hh):  # skip pure-number headers
                score_cols.append((i, hh))
        if not score_cols:
            # Fallback: treat each column after name as a score column
            score_cols = [(i, (h or f'Kolomu {i}').strip()) for i, h in enumerate(headers[1:], start=1)]

        students = []
        for r in rows:
            if len(r) < 2:
                continue
            name = (r[0] or '').strip()
            if not name:
                continue
            scores = {}
            for col_idx, col_name in score_cols:
                val = r[col_idx] if col_idx < len(r) else ''
                num = _extract_number(val)
                scores[col_name] = num if num is not None else None
            students.append({'name': name, 'scores': scores})

        if not students:
            return JsonResponse({'success': False, 'error': 'Hakuna wanafunzi walioonekana.'}, status=400)

        # ── Per-column (topic/subject) statistics ──
        col_stats = []
        for col_idx, col_name in score_cols:
            vals = [s['scores'].get(col_name) for s in students if s['scores'].get(col_name) is not None]
            if not vals:
                continue
            avg = round(sum(vals) / len(vals), 1)
            passed = sum(1 for v in vals if v >= pass_mark)
            col_stats.append({
                'name': col_name,
                'average': avg,
                'pass_count': passed,
                'total': len(vals),
                'pass_rate': round((passed / len(vals)) * 100) if vals else 0,
                'weak': avg < pass_mark,
            })
        col_stats.sort(key=lambda c: c['average'])

        # ── Student overall averages ──
        for s in students:
            vals = [v for v in s['scores'].values() if v is not None]
            s['average'] = round(sum(vals) / len(vals), 1) if vals else None
            s['passed'] = (s['average'] or 0) >= pass_mark
        students.sort(key=lambda s: (s['average'] or 0))

        class_avg = round(sum(s['average'] or 0 for s in students) / len(students), 1)
        below = [s for s in students if (s['average'] or 0) < pass_mark]
        pass_rate = round(((len(students) - len(below)) / len(students)) * 100)

        summary = {
            'subject': subject_name or (score_cols[0][1] if score_cols else ''),
            'class_name': class_name,
            'student_count': len(students),
            'class_average': class_avg,
            'pass_rate': pass_rate,
            'below_count': len(below),
            'weakest_columns': [c['name'] for c in col_stats[:5] if c['weak']],
            'strongest_columns': [c['name'] for c in reversed(col_stats[-3:])],
        }

        # Per-topic weak students list
        topic_students = []
        for c in col_stats:
            if c['weak']:
                weak_students = [s['name'] for s in students if (s['scores'].get(c['name']) or 0) < pass_mark]
                topic_students.append({'topic': c['name'], 'avg': c['average'], 'weak_students': weak_students[:15]})

        if client is None:
            return JsonResponse({'success': True, 'summary': summary, 'col_stats': col_stats,
                                 'students': students[:50], 'topic_students': topic_students,
                                 'advice': None, 'warning': 'Huduma ya AI haipo — unaona takwimu tu.'})

        weak_text = '\n'.join(f"- {c['name']}: wastani {c['average']} (faulu {c['pass_rate']}%)" for c in col_stats if c['weak']) or "(hakuna)"
        strong_text = '\n'.join(f"- {c['name']}: wastani {c['average']}" for c in reversed(col_stats[-3:])) or "(hakuna)"
        below_names = ', '.join(s['name'] for s in below[:20])

        prompt = f"""You are an experienced Tanzanian teacher-mentor and data analyst. Analyse these exam results and give practical advice.

SUBJECT: {subject_name or 'N/A'} | CLASS: {class_name or 'N/A'} | PASS MARK: {pass_mark}
STUDENTS: {len(students)} | CLASS AVERAGE: {class_avg} | PASS RATE: {pass_rate}% | BELOW PASS: {len(below)} students

WEAK TOPICS/AREAS (lowest averages):
{weak_text}

STRONG TOPICS/AREAS:
{strong_text}

STUDENTS WHO NEED MOST SUPPORT:
{below_names or '(hakuna)'}

Give advice (mix Kiswahili and English) covering:
1. Which topics MUST be re-taught and how (concrete strategies: remedial classes, group work, practical examples).
2. Specific students to give extra support and what kind (naming them).
3. How to group students for mixed-ability learning.
4. Suggestions for the next assessment (what to test, difficulty).
5. General teaching tips to lift the class average.
Use bullet points and short paragraphs."""

        response = client.models.generate_content(model=model_name, contents=prompt)
        return JsonResponse({
            'success': True,
            'summary': summary,
            'col_stats': col_stats,
            'students': students[:50],
            'topic_students': topic_students[:10],
            'advice': response.text,
        })
    except ValueError as ve:
        return JsonResponse({'success': False, 'error': str(ve)}, status=400)
    except Exception as e:
        logger.exception("[Results] Error")
        return JsonResponse({'success': False, 'error': str(e)[:300]}, status=500)
