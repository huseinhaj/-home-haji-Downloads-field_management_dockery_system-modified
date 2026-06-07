#!/usr/bin/env python3
"""
Script ya kutengeneza document ya Innovation Day kwa IMS project
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY  = RGBColor(0x0A, 0x2B, 0x5E)
GOLD  = RGBColor(0xC8, 0x90, 0x0A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0x4A, 0x55, 0x68)
LIGHT = RGBColor(0xEE, 0xF1, 0xF6)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        val = kwargs.get(edge, {})
        if val:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 6)))
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level=1, color=NAVY, size=16, bold=True, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def subheading(doc, text, color=GOLD, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'C8900A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def body(doc, text, size=11, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def numbered(doc, text, size=11):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    c1, c2 = row.cells[0], row.cells[1]
    c1.text = col1
    c2.text = col2
    for c, txt in [(c1, col1), (c2, col2)]:
        c.text = ''
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(11)
        if header:
            run.bold = True
            run.font.color.rgb = WHITE
            set_cell_bg(c, '0A2B5E')
        else:
            run.font.color.rgb = RGBColor(0x0F, 0x19, 0x23)
    return row

# ─── BUILD DOCUMENT ──────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t = title_p.add_run("INTERNSHIP MANAGEMENT SYSTEM")
t.bold = True
t.font.size = Pt(22)
t.font.color.rgb = NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = sub_p.add_run("IMS — Mfumo wa Usimamizi wa Mafunzo ya Vitendo")
s.bold = True
s.font.size = Pt(14)
s.font.color.rgb = GOLD

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.add_run("Innovation Day Presentation\n").font.size = Pt(12)
info_r = info_p.add_run("Wizara ya Elimu, Sayansi na Teknolojia — Tanzania")
info_r.font.size = Pt(11)
info_r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

line_p = doc.add_paragraph()
line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_r = line_p.add_run("━" * 45)
line_r.font.color.rgb = GOLD
line_r.font.size = Pt(10)

doc.add_page_break()

# ── 1. MUHTASARI / OVERVIEW ───────────────────────────────────────────────────
subheading(doc, "1. Muhtasari wa Mfumo (System Overview)")
doc.add_paragraph()
body(doc,
    "IMS ni mfumo wa kidijitali ulioundwa kusimamia mchakato mzima wa mafunzo ya vitendo (Teaching Practice) "
    "kwa wanafunzi-walimu wa vyuo vya ualimu Tanzania. Mfumo huu unabadilisha mchakato wa kawaida wa karatasi "
    "na kumfanya awe wa kidijitali, wa haraka, na unaoweza kufuatiliwa wakati wote.",
    size=11)
doc.add_paragraph()

# Quick facts table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_row(tbl, "Kipengele", "Maelezo", header=True)
facts = [
    ("Jina Kamili", "Internship Management System (IMS)"),
    ("Lengo", "Kusimamia mafunzo ya vitendo ya wanafunzi-walimu"),
    ("Watumiaji", "Wanafunzi-Walimu, Wakaguzi, DEO, Bodi, Wasimamizi"),
    ("Mwaka", "2024–2025"),
    ("Deployment", "Railway.app (Cloud Platform)"),
    ("Mwandishi", "Huseini Haji"),
]
for k, v in facts:
    add_table_row(tbl, k, v)

doc.add_paragraph()
doc.add_page_break()

# ── 2. MTIRIRIKO WA MFUMO ─────────────────────────────────────────────────────
subheading(doc, "2. Mtiririko wa Mfumo (System Flow)")
doc.add_paragraph()

flows = [
    ("HATUA 1 — Usajili & Idhini",
     "Mwanafunzi-mwalimu anajisajili kwenye mfumo kwa kutumia email na neno siri. "
     "Mtawala (Admin) anapitia maombi na kuidhinisha au kukataa. "
     "Baada ya idhini, mwanafunzi anaruhusiwa kuchagua shule."),

    ("HATUA 2 — Ugawaji wa Shule (DEO)",
     "Afisa Elimu wa Wilaya (DEO) anapakua na kupitia maombi ya wakuu wa shule. "
     "DEO anagawa idadi ya wanafunzi (quota) kwa kila shule. "
     "Mwanafunzi anachagua shule na masomo yake kwenye wilaya iliyoidhinishwa."),

    ("HATUA 3 — Mafunzo ya Vitendo",
     "Mwanafunzi-mwalimu anafundisha shuleni na kuandika logbook kila siku. "
     "Logbook inajumuisha shughuli za asubuhi, mchana, na jioni. "
     "Mfumo unakagua mahali (GPS) ili kuhakikisha mwanafunzi yuko shuleni."),

    ("HATUA 4 — Tathmini (Assessor)",
     "Mkaguzi (Assessor) anapewa shule moja au zaidi na Mtawala. "
     "Mkaguzi anatembelea shule, anaona maendeleo ya wanafunzi, na kuandika tathmini. "
     "Ripoti za mkaguzi zinaonekana kwa mwanafunzi na bodi."),

    ("HATUA 5 — Ufuatiliaji wa Bodi",
     "Bodi inaweza kuona ripoti za wilaya zote, shule zote, na wanafunzi wote. "
     "Wanaweza kuacha maoni (comments) kwa kila mwanafunzi. "
     "Ripoti zinaweza kupakiwa kama PDF moja kwa moja."),

    ("HATUA 6 — Barua za Uthibitisho",
     "Baada ya kukidhi masharti, mwanafunzi anapata barua ya uthibitisho (Individual Letter). "
     "Wakati shule inapofika idadi ya kiwango (quota), barua ya kikundi (Group Letter) inakuwa tayari. "
     "Barua zote zinaundwa kiotomatiki na mfumo."),
]

for title, desc in flows:
    p = doc.add_paragraph()
    r1 = p.add_run(f"  {title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(f"  {desc}")
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

doc.add_page_break()

# ── 3. TEKNOLOJIA ─────────────────────────────────────────────────────────────
subheading(doc, "3. Teknolojia Zilizotumika (Technology Stack)")
doc.add_paragraph()

tech_sections = {
    "Backend (Seva)": [
        ("Python 3.12", "Lugha ya programu — ya kisasa, yenye maktaba nyingi"),
        ("Django 5.x", "Web framework — MVC pattern, ORM, Auth, Admin"),
        ("PostgreSQL 15", "Database kuu — imara, inaweza kupanuka, ACID compliant"),
        ("Gunicorn", "WSGI server — inashughulikia maombi mengi kwa wakati mmoja"),
        ("Django ORM", "Object-Relational Mapper — usalama dhidi ya SQL Injection"),
    ],
    "Frontend (Upande wa Mteja)": [
        ("HTML5 / CSS3", "Muundo wa ukurasa — semantic markup"),
        ("Bootstrap 5", "Responsive design — inafanya kazi vizuri kwenye simu na kompyuta"),
        ("JavaScript (Vanilla)", "Interactivity — AJAX, form validation, GPS"),
        ("Font Awesome 6", "Icons za kitaalamu"),
        ("Google Fonts / Inter", "Fonti inayoonekana nzuri na ya kisasa"),
    ],
    "AI & Automation": [
        ("Claude AI (Anthropic)", "Kutengeneza Scheme of Work na Lesson Plans kiotomatiki"),
        ("Python-Docx", "Kutengeneza nyaraka za Word (.docx) kiotomatiki"),
        ("ReportLab / xHTML2PDF", "Kutengeneza PDF kiotomatiki"),
        ("Pandas", "Kusoma na kuchakata data za CSV (shule, masomo)"),
    ],
    "Infrastructure & DevOps": [
        ("Docker", "Containerization — mazingira yanafanana dev na production"),
        ("Railway.app", "Cloud hosting — auto-deploy, HTTPS, scaling"),
        ("GitHub", "Version control — kudhibiti mabadiliko ya code"),
        ("Gmail SMTP", "Kutuma emails za nywila kwa wakaguzi wapya"),
        ("Africa's Talking API", "Kutuma SMS kwa wakuu wa shule"),
    ],
}

for section_name, items in tech_sections.items():
    p = doc.add_paragraph()
    r = p.add_run(section_name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GOLD

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    add_table_row(tbl, "Teknolojia", "Kusudi", header=True)
    for k, v in items:
        add_table_row(tbl, k, v)
    doc.add_paragraph()

doc.add_page_break()

# ── 4. MUUNDO WA MFUMO (ARCHITECTURE) ────────────────────────────────────────
subheading(doc, "4. Muundo wa Mfumo (System Architecture)")
doc.add_paragraph()

body(doc, "IMS inatumia muundo wa MVT (Model-View-Template) — toleo la Django la MVC:", size=11)
doc.add_paragraph()

arch_items = [
    ("Model (Data Layer)",
     "Kuna models 25+ zinazosimamia data zote:\n"
     "• CustomUser, StudentTeacher — Watumiaji\n"
     "• School, District, Region — Jiografia\n"
     "• LogbookEntry, MonthlyReport — Rekodi za Kazi\n"
     "• Assessor, SchoolAssessment — Tathmini\n"
     "• DistrictAllocation, SchoolAllocation — Mgawanyo wa DEO\n"
     "• SchemeOfWork, LessonPlan — Mipango ya AI"),

    ("View (Business Logic)",
     "Views.py ina zaidi ya vitendo 100 vinavyosimamia:\n"
     "• Authentication na Authorization (Role-based access)\n"
     "• AJAX APIs kwa utafutaji wa shule bila kupakia ukurasa\n"
     "• Bulk operations (kuweka assessors wengi mara moja)\n"
     "• PDF na Word generation moja kwa moja"),

    ("Template (Presentation Layer)",
     "Templates zaidi ya 50 zinazofuata muundo wa serikali:\n"
     "• Rangi: Navy (#0A2B5E) + Gold (#C8900A) + Light (#EEF1F6)\n"
     "• Sidebar ya kando inayopanuka/kukunjwa\n"
     "• Responsive design — inafanya kazi vizuri simu na kompyuta"),

    ("Database Layer",
     "PostgreSQL na Django ORM:\n"
     "• Database indexes kwenye sehemu zinazotafutwa mara kwa mara\n"
     "• Foreign keys zinalinda uhusiano wa data\n"
     "• Migrations zinadhibiti mabadiliko ya schema salama"),

    ("Security Layer",
     "• CSRF Protection — kila form inalindwa\n"
     "• Role-based Access Control — kila mtumiaji anaona sehemu yake tu\n"
     "• Session management — vikao vya masaa 8 tu\n"
     "• HTTPS (SSL/TLS) — data inasafiri kwa usalama\n"
     "• HSTS — browser inakataa HTTP"),
]

for title, desc in arch_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.size = Pt(11.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

doc.add_page_break()

# ── 5. USALAMA WA DATA NYINGI (SCALABILITY) ───────────────────────────────────
subheading(doc, "5. Jinsi Mfumo Unavyoshughulikia Data Nyingi na Maombi Mengi")
doc.add_paragraph()

body(doc,
    "Swali: Je, mfumo unaweza kushughulikia wanafunzi 10,000+ na maombi mengi kwa wakati mmoja?",
    size=11, italic=True, color=GOLD)
body(doc, "Jibu: Ndiyo. Hivi ndivyo inavyofanya:", size=11)
doc.add_paragraph()

scalability = [
    ("1. Database Indexing",
     "Tumeweka index kwenye sehemu zote zinazotafutwa mara kwa mara kama vile jina la shule, "
     "nambari ya mwanafunzi, tarehe za logbook, na hali ya maombi. "
     "Hii inafanya utafutaji wa haraka hata kama database ina rekodi milioni nyingi."),

    ("2. Caching (Kuhifadhi Matokeo Kwenye Kumbukumbu ya Muda)",
     "Mfumo unatumia cache kwa data inayobadilika polepole kama vile mwaka wa masomo unaotumika sasa, "
     "orodha ya masomo kwa kila mwanafunzi, na orodha ya wilaya/mikoa. "
     "Hii inapunguza maswali ya database kwa 70-80%."),

    ("3. Bulk Operations",
     "Badala ya kuweka assessors mmoja mmoja (ambayo ingehitaji maswali 1000 kwa assessors 1000), "
     "mfumo unatumia bulk_create() na bulk_update() za Django — "
     "maswali 1000 yanafanywa kwa jibu moja tu la database."),

    ("4. Gunicorn Workers (Wafanyakazi Wengi)",
     "Railway inafanya kazi na Gunicorn ambayo inaweza kusimamia maombi mengi kwa wakati mmoja. "
     "Kila 'worker' inaweza kusimamia ombi la ziada. Kwa default, inafanyiwa kazi kwa formula: "
     "workers = (CPU cores × 2) + 1. Kwenye Railway, tunaweza kuongeza workers haraka."),

    ("5. Pagination (Kuonyesha Kidogo kwa Wakati Mmoja)",
     "Badala ya kupakia rekodi zote za wanafunzi 10,000 mara moja, "
     "mfumo unaonyesha kurasa 20-50 kwa wakati mmoja. "
     "Hii inapunguza mzigo wa database na kasi ya kupakia ukurasa."),

    ("6. Database Connection Pooling",
     "Django inasimamia muunganiko wa database kwa ufanisi. "
     "Muunganiko mmoja unaweza kutumika na maombi mengi badala ya kufungua "
     "muunganiko mpya kwa kila ombi — hii inapunguza muda wa kujibu kwa zaidi ya 50%."),

    ("7. Cloud Scaling (Railway.app)",
     "Wakati mzigo ukiwa mkubwa, Railway inaweza kuanzisha 'replicas' nyingi za mfumo kiotomatiki. "
     "Kila replica inashughulikia sehemu ya maombi. "
     "Hii inamaanisha mfumo unaweza kupanuka bila kuzima au kubadilisha code."),
]

for title, desc in scalability:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

doc.add_page_break()

# ── 6. UHIFADHI WA DATA BAADA YA INTAKE ──────────────────────────────────────
subheading(doc, "6. Data Zinakwenda Wapi Baada ya Intake Kuisha?")
doc.add_paragraph()

body(doc,
    "Swali: Baada ya intake (mzunguko wa mafunzo ya vitendo) kuisha, data za wanafunzi zinaenda wapi?",
    size=11, italic=True, color=GOLD)
doc.add_paragraph()

retention = [
    ("A — Academic Year System (Mwaka wa Masomo)",
     "Kila rekodi ya mwanafunzi (logbook, tathmini, maombi) imeunganishwa na 'Mwaka wa Masomo' (AcademicYear). "
     "Wakati mwaka mpya unapoanza, Mtawala anafungua mwaka mpya — "
     "data za mwaka uliopita HAZIANGUSHWI bali zinabaki kwenye database kama kumbukumbu ya kihistoria."),

    ("B — Backup System (Mfumo wa Chelezo)",
     "Kila wiki mfumo unatengeneza chelezo (backup) za database kwa muundo wa JSON. "
     "Chelezo hizi zinahifadhiwa kwenye:\n"
     "  • Seva ya Railway (cloud)\n"
     "  • Faili za .json kwenye directory ya /backups\n"
     "  • Mtawala anaweza kupakua chelezo wakati wowote anapohitaji"),

    ("C — Data za Mwanafunzi Huathiriwa Namna Gani?",
     "Mwanafunzi anayekwisha mafunzo ana haki ya:\n"
     "  • Kupakua logbook yake yote kama PDF\n"
     "  • Kupakua barua za uthibitisho wake\n"
     "  • Kuona rekodi zake za tathmini\n\n"
     "Data hizi HAZIFUTWE kamwe — zinabaki kwenye database kwa miaka yote inayofuata."),

    ("D — Data za Shule na DEO",
     "Mgawanyo wa wilaya (DistrictAllocation) na mgawanyo wa shule (SchoolAllocation) "
     "zimeunganishwa na mwaka wa masomo. "
     "Mwaka unaofuata, DEO anaanzisha mgawanyo mpya bila kuathiri takwimu za mwaka uliopita. "
     "Hii inawezesha kulinganisha takwimu kati ya miaka tofauti."),

    ("E — Compliance na Sheria ya Data",
     "Mfumo umejengwa kuzingatia ulinzi wa data:\n"
     "  • Nywila zimehifadhiwa kwa hashing (Django's PBKDF2 + SHA256)\n"
     "  • HTTPS inazuia kusikiwa kwa data wakati wa usafirishaji\n"
     "  • Role-based access — mwanafunzi anaona data yake tu\n"
     "  • Admin pekee ndiye anayeweza kufuta au kurekebisha rekodi nyeti"),
]

for title, desc in retention:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

doc.add_page_break()

# ── 7. MUUNDO WA UI ────────────────────────────────────────────────────────────
subheading(doc, "7. Muundo wa Kiolesura (UI/UX Design)")
doc.add_paragraph()

body(doc,
    "Mfumo umejengwa kuwa wa kitaalamu kama mifumo ya serikali ya Tanzania (TRA, NECTA, HESLB). "
    "Hii inaongeza imani ya watumiaji na kuifanya ikubalike rasmi.",
    size=11)
doc.add_paragraph()

ui_items = [
    ("Palette ya Rangi (Color Palette)",
     "• Navy Blue (#0A2B5E) — Imani, utaalamu, serikali\n"
     "• Gold (#C8900A) — Utukufu, Tanzania, umakini\n"
     "• Light Gray (#EEF1F6) — Usafi, uangalifu\n"
     "Rangi hizi zinafanana na bendera ya Tanzania na mifumo ya serikali."),

    ("Typography (Fonti)",
     "• Fonti: Inter — kisasa, inayosomeka vizuri kwenye skrini ndogo na kubwa\n"
     "• Maandishi ya headers: Uppercase + letter-spacing — yanaonyesha rasmiality\n"
     "• Ukubwa: 11–14px kwa maudhui, 18–22px kwa vichwa"),

    ("Layout (Mpangilio)",
     "• Sidebar inayopanuka/kukunjwa kwa click au hover\n"
     "• Cards zenye mipaka nyepesi na vivuli kidogo\n"
     "• Responsive — inafanya kazi vizuri kwenye simu, kompyuta kibao, na kompyuta\n"
     "• Mobile hamburger menu kwa watumiaji wa simu"),

    ("Uzoefu wa Mtumiaji (UX)",
     "• Loading states — mtumiaji anaona mfumo unafanya kazi\n"
     "• Flash messages — mafanikio/makosa yanaonyeshwa wazi\n"
     "• Hakuna uthibitisho wa karatasi — yote yanafanywa mtandaoni\n"
     "• Mtumiaji anaweza kubadilisha lugha (Kiswahili / English)"),

    ("Vipengele vya Kisasa (Modern Components)",
     "• Stat cards na takwimu za moja kwa moja\n"
     "• Progress bars zinaonyesha ukamilifu wa kazi\n"
     "• Badge labels za hali (Inasubiri / Imeidhinishwa / Imekataliwa)\n"
     "• Table ya kulingana na skrini (responsive tables)\n"
     "• Quick Actions grid kwa vitendo vya kawaida"),
]

for title, desc in ui_items:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

doc.add_page_break()

# ── 8. TAKWIMU ZA MFUMO ───────────────────────────────────────────────────────
subheading(doc, "8. Takwimu za Mfumo (System Statistics)")
doc.add_paragraph()

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
add_table_row(tbl2, "Kipengele", "Takwimu", header=True)
stats = [
    ("Mistari ya Code (Backend)", "~10,000+ mistari (Python/Django)"),
    ("Mistari ya Code (Frontend)", "~15,000+ mistari (HTML/CSS/JS)"),
    ("Idadi ya Models", "25+ models za database"),
    ("Idadi ya Views/APIs", "100+ views na API endpoints"),
    ("Idadi ya Templates", "50+ templates za HTML"),
    ("Idadi ya Migrations", "18 database migrations"),
    ("Moduli za Mfumo", "7 (Admin, Student, Assessor, DEO, Board, AI, API)"),
    ("Lugha Zinazotumika", "Python, JavaScript, SQL, HTML, CSS"),
    ("Lugha za Mtumiaji", "Kiswahili na Kiingereza (switchable)"),
    ("Wakati wa Maendeleo", "Miezi 6+"),
    ("Platform ya Hosting", "Railway.app (Cloud)"),
    ("Uwezo wa Kupanuka", "Hadi wanafunzi 100,000+ bila kubadilisha code"),
]
for k, v in stats:
    add_table_row(tbl2, k, v)

doc.add_paragraph()
doc.add_page_break()

# ── 9. FAIDA ZA MFUMO ─────────────────────────────────────────────────────────
subheading(doc, "9. Faida za Mfumo Huu Kuliko Mfumo wa Kawaida")
doc.add_paragraph()

comparison = [
    ("Muda", "Kawaida (karatasi): wiki 2-4 kupata barua. IMS: dakika 5 online"),
    ("Gharama", "Kawaida: karatasi, safari, posho. IMS: bure baada ya setup"),
    ("Ufuatiliaji", "Kawaida: haiwezekani. IMS: real-time kwa GPS + logbook"),
    ("Makosa", "Kawaida: mengi (kupotea kwa karatasi). IMS: database haiangukii"),
    ("Ripoti", "Kawaida: mwezi mzima. IMS: sekunde moja, PDF tayari"),
    ("Ufikiaji", "Kawaida: ofisi tu. IMS: popote duniani kwa simu au kompyuta"),
    ("Historia", "Kawaida: inapotea. IMS: miaka yote imehifadhiwa"),
]

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
hdr = tbl3.rows[0].cells
for c, txt in zip(hdr, ["Kipengele", "Mfumo wa Kawaida", "IMS"]):
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(11)
    set_cell_bg(c, '0A2B5E')

for k, v in comparison:
    row = tbl3.add_row()
    parts = v.split(". IMS: ")
    old = parts[0].replace("Kawaida (karatasi): ", "").replace("Kawaida: ", "")
    new = parts[1] if len(parts) > 1 else "—"
    for c, txt in zip(row.cells, [k, old, new]):
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10.5)

doc.add_paragraph()
doc.add_page_break()

# ── 10. HITIMISHO ─────────────────────────────────────────────────────────────
subheading(doc, "10. Hitimisho (Conclusion)")
doc.add_paragraph()

body(doc,
    "Internship Management System (IMS) ni mfumo kamili, salama, na wa kisasa ambao unajibu changamoto zote "
    "za usimamizi wa mafunzo ya vitendo Tanzania. Unajengwa kwa teknolojia za hali ya juu ambazo zinafanya kazi "
    "kwenye mazingira ya serikali na yanayohitaji uaminifu, usalama, na uwazi.",
    size=11)
doc.add_paragraph()
body(doc,
    "Mfumo huu unaweza kupanuka na kukua pamoja na mahitaji ya Wizara ya Elimu — "
    "kutoka wanafunzi 100 hadi 100,000+ bila kubadilisha muundo wa msingi.",
    size=11)
doc.add_paragraph()
body(doc,
    "Uvumbuzi mkubwa zaidi wa mfumo huu ni ujumuishaji wa AI (Claude AI) katika kuandaa "
    "Scheme of Work na Lesson Plans — kitu ambacho kinaweza kuokoa wakati wa masaa mengi kwa kila mwalimu.",
    size=11)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Asante kwa Usikilizaji Wenu")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Mwandishi: Huseini Haji  |  Innovation Day 2025")
r2.font.size = Pt(11)
r2.font.color.rgb = GOLD

# ── SAVE ───────────────────────────────────────────────────────────────────────
output = "/home/haji/Downloads/IMS_Innovation_Day_Presentation.docx"
doc.save(output)
print(f"✅ Document imehifadhiwa: {output}")
