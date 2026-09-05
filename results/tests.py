import json
from unittest.mock import patch

from django.test import Client, TestCase
from django.urls import reverse
import pandas as pd

from .models import ClassTimetableEntry, Exam, ExamResult, FormStudent, ProcessedResult, School, SchoolSubject, SpeechSubmissionSession, Student, Subject, TeacherAccount, TeachingAssignment, TimeSlot
from .services.class_timetable_service import TimetableConflict, generate_class_timetable, save_class_timetable, set_single_cell
from .services.scoresheet_ocr_service import (
    ScoreSheetOCRError,
    _clean_rows,
    _extract_json_array,
    _is_pdf,
    _load_page_images,
)
from .services.speech_submission_service import (
    create_or_get_session,
    extract_name_and_score,
    fuzzy_match_student_name,
	get_session_status,
	submit_speech_entries_batch,
	SpeechMatchReviewRequired,
	SpeechSubmissionError,
	parse_spoken_score,
    submit_speech_entry,
)
from .utils import (
	extract_subject_columns,
	normalize_gender,
	normalize_subject_name,
	parse_score,
)


class ResultsUtilsTests(TestCase):
	def test_normalize_subject_name_maps_common_aliases(self):
		self.assertEqual(normalize_subject_name('phy'), 'Physics')
		self.assertEqual(normalize_subject_name(' mathematics '), 'Mathematics')

	def test_extract_subject_columns_excludes_student_identity_fields(self):
		df = pd.DataFrame(columns=['First Name', 'Last Name', 'Gender', 'Physics', 'Chemistry'])
		self.assertEqual(extract_subject_columns(df), ['Physics', 'Chemistry'])

	def test_parse_score_handles_invalid_and_numeric_values(self):
		self.assertEqual(parse_score('78'), 78)
		self.assertEqual(parse_score(44.8), 44)
		self.assertIsNone(parse_score('not-a-number'))
		self.assertIsNone(parse_score(float('nan')))

	def test_normalize_gender_defaults_to_male_for_unknown_input(self):
		self.assertEqual(normalize_gender('Female'), 'F')
		self.assertEqual(normalize_gender('male'), 'M')
		self.assertEqual(normalize_gender(''), 'M')


class SpeechSubmissionServiceTests(TestCase):
	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=1)
		self.subject = Subject.objects.create(name='Mathematics')
		self.student_one = Student.objects.create(first_name='Amina', middle_name='', last_name='Juma', gender='F')
		self.student_two = Student.objects.create(first_name='Peter', middle_name='', last_name='Mushi', gender='M')

	def test_extract_name_and_score_from_transcript(self):
		name, score = extract_name_and_score('Amina Juma 78')
		self.assertEqual(name, 'amina juma')
		self.assertEqual(score, 78)

	def test_extract_name_and_score_handles_spaced_digits(self):
		name, score = extract_name_and_score('Amina Juma 8 0')
		self.assertEqual(name, 'amina juma')
		self.assertEqual(score, 80)

	def test_parse_spoken_score_handles_english_words(self):
		self.assertEqual(parse_spoken_score('Amina Juma eighty five'), 85)

	def test_parse_spoken_score_handles_english_filler_words(self):
		self.assertEqual(parse_spoken_score('Amina Juma score is eighty five'), 85)

	def test_parse_spoken_score_handles_swahili_words(self):
		self.assertEqual(parse_spoken_score('Amina Juma hamsini na tano'), 55)

	def test_parse_spoken_score_handles_swahili_phonetic_variants(self):
		self.assertEqual(parse_spoken_score('sistini'), 60)
		self.assertEqual(parse_spoken_score('Juma ali sistini'), 60)
		self.assertEqual(parse_spoken_score('Anasaiti, Nishinambili, Rasibari, Elateni, Yumaari, Tisinani.'), 90)
		self.assertEqual(parse_spoken_score('nishinambili'), 22)

	def test_extract_name_and_score_handles_non_trailing_spoken_score(self):
		name, score = extract_name_and_score('Juma ali sitini na mbili leo')
		self.assertEqual(score, 62)
		self.assertIn('juma', name)

	def test_fuzzy_match_student_name_returns_candidates(self):
		student, confidence, candidates = fuzzy_match_student_name('Amina Juma', Student.objects.all(), threshold=0.5)
		self.assertIsNotNone(student)
		self.assertEqual(student.id, self.student_one.id)
		self.assertGreaterEqual(confidence, 0.5)
		self.assertGreaterEqual(len(candidates), 2)

	def test_submit_speech_entry_finalizes_session_when_complete(self):
		session = create_or_get_session(exam=self.exam, subject=self.subject, teacher_name='Teacher One', expected_student_count=2)
		result_one = submit_speech_entry(
			session=session,
			transcript='Amina Juma 75',
			confirm_student_id=self.student_one.id,
		)
		self.assertEqual(result_one['score'], 75)
		result_two = submit_speech_entry(
			session=session,
			transcript='Peter Mushi 66',
			confirm_student_id=self.student_two.id,
		)
		session.refresh_from_db()
		self.assertEqual(result_two['score'], 66)
		self.assertEqual(session.status, SpeechSubmissionSession.STATUS_FINALIZED)

	def test_duplicate_submission_overwrites_existing_value(self):
		session = create_or_get_session(exam=self.exam, subject=self.subject, teacher_name='Teacher One', expected_student_count=2)
		submit_speech_entry(
			session=session,
			transcript='Amina Juma 75',
			confirm_student_id=self.student_one.id,
		)
		result = submit_speech_entry(
			session=session,
			transcript='Amina Juma 80',
			confirm_student_id=self.student_one.id,
		)
		self.assertEqual(result['score'], 80)

	def test_submit_speech_entries_batch_saves_multiple_students(self):
		session = create_or_get_session(
			exam=self.exam,
			subject=self.subject,
			teacher_name='Teacher One',
			expected_student_count=2,
			roster_student_ids=[self.student_one.id, self.student_two.id],
		)
		result = submit_speech_entries_batch(
			session=session,
			transcript='Amina Juma sabini na mbili. Peter Mushi themanini na moja.',
		)
		self.assertEqual(result['saved_count'], 2)
		self.assertEqual(result['skipped_count'], 0)

	def test_low_confidence_match_requires_review(self):
		session = create_or_get_session(exam=self.exam, subject=self.subject, teacher_name='Teacher One', expected_student_count=1)
		with self.assertRaises(SpeechMatchReviewRequired) as context:
			submit_speech_entry(
				session=session,
				transcript='completely different name 75',
			)
		self.assertGreaterEqual(len(context.exception.candidates), 1)

	def test_submit_speech_entry_recovers_noisy_name_from_transcript(self):
		session = create_or_get_session(exam=self.exam, subject=self.subject, teacher_name='Teacher One', expected_student_count=1)
		result = submit_speech_entry(
			session=session,
			transcript='ZZ Amina Jooma tisinani',
		)
		self.assertEqual(result['student']['id'], self.student_one.id)
		self.assertEqual(result['score'], 90)

	def test_get_session_status_includes_existing_subject_marks_in_debug_payload(self):
		session = create_or_get_session(
			exam=self.exam,
			subject=self.subject,
			teacher_name='Teacher One',
			expected_student_count=2,
			roster_student_ids=[self.student_one.id, self.student_two.id],
		)

		ExamResult.objects.create(exam=self.exam, student=self.student_two, subject=self.subject, score=64)
		submit_speech_entry(
			session=session,
			transcript='Amina Juma 88',
			confirm_student_id=self.student_one.id,
		)

		status = get_session_status(session, include_existing_marks=True)

		self.assertIn('existing_subject_marks', status)
		self.assertIn('saved_entries', status)
		self.assertEqual(len(status['existing_subject_marks']), 2)
		self.assertEqual(len(status['saved_entries']), 1)

		sources = {row['student_id']: row['source'] for row in status['existing_subject_marks']}
		self.assertEqual(sources[self.student_one.id], 'speech_session')
		self.assertEqual(sources[self.student_two.id], 'existing_result')


class ScoreSheetOCRParsingTests(TestCase):
	"""No network calls — only the response-parsing helpers, using canned
	AI-response text shapes (fenced, with surrounding prose, malformed)."""

	def test_extract_json_array_strips_markdown_fence(self):
		text = '```json\n[{"name": "Amina Juma", "score": 78}]\n```'
		self.assertEqual(_extract_json_array(text), [{"name": "Amina Juma", "score": 78}])

	def test_extract_json_array_ignores_surrounding_prose(self):
		text = 'Here are the results:\n[{"name": "Peter Mushi", "score": 55}]\nHope that helps!'
		self.assertEqual(_extract_json_array(text), [{"name": "Peter Mushi", "score": 55}])

	def test_extract_json_array_raises_on_no_array(self):
		with self.assertRaises(ScoreSheetOCRError):
			_extract_json_array('Sorry, I could not read the image.')

	def test_extract_json_array_raises_on_malformed_json(self):
		with self.assertRaises(ScoreSheetOCRError):
			_extract_json_array('[{"name": "Amina", "score": }]')

	def test_clean_rows_drops_out_of_range_and_garbage_fields(self):
		raw = [
			{"row": 1, "name": "Amina Juma", "score": 78},
			{"name": "", "score": 50},
			{"row": 3, "name": "No Score"},
			{"row": 4, "name": "Too High", "score": 150},
			{"row": 5, "name": "Too Low", "score": -5},
			{"row": 6, "name": "Not A Number", "score": "abc"},
		]
		self.assertEqual(_clean_rows(raw), [
			{"raw_name": "Amina Juma", "score": 78, "is_absent": False, "row": 1, "blank": False},
			{"raw_name": "No Score", "score": None, "is_absent": False, "row": 3, "blank": True},
		])

	def test_clean_rows_keeps_blank_rows_instead_of_dropping_them(self):
		"""A blank/dash score cell means 'student doesn't study this
		subject' — the row must be KEPT (blank=True, score=None) so a
		caller can tell it apart from a row the AI never reported at all
		(which usually means a mark was missed, not a non-taker)."""
		raw = [
			{"row": 1, "name": "Blank Cell", "score": "BLANK"},
			{"row": 2, "name": "Dash Cell", "score": "-"},
			{"row": 3, "name": "Null Cell", "score": None},
		]
		cleaned = _clean_rows(raw)
		self.assertEqual(len(cleaned), 3)
		self.assertTrue(all(r["blank"] and r["score"] is None and not r["is_absent"] for r in cleaned))

	def test_clean_rows_parses_row_number(self):
		raw = [{"row": "7", "name": "Amina Juma", "score": 78}]
		self.assertEqual(_clean_rows(raw)[0]["row"], 7)


def _build_pdf_bytes(num_pages=1):
	from io import BytesIO
	from reportlab.pdfgen import canvas

	buf = BytesIO()
	c = canvas.Canvas(buf)
	for i in range(num_pages):
		c.drawString(100, 700, f"Scoresheet page {i + 1}")
		c.showPage()
	c.save()
	return buf.getvalue()


class ScoreSheetDocumentLoadingTests(TestCase):
	"""Real PDF rendering via pypdfium2 (no network/AI call) — proves a
	scanned PDF is turned into page images the same way a photo is."""

	def test_is_pdf_detects_pdf_by_magic_bytes(self):
		from django.core.files.uploadedfile import SimpleUploadedFile
		pdf_file = SimpleUploadedFile('sheet.pdf', _build_pdf_bytes(), content_type='application/octet-stream')
		self.assertTrue(_is_pdf(pdf_file))

	def test_is_pdf_false_for_image(self):
		from django.core.files.uploadedfile import SimpleUploadedFile
		jpeg_file = SimpleUploadedFile('sheet.jpg', b'\xff\xd8\xff\xe0fake', content_type='image/jpeg')
		self.assertFalse(_is_pdf(jpeg_file))

	def test_load_page_images_renders_one_image_per_pdf_page(self):
		from django.core.files.uploadedfile import SimpleUploadedFile
		pdf_file = SimpleUploadedFile('sheet.pdf', _build_pdf_bytes(num_pages=2), content_type='application/pdf')
		images = _load_page_images(pdf_file)
		self.assertEqual(len(images), 2)
		for img in images:
			self.assertEqual(img.mode, 'RGB')

	def test_load_page_images_caps_at_max_pages(self):
		from django.core.files.uploadedfile import SimpleUploadedFile
		from .services import scoresheet_ocr_service
		pdf_file = SimpleUploadedFile('sheet.pdf', _build_pdf_bytes(num_pages=8), content_type='application/pdf')
		images = _load_page_images(pdf_file)
		self.assertEqual(len(images), scoresheet_ocr_service.MAX_PDF_PAGES)


class ScoreSheetMultiPageExtractionTests(TestCase):
	"""extract_scores_from_document reads a multi-page document's pages
	CONCURRENTLY (one slow ~180s vision call per page, so a 3-5 page
	scoresheet must not wait on them one at a time) -- these prove that
	stays correct: page order is preserved regardless of which finishes
	first, and one page's failure doesn't discard the others' rows.

	_load_page_images is patched to return plain sentinel strings instead
	of real rendered pages -- _read_page_with_ai is mocked too, so nothing
	downstream ever needs an actual PIL image, and using distinct sentinels
	lets fake_read tell pages apart deterministically instead of relying on
	call order, which the ThreadPoolExecutor doesn't guarantee."""

	def setUp(self):
		from .services import scoresheet_ocr_service
		self._service = scoresheet_ocr_service
		patcher = patch.object(scoresheet_ocr_service, 'OPENROUTER_API_KEY', 'test-key')
		patcher.start()
		self.addCleanup(patcher.stop)

	def _upload(self):
		from django.core.files.uploadedfile import SimpleUploadedFile
		return SimpleUploadedFile('sheet.pdf', b'%PDF-fake', content_type='application/pdf')

	def test_page_order_preserved_regardless_of_completion_order(self):
		import time as _time

		# 'page1' is slower than 'page2' -- if results were assembled in
		# completion order instead of page order, page 2's row would land
		# before page 1's.
		def fake_read(img):
			if img == 'page1':
				_time.sleep(0.15)
				return '[{"row": 1, "name": "Page One Student", "score": 50}]'
			return '[{"row": 2, "name": "Page Two Student", "score": 60}]'

		with patch.object(self._service, '_load_page_images', return_value=['page1', 'page2']), \
				patch.object(self._service, '_read_page_with_ai', side_effect=fake_read):
			rows = self._service.extract_scores_from_document(self._upload())

		names = [r['raw_name'] for r in rows]
		self.assertEqual(names, ['Page One Student', 'Page Two Student'])

	def test_one_failed_page_does_not_discard_others(self):
		def fake_read(img):
			if img == 'page1':
				raise RuntimeError('vision API timed out')
			return '[{"row": 2, "name": "Survivor", "score": 70}]'

		with patch.object(self._service, '_load_page_images', return_value=['page1', 'page2']), \
				patch.object(self._service, '_read_page_with_ai', side_effect=fake_read):
			rows = self._service.extract_scores_from_document(self._upload())

		self.assertEqual(len(rows), 1)
		self.assertEqual(rows[0]['raw_name'], 'Survivor')

	def test_all_pages_failing_raises(self):
		with patch.object(self._service, '_load_page_images', return_value=['page1', 'page2']), \
				patch.object(self._service, '_read_page_with_ai', side_effect=RuntimeError('down')):
			with self.assertRaises(ScoreSheetOCRError):
				self._service.extract_scores_from_document(self._upload())


class ScoreSheetPhotoExtractViewTests(TestCase):
	"""scoresheet_photo_extract now only dispatches a Celery task and
	returns a task_id — these force the task to run synchronously
	(task_always_eager) so the tests don't need a real worker/broker, then
	poll scoresheet_extract_status once for the actual result."""
	databases = {'default', 'results'}

	@classmethod
	def setUpClass(cls):
		super().setUpClass()
		from field_management.celery import app as celery_app
		cls._celery_app = celery_app
		cls._orig_conf = {
			'task_always_eager': celery_app.conf.task_always_eager,
			'task_eager_propagates': celery_app.conf.task_eager_propagates,
			'task_store_eager_result': celery_app.conf.task_store_eager_result,
		}
		celery_app.conf.task_always_eager = True
		celery_app.conf.task_eager_propagates = True
		celery_app.conf.task_store_eager_result = True

	@classmethod
	def tearDownClass(cls):
		cls._celery_app.conf.update(cls._orig_conf)
		super().tearDownClass()

	def setUp(self):
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=1)
		self.subject = Subject.objects.create(name='Mathematics')
		self.student_one = Student.objects.create(first_name='Amina', middle_name='', last_name='Juma', gender='F')
		self.student_two = Student.objects.create(first_name='Peter', middle_name='', last_name='Mushi', gender='M')
		self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER)
		self.teacher.subjects.set([self.subject])
		self.client = Client()
		self.client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')

	def _post(self, extracted_rows, roster):
		with patch('results.tasks.extract_scores_from_document', return_value=extracted_rows):
			from django.core.files.uploadedfile import SimpleUploadedFile
			photo = SimpleUploadedFile('sheet.jpg', b'fake-bytes', content_type='image/jpeg')
			kickoff = self.client.post(reverse('scoresheet_photo_extract'), {
				'photo': photo,
				'exam_id': self.exam.id,
				'subject_id': self.subject.id,
				'roster': json.dumps(roster),
			})
			self.assertEqual(kickoff.status_code, 202)
			task_id = kickoff.json()['task_id']
			return self.client.get(reverse('scoresheet_extract_status', args=[task_id]))

	def test_extracted_rows_are_fuzzy_matched_against_posted_roster(self):
		roster = [
			{'id': self.student_one.id, 'name': 'Amina Juma'},
			{'id': self.student_two.id, 'name': 'Peter Mushi'},
		]
		extracted = [
			{'raw_name': 'Amina Juma', 'score': 78},
			{'raw_name': 'Peter Mushi', 'score': 55},
		]
		response = self._post(extracted, roster)
		self.assertEqual(response.status_code, 200)
		data = response.json()
		self.assertEqual(len(data['matched']), 2)
		self.assertEqual(data['unmatched'], [])
		matched_by_id = {row['id']: row['score'] for row in data['matched']}
		self.assertEqual(matched_by_id[self.student_one.id], 78)
		self.assertEqual(matched_by_id[self.student_two.id], 55)

	def test_name_not_on_roster_creates_a_new_student(self):
		"""The photo IS the roster — a name that doesn't match anyone
		already loaded (or an empty/no roster at all) should create a new
		Student and come back in `matched` with is_new=True, not get
		silently dropped as 'unmatched'."""
		roster = [{'id': self.student_one.id, 'name': 'Amina Juma'}]
		extracted = [{'raw_name': 'Completely Different Person', 'score': 60}]
		response = self._post(extracted, roster)
		self.assertEqual(response.status_code, 200)
		data = response.json()
		self.assertEqual(data['unmatched'], [])
		self.assertEqual(len(data['matched']), 1)
		entry = data['matched'][0]
		self.assertTrue(entry['is_new'])
		self.assertEqual(entry['score'], 60)
		new_student = Student.objects.get(id=entry['id'])
		self.assertEqual(new_student.first_name, 'Completely')
		self.assertEqual(new_student.last_name, 'Person')

	def test_empty_roster_still_creates_students_from_photo(self):
		"""No roster uploaded beforehand is the common case — the photo
		alone should be enough to populate the table."""
		extracted = [
			{'raw_name': 'Amina Juma', 'score': 78},
			{'raw_name': 'Peter Mushi', 'score': 55},
		]
		response = self._post(extracted, roster=[])
		self.assertEqual(response.status_code, 200)
		data = response.json()
		self.assertEqual(data['unmatched'], [])
		self.assertEqual(len(data['matched']), 2)
		self.assertTrue(all(row['is_new'] for row in data['matched']))

	def test_unparseable_name_is_reported_as_unmatched(self):
		roster = [{'id': self.student_one.id, 'name': 'Amina Juma'}]
		extracted = [{'raw_name': 'x', 'score': 60}]
		response = self._post(extracted, roster)
		self.assertEqual(response.status_code, 200)
		data = response.json()
		self.assertEqual(data['matched'], [])
		self.assertEqual(len(data['unmatched']), 1)
		self.assertEqual(data['unmatched'][0]['raw_name'], 'x')

	def test_ocr_error_returns_400(self):
		with patch('results.tasks.extract_scores_from_document', side_effect=ScoreSheetOCRError('Hakuna alama iliyotambulika.')):
			from django.core.files.uploadedfile import SimpleUploadedFile
			photo = SimpleUploadedFile('sheet.jpg', b'fake-bytes', content_type='image/jpeg')
			kickoff = self.client.post(reverse('scoresheet_photo_extract'), {
				'photo': photo,
				'exam_id': self.exam.id,
				'subject_id': self.subject.id,
				'roster': '[]',
			})
			self.assertEqual(kickoff.status_code, 202)
			task_id = kickoff.json()['task_id']
			response = self.client.get(reverse('scoresheet_extract_status', args=[task_id]))
		self.assertEqual(response.status_code, 400)
		self.assertIn('error', response.json())


class DownloadScoresheetNamesPdfTests(TestCase):
	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=1)
		self.subject = Subject.objects.create(name='Mathematics')
		self.student_one = Student.objects.create(first_name='Amina', middle_name='', last_name='Juma', gender='F')
		self.student_two = Student.objects.create(first_name='Peter', middle_name='', last_name='Mushi', gender='M')
		# _resolve_class_roster's fallback tier picks up students via
		# existing ExamResult rows when there's no StoredRoster/FormStudent.
		ExamResult.objects.create(exam=self.exam, student=self.student_one, subject=self.subject, score=50)
		ExamResult.objects.create(exam=self.exam, student=self.student_two, subject=self.subject, score=60)
		self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER)
		self.teacher.subjects.set([self.subject])
		self.client = Client()
		self.client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')

	def test_returns_pdf_containing_registered_student_names(self):
		response = self.client.get(reverse('download_scoresheet_names_pdf'), {
			'exam_id': self.exam.id, 'subject_id': self.subject.id,
		})
		self.assertEqual(response.status_code, 200)
		self.assertEqual(response['Content-Type'], 'application/pdf')
		content = b''.join(response.streaming_content) if response.streaming else response.content
		self.assertTrue(content.startswith(b'%PDF'))

		import io
		import pdfplumber
		with pdfplumber.open(io.BytesIO(content)) as pdf:
			text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
		self.assertIn('Amina Juma', text)
		self.assertIn('Peter Mushi', text)

	def test_rejects_subject_teacher_is_not_assigned_to(self):
		other_subject = Subject.objects.create(name='Physics')
		response = self.client.get(reverse('download_scoresheet_names_pdf'), {
			'exam_id': self.exam.id, 'subject_id': other_subject.id,
		})
		self.assertEqual(response.status_code, 403)


class MarksEntryPreviewPdfTests(TestCase):
	"""The teacher's own review-before-save PDF — posted from the Marks
	Entry table's in-memory state after a photo/PDF scoresheet upload, so
	the teacher can check every row against the paper sheet before hitting
	'Hifadhi & Kagua'."""
	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=1)
		self.subject = Subject.objects.create(name='Mathematics')
		self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER)
		self.teacher.subjects.set([self.subject])
		self.client = Client()
		self.client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')

	def test_renders_pdf_with_scores_absent_and_flagged_rows(self):
		payload = {
			'exam_id': self.exam.id,
			'subject_name': 'Mathematics',
			'rows': [
				{'name': 'Amina Juma', 'score': 78, 'is_absent': False, 'flagged': False},
				{'name': 'Peter Mushi', 'score': None, 'is_absent': True, 'flagged': False},
				{'name': 'Zawadi Rama', 'score': None, 'is_absent': False, 'flagged': True},
			],
		}
		response = self.client.post(
			reverse('marks_entry_preview_pdf'),
			data=json.dumps(payload), content_type='application/json',
		)
		self.assertEqual(response.status_code, 200)
		self.assertEqual(response['Content-Type'], 'application/pdf')
		content = response.content
		self.assertTrue(content.startswith(b'%PDF'))

		import io
		import pdfplumber
		with pdfplumber.open(io.BytesIO(content)) as pdf:
			text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
		self.assertIn('AMINA JUMA', text.upper())
		self.assertIn('78', text)
		self.assertIn('ZAWADI RAMA', text.upper())

	def test_exam_or_subject_name_with_angle_bracket_is_not_silently_dropped(self):
		"""ReportLab's Paragraph treats unescaped '<...>' as markup and
		swallows it silently (no error) -- an exam/subject name containing
		one must still render in full, not vanish from the header."""
		exam = Exam.objects.create(name='Mock <Series 1>', year=2026, form=1)
		response = self.client.post(
			reverse('marks_entry_preview_pdf'),
			data=json.dumps({'exam_id': exam.id, 'subject_name': 'Civics & Ethics', 'rows': []}),
			content_type='application/json',
		)
		self.assertEqual(response.status_code, 200)

		import io
		import pdfplumber
		with pdfplumber.open(io.BytesIO(response.content)) as pdf:
			text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
		self.assertIn('Series 1', text)
		self.assertIn('Civics & Ethics', text)

	def test_missing_exam_returns_404(self):
		response = self.client.post(
			reverse('marks_entry_preview_pdf'),
			data=json.dumps({'exam_id': 999999, 'subject_name': 'Mathematics', 'rows': []}),
			content_type='application/json',
		)
		self.assertEqual(response.status_code, 404)


class StudentResultPdfTests(TestCase):
	"""The downloadable version of the public /matokeo/<token>/ page — no
	login required, same share token a parent already uses to view online."""
	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=2)
		self.subject = Subject.objects.create(name='Mathematics')
		self.student = Student.objects.create(first_name='Amina', middle_name='', last_name='Juma', gender='F')
		ExamResult.objects.create(exam=self.exam, student=self.student, subject=self.subject, score=78)
		self.result = ProcessedResult.objects.create(
			exam=self.exam, student=self.student,
			total_score=78, average_score=78, points=2, position=1, division='I',
		)

	def test_returns_pdf_with_student_name_and_division(self):
		response = self.client.get(reverse('student_result_pdf', args=[self.result.share_token]))
		self.assertEqual(response.status_code, 200)
		self.assertEqual(response['Content-Type'], 'application/pdf')
		content = response.content
		self.assertTrue(content.startswith(b'%PDF'))

		import io
		import pdfplumber
		with pdfplumber.open(io.BytesIO(content)) as pdf:
			text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
		self.assertIn('AMINA JUMA', text.upper())
		self.assertIn('78', text)

	def test_unknown_token_returns_404(self):
		import uuid
		response = self.client.get(reverse('student_result_pdf', args=[uuid.uuid4()]))
		self.assertEqual(response.status_code, 404)


class BulkStudentResultsPdfTests(TestCase):
	"""'Download all students' reports' — merges one result-slip page per
	student into a single PDF, restricted to the Academic Officer."""
	databases = {'default', 'results'}

	def setUp(self):
		self.school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
		self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=2, school=self.school)
		self.subject = Subject.objects.create(name='Mathematics')
		self.student_one = Student.objects.create(first_name='Amina', middle_name='', last_name='Juma', gender='F')
		self.student_two = Student.objects.create(first_name='Peter', middle_name='', last_name='Mushi', gender='M')
		ExamResult.objects.create(exam=self.exam, student=self.student_one, subject=self.subject, score=78)
		ExamResult.objects.create(exam=self.exam, student=self.student_two, subject=self.subject, score=55)
		ProcessedResult.objects.create(exam=self.exam, student=self.student_one, total_score=78, average_score=78, points=2, position=1, division='I')
		ProcessedResult.objects.create(exam=self.exam, student=self.student_two, total_score=55, average_score=55, points=5, position=2, division='III')
		self.academic = TeacherAccount.objects.create(email='academic@example.com', full_name='Academic One', role=TeacherAccount.ROLE_ACADEMIC, school=self.school)
		self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER, school=self.school)

	def test_merges_one_page_per_student_in_position_order(self):
		client = Client()
		client.force_login(self.academic, backend='results.backends.ResultsAuthBackend')
		response = client.get(reverse('generate_bulk_student_results_pdf', args=[self.exam.id]))
		self.assertEqual(response.status_code, 200)
		self.assertEqual(response['Content-Type'], 'application/pdf')
		content = response.content
		self.assertTrue(content.startswith(b'%PDF'))

		import io
		import pdfplumber
		with pdfplumber.open(io.BytesIO(content)) as pdf:
			self.assertEqual(len(pdf.pages), 2)
			text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
		self.assertIn('AMINA JUMA', text.upper())
		self.assertIn('PETER MUSHI', text.upper())

	def test_rejects_non_academic_teacher(self):
		client = Client()
		client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')
		response = client.get(reverse('generate_bulk_student_results_pdf', args=[self.exam.id]))
		self.assertEqual(response.status_code, 403)

	def test_no_processed_results_returns_404(self):
		empty_exam = Exam.objects.create(name='Empty Exam', year=2026, form=1, school=self.school)
		client = Client()
		client.force_login(self.academic, backend='results.backends.ResultsAuthBackend')
		response = client.get(reverse('generate_bulk_student_results_pdf', args=[empty_exam.id]))
		self.assertEqual(response.status_code, 404)


class BulkScoresheetPreviewPdfTests(TestCase):
    """Review-before-save PDF: the academic officer's bulk-upload review
    screen posts the CURRENT (possibly hand-corrected) table state here and
    gets back a PDF laid out like the paper scoresheet, so they can check
    every row landed on the right student before anything is saved."""
    databases = {'default', 'results'}

    def setUp(self):
        self.school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
        self.exam = Exam.objects.create(name='Midterm 1', year=2026, form=2, school=self.school)
        self.academic = TeacherAccount.objects.create(email='academic@example.com', full_name='Academic One', role=TeacherAccount.ROLE_ACADEMIC, school=self.school)
        self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER, school=self.school)
        self.client = Client()

    def test_renders_pdf_with_scores_absent_and_missing_rows(self):
        self.client.force_login(self.academic, backend='results.backends.ResultsAuthBackend')
        payload = {
            'subject_name': 'Mathematics',
            'rows': [
                {'name': 'Amina Juma', 'score': 78, 'is_absent': False, 'status': 'matched'},
                {'name': 'Peter Mushi', 'score': None, 'is_absent': True, 'status': 'matched'},
                {'name': 'Zawadi Rama', 'score': None, 'is_absent': False, 'status': 'missing'},
            ],
        }
        response = self.client.post(
            reverse('bulk_scoresheet_preview_pdf', args=[self.exam.id]),
            data=json.dumps(payload), content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        content = response.content
        self.assertTrue(content.startswith(b'%PDF'))

        import io
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        self.assertIn('AMINA JUMA', text.upper())
        self.assertIn('78', text)
        self.assertIn('PETER MUSHI', text.upper())
        self.assertIn('ZAWADI RAMA', text.upper())

    def test_rejects_non_academic_teacher(self):
        self.client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')
        response = self.client.post(
            reverse('bulk_scoresheet_preview_pdf', args=[self.exam.id]),
            data=json.dumps({'subject_name': 'Mathematics', 'rows': []}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 403)

    def test_bad_json_returns_400(self):
        self.client.force_login(self.academic, backend='results.backends.ResultsAuthBackend')
        response = self.client.post(
            reverse('bulk_scoresheet_preview_pdf', args=[self.exam.id]),
            data='not json', content_type='application/json',
        )
        self.assertEqual(response.status_code, 400)


class MergeDuplicateHistorySubjectsMigrationTests(TestCase):
	"""Data migration 0035: production had ended up with several
	differently-cased/typo'd 'Historia ya Tanzania na Maadili' Subject rows
	(created by an old path that bypassed normalize_subject_name), each
	with a blank `code` -- so on the results PDF both History and Historia
	fell back to the same 4-letter truncation ("HIST") and were
	indistinguishable. This exercises the actual migration function
	against real data, including a genuine FK conflict, to make sure the
	merge is safe to run against production."""
	databases = {'default', 'results'}

	def _run_migration(self):
		import importlib
		module = importlib.import_module('results.migrations.0035_merge_duplicate_history_subjects')
		from django.apps import apps
		module.merge_and_backfill(apps, None)

	def test_merges_variants_repoints_fks_and_backfills_codes(self):
		history = Subject.objects.create(name='History')
		canonical = Subject.objects.create(name='Historia ya Tanzania na Maadili')
		variant_lower_t = Subject.objects.create(name='Historia ya tanzania na Maadili')
		variant_upper = Subject.objects.create(name='HISTORIA YA TANZANIA NA MAADILI')
		variant_typo = Subject.objects.create(name='Historian ya Tanzania na maadili')

		school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
		exam = Exam.objects.create(name='Midterm 1', year=2026, form=1, school=school)
		student_a = Student.objects.create(first_name='Amina', last_name='Juma', gender='F')
		student_b = Student.objects.create(first_name='Peter', last_name='Mushi', gender='M')

		# Straightforward re-point: only exists under a variant.
		ExamResult.objects.create(exam=exam, student=student_a, subject=variant_lower_t, score=60)
		# Genuine conflict: student_b already has a result under the
		# CANONICAL subject for this exam -- the variant's row must be
		# dropped, not crash the migration with a unique-constraint error.
		ExamResult.objects.create(exam=exam, student=student_b, subject=canonical, score=70)
		ExamResult.objects.create(exam=exam, student=student_b, subject=variant_upper, score=99)

		SchoolSubject.objects.create(school=school, subject=variant_typo)

		self._run_migration()

		# Only the canonical Historia row (and plain History) survive.
		remaining = set(Subject.objects.filter(name__icontains='hist').values_list('name', flat=True))
		self.assertEqual(remaining, {'History', 'Historia ya Tanzania na Maadili'})

		# Codes backfilled so the PDF/Excel can tell them apart.
		history.refresh_from_db()
		canonical.refresh_from_db()
		self.assertEqual(history.code, 'HIST')
		self.assertEqual(canonical.code, 'HIST/M')

		# Straightforward case re-pointed to canonical.
		result_a = ExamResult.objects.get(exam=exam, student=student_a)
		self.assertEqual(result_a.subject_id, canonical.id)
		self.assertEqual(result_a.score, 60)

		# Conflict case: canonical's pre-existing result (70) survives;
		# the variant's conflicting result (99) was dropped, not merged.
		self.assertEqual(ExamResult.objects.filter(exam=exam, student=student_b).count(), 1)
		result_b = ExamResult.objects.get(exam=exam, student=student_b)
		self.assertEqual(result_b.subject_id, canonical.id)
		self.assertEqual(result_b.score, 70)

		# M2M-free single-FK model with no conflict risk re-pointed too.
		self.assertTrue(SchoolSubject.objects.filter(school=school, subject=canonical).exists())

	def test_merges_m2m_subject_references(self):
		canonical = Subject.objects.create(name='Historia ya Tanzania na Maadili')
		variant = Subject.objects.create(name='HISTORIA YA TANZANIA NA MAADILI')
		other_subject = Subject.objects.create(name='Geography')

		school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
		student = FormStudent.objects.create(school=school, form=5, first_name='Amina', last_name='Juma', gender='F')
		student.subjects.set([variant, other_subject])

		teacher = TeacherAccount.objects.create(email='t1@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER)
		teacher.subjects.set([variant])

		self._run_migration()

		self.assertEqual(set(student.subjects.values_list('name', flat=True)), {'Historia ya Tanzania na Maadili', 'Geography'})
		self.assertEqual(set(teacher.subjects.values_list('name', flat=True)), {'Historia ya Tanzania na Maadili'})
		self.assertFalse(Subject.objects.filter(id=variant.id).exists())

	def test_no_duplicates_is_a_safe_no_op(self):
		Subject.objects.create(name='History')
		Subject.objects.create(name='Historia ya Tanzania na Maadili')
		self._run_migration()  # must not raise
		self.assertEqual(Subject.objects.filter(name__icontains='hist').count(), 2)

	def test_no_history_subjects_at_all_creates_nothing(self):
		"""A school that doesn't teach History/Historia must not end up
		with a spurious Subject row just because the migration ran."""
		Subject.objects.create(name='Geography')
		self._run_migration()
		self.assertFalse(Subject.objects.filter(name__icontains='hist').exists())

	def test_promotes_oldest_row_when_no_row_is_exactly_canonical(self):
		"""If the correctly-spelled canonical name doesn't exist at all
		(every row is a case/typo variant), the oldest variant is renamed
		to canonical rather than being discarded."""
		only_variant = Subject.objects.create(name='HISTORIA YA TANZANIA NA MAADILI')
		self._run_migration()
		only_variant.refresh_from_db()
		self.assertEqual(only_variant.name, 'Historia ya Tanzania na Maadili')
		self.assertEqual(only_variant.code, 'HIST/M')


class ClassTimetableServiceTests(TestCase):
	"""The one rule that must never break: a teacher can never be double
	booked at the same time slot across two different classes."""
	databases = {'default', 'results'}

	def setUp(self):
		self.school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
		self.teacher = TeacherAccount.objects.create(email='t1@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER, school=self.school)
		self.math = Subject.objects.create(name='Mathematics')
		self.slot1 = TimeSlot.objects.create(school=self.school, day_of_week=0, order=0, start_time='08:00', end_time='08:40')
		self.slot2 = TimeSlot.objects.create(school=self.school, day_of_week=0, order=1, start_time='08:40', end_time='09:20')

	def test_generate_never_double_books_a_teacher(self):
		# One teacher, two classes, each wanting 2 periods/week — but only
		# 2 slots exist in total, so at most 2 of the 4 needed lessons can
		# ever be placed without the teacher being in two places at once.
		TeachingAssignment.objects.create(school=self.school, form=1, stream='A', subject=self.math, teacher=self.teacher, periods_per_week=2)
		TeachingAssignment.objects.create(school=self.school, form=1, stream='B', subject=self.math, teacher=self.teacher, periods_per_week=2)

		entries, unplaced = generate_class_timetable(self.school)

		teacher_slot_pairs = [(e['teacher_id'], e['time_slot_id']) for e in entries]
		self.assertEqual(len(teacher_slot_pairs), len(set(teacher_slot_pairs)))
		self.assertEqual(len(entries), 2)
		self.assertEqual(sum(u['missing'] for u in unplaced), 2)

	def test_generate_places_double_periods_on_consecutive_slots(self):
		# self.slot1/self.slot2 are back-to-back on Monday (setUp); this
		# Tuesday slot is not adjacent to either of them.
		tuesday_slot = TimeSlot.objects.create(school=self.school, day_of_week=1, order=0, start_time='08:00', end_time='08:40')
		TeachingAssignment.objects.create(
			school=self.school, form=1, stream='A', subject=self.math, teacher=self.teacher,
			periods_per_week=2, double_period=True,
		)
		entries, unplaced = generate_class_timetable(self.school)
		self.assertEqual(unplaced, [])
		slot_ids = {e['time_slot_id'] for e in entries}
		self.assertEqual(slot_ids, {self.slot1.id, self.slot2.id})
		self.assertNotIn(tuesday_slot.id, slot_ids)

	def test_double_period_goes_unplaced_without_a_consecutive_pair(self):
		# One slot per day across two separate days — no two slots are
		# ever adjacent, so a double session can never be seated.
		TimeSlot.objects.all().delete()
		TimeSlot.objects.create(school=self.school, day_of_week=0, order=0, start_time='08:00', end_time='08:40')
		TimeSlot.objects.create(school=self.school, day_of_week=1, order=0, start_time='08:00', end_time='08:40')
		TeachingAssignment.objects.create(
			school=self.school, form=1, stream='A', subject=self.math, teacher=self.teacher,
			periods_per_week=2, double_period=True,
		)
		entries, unplaced = generate_class_timetable(self.school)
		self.assertEqual(entries, [])
		self.assertEqual(sum(u['missing'] for u in unplaced), 2)

	def test_generate_raises_without_time_slots(self):
		TimeSlot.objects.all().delete()
		TeachingAssignment.objects.create(school=self.school, form=1, stream='A', subject=self.math, teacher=self.teacher, periods_per_week=1)
		with self.assertRaises(TimetableConflict):
			generate_class_timetable(self.school)

	def test_generate_raises_without_assignments(self):
		with self.assertRaises(TimetableConflict):
			generate_class_timetable(self.school)

	def test_save_class_timetable_only_replaces_touched_classes(self):
		untouched = ClassTimetableEntry.objects.create(
			school=self.school, form=9, stream='Z', time_slot=self.slot1, subject=self.math, teacher=self.teacher,
		)
		TeachingAssignment.objects.create(school=self.school, form=1, stream='A', subject=self.math, teacher=self.teacher, periods_per_week=1)
		entries, _ = generate_class_timetable(self.school, form_streams=[(1, 'A')])
		saved = save_class_timetable(self.school, entries, form_streams=[(1, 'A')])
		self.assertEqual(saved, 1)
		self.assertTrue(ClassTimetableEntry.objects.filter(school=self.school, form=1, stream='A').exists())
		self.assertTrue(ClassTimetableEntry.objects.filter(id=untouched.id).exists())

	def test_set_single_cell_rejects_double_booking(self):
		ClassTimetableEntry.objects.create(school=self.school, form=1, stream='A', time_slot=self.slot1, subject=self.math, teacher=self.teacher)
		with self.assertRaises(TimetableConflict):
			set_single_cell(self.school, form=1, stream='B', time_slot_id=self.slot1.id, subject_id=self.math.id, teacher_id=self.teacher.id)

	def test_set_single_cell_allows_same_class_update(self):
		entry = ClassTimetableEntry.objects.create(school=self.school, form=1, stream='A', time_slot=self.slot1, subject=self.math, teacher=self.teacher)
		other_subject = Subject.objects.create(name='English')
		set_single_cell(self.school, form=1, stream='A', time_slot_id=self.slot1.id, subject_id=other_subject.id, teacher_id=self.teacher.id)
		entry.refresh_from_db()
		self.assertEqual(entry.subject_id, other_subject.id)


class ClassTimetableViewTests(TestCase):
	databases = {'default', 'results'}

	def setUp(self):
		self.school = School.objects.create(name='Mfano Secondary', region='Dodoma', district='Dodoma')
		self.academic = TeacherAccount.objects.create(email='academic@example.com', full_name='Academic One', role=TeacherAccount.ROLE_ACADEMIC, school=self.school)
		self.teacher = TeacherAccount.objects.create(email='teacher@example.com', full_name='Teacher One', role=TeacherAccount.ROLE_TEACHER, school=self.school)
		self.subject = Subject.objects.create(name='Mathematics')
		self.client_academic = Client()
		self.client_academic.force_login(self.academic, backend='results.backends.ResultsAuthBackend')
		self.client_teacher = Client()
		self.client_teacher.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')

	def test_non_academic_cannot_manage_time_slots(self):
		response = self.client_teacher.get(reverse('time_slot_setup'))
		self.assertEqual(response.status_code, 403)

	def test_academic_can_add_and_delete_time_slot(self):
		response = self.client_academic.post(reverse('time_slot_setup'), {
			'day_of_week': '0', 'order': '0', 'start_time': '08:00', 'end_time': '08:40', 'is_teaching_slot': 'on',
		})
		self.assertEqual(response.status_code, 302)
		slot = TimeSlot.objects.get(school=self.school)
		self.assertTrue(slot.is_teaching_slot)

		response = self.client_academic.post(reverse('time_slot_setup'), {'action': 'delete', 'slot_id': slot.id})
		self.assertEqual(response.status_code, 302)
		self.assertFalse(TimeSlot.objects.filter(id=slot.id).exists())

	def test_academic_can_add_teaching_assignment(self):
		response = self.client_academic.post(reverse('teaching_assignment_manage'), {
			'teacher_id': self.teacher.id, 'subject_id': self.subject.id, 'form': '1', 'stream': 'A', 'periods_per_week': '3',
		})
		self.assertEqual(response.status_code, 302)
		self.assertTrue(TeachingAssignment.objects.filter(school=self.school, teacher=self.teacher, subject=self.subject).exists())

	def test_generate_view_shows_preview_and_save_persists(self):
		TimeSlot.objects.create(school=self.school, day_of_week=0, order=0, start_time='08:00', end_time='08:40')
		TeachingAssignment.objects.create(school=self.school, form=1, stream='A', subject=self.subject, teacher=self.teacher, periods_per_week=1)

		response = self.client_academic.post(reverse('generate_class_timetable'), {'action': 'generate'})
		self.assertEqual(response.status_code, 200)
		self.assertEqual(len(response.context['preview_rows']), 1)

		row = response.context['preview_rows'][0]
		save_response = self.client_academic.post(reverse('save_class_timetable'), {
			'form[]': [str(row['form'])],
			'stream[]': [row['stream']],
			'time_slot_id[]': [str(row['time_slot_id'])],
			'subject_id[]': [str(row['subject_id'])],
			'teacher_id[]': [str(row['teacher_id'])],
		})
		self.assertEqual(save_response.status_code, 302)
		self.assertTrue(ClassTimetableEntry.objects.filter(school=self.school, form=1, stream='A').exists())

	def test_class_timetable_view_renders_for_teacher_and_academic(self):
		slot = TimeSlot.objects.create(school=self.school, day_of_week=0, order=0, start_time='08:00', end_time='08:40')
		ClassTimetableEntry.objects.create(school=self.school, form=1, stream='A', time_slot=slot, subject=self.subject, teacher=self.teacher)
		TeachingAssignment.objects.create(school=self.school, form=1, stream='A', subject=self.subject, teacher=self.teacher, periods_per_week=1)

		for client in (self.client_academic, self.client_teacher):
			response = client.get(reverse('class_timetable_view'))
			self.assertEqual(response.status_code, 200)
			self.assertContains(response, 'Mathematics')

	def test_cell_edit_rejects_conflicting_teacher(self):
		slot = TimeSlot.objects.create(school=self.school, day_of_week=0, order=0, start_time='08:00', end_time='08:40')
		ClassTimetableEntry.objects.create(school=self.school, form=1, stream='A', time_slot=slot, subject=self.subject, teacher=self.teacher)

		response = self.client_academic.post(reverse('class_timetable_cell_edit'), {
			'form': '1', 'stream': 'B', 'time_slot_id': slot.id, 'subject_id': self.subject.id, 'teacher_id': self.teacher.id,
		})
		self.assertEqual(response.status_code, 409)

	def test_default_template_creates_slots_only_when_none_exist(self):
		response = self.client_academic.post(reverse('time_slot_setup'), {'action': 'default_template'})
		self.assertEqual(response.status_code, 302)
		count_after_first = TimeSlot.objects.filter(school=self.school).count()
		self.assertGreater(count_after_first, 0)

		# Second call must not duplicate — refuses instead of stacking a second template.
		response = self.client_academic.post(reverse('time_slot_setup'), {'action': 'default_template'})
		self.assertEqual(response.status_code, 302)
		self.assertEqual(TimeSlot.objects.filter(school=self.school).count(), count_after_first)

	def test_auto_populate_bootstraps_from_teacher_form_assignment(self):
		from .models import TeacherFormAssignment
		TeacherFormAssignment.objects.create(teacher=self.teacher, form=2, subject=self.subject, school=self.school)

		response = self.client_academic.post(reverse('teaching_assignment_manage'), {'action': 'auto_populate'})
		self.assertEqual(response.status_code, 302)
		ta = TeachingAssignment.objects.get(school=self.school, form=2, subject=self.subject)
		self.assertEqual(ta.teacher_id, self.teacher.id)
		self.assertEqual(ta.stream, '')
		self.assertEqual(ta.periods_per_week, 5)

		# Re-running must not clobber a since-edited row.
		ta.periods_per_week = 3
		ta.save(update_fields=['periods_per_week'])
		self.client_academic.post(reverse('teaching_assignment_manage'), {'action': 'auto_populate'})
		ta.refresh_from_db()
		self.assertEqual(ta.periods_per_week, 3)


class HistoriaYaTanzaniaNaMaadiliTests(TestCase):
	"""History and Historia ya Tanzania na Maadili must stay distinct."""

	databases = {'default', 'results'}

	def test_normalize_keeps_the_two_history_subjects_apart(self):
		from .utils import normalize_subject_name
		for raw in ('History', 'HIST', 'hist'):
			self.assertEqual(normalize_subject_name(raw), 'History')
		for raw in ('HIST/M', 'hist-m', 'HISTM', 'Historia ya Tanzania na Maadili',
					'Maadili', 'History of Tanzania and Ethics', 'HTE'):
			self.assertEqual(normalize_subject_name(raw), 'Historia ya Tanzania na Maadili')

	def test_canon_subject_keeps_them_apart(self):
		from .combinations import canon_subject
		self.assertEqual(canon_subject('HIST'), 'History')
		self.assertEqual(canon_subject('HIST/M'), 'Historia ya Tanzania na Maadili')
		self.assertNotEqual(canon_subject('History'), canon_subject('Historia ya Tanzania na Maadili'))

	def test_safe_get_or_create_stamps_short_code(self):
		from .utils import safe_get_or_create_subject
		self.assertEqual(safe_get_or_create_subject('History').code, 'HIST')
		self.assertEqual(
			safe_get_or_create_subject('Historia ya Tanzania na Maadili').code, 'HIST/M')

	def test_code_is_backfilled_on_a_pre_existing_row(self):
		from .utils import safe_get_or_create_subject
		Subject.objects.create(name='Historia ya Tanzania na Maadili')  # no code
		subject = safe_get_or_create_subject('Historia ya Tanzania na Maadili')
		self.assertEqual(subject.code, 'HIST/M')

	def test_safe_get_or_create_normalises_raw_aliases(self):
		from .utils import safe_get_or_create_subject
		a = safe_get_or_create_subject('HIST/M')
		b = safe_get_or_create_subject('Maadili')
		self.assertEqual(a.pk, b.pk)
		self.assertEqual(a.name, 'Historia ya Tanzania na Maadili')

	def test_resolve_subject_columns_splits_two_history_columns(self):
		from .utils import resolve_subject_columns
		# pandas renames a duplicate "HIST" header to "HIST.1"
		resolved = resolve_subject_columns(['GEOG', 'HIST', 'HIST.1'])
		self.assertEqual(resolved, [
			('GEOG', 'Geography'),
			('HIST', 'History'),
			('HIST.1', 'Historia ya Tanzania na Maadili'),
		])

	def test_resolve_subject_columns_leaves_distinct_headers_alone(self):
		from .utils import resolve_subject_columns
		resolved = resolve_subject_columns(['History', 'HIST/M'])
		self.assertEqual(resolved, [
			('History', 'History'),
			('HIST/M', 'Historia ya Tanzania na Maadili'),
		])

	def test_upload_with_two_history_columns_creates_two_subjects(self):
		import io
		from django.core.files.uploadedfile import SimpleUploadedFile
		from .services.upload_processing_service import process_uploaded_results
		exam = Exam.objects.create(name='Mock', year=2026, form=4)
		csv = (
			"First Name,Last Name,Gender,HIST,HIST\n"
			"Asha,Juma,F,80,55\n"
		)
		process_uploaded_results(exam, SimpleUploadedFile(
			'r.csv', csv.encode('utf-8'), content_type='text/csv'))
		names = set(ExamResult.objects.filter(exam=exam).values_list('subject__name', flat=True))
		self.assertEqual(names, {'History', 'Historia ya Tanzania na Maadili'})

	def test_historia_ya_tanzania_never_enters_an_acsee_combination(self):
		from .combinations import detect_acsee_combination
		# Student's real combination is HGK; HIST/M is just an extra subject.
		code, subs = detect_acsee_combination(
			['History', 'Geography', 'Kiswahili', 'Historia ya Tanzania na Maadili'],
			lambda n: 1,
		)
		self.assertEqual(code, 'HGK')
		self.assertNotIn('Historia ya Tanzania na Maadili', subs)


class AcseeSubsidiarySubjectTests(TestCase):
	def test_general_studies_and_bam_are_subsidiary(self):
		from .utils import is_acsee_subsidiary_subject
		for name in (
			'General Studies', 'general studies', 'GS', 'G/Studies', 'G Studies',
			'Basic Applied Mathematics', 'basic applied maths', 'BAM',
			'Applied Mathematics', '  General   Studies  ',
		):
			self.assertTrue(is_acsee_subsidiary_subject(name), name)

	def test_principal_subjects_are_not_subsidiary(self):
		from .utils import is_acsee_subsidiary_subject
		for name in ('Physics', 'Advanced Mathematics', 'History', 'Economics', 'Mathematics', ''):
			self.assertFalse(is_acsee_subsidiary_subject(name), name)


class AcseeCombinationDetectionTests(TestCase):
	def test_canon_subject_normalises_aliases(self):
		from .combinations import canon_subject
		self.assertEqual(canon_subject('Maths'), 'Advanced Mathematics')
		self.assertEqual(canon_subject('  mathematics '), 'Advanced Mathematics')
		self.assertEqual(canon_subject('ENGLISH'), 'English Language')
		self.assertEqual(canon_subject('Literature in English'), 'Literature in English')
		self.assertEqual(canon_subject('Literature'), 'Literature in English')
		self.assertEqual(canon_subject('Phy'), 'Physics')
		self.assertEqual(canon_subject('BAM'), 'Basic Applied Mathematics')
		self.assertEqual(canon_subject('Nutmeg Studies'), 'Nutmeg Studies')

	def test_detects_unambiguous_combination(self):
		from .combinations import detect_acsee_combination
		code, subs = detect_acsee_combination(
			['Physics', 'Chemistry', 'Biology'], lambda n: 1)
		self.assertEqual(code, 'PCB')
		self.assertEqual(set(subs), {'Physics', 'Chemistry', 'Biology'})
		code, _ = detect_acsee_combination(
			['History', 'Geography', 'Kiswahili'], lambda n: 1)
		self.assertEqual(code, 'HGK')

	def test_no_match_returns_none(self):
		from .combinations import detect_acsee_combination
		self.assertIsNone(detect_acsee_combination(
			['Physics', 'History', 'Kiswahili'], lambda n: 1))

	def test_hgl_and_hgli_are_distinguished(self):
		from .combinations import detect_acsee_combination
		code, _ = detect_acsee_combination(
			['History', 'Geography', 'English Language'], lambda n: 1)
		self.assertEqual(code, 'HGL')
		code, subs = detect_acsee_combination(
			['History', 'Geography', 'Literature'], lambda n: 1)
		self.assertEqual(code, 'HGLi')
		self.assertIn('Literature in English', subs)

	def test_ambiguous_match_drops_the_best_extra_subject(self):
		from .combinations import detect_acsee_combination
		# Physics/Chemistry/Biology/Advanced Mathematics satisfies PCB, PCM
		# and CBM. Maths is the student's best subject (1 pt), the rest 5.
		# The conservative pick is the combination totalling the MOST
		# points — PCB — so the strong Maths is left uncounted.
		points = {'Advanced Mathematics': 1, 'Physics': 5, 'Chemistry': 5, 'Biology': 5}
		code, subs = detect_acsee_combination(
			['Physics', 'Chemistry', 'Biology', 'Mathematics'], lambda n: points[n])
		self.assertEqual(code, 'PCB')
		self.assertNotIn('Advanced Mathematics', subs)


class RecomputeAcseeDivisionTests(TestCase):
	"""ACSEE (Form 5-6) division: the student's COMBINATION subjects only,
	extras dropped, and a combination short of 3 padded with F (7 pts)."""

	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Mock ACSEE', year=2026, form=5)
		self._subjects = {}

	def _subject(self, name):
		if name not in self._subjects:
			self._subjects[name] = Subject.objects.create(name=name)
		return self._subjects[name]

	def _student(self, first, last):
		return Student.objects.create(first_name=first, middle_name='', last_name=last, gender='M')

	def _enter(self, student, **scores):
		for subject_name, score in scores.items():
			ExamResult.objects.create(
				exam=self.exam, student=student,
				subject=self._subject(subject_name.replace('_', ' ')), score=score,
			)

	def _processed(self, student):
		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		return ProcessedResult.objects.get(exam=self.exam, student=student)

	def test_general_studies_and_bam_excluded_from_division(self):
		# PCB combination: Physics D(4), Chemistry D(4), Biology E(5) -> 13
		# -> Div III. BAM C(3) is better than Biology but must NOT count.
		student = self._student('Asha', 'Kimaro')
		self._enter(student, Physics=52, Chemistry=55, Biology=45,
			Basic_Applied_Mathematics=65, General_Studies=42)
		result = self._processed(student)
		self.assertEqual(result.points, 13)
		self.assertEqual(result.division, 'III')
		self.assertTrue(result.counted_subjects.startswith('PCB:'), result.counted_subjects)
		self.assertNotIn('Basic Applied Mathematics', result.counted_subjects)
		self.assertNotIn('General Studies', result.counted_subjects)

	def test_extra_fourth_principal_subject_is_not_counted_even_if_better(self):
		# PCB student who also sat Advanced Mathematics and aced it.
		# Combination subjects: Physics C(3), Chemistry C(3), Biology E(5)
		# -> 11 -> Div II. A naive best-3 would take Maths A(1) instead of
		# Biology -> 7 -> Div I, which is wrong.
		student = self._student('Deo', 'Marwa')
		self._enter(student, Physics=62, Chemistry=62, Biology=42,
			Advanced_Mathematics=95, General_Studies=80)
		result = self._processed(student)
		self.assertEqual(result.points, 11)
		self.assertEqual(result.division, 'II')
		self.assertTrue(result.counted_subjects.startswith('PCB:'), result.counted_subjects)
		self.assertNotIn('Advanced Mathematics', result.counted_subjects)

	def test_arts_combination_is_detected(self):
		# HGL: English C(3), History D(4), Geography E(5) -> 12 -> Div II.
		student = self._student('Rehema', 'Kato')
		self._enter(student, History=55, Geography=45, English_Language=62)
		result = self._processed(student)
		self.assertEqual(result.points, 12)
		self.assertEqual(result.division, 'II')
		self.assertTrue(result.counted_subjects.startswith('HGL:'), result.counted_subjects)

	def test_three_subject_combination_with_subsidiaries_is_division_I(self):
		student = self._student('Frank', 'Noel')
		self._enter(student, Physics=85, Chemistry=78, Biology=70,
			General_Studies=55, Basic_Applied_Mathematics=60)
		result = self._processed(student)
		self.assertEqual(result.points, 5)          # A(1) + B(2) + B(2)
		self.assertEqual(result.division, 'I')
		self.assertTrue(result.counted_subjects.startswith('PCB:'), result.counted_subjects)
		self.assertNotIn('General Studies', result.counted_subjects)

	def test_ranking_tiebreak_uses_combination_total_not_general_studies(self):
		# Both are PCB with 3 points. A's combination marks total less than
		# B's, but A also has a huge General Studies mark. Ranking must go
		# by the combination total only -> B first.
		a = self._student('Aled', 'One')
		self._enter(a, Physics=80, Chemistry=80, Biology=80, General_Studies=99)
		b = self._student('Bex', 'Two')
		self._enter(b, Physics=83, Chemistry=83, Biology=83)
		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		a_r = ProcessedResult.objects.get(exam=self.exam, student=a)
		b_r = ProcessedResult.objects.get(exam=self.exam, student=b)
		self.assertEqual((a_r.points, b_r.points), (3, 3))
		self.assertLess(b_r.position, a_r.position)

	def test_student_absent_in_all_subjects_ranks_last(self):
		good = self._student('Good', 'Student')
		self._enter(good, Physics=90, Chemistry=90, Biology=90)
		absent = self._student('Absent', 'Student')
		ExamResult.objects.create(
			exam=self.exam, student=absent, subject=self._subject('Physics'),
			score=None, is_absent=True,
		)
		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		good_r = ProcessedResult.objects.get(exam=self.exam, student=good)
		absent_r = ProcessedResult.objects.get(exam=self.exam, student=absent)
		self.assertEqual(absent_r.division, '0')
		self.assertGreater(absent_r.position, good_r.position)

	def test_single_principal_subject_A_is_padded_to_division_III(self):
		# One A (1 pt) + two missing slots scored F (7) each = 15 -> Div III.
		student = self._student('Baraka', 'Mushi')
		self._enter(student, Physics=85)
		result = self._processed(student)
		self.assertEqual(result.points, 15)
		self.assertEqual(result.division, 'III')

	def test_single_principal_subject_F_is_division_0(self):
		student = self._student('Neema', 'Paul')
		self._enter(student, Physics=20)
		result = self._processed(student)
		self.assertEqual(result.points, 21)
		self.assertEqual(result.division, '0')

	def test_two_principal_subjects_are_padded_once(self):
		# A(1) + B(2) + one F(7) = 10 -> Div II.
		student = self._student('Juma', 'Ally')
		self._enter(student, Physics=85, Chemistry=75)
		result = self._processed(student)
		self.assertEqual(result.points, 10)
		self.assertEqual(result.division, 'II')

	def test_padded_candidate_ranks_below_a_full_three_subject_candidate(self):
		full = self._student('Full', 'Combination')
		self._enter(full, Physics=85, Chemistry=85, Biology=85)  # 3 pts, Div I
		partial = self._student('One', 'Subject')
		self._enter(partial, History=85)  # padded to 15 pts, Div III
		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		full_r = ProcessedResult.objects.get(exam=self.exam, student=full)
		partial_r = ProcessedResult.objects.get(exam=self.exam, student=partial)
		self.assertEqual(full_r.division, 'I')
		self.assertLess(full_r.position, partial_r.position)

	def test_three_full_principals_are_not_padded(self):
		student = self._student('Grace', 'Mena')
		self._enter(student, Physics=62, Chemistry=55, Geography=45)  # C,D,E = 3+4+5
		result = self._processed(student)
		self.assertEqual(result.points, 12)
		self.assertEqual(result.division, 'II')

	def test_only_subsidiary_subjects_falls_back_without_crashing(self):
		student = self._student('Said', 'Omary')
		self._enter(student, General_Studies=65, Basic_Applied_Mathematics=70)
		result = self._processed(student)
		# B(2) + C(3) + one padded F(7) = 12 -> Div II
		self.assertEqual(result.points, 12)
		self.assertEqual(result.division, 'II')


class RecomputeCseeMinimumSubjectRuleTests(TestCase):
	"""The existing CSEE (<7 subjects) rule must be untouched by the
	ACSEE changes."""

	databases = {'default', 'results'}

	def setUp(self):
		self.exam = Exam.objects.create(name='Terminal', year=2026, form=4)

	def _run(self, **scores):
		student = Student.objects.create(first_name='Test', middle_name='', last_name='Pupil', gender='F')
		for name, score in scores.items():
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=self.exam, student=student, subject=subject, score=score)
		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		return ProcessedResult.objects.get(exam=self.exam, student=student)

	def test_fewer_than_seven_with_a_pass_is_capped_at_division_IV(self):
		result = self._run(Physics=90, Chemistry=88, Biology=85)  # 3x A
		self.assertEqual(result.division, 'IV')

	def test_fewer_than_seven_without_a_real_pass_is_division_0(self):
		result = self._run(Physics=20, Chemistry=25)  # 2x F
		self.assertEqual(result.division, '0')

	def test_fewer_subjects_straight_As_does_not_outrank_full_subject_straight_As(self):
		"""Raw points alone would rank a 4-subject straight-A student (4
		points) above a 7-subject straight-A student (7 points), even
		though NECTA's own rule caps the 4-subject student at Division IV.
		Position must reflect division first so 'top performers' are the
		students who are actually strongest, not just the ones who sat
		fewer exams."""
		partial = Student.objects.create(first_name='Partial', last_name='Subjects', gender='F')
		for name, score in [('Physics', 90), ('Chemistry', 88), ('Biology', 85), ('Mathematics', 92)]:
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=self.exam, student=partial, subject=subject, score=score)

		full = Student.objects.create(first_name='Full', last_name='Subjects', gender='M')
		for name, score in [
			('Physics', 90), ('Chemistry', 88), ('Biology', 85), ('Mathematics', 92),
			('English', 91), ('Kiswahili', 90), ('Civics', 93),
		]:
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=self.exam, student=full, subject=subject, score=score)

		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		partial_r = ProcessedResult.objects.get(exam=self.exam, student=partial)
		full_r = ProcessedResult.objects.get(exam=self.exam, student=full)

		self.assertEqual(partial_r.points, 4)
		self.assertEqual(partial_r.division, 'IV')
		self.assertEqual(full_r.points, 7)
		self.assertEqual(full_r.division, 'I')
		self.assertLess(full_r.position, partial_r.position)

	def test_within_same_division_lower_points_still_ranks_first(self):
		better = Student.objects.create(first_name='Better', last_name='Points', gender='F')
		worse = Student.objects.create(first_name='Worse', last_name='Points', gender='M')
		subjects = ['Physics', 'Chemistry', 'Biology', 'Mathematics', 'English', 'Kiswahili', 'Civics']
		# Both Division I (points <= 17), but `better` scores higher overall.
		for name, score in zip(subjects, [90, 90, 90, 90, 90, 90, 90]):  # 7x A = 7 pts
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=self.exam, student=better, subject=subject, score=score)
		for name, score in zip(subjects, [70, 70, 70, 70, 70, 70, 70]):  # 7x B = 14 pts
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=self.exam, student=worse, subject=subject, score=score)

		from .services.upload_processing_service import recompute_processed_results_for_exam
		recompute_processed_results_for_exam(self.exam)
		better_r = ProcessedResult.objects.get(exam=self.exam, student=better)
		worse_r = ProcessedResult.objects.get(exam=self.exam, student=worse)

		self.assertEqual((better_r.division, worse_r.division), ('I', 'I'))
		self.assertLess(better_r.position, worse_r.position)


class RecomputeCseePositionRankingMigrationTests(TestCase):
	"""Data migration 0036 backfills the division-first ranking fix onto
	every existing CSEE exam's already-cached ProcessedResult rows."""
	databases = {'default', 'results'}

	def _run_migration(self):
		import importlib
		module = importlib.import_module('results.migrations.0036_recompute_csee_position_ranking')
		from django.apps import apps
		module.recompute_csee_exams(apps, None)

	def test_recomputes_stale_position_for_existing_csee_exam(self):
		exam = Exam.objects.create(name='Terminal', year=2026, form=4)
		partial = Student.objects.create(first_name='Partial', last_name='Subjects', gender='F')
		full = Student.objects.create(first_name='Full', last_name='Subjects', gender='M')
		for name, score in [('Physics', 90), ('Chemistry', 88), ('Biology', 85), ('Mathematics', 92)]:
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=exam, student=partial, subject=subject, score=score)
		for name, score in [
			('Physics', 90), ('Chemistry', 88), ('Biology', 85), ('Mathematics', 92),
			('English', 91), ('Kiswahili', 90), ('Civics', 93),
		]:
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=exam, student=full, subject=subject, score=score)

		# Simulate the OLD (buggy) cached ranking: partial (fewer subjects,
		# lower raw points) was stored as position 1.
		ProcessedResult.objects.create(exam=exam, student=partial, total_score=355, average_score=88.75, points=4, position=1, division='IV')
		ProcessedResult.objects.create(exam=exam, student=full, total_score=629, average_score=89.86, points=7, position=2, division='I')

		self._run_migration()

		partial_r = ProcessedResult.objects.get(exam=exam, student=partial)
		full_r = ProcessedResult.objects.get(exam=exam, student=full)
		self.assertLess(full_r.position, partial_r.position)

	def test_non_csee_exams_are_left_alone(self):
		"""ACSEE (form 5-6) exams aren't touched by this backfill."""
		exam = Exam.objects.create(name='Mock ACSEE', year=2026, form=5)
		student = Student.objects.create(first_name='Asha', last_name='Kimaro', gender='F')
		for name, score in [('Physics', 85), ('Chemistry', 85), ('Biology', 85)]:
			subject, _ = Subject.objects.get_or_create(name=name)
			ExamResult.objects.create(exam=exam, student=student, subject=subject, score=score)
		self._run_migration()  # must not raise on an ACSEE exam
		self.assertFalse(ProcessedResult.objects.filter(exam=exam).exists())


def _build_docx_bytes(*, table_rows=None, paragraph_lines=None):
	from io import BytesIO
	from docx import Document

	doc = Document()
	if table_rows:
		table = doc.add_table(rows=0, cols=len(table_rows[0]))
		for row in table_rows:
			cells = table.add_row().cells
			for i, value in enumerate(row):
				cells[i].text = value
	if paragraph_lines:
		for line in paragraph_lines:
			doc.add_paragraph(line)
	buf = BytesIO()
	doc.save(buf)
	return buf.getvalue()


class DocxRosterUploadTests(TestCase):
	"""Roster upload also accepts a .docx (Word) class list -- a table
	first, falling back to one student per paragraph line, sharing the
	exact same row/line parsing as the PDF roster path."""

	databases = {'default', 'results'}

	def _docx_file(self, name='roster.docx', **kwargs):
		from django.core.files.uploadedfile import SimpleUploadedFile
		content = _build_docx_bytes(**kwargs)
		return SimpleUploadedFile(
			name, content,
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
		)

	def test_parses_students_from_a_docx_table(self):
		from .views import _parse_docx_roster
		docx_file = self._docx_file(table_rows=[
			['Jina la Mwanafunzi', 'Jinsia'],
			['Amina Juma', 'F'],
			['Peter Mushi', 'M'],
		])
		collected = []
		_parse_docx_roster(docx_file, on_student=lambda f, m, l, g: collected.append((f, m, l, g)))
		names = {(row[0], row[2]) for row in collected}
		self.assertEqual(names, {('Amina', 'Juma'), ('Peter', 'Mushi')})

	def test_parses_students_from_docx_paragraph_lines_when_no_table(self):
		from .views import _parse_docx_roster
		docx_file = self._docx_file(paragraph_lines=[
			'Halima Ally Mohamed F',
			'John Peter Komba M',
		])
		collected = []
		_parse_docx_roster(docx_file, on_student=lambda f, m, l, g: collected.append((f, m, l, g)))
		self.assertEqual(len(collected), 2)
		self.assertIn(('Halima', 'Ally', 'Mohamed', 'F'), collected)
		self.assertIn(('John', 'Peter', 'Komba', 'M'), collected)

	def test_collect_roster_rows_dispatches_docx_by_flag(self):
		from .views import _collect_roster_rows
		docx_file = self._docx_file(paragraph_lines=['Amina Juma F'])
		rows = _collect_roster_rows(docx_file, is_pdf=False, is_docx=True)
		self.assertEqual(rows, [('Amina', '', 'Juma', 'F')])

	def test_upload_roster_view_accepts_docx(self):
		teacher = TeacherAccount.objects.create(email='t2@example.com', full_name='Teacher Two', role=TeacherAccount.ROLE_TEACHER)
		client = Client()
		client.force_login(teacher, backend='results.backends.ResultsAuthBackend')
		docx_file = self._docx_file(paragraph_lines=['Amina Juma F', 'Peter Mushi M'])
		response = client.post(reverse('upload_roster'), {'file': docx_file})
		self.assertEqual(response.status_code, 200)
		data = response.json()
		names = {s['name'] for s in data.get('students', [])}
		self.assertEqual(names, {'Amina Juma', 'Peter Mushi'})


class MarksEntryAddStudentTests(TestCase):
	"""A student added inline on the Marks Entry page must stick — the
	review page filters marks to the roster, so the new student has to be
	written into the teacher's StoredRoster or the row vanishes."""

	databases = {'default', 'results'}

	def setUp(self):
		from .models import FormStudent, StoredRoster
		self.StoredRoster = StoredRoster
		self.school = School.objects.create(name='Mfano Sekondari', region='Dodoma', district='Dodoma')
		self.exam = Exam.objects.create(name='Terminal', year=2026, form=4, school=self.school)
		self.subject = Subject.objects.create(name='History')
		self.teacher = TeacherAccount.objects.create(
			email='t@example.com', full_name='Mwalimu', role=TeacherAccount.ROLE_TEACHER,
			school=self.school)
		self.teacher.subjects.set([self.subject])
		FormStudent.objects.create(
			school=self.school, form=4, first_name='Asha', last_name='Kimaro', gender='F')
		self.client = Client()
		self.client.force_login(self.teacher, backend='results.backends.ResultsAuthBackend')

	def _add(self, first, last):
		return self.client.post(
			reverse('marks_entry_add_student'),
			data=json.dumps({
				'exam_id': self.exam.id, 'subject_id': self.subject.id,
				'first_name': first, 'last_name': last, 'gender': 'M',
			}),
			content_type='application/json',
		)

	def test_added_student_is_written_into_the_stored_roster(self):
		resp = self._add('Baraka', 'Mushi')
		self.assertEqual(resp.status_code, 200)
		new_id = resp.json()['id']

		stored = self.StoredRoster.objects.get(
			teacher=self.teacher, exam=self.exam, subject=self.subject)
		names = {s['name'] for s in stored.students}
		ids = {s['id'] for s in stored.students}
		self.assertIn(new_id, ids)                       # the new student
		self.assertIn('Asha Kimaro', names)              # seeded from FormStudent
		self.assertEqual(stored.student_count, len(stored.students))

	def test_added_student_survives_into_the_review_page(self):
		new_id = self._add('Baraka', 'Mushi').json()['id']

		save = self.client.post(
			reverse('marks_entry_save'),
			data=json.dumps({
				'exam_id': self.exam.id, 'subject_id': self.subject.id,
				'entries': [{'student_id': new_id, 'score': 72}],
			}),
			content_type='application/json',
		)
		self.assertEqual(save.status_code, 200)

		review = self.client.get(
			reverse('marks_entry') + f'?exam={self.exam.id}&subject={self.subject.id}&review=1')
		self.assertContains(review, 'Baraka')

	def test_adding_twice_does_not_duplicate_the_roster_row(self):
		first_id = self._add('Baraka', 'Mushi').json()['id']
		self._add('Baraka', 'Mushi')
		stored = self.StoredRoster.objects.get(
			teacher=self.teacher, exam=self.exam, subject=self.subject)
		self.assertEqual([s['id'] for s in stored.students].count(first_id), 1)


class CentreCountedSubjectsTests(TestCase):
	"""The class-results PDF 'Subjects Counted' / centre GPA divisor."""

	def _counted(self, form, *counted_subjects):
		from types import SimpleNamespace
		from .services.pdf_export_service import _centre_counted_subjects
		rows = [SimpleNamespace(counted_subjects=cs) for cs in counted_subjects]
		return _centre_counted_subjects(rows, form)

	def test_olevel_falls_back_to_seven_when_rows_have_no_counted_subjects(self):
		self.assertEqual(self._counted(4, '', None), 7)
		self.assertEqual(self._counted(4), 7)          # no results at all

	def test_alevel_falls_back_to_three(self):
		self.assertEqual(self._counted(5, ''), 3)

	def test_uses_the_widest_real_count_when_present(self):
		# one partial candidate (4 subjects), one full (7)
		self.assertEqual(self._counted(
			4,
			'Maths, Eng, Kisw, Bio',
			'Maths, Eng, Kisw, Bio, Chem, Phys, Geo',
		), 7)

	def test_alevel_combination_prefix_still_counts_three(self):
		self.assertEqual(self._counted(5, 'PCB: Physics, Chemistry, Biology'), 3)
