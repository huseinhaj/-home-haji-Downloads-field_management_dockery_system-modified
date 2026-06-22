from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0A, 0x1F, 0x44)   # dark navy
GOLD  = RGBColor(0xC8, 0xA0, 0x32)   # gold
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xF5, 0xF7, 0xFA)
GREY  = RGBColor(0x55, 0x55, 0x55)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val', 'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz', '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '0A1F44'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading_para(doc, text, size=13, bold=True, color=NAVY, center=False, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def body_para(doc, text, size=10.5, color=GREY, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return p

def divider(doc, color_hex='C8A032'):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr= OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# HEADER BANNER TABLE
# ══════════════════════════════════════════════════════════════════════════════
hdr = doc.add_table(rows=1, cols=1)
hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = hdr.cell(0, 0)
set_cell_bg(cell, '0A1F44')
set_cell_border(cell, val='none', sz='0', color='0A1F44')

for line, sz, bold in [
    ('CHUO KIKUU CHA DODOMA (UDOM)', 15, True),
    ('COLLEGE OF INFORMATICS AND VIRTUAL EDUCATION (CIVE)', 11, True),
    ('Department of Computer Science and Engineering (CSE)', 10, False),
    ('CIVE DAY 2026 — PROJECT PRESENTATION', 10, False),
]:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(line)
    run.bold           = bold
    run.font.size      = Pt(sz)
    run.font.color.rgb = WHITE if not bold else GOLD
    if line.startswith('CIVE DAY'):
        run.font.color.rgb = WHITE

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DEPARTMENT & PROJECT INFO
# ══════════════════════════════════════════════════════════════════════════════
heading_para(doc, '1. DEPARTMENT & PROJECT IDENTIFICATION', size=12, color=NAVY, space_before=6)
divider(doc)

info_tbl = doc.add_table(rows=4, cols=2)
info_tbl.style = 'Table Grid'
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

labels = ['Department', 'Project Title', 'Team Members', 'Supervisor']
values = [
    'Computer Science and Engineering (CSE)',
    'IMS — Web-Based Internship (Field Practice) Management System\nfor Student Teachers in Tanzania',
    '1. Haji Hamisi Huseni\n2. Hamisi Selemani',
    'Mr. Barongo',
]

for i, (lbl, val) in enumerate(zip(labels, values)):
    row = info_tbl.rows[i]
    # label cell
    lc = row.cells[0]
    set_cell_bg(lc, '0A1F44')
    set_cell_border(lc, color='C8A032')
    lc.width = Cm(4)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lrun = lp.add_run(lbl.upper())
    lrun.bold = True
    lrun.font.size = Pt(9)
    lrun.font.color.rgb = GOLD
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # value cell
    vc = row.cells[1]
    set_cell_bg(vc, 'F5F7FA')
    set_cell_border(vc, color='0A1F44')
    vp = vc.paragraphs[0]
    vrun = vp.add_run(val)
    vrun.font.size = Pt(10)
    vrun.font.color.rgb = NAVY
    vrun.bold = (i in [1])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading_para(doc, '2. INTRODUCTION', size=12, color=NAVY, space_before=8)
divider(doc)

heading_para(doc, 'a. Background', size=10.5, color=GOLD, bold=True, space_before=6, space_after=2)
body_para(doc,
    'In Tanzania, student teachers from universities and colleges are required to complete practical '
    'field training (Mafunzo ya Vitendo) in secondary and primary schools across the country. '
    'This involves placing students in schools, assigning supervisors (assessors), monitoring '
    'teaching activities through daily logbooks, and issuing completion certificates at the end '
    'of the internship period. Coordination spans multiple stakeholders: students, school assessors, '
    'head teachers, District Education Officers (DEOs), and university administration. '
    'Traditionally, this entire process has been handled manually — through paper logbooks, '
    'physical assessment forms, and in-person coordination — leading to significant inefficiencies.'
)

heading_para(doc, 'b. Problem Statement', size=10.5, color=GOLD, bold=True, space_before=6, space_after=2)
body_para(doc, 'The manual management of student teacher internships creates the following critical challenges:')
problems = [
    'No centralized system to track hundreds of student teachers across regions and districts.',
    'Paper logbooks are prone to loss, forgery, and inconsistency in record-keeping.',
    'Assessors have no efficient digital tool to record evaluations and observations.',
    'DEOs and head teachers lack real-time visibility into student progress.',
    'Certificate issuance is delayed due to multi-level manual approval processes.',
    'Communication between university admin, schools, and students is slow and unreliable.',
]
for prob in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(prob)
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

heading_para(doc, 'c. Main Objective', size=10.5, color=GOLD, bold=True, space_before=6, space_after=2)
body_para(doc,
    'To design and develop a web-based Internship Management System (IMS) that fully digitizes '
    'and automates the field practice workflow — from student registration and school placement '
    'to logbook management, assessor evaluation, and certificate issuance — providing real-time '
    'monitoring and transparency for all stakeholders.'
)

heading_para(doc, 'Specific Objectives:', size=10, color=NAVY, bold=True, space_before=4, space_after=2)
objectives = [
    'Automate student registration, school selection, and subject assignment.',
    'Enable digital logbook submission with assessor feedback and remarks.',
    'Provide assessors with tools to evaluate and grade student performance.',
    'Give DEOs and head teachers a dashboard for district-wide progress monitoring.',
    'Generate tamper-proof, digitally-signed completion certificates and official letters.',
    'Integrate AI tools to assist students in creating Schemes of Work and Lesson Plans.',
]
for obj in objectives:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(obj)
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE CORE INNOVATION
# ══════════════════════════════════════════════════════════════════════════════
heading_para(doc, '3. THE CORE INNOVATION', size=12, color=NAVY, space_before=10)
divider(doc)

# ── System Architecture Diagram (table-based) ─────────────────────────────────
heading_para(doc, 'a. System Architecture Diagram', size=10.5, color=GOLD, bold=True, space_before=6, space_after=4)

arch = doc.add_table(rows=7, cols=5)
arch.alignment = WD_TABLE_ALIGNMENT.CENTER

def arch_cell(cell, text, bg='0A1F44', fg=WHITE, sz=8, bold=True, center=True):
    set_cell_bg(cell, bg)
    set_cell_border(cell, color='C8A032', sz='4')
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz)
    r.font.color.rgb = fg

# Row 0 — User layer
arch.cell(0,0).merge(arch.cell(0,4))
arch_cell(arch.cell(0,0),
    '[ WEB BROWSER — HTTPS ]\nStudent Teacher  |  Assessor  |  Admin  |  DEO  |  Head Teacher',
    bg='0A1F44', fg=GOLD, sz=8.5)

# Row 1 — arrow
arch.cell(1,0).merge(arch.cell(1,4))
arch_cell(arch.cell(1,0), '▼  Django Web Application (Python 3.12 / Django 5.2)  ▼',
    bg='C8A032', fg=WHITE, sz=8.5, bold=True)

# Row 2 — modules row 1
for col, (txt, bg) in enumerate([
    ('AUTH MODULE\nLogin / Register\nPassword Reset', '0D2A5E'),
    ('STUDENT MODULE\nSchool Selection\nLogbook / Reports', '0D2A5E'),
    ('ASSESSOR MODULE\nEvaluations\nRemarks / Grades', '0D2A5E'),
    ('BOARD / DEO\nDashboard\nApprovals', '0D2A5E'),
    ('ADMIN MODULE\nUser Mgmt\nAssignments', '0D2A5E'),
]):
    arch_cell(arch.cell(2,col), txt, bg=bg, fg=WHITE, sz=7.5)

# Row 3 — modules row 2
for col, (txt, bg) in enumerate([
    ('AI TOOLS\nScheme of Work\nLesson Plans', '152B5C'),
    ('PDF ENGINE\nCertificates\nLogbook PDFs', '152B5C'),
    ('SMS / EMAIL\nAfrica\'s Talking\nGmail SMTP', '152B5C'),
    ('CACHING\nRedis\nSession Store', '152B5C'),
    ('REST APIs\nAJAX Endpoints\nJSON Responses', '152B5C'),
]):
    arch_cell(arch.cell(3,col), txt, bg=bg, fg=GOLD, sz=7.5)

# Row 4 — arrow down to db
arch.cell(4,0).merge(arch.cell(4,4))
arch_cell(arch.cell(4,0), '▼  Data Layer  ▼', bg='C8A032', fg=WHITE, sz=8)

# Row 5 — databases + external
arch.cell(5,0).merge(arch.cell(5,1))
arch_cell(arch.cell(5,0),
    'PostgreSQL 16\nDatabase\n(Schools / Students /\nAssessments / Logbooks)',
    bg='0A1F44', fg=WHITE, sz=7.5)
arch.cell(5,2).merge(arch.cell(5,2))
arch_cell(arch.cell(5,2),
    'Docker\nContainer\n(Web + DB\nIsolated)',
    bg='1A3A6B', fg=GOLD, sz=7.5)
arch.cell(5,3).merge(arch.cell(5,4))
arch_cell(arch.cell(5,3),
    'External Services\n• OpenAI / Google Gemini (AI)\n• Gmail SMTP (Email)\n• Africa\'s Talking (SMS)',
    bg='0A1F44', fg=WHITE, sz=7.5)

# Row 6 — legend
arch.cell(6,0).merge(arch.cell(6,4))
arch_cell(arch.cell(6,0),
    'LEGEND:  Navy = Core App Modules  |  Gold = Support Layers  |  Dark Blue = Data & External',
    bg='F5F7FA', fg=GREY, sz=7.5, bold=False)
for c in arch.cell(6,0).paragraphs:
    for r in c.runs:
        r.font.color.rgb = GREY

doc.add_paragraph()

# ── Methodology ───────────────────────────────────────────────────────────────
heading_para(doc, 'b. Methodology', size=10.5, color=GOLD, bold=True, space_before=6, space_after=4)

meth_tbl = doc.add_table(rows=8, cols=2)
meth_tbl.style = 'Table Grid'
meth_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

meth_rows = [
    ('Programming Language', 'Python 3.12'),
    ('Web Framework',        'Django 5.2 (MVT Architecture)'),
    ('Database',             'PostgreSQL 16 (via Docker)'),
    ('AI Integration',       'OpenAI GPT-4 & Google Gemini — AI-powered Scheme of Work & Lesson Plan generation'),
    ('PDF Generation',       'ReportLab — Certificates, Logbooks, Official Letters with watermarks & signatures'),
    ('Containerization',     'Docker & Docker Compose — isolated, reproducible deployment'),
    ('Frontend',             'HTML5, CSS3, JavaScript — Responsive multi-role web interface'),
    ('Communication',        'Gmail SMTP (email credentials) · Africa\'s Talking API (SMS to head teachers)'),
]

for i, (label, value) in enumerate(meth_rows):
    row = meth_tbl.rows[i]
    lc = row.cells[0]
    vc = row.cells[1]
    set_cell_bg(lc, '0A1F44')
    set_cell_border(lc, color='C8A032')
    lc.width = Cm(4.5)
    lp = lc.paragraphs[0]
    lrun = lp.add_run(label)
    lrun.bold = True
    lrun.font.size = Pt(9)
    lrun.font.color.rgb = GOLD
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    set_cell_bg(vc, 'F5F7FA')
    set_cell_border(vc, color='0A1F44')
    vp = vc.paragraphs[0]
    vrun = vp.add_run(value)
    vrun.font.size = Pt(9.5)
    vrun.font.color.rgb = NAVY

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
divider(doc, 'C8A032')
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(4)
fr = fp.add_run('CIVE DAY 2026  ·  UDOM  ·  CSE Department  ·  NI ELIMU TU ✍  ·  MUNGU IBARIKI UDOM')
fr.font.size = Pt(8.5)
fr.font.color.rgb = NAVY
fr.bold = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/haji/Downloads/field_management_dockery_system-modifiedied/CIVE_DAY_2026_IMS_Report.docx'
doc.save(out)
print(f'SAVED: {out}')
